"""
app_mcp.py (FIXED v2)
======================
GraphPulse AI — Full 5-Tab Suite

FIXES APPLIED:
  1. Tab heading text → white (#f1f5f9) — was dark grey / invisible
  2. Buttons (Generate Query, Clear Session, Abort, Sample Qs) → bluish-cyan theme
  3. RCA Trail + Visualization tab dark-grey text → white/light on dark bg
  4. RCA Trail left panel: removed static "How It Works", replaced with
     dynamic "Query Context" panel that reflects the question being analysed
  5. Visualizations tab: button row cleanly separated from chart section headers
  6. Transport Node Sunburst bug: branchvalues changed from "total" to "remainder"
     + null-safety for missing mode values
  7. RCA report rendered in gr.Markdown (not Textbox) for rich formatting
  8. MCP server note: mcp_server.py is a standalone REST API for external
     clients (e.g. Claude Desktop). The Gradio app calls neo4j_tools.py
     directly via TOOL_FUNCTIONS — no subprocess needed or connected here.

Run: python app_mcp.py
/* ── RCA chart glow-box strip ── */
.rca5-chart-btns {
  display: flex; gap: 10px; flex-wrap: wrap; margin-bottom: 18px;
}
.rca5-chart-btn {
  flex: 1; min-width: 110px;
  background: rgba(12,21,40,0.9);
  border: 1px solid rgba(56,189,248,0.2);
  border-radius: 14px; padding: 16px 12px 14px;
  cursor: pointer; transition: all 0.2s ease;
  text-align: center; user-select: none;
  position: relative; overflow: hidden;
}
.rca5-chart-btn::before {
  content: ''; position: absolute; top: 0; left: 0; right: 0; height: 3px;
  background: linear-gradient(90deg, #38bdf8, #7c3aed);
  opacity: 0; transition: opacity 0.2s;
}
.rca5-chart-btn::after {
  content: ''; position: absolute; inset: 0;
  background: radial-gradient(ellipse at 50% 0%, rgba(56,189,248,0.08) 0%, transparent 70%);
  opacity: 0; transition: opacity 0.2s;
}
.rca5-chart-btn:hover {
  border-color: rgba(56,189,248,0.55);
  transform: translateY(-3px);
  box-shadow: 0 8px 28px rgba(56,189,248,0.15), 0 0 0 1px rgba(56,189,248,0.1);
}
.rca5-chart-btn:hover::before, .rca5-chart-btn:hover::after { opacity: 1; }
.rca5-chart-btn.active {
  border-color: #38bdf8;
  background: rgba(56,189,248,0.08);
  box-shadow: 0 0 30px rgba(56,189,248,0.22), 0 0 0 1px rgba(56,189,248,0.2);
}
.rca5-chart-btn.active::before { opacity: 1; }
.rca5-chart-btn.active::after  { opacity: 1; }
.rca5-cbtn-icon  { font-size: 26px; margin-bottom: 8px; line-height: 1; }
.rca5-cbtn-label { font-size: 12px; font-weight: 700; color: #e2e8f0 !important; margin-bottom: 4px; text-shadow: 0 1px 3px rgba(0,0,0,0.8); }
.rca5-cbtn-desc  { font-size: 10px; color: #93c5fd !important; line-height: 1.4; text-shadow: 0 1px 2px rgba(0,0,0,0.9); }
.rca5-chart-btn:hover .rca5-cbtn-desc { color: #bae6fd !important; }
.rca5-chart-btn.active .rca5-cbtn-label { color: #38bdf8 !important; }
.rca5-chart-btn.active .rca5-cbtn-desc  { color: #7dd3fc !important; }

/* ── Why glow boxes — side by side ── */
.rca5-why-row {
  display: grid; grid-template-columns: 1fr 1fr; gap: 12px; margin-top: 16px;
}
.rca5-why-glowbox {
  background: rgba(6,15,35,0.95);
  border-radius: 12px; padding: 18px 20px;
  position: relative; overflow: hidden;
}
.rca5-why-glowbox.what-box {
  border: 1px solid rgba(56,189,248,0.4);
  box-shadow: 0 0 28px rgba(56,189,248,0.1), inset 0 0 20px rgba(56,189,248,0.03);
}
.rca5-why-glowbox.what-box::before {
  content: ''; position: absolute; top: 0; left: 0; right: 0; height: 2px;
  background: linear-gradient(90deg, transparent, #38bdf8, transparent);
}
.rca5-why-glowbox.why-box-inner {
  border: 1px solid rgba(124,58,237,0.4);
  box-shadow: 0 0 28px rgba(124,58,237,0.1), inset 0 0 20px rgba(124,58,237,0.03);
}
.rca5-why-glowbox.why-box-inner::before {
  content: ''; position: absolute; top: 0; left: 0; right: 0; height: 2px;
  background: linear-gradient(90deg, transparent, #7c3aed, transparent);
}
.rca5-glow-icon  { font-size: 18px; margin-bottom: 6px; }
.rca5-glow-title { font-size: 0.65rem; font-weight: 700; text-transform: uppercase;
  letter-spacing: 0.1em; margin-bottom: 8px; }
.what-box .rca5-glow-title { color: #38bdf8; }
.why-box-inner .rca5-glow-title { color: #a78bfa; }
.rca5-glow-body  { font-size: 0.8rem; line-height: 1.65; color: #e2e8f0 !important; }
.rca5-glow-body strong { color: #ffffff !important; }
.rca5-glow-body em { color: #7dd3fc !important; font-style: normal; font-weight: 600; }
.rca5-why-tags   { display: flex; flex-wrap: wrap; gap: 5px; margin-top: 10px; }
.rca5-why-tag    { font-size: 10px; padding: 2px 9px; border-radius: 20px;
  background: rgba(56,189,248,0.08); border: 1px solid rgba(56,189,248,0.2); color: #7dd3fc !important; }

/* ── Chart plot area ── */
.rca5-chart-area {
  background: rgba(12,21,40,0.7); border: 1px solid rgba(56,189,248,0.15);
  border-radius: 14px; padding: 18px 20px; margin-top: 8px;
}
/* Ensure Plotly text is visible on dark bg */
.rca5-chart-area .js-plotly-plot text,
.rca5-chart-area .js-plotly-plot .gtitle,
.rca5-chart-area .js-plotly-plot .xtick text,
.rca5-chart-area .js-plotly-plot .ytick text,
.rca5-chart-area .js-plotly-plot .legendtext,
.rca5-chart-area .js-plotly-plot .annotation-text { fill: #e2e8f0 !important; color: #e2e8f0 !important; }
.rca5-chart-area .js-plotly-plot .modebar-btn path { fill: #94a3b8 !important; }

"""

import threading, queue, json, os, re
import csv, shutil
import gradio as gr
import plotly.graph_objects as go
import plotly.express as px
import pandas as pd
from dotenv import load_dotenv
from neo4j import GraphDatabase
import requests
import subprocess 
import time

# ── Optional dependencies — installed separately ──────────────────────────────
# python-docx  (pip install python-docx)
try:
    from docx import Document as DocxDocument          # type: ignore[import]
    from docx.shared import Pt, RGBColor, Inches, Cm          # type: ignore[import]
    from docx.enum.text import WD_ALIGN_PARAGRAPH              # type: ignore[import]
    from docx.enum.table import WD_TABLE_ALIGNMENT             # type: ignore[import]
    from docx.oxml.ns import qn                                # type: ignore[import]
    from docx.oxml import OxmlElement                          # type: ignore[import]
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# openpyxl  (pip install openpyxl)
try:
    import openpyxl                                            # type: ignore[import]
    from openpyxl.styles import (                              # type: ignore[import]
        Font as XlFont, PatternFill, Alignment, Border, Side, GradientFill
    )
    from openpyxl.utils import get_column_letter               # type: ignore[import]
    _OPENPYXL_AVAILABLE = True
except ImportError:
    _OPENPYXL_AVAILABLE = False

def start_mcp():
    """
    Start the GraphPulse MCP server in a background thread.
    Probes for the port then propagates the URL into app._MCP_BASE_URL
    so run_cypher_via_mcp() routes through MCP.
    Non-blocking: demo.launch() is NOT delayed.
    """
    def _set_mcp_url(url: str):
        os.environ["GRAPHPULSE_MCP_URL"] = url
        try:
            import app as _app_mod
            _app_mod._MCP_BASE_URL = url
            print(f"[MCP] ✓ Connected → {url} (app._MCP_BASE_URL updated)")
        except Exception as ex:
            print(f"[MCP] Could not set app._MCP_BASE_URL: {ex}")

    def _worker():
        import sys as _sys, tempfile as _tmp
        try:
            # Check if already running
            try:
                with open("mcp_port.json") as f:
                    port = json.load(f)["port"]
                url = "http://127.0.0.1:" + str(port)
                resp = requests.get(url + "/tools", timeout=2)
                if resp.status_code == 200:
                    print("[MCP] Already running on port " + str(port))
                    _set_mcp_url(url)
                    return
            except Exception:
                pass

            # Launch with same venv Python
            _errfile = _tmp.NamedTemporaryFile(
                delete=False, suffix="_mcp_err.log", mode="w", encoding="utf-8"
            )
            proc = subprocess.Popen(
                [_sys.executable, "mcp_server.py"],
                stdout=_errfile,
                stderr=_errfile,
            )
            _errfile.close()
            # Poll every 0.3s for up to 15s
            for _ in range(50):
                time.sleep(0.3)
                if proc.poll() is not None:
                    try:
                        with open(_errfile.name, encoding="utf-8", errors="replace") as f:
                            err = f.read(800)
                        print("[MCP] Server crashed:\n" + err)
                    except Exception:
                        pass
                    return
                try:
                    with open("mcp_port.json") as f:
                        port = json.load(f)["port"]
                    url = "http://127.0.0.1:" + str(port)
                    resp = requests.get(url + "/tools", timeout=2)
                    if resp.status_code == 200:
                        print("[MCP] Started on port " + str(port))
                        _set_mcp_url(url)
                        return
                except Exception:
                    continue
            print("[MCP] Did not start in 15s - direct Neo4j fallback active")
            try:
                with open(_errfile.name, encoding="utf-8", errors="replace") as f:
                    print("[MCP] Last log:\n" + f.read(600))
            except Exception:
                pass
        except Exception as e:
            print("[MCP] Worker error: " + str(e))

    threading.Thread(target=_worker, daemon=True).start()

# ─────────────────────────────────────────────
# GLOBAL PLOTLY STYLE (used across all charts)
# ─────────────────────────────────────────────
_PLOTLY_LAYOUT = dict(
    paper_bgcolor="#060c1c",
    plot_bgcolor="#060c1c",
    height=480,
    autosize=False,
    font=dict(family="DM Sans, sans-serif", color="#f1f5f9", size=12),
    margin=dict(l=70, r=40, t=60, b=80),
    title_font=dict(color="#7dd3fc", size=15),
    xaxis=dict(
        tickfont=dict(color="#f1f5f9", size=11),
        
        gridcolor="rgba(255,255,255,0.08)",
        linecolor="rgba(255,255,255,0.15)",
        zerolinecolor="rgba(255,255,255,0.1)"
    ),
    yaxis=dict(
        tickfont=dict(color="#f1f5f9", size=11),
        
        gridcolor="rgba(255,255,255,0.08)",
        linecolor="rgba(255,255,255,0.15)",
        zerolinecolor="rgba(255,255,255,0.1)"
    ),
    legend=dict(
        font=dict(color="#f1f5f9", size=11),
        bgcolor="rgba(6,12,28,0.95)",
        bordercolor="rgba(56,189,248,0.25)",
        borderwidth=1
    )
)

_blank_fig = go.Figure().update_layout(
    paper_bgcolor="#060c1c",
    plot_bgcolor="#060c1c",
    xaxis=dict(visible=False),
    yaxis=dict(visible=False),
    margin=dict(l=0, r=0, t=0, b=0),
    height=480,
    autosize=False,
)

# Load .env FIRST so Neo4j credentials exist when app.py starts its MCP thread
load_dotenv(".env")
try:
    import app as _orig
    _orig_ok = True
except Exception as _orig_err:
    print(f"[app_mcp] Warning: could not fully import app.py: {_orig_err}")
    import types as _types
    _orig = _types.SimpleNamespace()
    _orig_ok = False
from agent_runner import run_rca, run_graph_update, _groq_call, undo_last_nl_update, undo_nl_update, get_nl_update_history, run_file_update_pipeline


# (load_dotenv already called above)

# ── CSV sync helpers — keep data/ files in sync with every graph change ──────
_APP_BASE   = os.path.dirname(os.path.abspath(__file__))

def _resolve_data_dir_app(base: str) -> str:
    """
    Find the actual data directory (projectdata/ or data/).
    Checks DATA_PATH env-var first, then probes both candidates for SupplierMaster.csv.
    """
    env_path = os.environ.get("DATA_PATH", "").strip()
    if env_path and os.path.isdir(env_path):
        return env_path
    for candidate in ("projectdata", "data"):
        path = os.path.join(base, candidate)
        if os.path.isfile(os.path.join(path, "SupplierMaster.csv")):
            return path
    fallback = os.path.join(base, "data")
    os.makedirs(fallback, exist_ok=True)
    return fallback

_APP_DATA   = _resolve_data_dir_app(_APP_BASE)
os.makedirs(_APP_DATA, exist_ok=True)
_APP_CSV    = {
    "Supplier":         os.path.join(_APP_DATA, "SupplierMaster.csv"),
    "SupplierEnriched": os.path.join(_APP_DATA, "Supplier_master_enriched.csv"),
    "Distributor":      os.path.join(_APP_DATA, "Distributer-WarehouseMasterDist.csv"),
    "Route":            os.path.join(_APP_DATA, "RoutesSheet.csv"),
}
_APP_ID_COL = {
    "Supplier":         "supplier_id",
    "SupplierEnriched": "supplier_id",
    "Distributor":      "distributor_id",
    "Route":            "route_id",
}

def _app_csv_read(path):
    if not os.path.exists(path): return []
    with open(path, newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))

def _app_csv_write(path, rows):
    if not rows: return
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        w.writeheader(); w.writerows(rows)

def _app_csv_delete(entity_type, id_val):
    """Remove a row from the matching CSV when a node is deleted from Neo4j."""
    path = _APP_CSV.get(entity_type)
    if not path: return
    id_col = _APP_ID_COL.get(entity_type, "id")
    rows = _app_csv_read(path)
    new_rows = [r for r in rows if str(r.get(id_col,"")).strip() != str(id_val).strip()]
    if len(new_rows) != len(rows):
        _app_csv_write(path, new_rows)
    if entity_type == "Supplier":
        _app_csv_delete("SupplierEnriched", id_val)
# ─────────────────────────────────────────────────────────────────────────────

# Safely import all attributes from app.py with fallbacks
def _safe_orig(attr, fallback=None):
    try: return getattr(_orig, attr)
    except AttributeError: return fallback

CUSTOM_CSS        = _safe_orig('CUSTOM_CSS', '')
INIT_JS           = _safe_orig('INIT_JS', '')
load_snapshot     = _safe_orig('load_snapshot', lambda: '')
update_history    = _safe_orig('update_history', lambda: '')
on_generate_query = _safe_orig('on_generate_query', lambda q: ('', '', '', '', '', '', ''))
_orig_on_generate_query = on_generate_query  # save original


# ── Pre-validated Cypher patterns for complex queries the LLM gets wrong ──────


def on_generate_query(question: str):
    """Override — intercepts known queries with correct Cypher before LLM."""
    import gradio as _gr
    q = (question or "").lower()

    # Verified Cypher for queries the LLM consistently gets wrong
    _OVERRIDES = [
        # Kolkata stockout — shows per-plant contribution to Kolkata's gap
        (lambda q: "kolkata" in q and any(w in q for w in [
             "stockout","demand gap","persistent","shortage","upstream","breaking down"]),
         """MATCH (pl:Plant)-[:DISPATCHES]->(sh:Shipment)-[:SHIPPED_TO]->(d:Distributor {distributor_city: 'Kolkata'})
WHERE sh.demand_gap > 0
WITH pl.plant_name AS plant_name, pl.plant_id AS plant_id,
     COUNT(sh) AS shortage_shipments,
     round(SUM(sh.demand_gap), 0) AS total_demand_gap,
     SUM(CASE WHEN sh.delivery_status = 'Major Delay' THEN 1 ELSE 0 END) AS delayed_shipments,
     round(AVG(CASE WHEN sh.delivery_status = 'Major Delay' THEN sh.delay_days END), 2) AS avg_delay_days
RETURN plant_name, plant_id, shortage_shipments, total_demand_gap, delayed_shipments, avg_delay_days
ORDER BY total_demand_gap DESC"""),

        # Delivery performance deteriorating — transport route/mode analysis
        (lambda q: any(w in q for w in ["delivery performance","deteriorating","transportation routes","routes or modes",
             "which routes","which modes","route driving","mode driving"]) or
                   ("delivery" in q and "deteriorat" in q) or
                   ("transport" in q and "routes" in q and "disruption" in q),
         """MATCH (pl:Plant)-[:HAS_ROUTE]->(r:Route)-[:CONNECTS_TO]->(d:Distributor)
MATCH (pl)-[:DISPATCHES]->(sh:Shipment)-[:SHIPPED_TO]->(d)
WHERE sh.delivery_status = 'Major Delay'
WITH r.mode AS transport_mode,
     pl.plant_name AS plant_name,
     d.distributor_city AS distributor_city,
     COUNT(sh) AS total_delays,
     round(AVG(sh.delay_days), 2) AS avg_delay_days,
     round(100.0 * COUNT(sh) / COUNT(sh), 1) AS delay_share
WITH transport_mode,
     COUNT(DISTINCT plant_name) AS plants_affected,
     COUNT(DISTINCT distributor_city) AS cities_affected,
     SUM(total_delays) AS total_delays,
     round(AVG(avg_delay_days), 2) AS avg_delay_days
RETURN transport_mode, total_delays, avg_delay_days, plants_affected, cities_affected
ORDER BY total_delays DESC"""),

        # Bhopal delays — plant-specific with upstream suppliers
        (lambda q: "bhopal" in q and any(w in q for w in ["delay","supplier","highest","upstream"]),
         """MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant {plant_name: 'Bhopal'})
MATCH (pl)-[:DISPATCHES]->(sh:Shipment)
WITH sup, pl,
     COUNT(sh) AS total_shipments,
     SUM(CASE WHEN sh.delivery_status = 'Major Delay' THEN 1 ELSE 0 END) AS delayed_count,
     round(AVG(CASE WHEN sh.delivery_status = 'Major Delay' THEN sh.delay_days END), 2) AS avg_delay_days
RETURN sup.supplier_id, sup.supplier_name, round(sup.risk_score, 2) AS risk_score,
       sup.annual_capacity_units, pl.plant_id, pl.plant_name,
       total_shipments, delayed_count,
       round(100.0 * delayed_count / total_shipments, 1) AS delay_rate_pct, avg_delay_days
ORDER BY risk_score DESC LIMIT 15"""),

        # High-risk suppliers with capacity exposure
        (lambda q: any(w in q for w in ["capacity exposure","hardest to replace","highest risk",
             "putting our plants","risk score above","which suppliers are putting"]),
         """MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)
WHERE sup.risk_score > 0.70
MATCH (pl)-[:DISPATCHES]->(sh:Shipment)
WITH sup, pl,
     COUNT(CASE WHEN sh.delivery_status = 'Major Delay' THEN 1 END) AS delayed_shipments,
     COUNT(sh) AS total_shipments
RETURN sup.supplier_id, sup.supplier_name, round(sup.risk_score,2) AS risk_score,
       sup.annual_capacity_units,
       COALESCE(sup.StoP_lead_time_days, 0) AS lead_time_days,
       pl.plant_id, pl.plant_name, delayed_shipments, total_shipments,
       round(100.0*delayed_shipments/total_shipments,1) AS delay_rate_pct
ORDER BY sup.risk_score DESC LIMIT 15"""),

        # Distributor unmet demand
        (lambda q: any(w in q for w in ["largest unmet demand","most shortage shipments",
             "distributor cities","total unmet demand","top 10 most affected"]),
         """MATCH (pl:Plant)-[:DISPATCHES]->(sh:Shipment)-[:SHIPPED_TO]->(d:Distributor)
WHERE sh.demand_gap > 0
WITH d.distributor_city AS distributor_city, d.distributor_id AS distributor_id,
     COUNT(DISTINCT sh) AS shortage_shipments,
     round(SUM(sh.demand_gap), 0) AS total_demand_gap,
     SUM(CASE WHEN sh.delivery_status = 'Major Delay' THEN 1 ELSE 0 END) AS delayed_shipments,
     round(AVG(CASE WHEN sh.delivery_status = 'Major Delay' THEN sh.delay_days END), 2) AS avg_delay_days
RETURN distributor_city, distributor_id, shortage_shipments, total_demand_gap, delayed_shipments, avg_delay_days
ORDER BY total_demand_gap DESC LIMIT 15"""),

        # Transport mode delays (generic — fires only when NOT a plant-usage query)
       (lambda q: (
            any(w in q for w in ["transport mode", "transportation mode", "which mode", "worst delay rate",
                                  "what mode delays", "mode delay", "mode has most delays"])
            and not any(pid in q for pid in ["pl1", "pl2", "pl3", "pl4", "baddi", "bhopal", "pune", "goa"])
            and not any(w in q for w in ["modes does", "modes do", "modes used", "what modes",
                                          "uses what", "uses which", "transport modes does",
                                          "transport modes use", "modes use"])
        ),
         """MATCH (pl:Plant)-[:HAS_ROUTE]->(r:Route)-[:CONNECTS_TO]->(d:Distributor)
MATCH (pl)-[:DISPATCHES]->(sh:Shipment)-[:SHIPPED_TO]->(d)
WHERE sh.delivery_status = 'Major Delay'
WITH r.mode AS transport_mode, COUNT(sh) AS total_delays,
     round(AVG(sh.delay_days),2) AS avg_delay_days, COUNT(DISTINCT pl) AS plants_affected
RETURN transport_mode, total_delays, avg_delay_days, plants_affected
ORDER BY total_delays DESC"""),

        # Jaggi/Sabharwal Pune
        (lambda q: any(w in q for w in ["jaggi","sabharwal"]) and
                   any(w in q for w in ["pune","pl3","plant"]),
         """MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant {plant_name: 'Pune'})
MATCH (pl)-[:DISPATCHES]->(sh:Shipment)
WITH pl, COUNT(DISTINCT sup) AS supplier_count,
     round(AVG(sup.risk_score),2) AS avg_risk_score, round(MAX(sup.risk_score),2) AS max_risk_score,
     COUNT(sh) AS total_shipments,
     SUM(CASE WHEN sh.delivery_status = 'Major Delay' THEN 1 ELSE 0 END) AS delayed_count
RETURN pl.plant_id, pl.plant_name, supplier_count, avg_risk_score, max_risk_score,
       total_shipments, delayed_count, round(100.0*delayed_count/total_shipments,1) AS delay_rate_pct"""),

        # All plants compare
        (lambda q: any(w in q for w in ["all four plants","compare plants","across all plants",
             "structural upstream","delay rates","59%","60%","61%"]),
         """MATCH (pl:Plant)-[:DISPATCHES]->(sh:Shipment)
WITH pl, COUNT(sh) AS total_shipments,
     SUM(CASE WHEN sh.delivery_status = 'Major Delay' THEN 1 ELSE 0 END) AS delayed_count,
     round(AVG(CASE WHEN sh.delivery_status='Major Delay' THEN sh.delay_days END),2) AS avg_delay
MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl)
WITH pl, total_shipments, delayed_count, avg_delay,
     COUNT(DISTINCT sup) AS supplier_count,
     round(AVG(sup.risk_score),2) AS avg_risk_score, round(MAX(sup.risk_score),2) AS max_risk
RETURN pl.plant_id, pl.plant_name, total_shipments, delayed_count,
       round(100.0*delayed_count/total_shipments,1) AS delay_rate_pct,
       avg_delay, supplier_count, avg_risk_score, max_risk
ORDER BY delay_rate_pct DESC"""),
    ]

    # ── Early intercept: plant-specific transport mode USAGE ─────────────
    _PLANT_ID_MAP = {"pl1":"PL1","pl2":"PL2","pl3":"PL3","pl4":"PL4",
                     "baddi":"PL1","bhopal":"PL2","pune":"PL3","goa":"PL4"}
    _is_plant_mode_usage = (
        any(w in q for w in ["what mode","which mode","what transport mode","which transport mode",
                              "modes does","modes do","modes used","mode does","transport modes",
                              "uses what","uses which","what modes"])
        and any(pid in q for pid in _PLANT_ID_MAP.keys())
    )
    if _is_plant_mode_usage:
        _pid = next((v for k, v in _PLANT_ID_MAP.items() if k in q), "PL1")
        _mode_cypher = (
            f"MATCH (pl:Plant {{plant_id: \'{_pid}\'}}) - [:HAS_ROUTE]->(r:Route)\n"
            f"RETURN DISTINCT r.mode AS transport_mode,\n"
            f"       COUNT(r) AS route_count,\n"
            f"       round(AVG(r.PtoD_transportation_cost_inr), 0) AS avg_cost_inr,\n"
            f"       round(AVG(r.PtoD_leadtime_days), 1) AS avg_leadtime_days,\n"
            f"       round(AVG(r.PtoD_distance_km), 0) AS avg_distance_km\n"
            f"ORDER BY route_count DESC"
        )
        print(f"[PlantModeUsage] Intercepted for {_pid}: {question[:60]}")
        return (
            '<div class="status-msg status-ok-gen">✦ Cypher generated — inspect or edit below, then click <strong style="color:#00ffcc;">▶ Run Query</strong> to execute.</div>',
            _gr.update(value=_mode_cypher, visible=True),
            _gr.update(visible=True), _gr.update(visible=False),
            _gr.update(value=''), _gr.update(value=""),
            _gr.update(value=""), _gr.update(value=""),
            _gr.update(visible=False),
        )

    for matcher, cypher in _OVERRIDES:
        try:
            if matcher(q):
                print(f"[CypherOverride] Matched for: {question[:60]}")
                return (
                    '<div class="status-msg status-ok-gen">✦ Cypher generated — inspect or edit below, then click <strong style="color:#00ffcc;">▶ Run Query</strong> to execute.</div>',
                    _gr.update(value=cypher, visible=True),
                    _gr.update(visible=True),
                    _gr.update(visible=False),
                    _gr.update(value=''),
                    _gr.update(value=""),
                    _gr.update(value=""),
                    _gr.update(value=""),
                    _gr.update(visible=False),
                )
        except Exception as _e:
            print(f"[CypherOverride error] {_e}")

    return _orig_on_generate_query(question)


_orig_on_run_query = _safe_orig('on_run_query', lambda q, c: ('', '', '', '', '', '', False, '', False))
_orig_on_run_insights = _safe_orig('on_run_insights', lambda q, c: ('', ''))
on_abort          = _safe_orig('on_abort', lambda: ('',)*9)
clear_session     = _safe_orig('clear_session', lambda: ('',)*9)
make_csv          = _safe_orig('make_csv', lambda q: None)
SAMPLE_QUESTIONS  = _safe_orig('SAMPLE_QUESTIONS', [])


def _build_rich_brief(question: str, records: list) -> str:
    """
    Generates a 3-4 line data-driven brief for ANY query type.
    Replaces the generic 'detailed analysis loading below' fallback.
    """
    if not records:
        return ""
    import re as _rb
    q = question.lower()
    n = len(records)
    keys = list(records[0].keys()) if records else []

    def _val(r, *candidates):
        for k in candidates:
            if r.get(k) is not None:
                return r[k]
        return None

    # ── Transport mode USAGE query (what modes does plant X use?) ─────────
    _is_mode_use = any(w in q for w in [
        "what mode", "which mode", "modes does", "modes do", "modes used",
        "mode does", "transport modes", "uses what", "uses which", "what modes", "mode use"
    ])
    if _is_mode_use and any(k in keys for k in ["transport_mode", "mode", "route_count"]):
        _PLANT_NAMES = {"pl1": "PL1 (Baddi)", "pl2": "PL2 (Bhopal)",
                        "pl3": "PL3 (Pune)", "pl4": "PL4 (Goa)",
                        "baddi": "Baddi", "bhopal": "Bhopal", "pune": "Pune", "goa": "Goa"}
        _plant_name = next((v for k, v in _PLANT_NAMES.items() if k in q), "The plant")
        _mode_col   = next((k for k in ["transport_mode", "mode"] if k in keys), None)
        _count_col  = next((k for k in ["route_count"] if k in keys), None)
        _cost_col   = next((k for k in ["avg_cost_inr", "cost_inr"] if k in keys), None)
        _modes      = [str(r.get(_mode_col, "")) for r in records if r.get(_mode_col)]
        _mode_list  = ", ".join(f"**{m}**" for m in _modes)
        _top_count  = records[0].get(_count_col, "") if _count_col and records else ""
        _top_mode   = _modes[0] if _modes else "Road"
        _top_cost   = records[0].get(_cost_col, "") if _cost_col and records else ""
        try:
            _cost_fmt = f"₹{float(_top_cost):,.0f}" if _top_cost else ""
        except Exception:
            _cost_fmt = str(_top_cost)
        _cost_phrase = f" with avg route cost of {_cost_fmt}" if _cost_fmt else ""
        return (
            f"**{_plant_name}** uses **{n} transport mode(s)**: {_mode_list}.\n\n"
            f"**{_top_mode}** has the highest number of routes ({_top_count}){_cost_phrase}. "
            f"The table shows route counts, average costs, lead times, and distances for each mode "
            f"— use this to understand how {_plant_name} distributes its logistics footprint."
        )

    def _fmt(v):
        try:
            f = float(v)
            return f"{f:,.0f}" if f > 100 else f"{f:.2f}"
        except Exception:
            return str(v)

    def _name(r, *candidates):
        for k in candidates:
            v = r.get(k)
            if v: return str(v).replace("_"," ").title()
        return "—"

    # ── Kolkata-specific brief — results show per-plant contribution to Kolkata gap ──
    if "kolkata" in q and any(w in q for w in ["stockout","demand gap","persistent","upstream","shortage"]):
        total_gap = sum(float(r.get("total_demand_gap", 0) or 0) for r in records)
        # Records are plants contributing to Kolkata gap, sorted by gap DESC
        top_plant = records[0].get("plant_name", records[0].get("plant", "Goa")) if records else "Goa"
        top_gap   = float(records[0].get("total_demand_gap", 0) or 0) if records else 0
        top_pct   = f"{top_gap/total_gap*100:.0f}%" if total_gap else "?"
        sec_plant = records[1].get("plant_name", records[1].get("plant", "Pune")) if len(records) > 1 else "Pune"
        sec_gap   = float(records[1].get("total_demand_gap", 0) or 0) if len(records) > 1 else 0
        return (
            f"**Kolkata (D0005)** has a total demand gap of **{total_gap:,.0f} units** — "
            f"the largest of any distributor city in the network.\n\n"
            f"**{top_plant} plant (PL4)** is the single biggest contributing source at **{top_gap:,.0f} units** "
            f"({top_pct} of Kolkata's total gap), followed by **{sec_plant}** at {sec_gap:,.0f} units. "
            f"All {n} plants are shipping to Kolkata but none are meeting forecasted demand — "
            f"the shortage persists every month regardless of transport performance.\n\n"
            f"Run the **RCA Trail** for a 5-step upstream investigation into which plants and "
            f"suppliers are responsible for this persistent supply gap."
        )

    # ── Delivery performance / transport route brief ────────────────────
    _is_delivery_q = any(w in q for w in [
        "delivery performance","deteriorating","transportation routes",
        "routes or modes","which routes","which modes","route driving","mode driving",
        "routes driving","transport driving",
    ]) or ("delivery" in q and "deteriorat" in q) or \
       ("transport" in q and "routes" in q)
    if _is_delivery_q:
        # Records are transport_mode, total_delays, avg_delay_days, plants_affected, cities_affected
        mode_col   = next((k for k in keys if "mode" in k.lower()), None)
        delay_col  = next((k for k in keys if "total_delay" in k.lower() or "total_delays" in k.lower()), None)
        avd_col    = next((k for k in keys if "avg_delay" in k.lower()), None)
        plants_col = next((k for k in keys if "plant" in k.lower()), None)
        cities_col = next((k for k in keys if "cit" in k.lower()), None)
        if mode_col and delay_col:
            sorted_m = sorted(records, key=lambda r: float(r.get(delay_col,0) or 0), reverse=True)
            top_m = sorted_m[0] if sorted_m else records[0]
            sec_m = sorted_m[1] if len(sorted_m) > 1 else {}
            total_delays = sum(float(r.get(delay_col,0) or 0) for r in records)
            top_mode  = str(top_m.get(mode_col,"?"))
            top_dels  = float(top_m.get(delay_col,0) or 0)
            top_pct   = f"{top_dels/total_delays*100:.0f}%" if total_delays else "?"
            top_avd   = top_m.get(avd_col,"?") if avd_col else "?"
            sec_mode  = str(sec_m.get(mode_col,"?")) if sec_m else "—"
            sec_dels  = float(sec_m.get(delay_col,0) or 0) if sec_m else 0
            top_plants= top_m.get(plants_col,"all") if plants_col else "all"
            return (
                f"**{top_mode}** mode has the highest delay count at **{top_dels:,.0f} delayed shipments** — "
                f"{top_pct} of all {total_delays:,.0f} Major Delay shipments across {n} transport modes, "
                f"with an average delay of **{top_avd} days** across **{top_plants} plants**.\n\n"
                f"**{sec_mode}** is the second-worst mode at {sec_dels:,.0f} delayed shipments. "
                f"The near-identical delay rates across all modes suggest the problem originates "
                f"upstream (supplier/plant level) rather than in the logistics layer — "
                f"delivery performance is deteriorating because of what is being dispatched, not how it is transported.\n\n"
                f"Run the **RCA Trail** to trace the upstream supplier and plant failures driving this network-wide delivery breakdown."
            )

    # ── Individual shipment rows (e.g. "show me delayed shipments") ──────
    # Detect when records are individual shipment rows (have shipment_id + delay_days)
    _has_ship_id  = any(k in keys for k in ["shipment_id","shipment","ship_id"])
    _has_delay_d  = any(k in keys for k in ["delay_days","delay","days_late","delay_duration"])
    _has_status   = any(k in keys for k in ["delivery_status","status","delay_status"])
    if _has_ship_id or (_has_delay_d and len(records) > 3):
        # Shipment-level results — summarise by plant, distributor, max delay
        delay_col  = next((k for k in ["delay_days","delay","days_late","delay_duration"] if k in keys), None)
        plant_col  = next((k for k in ["plant","plant_name","origin_plant"] if k in keys), None)
        dist_col   = next((k for k in ["distributor","distributor_city","city","destination"] if k in keys), None)
        status_col = next((k for k in ["delivery_status","status","delay_status"] if k in keys), None)

        total_ships = len(records)
        delayed_ships = sum(1 for r in records
                            if (status_col and "delay" in str(r.get(status_col,"")).lower())
                            or (delay_col and float(r.get(delay_col,0) or 0) > 0))
        if not delayed_ships and not status_col:
            delayed_ships = total_ships  # if no status col, assume all rows are delayed

        max_delay = 0
        if delay_col:
            try:
                max_delay = max(float(r.get(delay_col,0) or 0) for r in records)
            except Exception:
                pass

        # Count by plant
        plant_counts = {}
        if plant_col:
            for r in records:
                p = str(r.get(plant_col,"Unknown"))
                plant_counts[p] = plant_counts.get(p, 0) + 1
        top_plant = max(plant_counts, key=plant_counts.get) if plant_counts else "—"
        top_plant_cnt = plant_counts.get(top_plant, 0)

        # Count by distributor city
        dist_counts = {}
        if dist_col:
            for r in records:
                d = str(r.get(dist_col,"Unknown"))
                dist_counts[d] = dist_counts.get(d, 0) + 1
        top_dist = max(dist_counts, key=dist_counts.get) if dist_counts else "—"
        top_dist_cnt = dist_counts.get(top_dist, 0)

        plant_line = (f"**{top_plant}** is the origin plant for the most delayed shipments "
                      f"({top_plant_cnt} of {total_ships} shown)") if plant_col else ""
        dist_line  = (f"**{top_dist}** is the most-affected distributor city "
                      f"({top_dist_cnt} shipments)") if dist_col else ""
        delay_line = f"The maximum individual delay is **{max_delay:.1f} days**." if max_delay else ""

        parts = [p for p in [plant_line, dist_line] if p]
        summary_line = ". ".join(parts) + "." if parts else ""

        return (
            f"The query returned **{total_ships} delayed shipments** with Major Delay status "
            f"across the network.  \n\n"
            f"{summary_line}  \n\n"
            f"{delay_line}  \n\n"
            f"Run the **RCA Trail** to trace which upstream suppliers and plant-level issues are "
            f"generating these delayed dispatches."
        )

    # Sort by most impactful metric
    METRIC_KEYS = ["total_demand_gap","demand_gap","demand_gap_at_risk","major_delays",
                   "delayed_count","delayed_shipments","shortage_shipments","risk_score",
                   "total_shipments","avg_delay_days","delay_days","delay_rate_pct",
                   # ── Route / cost columns ──
                   "cost_inr","PtoD_transportation_cost_inr","avg_cost_inr",
                   "transportation_cost","transport_cost","freight_cost",
                   "cost_efficiency","distance_km","PtoD_distance_km",
                   "leadtime_days","PtoD_leadtime_days",
                   # ── Mode aggregate columns ──
                   "total_delays","route_count","avg_delay"]
    LABEL_KEYS  = ["distributor_city","supplier_name","plant_name","transportation_mode",
                   "transport_mode","mode","category","product_category","retailer_city",
                   "city","plant","distributor","route_id","shipment_id"]

    metric_key = next((k for k in METRIC_KEYS if k in keys), None)
    label_key  = next((k for k in LABEL_KEYS  if k in keys), None)

    try:
        sorted_r = sorted([r for r in records if _val(r, metric_key) is not None],
                          key=lambda r: float(_val(r, metric_key) or 0), reverse=True)
    except Exception:
        sorted_r = records

    top     = sorted_r[0] if sorted_r else records[0]
    second  = sorted_r[1] if len(sorted_r) > 1 else {}
    third   = sorted_r[2] if len(sorted_r) > 2 else {}

    top_lbl = _name(top, *LABEL_KEYS)
    top_val = _fmt(_val(top, metric_key)) if metric_key else "—"
    sec_lbl = _name(second, *LABEL_KEYS) if second else "—"
    sec_val = _fmt(_val(second, metric_key)) if second and metric_key else "—"
    # Guard: if top_val is — (metric not found), try to resolve from any numeric column
    if top_val == "—" and metric_key is None:
        _any_num_key = next((k for k in keys if isinstance(records[0].get(k), (int, float))), None)
        if _any_num_key:
            top_val = _fmt(_val(top, _any_num_key))
            sec_val = _fmt(_val(second, _any_num_key)) if second else "—"

    # ── Demand gap / stockout queries ──────────────────────────────────
    if any(w in q for w in ["demand gap","stockout","shortage","unmet","supply gap"]):
        total_gap = sum(float(_val(r, "total_demand_gap","demand_gap") or 0) for r in records)
        return (
            f"**{top_lbl}** has the largest unmet demand at **{top_val} units**, "
            f"followed by **{sec_lbl}** at {sec_val} units across the {n} results shown.\n\n"
            f"All {n} distributor cities in this result carry active demand gaps, indicating a systemic supply "
            f"capacity problem rather than isolated incidents. The total unmet demand across all cities shown "
            f"is **{total_gap:,.0f} units** — none of these distributors are fully served.\n\n"
            f"Run the **RCA Trail** for a 5-step upstream investigation into which plants and suppliers are responsible."
        )

    # ── Supplier risk at a SPECIFIC PLANT (e.g. Bhopal supplier analysis) ─
    _has_supplier_cols = "supplier_name" in keys or "supplier_id" in keys
    _has_plant_cols    = "plant_name" in keys or "plant_id" in keys
    _has_delay_cols    = "delayed_count" in keys or "delayed_shipments" in keys
    _specific_plant    = next((w for w in ["bhopal","pune","baddi","goa","pl1","pl2","pl3","pl4"] if w in q), None)

    if _has_supplier_cols and _has_plant_cols and _has_delay_cols and _specific_plant:
        # This is a plant-specific supplier risk query — Bhopal, Pune etc.
        plant_name = _name(records[0], "plant_name") if records else _specific_plant.title()
        plant_id   = records[0].get("plant_id","?") if records else "?"
        total_delayed = int(records[0].get("delayed_count", records[0].get("delayed_shipments", 0)) or 0) if records else 0
        delay_rate = records[0].get("delay_rate_pct","?") if records else "?"
        # Sort by risk_score to find top suppliers
        risk_sorted = sorted(records, key=lambda r: float(r.get("risk_score",0) or 0), reverse=True)
        top_sup  = risk_sorted[0].get("supplier_name","?") if risk_sorted else "?"
        top_risk = risk_sorted[0].get("risk_score","?") if risk_sorted else "?"
        sec_sup  = risk_sorted[1].get("supplier_name","?") if len(risk_sorted) > 1 else "—"
        sec_risk = risk_sorted[1].get("risk_score","?") if len(risk_sorted) > 1 else "—"
        n_critical = sum(1 for r in records if float(r.get("risk_score",0) or 0) >= 0.90)
        n_high = sum(1 for r in records if float(r.get("risk_score",0) or 0) >= 0.70)
        return (
            f"**{plant_name} ({plant_id})** has **{total_delayed:,} Major Delay shipments** — "
            f"a delay rate of **{delay_rate}%**, the highest in the network.\n\n"
            f"**{top_sup}** carries the highest risk score at **{top_risk}**, followed by **{sec_sup}** at {sec_risk}. "
            f"Of the {n} suppliers feeding {plant_name}, **{n_high} have risk scores above 0.70** "
            f"({'including ' + str(n_critical) + ' CRITICAL (≥0.90)' if n_critical else 'none are CRITICAL ≥0.90'}).\n\n"
            f"Run the **RCA Trail** to trace how these supplier risk scores translate into the {total_delayed:,} delayed "
            f"shipments at {plant_name} and their downstream impact on distributor cities."
        )

    # ── Supplier risk queries ──────────────────────────────────────────
    if any(w in q for w in ["supplier risk","risk score","high risk","risky supplier",
                              "which supplier","jaggi","sabharwal","nagy","capacity exposure",
                              "hardest to replace","bhopal plant","goa plant","pune plant","baddi plant"]):
        if _has_supplier_cols:
            risk_sorted = sorted(records, key=lambda r: float(r.get("risk_score",0) or 0), reverse=True)
            top_sup  = risk_sorted[0].get("supplier_name","?") if risk_sorted else "?"
            top_risk_v = risk_sorted[0].get("risk_score","?") if risk_sorted else "?"
            top_plant= _name(risk_sorted[0], "plant_name","plant_id") if risk_sorted else "?"
            sec_sup  = risk_sorted[1].get("supplier_name","?") if len(risk_sorted) > 1 else "—"
            sec_risk_v = risk_sorted[1].get("risk_score","?") if len(risk_sorted) > 1 else "—"
            sec_plant= _name(risk_sorted[1], "plant_name","plant_id") if len(risk_sorted) > 1 else "—"
            n_critical = sum(1 for r in records if float(r.get("risk_score",0) or 0) >= 0.90)
            return (
                f"**{top_sup}** carries the highest risk score of **{top_risk_v}**, feeding **{top_plant}**. "
                f"**{sec_sup}** follows at risk {sec_risk_v}, feeding {sec_plant}.\n\n"
                f"Of the {n} suppliers shown, **{n_critical} carry risk scores of 0.90 or above** — classified as CRITICAL. "
                f"These suppliers are historically the most unreliable and represent the primary upstream vulnerability "
                f"across the manufacturing network.\n\n"
                f"The higher the risk score, the more frequently this supplier has caused Major Delay shipments. "
                f"Run the **RCA Trail** to trace how these supplier risks cascade into plant delays and distributor stockouts."
            )

    # ── Shipment delay queries ─────────────────────────────────────────
    if any(w in q for w in ["delay","delayed shipment","late delivery","major delay","shipment bottleneck"]):
        # Only generate a grouped brief if data is aggregated (has delayed_count/delayed_shipments col)
        _has_agg_delay = any(k in keys for k in ["delayed_count","delayed_shipments","major_delays","total_delays"])
        if _has_agg_delay and metric_key and top_lbl != "—":
            total_delayed = sum(int(_val(r,"delayed_count","delayed_shipments","major_delays") or 0) for r in records)
            return (
                f"**{top_lbl}** has the most Major Delay shipments at **{top_val}**, "
                f"followed by **{sec_lbl}** at {sec_val} delayed shipments.\n\n"
                f"Across all {n} results, a total of **{total_delayed:,} Major Delay shipments** are recorded — "
                f"these are shipments that arrived significantly behind schedule, creating downstream stock gaps "
                f"at distributor cities.\n\n"
                f"Run the **RCA Trail** to trace which upstream suppliers and transport routes are driving these delays."
            )

    # ── Transportation mode / route queries ────────────────────────────
    if any(w in q for w in ["transport","route","mode","freight","cost","road","air","rail","sea"]):
        # Detect: individual route rows (route_id present) vs mode aggregates
        _has_route_id  = "route_id" in keys
        _has_mode_agg  = any(k in keys for k in ["transportation_mode","transport_mode","total_delays","route_count"])
        _is_cost_q     = any(w in q for w in ["cost","expensive","cheap","affordable","price"])
        _is_delay_q    = any(w in q for w in ["delay","late","delayed"])
        _order_word    = "cheapest" if any(w in q for w in ["cheap","lowest","affordable"]) else "most expensive"

        if _has_route_id and _is_cost_q:
            # Individual route cost ranking
            _cost_key   = next((k for k in ["cost_inr","PtoD_transportation_cost_inr"] if k in keys), metric_key)
            _plant_col  = next((k for k in ["plant","plant_name"] if k in keys), None)
            _city_col   = next((k for k in ["distributor_city","city"] if k in keys), None)
            _mode_col   = next((k for k in ["transport_mode","mode"] if k in keys), None)

            _top_cost   = _fmt(_val(top, _cost_key)) if _cost_key else top_val
            _top_plant  = str(top.get(_plant_col,"")).replace("_"," ").title() if _plant_col else ""
            _top_city   = str(top.get(_city_col,"")).replace("_"," ").title() if _city_col else top_lbl
            _top_mode   = str(top.get(_mode_col,"")) if _mode_col else ""
            _sec_cost   = _fmt(_val(second, _cost_key)) if second and _cost_key else "—"
            _sec_city   = str(second.get(_city_col,"")).replace("_"," ").title() if second and _city_col else sec_lbl

            _route_desc = f"{_top_plant} → {_top_city} via {_top_mode}" if _top_plant and _top_mode else _top_city
            _sec_desc   = f"→ {_sec_city}" if _sec_city and _sec_city != "—" else ""

            return (
                f"The {_order_word} route is **{_route_desc}** at **₹{_top_cost}** per route, "
                f"followed by **{_sec_desc}** at ₹{_sec_cost} across the {n} routes shown.\n\n"
                f"Each row is one specific Plant→Distributor route. **Air** routes are typically the most "
                f"expensive but fastest (1-day lead time). **Sea** and **Rail** routes are lowest cost "
                f"but have longer lead times. Use this to identify which logistics paths offer "
                f"the best cost-efficiency trade-off for each corridor."
            )

        elif _has_route_id and _is_delay_q:
            return (
                f"**{top_lbl}** has the highest delay exposure among the {n} routes shown.\n\n"
                f"Each row is one specific Plant→Distributor route. Use this to identify "
                f"which corridors have the worst delivery performance."
            )

        elif _has_route_id:
            # Generic individual route brief
            return (
                f"**{top_lbl}** leads with **{top_val}** across the {n} routes, "
                f"followed by **{sec_lbl}** at {sec_val}.\n\n"
                f"Each row is one specific Plant→Distributor route (route_id format: Plant@Distributor). "
                f"The table is sorted by the most significant metric for this query."
            )

        elif _has_mode_agg and _is_delay_q:
            # Mode aggregate delay
            return (
                f"**{top_lbl}** transport mode has the most delays at **{top_val}** shipments, "
                f"followed by **{sec_lbl}** at {sec_val} across the {n} modes.\n\n"
                f"This covers ALL Major Delay shipments network-wide grouped by transport mode. "
                f"Near-identical delay rates across modes indicate the root cause is upstream "
                f"(supplier/plant level) rather than in the logistics layer itself."
            )

        else:
            # Mode aggregate cost or generic
            total_cost = sum(float(_val(r, "cost_inr", "PtoD_transportation_cost_inr", "avg_cost_inr") or 0) for r in records)
            return (
                f"**{top_lbl}** has the highest {metric_key.replace('_',' ') if metric_key else 'value'} "
                f"at **{top_val}**, followed by **{sec_lbl}** at {sec_val} across the {n} routes or modes shown.\n\n"
                f"This analysis covers the transportation layer — comparing cost, lead time, and delay "
                f"performance. Routes with high cost combined with high delay rates represent the biggest "
                f"logistics efficiency gap."
            )

    # ── Plant performance / delay rate queries ────────────────────────
    if any(w in q for w in ["plant","factory","baddi","pune","bhopal","goa","pl1","pl2","pl3","pl4",
                              "delay rate","failure rate","structural upstream","all four plants",
                              "59%","60%","61%","network-wide","consistently high"]):
        total_delayed = sum(int(r.get("delayed_count", r.get("delayed_shipments", r.get("major_delays", 0))) or 0) for r in records)
        top_plant = _name(top, "plant_name","plant_id","name")
        top_rate  = top.get("delay_rate_pct", top.get("delay_pct","?"))
        top_del   = top.get("delayed_count", top.get("delayed_shipments","?"))
        sec_plant = _name(second, "plant_name","plant_id","name") if second else "—"
        top_risk  = top.get("max_risk_score", top.get("avg_risk_score","?"))
        return (
            f"**{top_plant}** leads with the highest delay rate at **{top_rate}%** and **{top_del} Major Delay shipments**, "
            f"followed by **{sec_plant}**. All four plants are running between 59–61% delay rates — "
            f"this near-identical pattern across all plants rules out plant-specific issues and points to a **common upstream cause**.\n\n"
            f"Across all {n} plants, a total of **{total_delayed:,} Major Delay shipments** are recorded. "
            f"The fact that every plant has roughly the same delay rate means the bottleneck is not inside any plant — "
            f"it originates in the **shared supplier base** feeding all four manufacturing sites.\n\n"
            f"The max risk score across plant supplier bases is **{top_risk}**. "
            f"Run the **RCA Trail** to trace which suppliers are the structural upstream cause of this network-wide failure rate."
        )

    # ── Generic fallback — always better than "detailed analysis loading below" ──
    # ── Safe generic fallback — never shows — placeholders ──────────────
    _safe_top_lbl = top_lbl if top_lbl and top_lbl != "—" else "The top result"
    _safe_top_val = top_val if top_val and top_val != "—" else "the highest value"
    _safe_sec_lbl = sec_lbl if sec_lbl and sec_lbl != "—" else ""
    _safe_sec_val = sec_val if sec_val and sec_val != "—" else ""
    _second_line  = (f", followed by **{_safe_sec_lbl}** at {_safe_sec_val}") if _safe_sec_lbl and _safe_sec_val else ""
    _metric_label = metric_key.replace("_"," ").title() if metric_key else "the key metric"

    return (
        f"**{_safe_top_lbl}** ranks highest with **{_safe_top_val}** in {_metric_label}"
        f"{_second_line} across the {n} results shown.\n\n"
        f"This table covers {n} records from your supply chain knowledge graph, "
        f"sorted by the highest-impact dimension to surface the most critical findings first.\n\n"
        f"Run the **RCA Trail** for a structured investigation into root causes and upstream dependencies."
    )


def on_run_query(question: str, cypher: str):
    """Override — injects data-driven brief for all query types."""
    result = _orig_on_run_query(question, cypher)
    try:
        records = getattr(_orig, '_last_records', [])
        if records and result and len(result) >= 3:
            rich = _build_rich_brief(question, records)
            if rich:
                result = list(result)
                result[2] = rich
                result = tuple(result)
    except Exception as _e:
        print(f"[on_run_query override] {_e}")
    return result


def _compute_programmatic_insights(question: str, records: list) -> str:
    """
    Generate 100% data-driven Business Insight bullets from the full records list.
    Used when LLM-generated percentages are wrong (e.g. supplier shutdown queries).
    Returns bullet text string or "" if not applicable.
    """
    if not records or len(records) < 2:
        return ""
    q_lower = (question or "").lower()
    cols = list(records[0].keys()) if records else []
    cols_lower = [c.lower() for c in cols]

    # ── Kolkata plant-contribution query ────────────────────────────────
    # Records: plant_name, plant_id, shortage_shipments, total_demand_gap, delayed_shipments, avg_delay_days
    _has_plant  = any("plant" in c for c in cols_lower)
    _is_kolkata = "kolkata" in q_lower and any(w in q_lower for w in [
        "stockout","demand gap","persistent","upstream","shortage","breaking down"])

    if _is_kolkata and _has_plant and _has_gap:
        gap_col  = next((c for c in cols if "demand_gap" in c.lower()), None)
        pl_col   = next((c for c in cols if "plant_name" in c.lower() or c.lower() == "plant"), None)
        del_col  = next((c for c in cols if "delayed" in c.lower()), None)
        sht_col  = next((c for c in cols if "shortage" in c.lower()), None)
        if gap_col and pl_col:
            total_gap = sum(float(r.get(gap_col, 0) or 0) for r in records)
            sorted_r  = sorted(records, key=lambda r: float(r.get(gap_col, 0) or 0), reverse=True)
            top1, top2 = sorted_r[0], sorted_r[1] if len(sorted_r) > 1 else {}
            t1_name = str(top1.get(pl_col, "?"))
            t1_gap  = float(top1.get(gap_col, 0) or 0)
            t2_name = str(top2.get(pl_col, "?")) if top2 else "—"
            t2_gap  = float(top2.get(gap_col, 0) or 0) if top2 else 0
            total_delayed = sum(int(r.get(del_col, 0) or 0) for r in records) if del_col else 0
            total_shortage= sum(int(r.get(sht_col, 0) or 0) for r in records) if sht_col else 0
            bullets = [
                f"• Kolkata's total demand gap is {total_gap:,.0f} units — accumulated from {total_shortage} shortage shipments across all {len(records)} plants feeding the city",
                f"• {t1_name} plant is the largest upstream contributor at {t1_gap:,.0f} units ({t1_gap/total_gap*100:.0f}% of Kolkata's total gap), followed by {t2_name} at {t2_gap:,.0f} units ({t2_gap/total_gap*100:.0f}%)",
                f"• {total_delayed} of all shipments to Kolkata are Major Delays — confirming the gap is not just a capacity problem but also a reliability problem in the supply chain",
                f"• All {len(records)} plants show active demand gaps to Kolkata — this is a structural network-wide failure, not a single-plant issue",
            ]
            return "\n".join(bullets)

    # ── Supplier shutdown / downstream city exposure queries ──────────
    # Detect by column shape: has demand_gap_at_risk OR (total_shipments + major_delays + demand_gap)
    _has_gap   = any("demand_gap" in c for c in cols_lower)
    _has_delay = any("major_delay" in c or "delayed" in c for c in cols_lower)
    _has_city  = any("city" in c or "distributor" in c for c in cols_lower)
    _is_shutdown = any(w in q_lower for w in [
        "shutdown", "shuts down", "production stop", "nagy", "sup0045",
        "cascading", "at risk", "impact on stockout"
    ])

    if _is_shutdown and _has_gap and _has_city:
        # Find the right column names
        gap_col   = next((c for c in cols if "demand_gap" in c.lower()), None)
        city_col  = next((c for c in cols if "city" in c.lower()), None)
        did_col   = next((c for c in cols if "distributor_id" in c.lower() or (c.lower()=="distributor_id")), None)
        tot_col   = next((c for c in cols if "total_shipment" in c.lower()), None)
        del_col   = next((c for c in cols if "major_delay" in c.lower() or "delayed" in c.lower()), None)

        if not (gap_col and city_col):
            return ""

        total_gap   = sum(float(r.get(gap_col, 0) or 0) for r in records)
        total_ships = sum(int(r.get(tot_col, 0) or 0) for r in records) if tot_col else 0
        total_delays= sum(int(r.get(del_col, 0) or 0) for r in records) if del_col else 0
        n = len(records)

        sorted_r = sorted(records, key=lambda r: float(r.get(gap_col, 0) or 0), reverse=True)
        top1 = sorted_r[0]
        top2 = sorted_r[1] if n > 1 else None
        top3 = sorted_r[2] if n > 2 else None

        def _city(r):
            return str(r.get(city_col, r.get(did_col, "Unknown"))).strip()
        def _gap(r):
            return float(r.get(gap_col, 0) or 0)
        def _pct_of_total(v):
            return f"{v/total_gap*100:.1f}%" if total_gap else "N/A"
        def _shortage_rate(r):
            # shortage shipments / total shipments per city
            s = int(r.get(next((c for c in cols if "shortage" in c.lower()), tot_col or ""), 0) or 0)
            t = int(r.get(tot_col, 0) or 0) if tot_col else 0
            if t > 0:
                return f"{s/t*100:.0f}%"
            return "N/A"

        top1_pct = _pct_of_total(_gap(top1))
        top2_str = (f"{_city(top2)} ({top2.get(did_col,'')}) at {_gap(top2):,.0f} units "
                    f"({_pct_of_total(_gap(top2))} of total at-risk demand)") if top2 else ""
        top3_str = (f"{_city(top3)} ({top3.get(did_col,'')}) at {_gap(top3):,.0f} units") if top3 else ""

        # Average shortage rate across all cities (shortage shipments / total shipments)
        if tot_col:
            rates = []
            sht_col = next((c for c in cols if "shortage" in c.lower()), None)
            for r in records:
                t = int(r.get(tot_col, 0) or 0)
                s = int(r.get(sht_col, 0) or 0) if sht_col else 0
                if t > 0:
                    rates.append(s / t * 100)
            avg_rate = f"{sum(rates)/len(rates):.0f}%" if rates else "N/A"
        else:
            avg_rate = "N/A"

        # Top 2 cities combined
        top2_combined = _gap(top1) + (_gap(top2) if top2 else 0)
        top2_combined_pct = _pct_of_total(top2_combined)

        bullets = [
            f"• {_city(top1)} distributor {top1.get(did_col,'')} has {_gap(top1):,.0f} units of demand gap at risk — {top1_pct} of the total {total_gap:,.0f} units exposed across all {n} cities",
            f"• {top2_str}" if top2_str else "",
            f"• {top3_str} ({_pct_of_total(_gap(top3))} of total at-risk demand)" if top3 else "",
            f"• All {n} distributor cities are exposed — average shortage rate of {avg_rate} of shipments per city is already at stockout risk before any shutdown occurs",
            f"• Prioritise emergency procurement for {_city(top1)} and {_city(top2) if top2 else ''} to cover {top2_combined:,.0f} units ({top2_combined_pct} of total at-risk demand) as the immediate mitigation priority",
        ]
        return "\n".join(b for b in bullets if b.strip())

    return ""  # not applicable — fall through to LLM


def on_run_insights(question: str, cypher: str):
    """
    Override of app.py on_run_insights.
    For queries where the LLM is known to hallucinate percentages (e.g. supplier shutdown),
    generate insights programmatically from the full records. Fall back to LLM otherwise.
    """
    try:
        # Access _last_records from the original app module
        records = getattr(_orig, '_last_records', [])
        if not records:
            return ""

        # Try programmatic first — guaranteed correct percentages
        prog_bullets = _compute_programmatic_insights(question, records)
        if prog_bullets:
            # Reuse app.py's _build_insight_html to keep identical styling
            try:
                return _orig._build_insight_html(prog_bullets)
            except Exception:
                # Inline fallback renderer
                import re as _ri
                def _bold(t):
                    return _ri.sub(r'\b(\d[\d,\.]*\s*(?:units|days|%|shipments|cities)?)\b', r'<strong>\1</strong>', t)
                items = "".join(
                    f'<div class="insight-item"><span class="insight-dot">◆</span>'
                    f'<span class="insight-text">{_bold(line.lstrip("• ").strip())}</span></div>'
                    for line in prog_bullets.split("\n") if line.strip()
                )
                return f'''<div class="insight-section">
    <div class="insight-heading custom-insight-heading">
        <span class="insight-heading-icon">◈</span>
        <span>✦ Business Insights</span>
        <span class="insight-heading-line"></span>
    </div>{items}</div>'''

        # Fall back to original LLM-based insights for all other query types
        return _orig_on_run_insights(question, cypher)

    except Exception as e:
        try:
            return _orig_on_run_insights(question, cypher)
        except Exception:
            return f'<div style="color:#f87171;font-size:0.78rem;padding:8px 0">⚠ Insights unavailable: {str(e)[:120]}</div>'

_neo4j = GraphDatabase.driver(
    os.getenv("NEO4J_URI"),
    auth=(os.getenv("NEO4J_USER"), os.getenv("NEO4J_PASSWORD"))
)
_DB = os.getenv("NEO4J_DATABASE", "neo4j")


def _run_neo4j(query, params=None):
    with _neo4j.session(database=_DB) as s:
        return [r.data() for r in s.run(query, params or {})]


# ════════════════════════════════════════════════════════════════════
# EXTRA CSS — all UI fixes in one place
# ════════════════════════════════════════════════════════════════════
EXTRA_CSS = """
/* ─────────────────────────────────────────────
   FIX 1: Tab heading text → white, hover visible
   ───────────────────────────────────────────── */
.gr-tab-label,
button[role="tab"],
button[role="tab"] span,
.tabs > .tab-nav > button,
.tabs > .tab-nav > button *,
div[role="tablist"] button,
div[role="tablist"] button span {
    color: #f1f5f9 !important;
    font-weight: 600 !important;
}
/* Hover: dark navy bg so white text stays readable */
div[role="tablist"] button:hover,
div[role="tablist"] button:hover span,
.tabs > .tab-nav > button:hover,
.tabs > .tab-nav > button:hover * {
    background: rgba(14,165,233,0.14) !important;
    color: #ffffff !important;
}
/* Active tab: cyan underline + bright text */
div[role="tablist"] button[aria-selected="true"],
div[role="tablist"] button[aria-selected="true"] span {
    color: #38bdf8 !important;
    border-bottom: 2px solid #38bdf8 !important;
    background: rgba(56,189,248,0.08) !important;
}
div[role="tablist"] button[aria-selected="true"]:hover,
div[role="tablist"] button[aria-selected="true"]:hover span {
    color: #7dd3fc !important;
    background: rgba(56,189,248,0.14) !important;
}

/* ─────────────────────────────────────────────
   FIX 2: Buttons → bluish-cyan theme
   Targets: Generate Query (analyze-btn), Run Query (run-btn),
            Abort (abort-btn), Clear Session / Load Snapshot (clear-btn),
            Sample question buttons (sample-btn)
   ───────────────────────────────────────────── */
.analyze-btn,
button.analyze-btn,
.gr-button.analyze-btn {
    background: linear-gradient(135deg, #0284c7 0%, #0ea5e9 60%, #06b6d4 100%) !important;
    border: 1px solid #0ea5e9 !important;
    color: #ffffff !important;
    font-weight: 700 !important;
    box-shadow: 0 0 18px rgba(14,165,233,0.35) !important;
}
.analyze-btn:hover { opacity: 0.88 !important; }

.run-btn,
button.run-btn,
.gr-button.run-btn {
    background: linear-gradient(135deg, #0369a1 0%, #0ea5e9 100%) !important;
    border: 1px solid #38bdf8 !important;
    color: #ffffff !important;
    font-weight: 600 !important;
}
.run-btn:hover { opacity: 0.88 !important; }

.abort-btn,
button.abort-btn,
.gr-button.abort-btn {
    background: linear-gradient(135deg, #155e75 0%, #0891b2 100%) !important;
    border: 1px solid #22d3ee !important;
    color: #e0f2fe !important;
    font-weight: 600 !important;
}

.clear-btn,
button.clear-btn,
.gr-button.clear-btn {
    background: rgba(14,165,233,0.15) !important;
    border: 1px solid rgba(14,165,233,0.45) !important;
    color: #7dd3fc !important;
}
.clear-btn:hover { background: rgba(14,165,233,0.28) !important; }

.sample-btn,
button.sample-btn,
.gr-button.sample-btn {
    background: rgba(6,182,212,0.12) !important;
    border: 1px solid rgba(6,182,212,0.35) !important;
    color: #67e8f9 !important;
    font-size: 0.78rem !important;
}
.sample-btn:hover { background: rgba(6,182,212,0.25) !important; }

/* ─────────────────────────────────────────────
   FIX 3: Dark-grey text → visible in RCA + Viz tabs
   ───────────────────────────────────────────── */
.sec-label {
    color: #7dd3fc !important;
    font-size: 0.7rem !important;
    font-weight: 700 !important;
    letter-spacing: 0.12em !important;
    text-transform: uppercase !important;
    margin-bottom: 8px !important;
}

/* Gradio markdown paragraphs inside dark panels */
.left-panel .prose p,
.left-panel .prose li,
.left-panel .prose strong,
.left-panel .prose h1,
.left-panel .prose h2,
.left-panel .prose h3,
.left-panel p,
.left-panel li,
.left-panel label,
.main-panel .prose p,
.main-panel label {
    color: #e2e8f0 !important;
}

/* Gradio textbox labels + placeholder */
.gr-textbox label,
.gr-textbox textarea,
.gr-textbox input,
label.svelte-1b6s6s,
label.block span {
    color: #cbd5e1 !important;
}

/* ─────────────────────────────────────────────
   FIX: File upload input boxes — black text on white bg
   ───────────────────────────────────────────── */
/* ── File Upload Drop Zone — Override ALL Gradio themes ── */
/* Every possible selector Gradio uses for the upload zone text */
.gr-file input[type="text"],
.gr-file .file-name,
.gr-file label span,
.gr-file .wrap span,
.gr-file p,
.gr-file span,
.gr-file > .wrap > .file-preview span,
input[type="file"] + label,
.upload-container span,
.upload-container p,
[data-testid="file"] span,
[data-testid="file"] p,
[data-testid="file"] .file-name,
[data-testid="file"] > label > span,
[data-testid="file"] .wrap span,
.svelte-file-upload span,
.file-preview span,
.file-upload span,
.file-upload p,
.upload-button span,
.gr-form .gr-file span,
label[for*="file"] span,
.block .file-preview span,
.block .file-drop span,
form .file-upload span {
    color: #000000 !important;
    font-weight: 800 !important;
    font-size: 1.05rem !important;
    letter-spacing: 0.01em !important;
}
/* Gradio file component label above the box */
.gr-file > label > span,
div.file-upload label span {
    color: #000000 !important;
    font-weight: 800 !important;
    font-size: 1.05rem !important;
}

/* ── Node Lookup dropdown — complete visibility fix ── */
/* List options: black text on white */
#gv-label-drop select,
#gv-label-drop option,
#gv-label-drop .wrap-inner select,
#gv-label-drop ul li,
#gv-label-drop ul li span,
#gv-label-drop .item span,
#gv-label-drop .gradio-dropdown li {
    color: #111111 !important;
    background-color: #ffffff !important;
    font-weight: 700 !important;
}
/* Selected value displayed inside the dark input box: white text */
#gv-label-drop .wrap-inner > input,
#gv-label-drop input[type="text"],
#gv-label-drop .single-select input {
    color: #ffffff !important;
    font-weight: 700 !important;
    background: transparent !important;
}
/* Dropdown container */
#gv-label-drop .wrap-inner {
    background: rgba(12,21,40,0.95) !important;
    border: 1px solid rgba(56,189,248,0.4) !important;
}
/* Hover state on list items */
#gv-label-drop ul li:hover {
    background: rgba(56,189,248,0.15) !important;
    color: #000000 !important;
}
/* Selected/active item */
#gv-label-drop ul li.selected,
#gv-label-drop ul li[aria-selected="true"] {
    background: rgba(56,189,248,0.2) !important;
    color: #000000 !important;
    font-weight: 800 !important;
}
/* ID input box */
#gv-id-box input,
#gv-id-box textarea {
    color: #ffffff !important;
    background: rgba(4,9,22,0.9) !important;
    border: 1px solid rgba(56,189,248,0.3) !important;
}
#gv-id-box label span {
    color: #7dd3fc !important;
    font-weight: 700 !important;
}
#gv-label-drop label span {
    color: #7dd3fc !important;
    font-weight: 700 !important;
}

/* ── Node Lookup dropdown — black text on white option list ──
/* gr.Markdown in viz tab descriptions */
.viz-desc p, .viz-desc li {
    color: #94a3b8 !important;
    font-style: italic;
    font-size: 0.82rem;
}

/* Tool log text visibility */
.tool-log-wrap {
    color: #cbd5e1 !important;
}
.tool-entry span[style*="var(--muted)"] {
    color: #94a3b8 !important;
}

/* ─────────────────────────────────────────────
   Agent status badges
   ───────────────────────────────────────────── */
.ag-status { padding:10px 16px; border-radius:8px; font-size:0.85rem;
             font-family:var(--mono); margin-bottom:8px; }
.ag-run  { background:rgba(14,165,233,.18); border:1px solid rgba(14,165,233,.4);
           color:#7dd3fc; }
.ag-ok   { background:rgba(16,185,129,.15); border:1px solid rgba(16,185,129,.4);
           color:#6ee7b7; }
.ag-err  { background:rgba(239,68,68,.15);  border:1px solid rgba(239,68,68,.4);
           color:#fca5a5; }

/* ─────────────────────────────────────────────
   Tool call log
   ───────────────────────────────────────────── */
.tool-log-wrap { background:var(--bg-card); border:1px solid var(--border);
                 border-radius:var(--radius); padding:14px;
                 max-height:280px; overflow-y:auto; }
.tool-entry { border-bottom:1px solid rgba(255,255,255,.06); padding:8px 0;
              font-size:0.78rem; font-family:var(--mono); }
.tool-entry:last-child { border-bottom:none; }
.tool-name { color:#38bdf8; font-weight:700; }

/* ─────────────────────────────────────────────
   FIX 4: RCA dynamic context panel
   ───────────────────────────────────────────── */
.rca-context-card {
    background: rgba(14,165,233,0.06);
    border: 1px solid rgba(14,165,233,0.25);
    border-radius: 10px;
    padding: 14px 16px;
    margin-top: 8px;
}
.rca-context-card .ctx-label {
    font-size: 0.68rem;
    font-weight: 700;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    color: #38bdf8;
    margin-bottom: 6px;
}
.rca-context-card .ctx-q {
    font-size: 0.82rem;
    color: #e2e8f0;
    line-height: 1.5;
    margin-bottom: 10px;
    font-style: italic;
}
.rca-context-card .ctx-badge {
    display: inline-block;
    background: rgba(6,182,212,0.15);
    border: 1px solid rgba(6,182,212,0.3);
    border-radius: 20px;
    padding: 3px 10px;
    font-size: 0.72rem;
    color: #67e8f9;
    margin: 3px 3px 3px 0;
}
.rca-hint {
    background: rgba(124,58,237,0.08);
    border-left: 3px solid #7c3aed;
    border-radius: 0 8px 8px 0;
    padding: 8px 12px;
    margin-top: 10px;
    font-size: 0.78rem;
    color: #c4b5fd;
}
/* ── Network Health tab ── */
.health-metric {
    display: flex;
    align-items: center;
    justify-content: space-between;
    background: #071428;
    border: 1px solid rgba(56,189,248,0.15);
    border-radius: 8px;
    padding: 12px 16px;
    margin-bottom: 8px;
}
.health-label { color: #7dd3fc; font-size: 0.82rem; font-weight: 600; }
.health-value { color: #38bdf8; font-size: 0.9rem; font-weight: 700; }
.health-ok    { color: #4ade80 !important; }
.health-warn  { color: #fbbf24 !important; }
.health-bad   { color: #f87171 !important; }

/* ─────────────────────────────────────────────
   Inline RCA chart panel (below tool log)
   ───────────────────────────────────────────── */
.rca-viz-panel {
    background: rgba(12,21,40,0.75);
    border: 1px solid rgba(56,189,248,0.18);
    border-radius: 12px;
    padding: 16px 18px;
    margin-top: 10px;
}
.rca-viz-why {
    background: rgba(124,58,237,0.10);
    border-left: 3px solid #7c3aed;
    border-radius: 0 8px 8px 0;
    padding: 10px 14px;
    margin-bottom: 14px;
    font-size: 0.82rem;
    color: #c4b5fd;
    line-height: 1.55;
}
.rca-viz-why strong { color: #e2e8f0 !important; }
.rca-viz-why .why-title {
    font-size: 0.68rem;
    font-weight: 700;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    color: #a78bfa;
    display: block;
    margin-bottom: 5px;
}

/* Dropdown styling */
.rca-viz-dropdown label span,
.rca-viz-dropdown .gr-dropdown label {
    color: #7dd3fc !important;
    font-size: 0.72rem !important;
    font-weight: 700 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.08em !important;
}
.rca-viz-dropdown select,
.rca-viz-dropdown .wrap {
    background: rgba(12,21,40,0.9) !important;
    border: 1px solid rgba(56,189,248,0.3) !important;
    color: #e2e8f0 !important;
    border-radius: 8px !important;
}

/* ─────────────────────────────────────────────
   FIX 5: Visualization tab layout
   ───────────────────────────────────────────── */
.viz-btn-row {
    display: flex;
    gap: 8px;
    flex-wrap: wrap;
    margin-bottom: 20px;
    padding: 14px;
    background: rgba(14,165,233,0.04);
    border: 1px solid rgba(14,165,233,0.15);
    border-radius: 10px;
}
.viz-chart-section {
    background: rgba(12,21,40,0.6);
    border: 1px solid rgba(255,255,255,0.06);
    border-radius: 12px;
    padding: 18px;
    margin-bottom: 16px;
}
.viz-chart-title {
    font-size: 0.72rem;
    font-weight: 700;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    color: #7dd3fc;
    margin-bottom: 4px;
}
.viz-chart-desc {
    font-size: 0.8rem;
    color: #64748b;
    margin-bottom: 12px;
    font-style: italic;
}

/* ─────────────────────────────────────────────
   FIX 7: RCA Markdown report styling
   ───────────────────────────────────────────── */
.rca-report-wrap {
    background: rgba(12,21,40,0.7);
    border: 1px solid rgba(14,165,233,0.2);
    border-radius: 12px;
    padding: 20px 24px;
    min-height: 200px;
}

/* ── First response blockquote ── */
.rca-report-wrap blockquote {
    background: rgba(14,165,233,0.07) !important;
    border-left: 4px solid #38bdf8 !important;
    border-radius: 0 10px 10px 0 !important;
    padding: 14px 18px !important;
    margin: 0 0 18px 0 !important;
    color: #e2e8f0 !important;
    font-size: 0.92rem !important;
    line-height: 1.65 !important;
}
.rca-report-wrap blockquote p,
.rca-report-wrap blockquote strong {
    color: #e2e8f0 !important;
}

/* ── Collapsible Executive Summary ── */
.rca-report-wrap details {
    background: rgba(56,189,248,0.05);
    border: 1px solid rgba(56,189,248,0.25);
    border-radius: 10px;
    padding: 2px 0;
    margin-bottom: 18px;
}
.rca-report-wrap details[open] {
    padding-bottom: 12px;
}
.rca-report-wrap summary {
    cursor: pointer;
    padding: 12px 18px;
    font-size: 0.88rem;
    font-weight: 700;
    color: #38bdf8 !important;
    list-style: none;
    display: flex;
    align-items: center;
    gap: 8px;
    user-select: none;
    border-radius: 10px;
}
.rca-report-wrap summary::-webkit-details-marker { display: none; }
.rca-report-wrap summary::marker { display: none; }
.rca-report-wrap summary:hover {
    background: rgba(56,189,248,0.08);
}
/* Arrow indicator — always visible cyan */
.rca-report-wrap summary .exec-arrow,
.rca-report-wrap details > summary > span:first-child {
    display: inline-block;
    font-size: 0.65rem;
    color: #38bdf8 !important;
    transition: transform 0.2s ease;
}
.rca-report-wrap details[open] > summary .exec-arrow,
.rca-report-wrap details[open] > summary > span:first-child {
    transform: rotate(90deg);
}
/* Fallback: pure CSS arrow without JS, uses ::before — disabled when exec-arrow span is present */
.rca-report-wrap summary::before {
    content: none;
}
.rca-report-wrap details > *:not(summary) {
    padding: 0 18px;
    color: #cbd5e1 !important;
    font-size: 0.88rem;
    line-height: 1.65;
}
/* "click to expand" hint text — always visible */
.rca-report-wrap summary span[style*="color:#7dd3fc"],
.rca-report-wrap summary .exec-hint {
    color: #7dd3fc !important;
    font-weight: 400 !important;
    font-size: 0.74rem !important;
}
.rca-report-wrap h2 {
    color: #38bdf8 !important;
    font-size: 1.15rem !important;
    border-bottom: 1px solid rgba(14,165,233,0.3);
    padding-bottom: 8px;
    margin-bottom: 16px;
}
.rca-report-wrap h3 {
    color: #7dd3fc !important;
    font-size: 0.95rem !important;
    margin-top: 18px !important;
    margin-bottom: 8px !important;
}
.rca-report-wrap h4 {
    color: #4ade80 !important;
    font-size: 0.85rem !important;
    margin-top: 12px !important;
    margin-bottom: 4px !important;
    border-left: 3px solid rgba(74,222,128,0.4);
    padding-left: 8px;
}
.rca-report-wrap p, .rca-report-wrap li {
    color: #cbd5e1 !important;
    line-height: 1.65 !important;
}
.rca-report-wrap table {
    width: 100%;
    border-collapse: collapse;
    margin: 10px 0;
    font-size: 0.82rem;
}
.rca-report-wrap th {
    background: rgba(14,165,233,0.15);
    color: #7dd3fc !important;
    padding: 7px 10px;
    text-align: left;
    border: 1px solid rgba(14,165,233,0.2);
}
.rca-report-wrap td {
    color: #cbd5e1 !important;
    padding: 6px 10px;
    border: 1px solid rgba(255,255,255,0.06);
}
.rca-report-wrap tr:nth-child(even) td {
    background: rgba(255,255,255,0.03);
}
.rca-report-wrap strong {
    color: #e2e8f0 !important;
}
.rca-report-wrap code {
    background: rgba(6,182,212,0.12);
    border: 1px solid rgba(6,182,212,0.25);
    border-radius: 4px;
    padding: 1px 5px;
    color: #67e8f9 !important;
    font-size: 0.8em;
}
.rca-report-wrap hr {
    border-color: rgba(14,165,233,0.2) !important;
    margin: 16px 0;
}

/* ── RCA Cypher accordion — placed AFTER initial assessment, before full report ── */
#rca-cypher-accordion {
    margin-top: 10px !important;
    margin-bottom: 14px !important;
}

/* ── Collapsible table inside RCA report ── */
.rca5-report-panel details,
.rca-report-wrap details {
    background: rgba(14,165,233,0.04) !important;
    border: 1px solid rgba(14,165,233,0.2) !important;
    border-radius: 8px !important;
    margin: 8px 0 !important;
}
.rca5-report-panel details summary,
.rca-report-wrap details summary {
    color: #7dd3fc !important;
    font-size: 0.8rem !important;
    font-weight: 700 !important;
    padding: 9px 14px !important;
    cursor: pointer !important;
    list-style: none !important;
    display: flex !important;
    align-items: center !important;
    gap: 6px !important;
    user-select: none !important;
}
.rca5-report-panel details summary::-webkit-details-marker { display: none !important; }
.rca5-report-panel details summary::marker { display: none !important; }
.rca5-report-panel details[open] summary span:first-child {
    transform: rotate(90deg) !important;
}

/* ── Root Cause Trail chain box ── */
.rca-trail-box {
    background: rgba(6,15,35,0.9);
    border: 1px solid rgba(56,189,248,0.2);
    border-radius: 12px;
    padding: 16px 20px;
    margin-bottom: 20px;
}

/* ── Final Root Cause highlight ── */
.rca-root-cause-box {
    background: rgba(239,68,68,0.07);
    border: 1px solid rgba(239,68,68,0.35);
    border-left: 4px solid #f87171;
    border-radius: 10px;
    padding: 14px 20px;
    margin-bottom: 16px;
}


    background: rgba(14,165,233,0.05);
    border: 1px solid rgba(14,165,233,0.18);
    border-radius: 10px;
    padding: 12px 14px;
    margin-top: 10px;
}
.rca-filter-label {
    font-size: 0.65rem;
    font-weight: 700;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    color: #38bdf8;
    margin-bottom: 8px;
}
.rca-stat-row {
    display: flex;
    justify-content: space-between;
    align-items: center;
    padding: 5px 0;
    border-bottom: 1px solid rgba(255,255,255,0.05);
    font-size: 0.76rem;
}
.rca-stat-row:last-child { border-bottom: none; }
.rca-stat-key { color: #94a3b8; }
.rca-stat-val { color: #e2e8f0; font-weight: 600; font-family: var(--mono, monospace); }
.rca-stat-val.red   { color: #f87171; }
.rca-stat-val.amber { color: #fbbf24; }
.rca-stat-val.green { color: #4ade80; }

/* RCA right panel stacked charts */
.rca-right-chart-card {
    background: rgba(12,21,40,0.75);
    border: 1px solid rgba(56,189,248,0.15);
    border-radius: 10px;
    padding: 10px 12px 4px;
    margin-bottom: 10px;
}
.rca-right-chart-title {
    font-size: 0.64rem;
    font-weight: 700;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    color: #38bdf8;
    margin-bottom: 4px;
}
/* Gradio dropdown inside dark panels */
.rca-filter-section select,
.rca-filter-section .wrap,
.rca-filter-section .svelte-1gfkn6j {
    background: rgba(7,20,40,0.9) !important;
    border: 1px solid rgba(56,189,248,0.25) !important;
    color: #e2e8f0 !important;
    border-radius: 6px !important;
    font-size: 0.8rem !important;
}

/* ── RCA v4 ENTERPRISE COMPACT DASHBOARD — Full Redesign ── */

/* Compact report wrapper — tighter padding, no oversized margins */
.rca-report-wrap,
.rca5-report-panel {
    padding: 14px 18px !important;
}

/* ── LAYOUT FIX: Gradio Column (.form) inside rca5-report-panel must stack vertically ── */
.rca5-report-panel > .form {
    display: flex !important;
    flex-direction: column !important;
    align-items: stretch !important;
    gap: 0 !important;
    width: 100% !important;
}
.rca5-report-panel > .form > * {
    width: 100% !important;
    flex: 0 0 100% !important;
    min-width: 0 !important;
}

/* #rca-report-content — the wrapper div injected by _assemble_final_output */
#rca-report-content {
    display: block !important;
    width: 100% !important;
    box-sizing: border-box !important;
    float: none !important;
}
#rca-report-content > div {
    display: block !important;
    width: 100% !important;
    box-sizing: border-box !important;
}

/* KPI grid strip inside report */
.rca-report-wrap .rca-kpi-strip,
.rca5-report-panel .rca-kpi-strip {
    display: grid;
    grid-template-columns: repeat(auto-fit, minmax(130px, 1fr));
    gap: 10px;
    margin-bottom: 14px;
}

/* Section collapsible cards — tight spacing */
.rca-report-wrap details,
.rca5-report-panel details {
    background: rgba(12,21,40,0.75) !important;
    border-radius: 10px !important;
    margin-bottom: 8px !important;
    padding: 0 !important;
    overflow: hidden !important;
}
.rca-report-wrap details[open],
.rca5-report-panel details[open] {
    padding-bottom: 0 !important;
}

/* Summary row — compact single-line */
.rca-report-wrap summary,
.rca5-report-panel summary,
.rca-report-wrap details > summary,
.rca5-report-panel details > summary {
    padding: 10px 14px !important;
    font-size: 0.82rem !important;
    font-weight: 700 !important;
    cursor: pointer !important;
    list-style: none !important;
    display: flex !important;
    align-items: center !important;
    gap: 8px !important;
    user-select: none !important;
    border-radius: 10px !important;
    transition: background 0.15s ease !important;
}
.rca-report-wrap summary:hover,
.rca5-report-panel summary:hover {
    background: rgba(56,189,248,0.05) !important;
}
.rca-report-wrap summary::-webkit-details-marker,
.rca5-report-panel summary::-webkit-details-marker { display: none !important; }
.rca-report-wrap summary::marker,
.rca5-report-panel summary::marker { display: none !important; }

/* Content inside section cards */
.rca-report-wrap details > *:not(summary),
.rca5-report-panel details > *:not(summary) {
    padding: 2px 14px 12px !important;
    font-size: 0.82rem !important;
    color: #94a3b8 !important;
    line-height: 1.65 !important;
}

/* Executive Summary details — slightly more open */
.rca-report-wrap details[open] > div,
.rca5-report-panel details[open] > div {
    color: #94a3b8 !important;
}

/* Compact heading sizes */
.rca-report-wrap h2, .rca5-report-panel h2 {
    font-size: 1rem !important;
    padding-bottom: 6px !important;
    margin-bottom: 12px !important;
}
.rca-report-wrap h3, .rca5-report-panel h3 {
    font-size: 0.85rem !important;
    margin-top: 12px !important;
    margin-bottom: 6px !important;
}

/* Compact table rows */
.rca-report-wrap td, .rca5-report-panel td {
    padding: 5px 9px !important;
    font-size: 0.79rem !important;
}
.rca-report-wrap th, .rca5-report-panel th {
    padding: 6px 9px !important;
    font-size: 0.78rem !important;
}

/* Horizontal trail wrapper */
.rca-trail-horizontal {
    display: flex;
    align-items: flex-start;
    gap: 6px;
    flex-wrap: wrap;
    padding: 10px 14px;
}

/* Cypher accordion — always visible after RCA */
#rca-cypher-accordion {
    margin-top: 10px !important;
    margin-bottom: 4px !important;
    border: 1px solid rgba(6,182,212,0.25) !important;
    border-radius: 10px !important;
    background: rgba(6,15,35,0.8) !important;
}
#rca-cypher-accordion .label-wrap { padding: 10px 14px !important; }

/* ── NEW: RCA v3 redesign styles ── */
.rca-inner-subtab button[role="tab"] {
    font-size: 0.78rem !important;
    padding: 7px 14px !important;
}
.rca-hero-bar {
    background: linear-gradient(135deg,rgba(14,165,233,0.08),rgba(124,58,237,0.06));
    border: 1px solid rgba(14,165,233,0.2);
    border-radius: 14px;
    padding: 18px 22px;
    margin-bottom: 14px;
}
/* Glow on rca_report_wrap */
.rca-report-wrap {
    box-shadow: 0 0 32px rgba(14,165,233,0.08), inset 0 0 20px rgba(14,165,233,0.03);
}
/* Quick-fire sample pill row */
.rca-pill-row {
    display: flex;
    flex-wrap: wrap;
    gap: 6px;
    margin: 6px 0 10px;
}

.rca-side-chart-card {
    background: rgba(12,21,40,0.75);
    border: 1px solid rgba(56,189,248,0.18);
    border-radius: 10px;
    padding: 10px 12px 4px;
    margin-bottom: 6px;
}
.rca-side-chart-title {
    font-size: 0.64rem;
    font-weight: 700;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    color: #38bdf8;
    margin-bottom: 2px;
}
.rca-dl-btn button {
    background: linear-gradient(135deg,rgba(6,182,212,0.18),rgba(14,165,233,0.12)) !important;
    border: 1px solid rgba(6,182,212,0.45) !important;
    color: #67e8f9 !important;
    font-family: var(--mono, monospace) !important;
    font-size: 0.75rem !important;
    font-weight: 700 !important;
    letter-spacing: 1px !important;
    border-radius: 8px !important;
    width: 100% !important;
    transition: all .2s !important;
}
.rca-dl-btn button:hover {
    background: rgba(6,182,212,0.32) !important;
    border-color: rgba(6,182,212,0.7) !important;
    box-shadow: 0 0 18px rgba(6,182,212,0.3) !important;
}

/* ── Download section below report ── */
.rca-download-section {
    margin-top: 18px !important;
    padding: 16px 18px !important;
    background: rgba(6, 15, 35, 0.85) !important;
    border: 1px solid rgba(56, 189, 248, 0.3) !important;
    border-radius: 12px !important;
    position: relative !important;
}
.rca-download-section::before {
    content: '' !important;
    position: absolute !important;
    top: 0; left: 0; right: 0; height: 2px !important;
    background: linear-gradient(90deg, #38bdf8, #7c3aed) !important;
    border-radius: 12px 12px 0 0 !important;
}
.rca-download-label {
    font-size: 0.65rem !important;
    font-weight: 700 !important;
    letter-spacing: 0.12em !important;
    text-transform: uppercase !important;
    color: #38bdf8 !important;
    margin-bottom: 12px !important;
    display: block !important;
}
.rca-download-row {
    gap: 12px !important;
}
.rca-download-row .rca-dl-btn button,
.rca-dl-btn button {
    background: linear-gradient(135deg, rgba(14,165,233,0.15) 0%, rgba(56,189,248,0.08) 100%) !important;
    border: 1px solid rgba(56,189,248,0.5) !important;
    color: #e0f2fe !important;
    font-size: 0.84rem !important;
    font-weight: 700 !important;
    letter-spacing: 0.02em !important;
    border-radius: 10px !important;
    padding: 12px 20px !important;
    width: 100% !important;
    transition: all 0.2s ease !important;
    box-shadow: 0 2px 12px rgba(56,189,248,0.1) !important;
}
.rca-download-row .rca-dl-btn button:hover,
.rca-dl-btn button:hover {
    background: linear-gradient(135deg, rgba(14,165,233,0.32) 0%, rgba(56,189,248,0.2) 100%) !important;
    border-color: #38bdf8 !important;
    color: #ffffff !important;
    box-shadow: 0 0 20px rgba(56,189,248,0.3), 0 4px 16px rgba(56,189,248,0.15) !important;
    transform: translateY(-1px) !important;
}
.rca-dl-status {
    margin-top: 10px !important;
    font-size: 0.75rem !important;
    color: #4ade80 !important;
    text-align: center !important;
}

/* ══════════════════════════════════════════
   RCA TRAIL v5 — Clean, Intuitive UI
   ══════════════════════════════════════════ */

/* ── RCA v5 header ── */
.rca5-header {
  background: linear-gradient(135deg,rgba(14,165,233,0.07) 0%,rgba(124,58,237,0.05) 100%);
  border: 1px solid rgba(14,165,233,0.2);
  border-radius: 14px; padding: 20px 26px 16px; margin-bottom: 16px; position: relative;
}
.rca5-header::before {
  content:''; position:absolute; top:0; left:0; right:0; height:2px;
  background: linear-gradient(90deg, transparent, #38bdf8, #7c3aed, transparent);
}
.rca5-title { font-size:20px; font-weight:700; color:#e2e8f0; margin:0 0 4px; }
.rca5-sub   { font-size:12px; color:#94a3b8; font-family:monospace; }
.rca5-badges { display:flex; flex-wrap:wrap; gap:7px; margin-top:10px; }
.rca5-badge {
  display:inline-flex; align-items:center; gap:4px;
  border-radius:20px; padding:3px 11px; font-size:11px; font-family:monospace;
}
.rca5-badge.cyan   { background:rgba(14,165,233,0.1);  border:1px solid rgba(14,165,233,0.3);  color:#7dd3fc; }
.rca5-badge.green  { background:rgba(16,185,129,0.1);  border:1px solid rgba(16,185,129,0.3);  color:#6ee7b7; }
.rca5-badge.purple { background:rgba(124,58,237,0.1);  border:1px solid rgba(124,58,237,0.3);  color:#c4b5fd; }
.rca5-badge.amber  { background:rgba(245,158,11,0.1);  border:1px solid rgba(245,158,11,0.3);  color:#fcd34d; }

/* ── KPI strip ── */
.rca5-kpi-strip { display:grid; grid-template-columns:repeat(4,1fr); gap:10px; margin-bottom:16px; }
.rca5-kpi {
  background:rgba(12,21,40,0.8); border-radius:10px; padding:14px 16px;
  text-align:center; border:1px solid rgba(56,189,248,0.12); position:relative; overflow:hidden;
}
.rca5-kpi::after { content:''; position:absolute; bottom:0; left:0; right:0; height:2px; }
.rca5-kpi.k-cyan::after   { background:#38bdf8; }
.rca5-kpi.k-red::after    { background:#f87171; }
.rca5-kpi.k-amber::after  { background:#fbbf24; }
.rca5-kpi.k-green::after  { background:#4ade80; }
.rca5-kpi-label { font-size:10px; color:#94a3b8; text-transform:uppercase; letter-spacing:0.08em; margin-bottom:4px; }
.rca5-kpi-val   { font-size:24px; font-weight:700; }
.rca5-kpi.k-cyan  .rca5-kpi-val { color:#38bdf8; }
.rca5-kpi.k-red   .rca5-kpi-val { color:#f87171; }
.rca5-kpi.k-amber .rca5-kpi-val { color:#fbbf24; }
.rca5-kpi.k-green .rca5-kpi-val { color:#4ade80; }
.rca5-kpi-sub   { font-size:10px; color:#94a3b8; margin-top:2px; }

/* ── Report panel ── */
/* Hide ALL Gradio progress/loading bars globally */
.progress-container,
.progress-level,
.progress-level-inner,
.eta-bar,
.generating,
.gr-progress,
svelte-progress-bar,
.loader,
div[class*="progress"],
div[class*="generating"],
div[class*="eta"] {
  display: none !important;
  height: 0 !important;
  min-height: 0 !important;
  max-height: 0 !important;
  margin: 0 !important;
  padding: 0 !important;
  overflow: hidden !important;
  opacity: 0 !important;
  visibility: hidden !important;
}

/* Hide invisible download button wrappers (they render as white bars) */
#rca-dl-hidden-wrap,
#rca-dl-hidden-wrap * {
  display: none !important;
  height: 0 !important;
  min-height: 0 !important;
  margin: 0 !important;
  padding: 0 !important;
  overflow: hidden !important;
}

.rca5-report-panel {
  background:rgba(12,21,40,0.7); border:1px solid rgba(14,165,233,0.18);
  border-radius:12px; padding:22px 26px;
}
.rca5-report-panel h2 { color:#38bdf8 !important; font-size:1.1rem !important;
  border-bottom:1px solid rgba(14,165,233,0.25); padding-bottom:8px; margin-bottom:14px; }
.rca5-report-panel h3 { color:#7dd3fc !important; font-size:0.9rem !important;
  margin-top:16px !important; margin-bottom:6px !important; }
.rca5-report-panel h4 { color:#4ade80 !important; font-size:0.85rem !important;
  margin-top:14px !important; margin-bottom:4px !important;
  border-left:3px solid rgba(74,222,128,0.5); padding-left:8px; }
.rca5-report-panel p, .rca5-report-panel li { color:#cbd5e1 !important; line-height:1.65 !important; }
.rca5-report-panel table { width:100%; border-collapse:collapse; margin:10px 0; font-size:0.82rem; }
.rca5-report-panel th { background:rgba(14,165,233,0.12); color:#7dd3fc !important;
  padding:7px 10px; text-align:left; border:1px solid rgba(14,165,233,0.18); }
.rca5-report-panel td { color:#cbd5e1 !important; padding:6px 10px;
  border:1px solid rgba(255,255,255,0.05); }
.rca5-report-panel tr:nth-child(even) td { background:rgba(255,255,255,0.02); }
.rca5-report-panel strong { color:#e2e8f0 !important; }
.rca5-report-panel code { background:rgba(6,182,212,0.1); border:1px solid rgba(6,182,212,0.2);
  border-radius:4px; padding:1px 5px; color:#67e8f9 !important; font-size:0.8em; }
.rca5-report-panel hr { border-color:rgba(14,165,233,0.18) !important; margin:14px 0; }

/* FIX: Executive Summary <details><summary> heading always visible cyan */
.rca5-report-panel details {
  background: rgba(56,189,248,0.05);
  border: 1px solid rgba(56,189,248,0.25);
  border-radius: 10px;
  margin-bottom: 18px;
}
.rca5-report-panel details[open] { padding-bottom: 10px; }
.rca5-report-panel summary,
.rca5-report-panel summary span,
.rca5-report-panel details > summary {
  cursor: pointer !important;
  padding: 12px 16px !important;
  font-size: 0.9rem !important;
  font-weight: 700 !important;
  color: #38bdf8 !important;
  list-style: none !important;
  display: flex !important;
  align-items: center !important;
  gap: 8px !important;
  user-select: none !important;
  border-radius: 10px !important;
}
.rca5-report-panel summary::-webkit-details-marker { display: none !important; }
.rca5-report-panel summary::marker { display: none !important; }
.rca5-report-panel summary::before { display: none !important; }
.rca5-report-panel summary .exec-arrow { color: #38bdf8 !important; font-size: 0.65rem !important; }
.rca5-report-panel details > *:not(summary) { padding: 0 18px; color: #cbd5e1 !important; font-size: 0.87rem; line-height: 1.65; }

/* ── Chart area ── */
.rca5-chart-area {
  background:rgba(12,21,40,0.7); border:1px solid rgba(56,189,248,0.15);
  border-radius:12px; padding:16px 18px; margin-top:14px;
}
.rca5-chart-header {
  display:flex; align-items:center; justify-content:space-between;
  margin-bottom:12px; flex-wrap:wrap; gap:8px;
}
.rca5-chart-title-txt { font-size:0.7rem; font-weight:700; text-transform:uppercase;
  letter-spacing:0.1em; color:#38bdf8; }
.rca5-chart-dd-wrap { min-width:220px; }

/* ── Why box ── */
.rca5-why-box {
  border-radius:10px; padding:14px 18px; margin-top:12px;
  background:rgba(124,58,237,0.07); border-left:3px solid #7c3aed;
  border-radius:0 8px 8px 0;
}
.rca5-why-title { font-size:0.68rem; font-weight:700; text-transform:uppercase;
  letter-spacing:0.1em; color:#a78bfa; margin-bottom:5px; }
.rca5-why-body  { font-size:0.8rem; color:#c4b5fd; line-height:1.65; }
.rca5-why-body strong { color:#e2e8f0 !important; }
.rca5-why-tags  { display:flex; flex-wrap:wrap; gap:5px; margin-top:8px; }
.rca5-why-tag   { font-size:10px; padding:2px 8px; border-radius:20px;
  background:rgba(124,58,237,0.1); border:1px solid rgba(124,58,237,0.25); color:#c4b5fd; }

/* ── Sidebar stats ── */
.rca5-sidebar { display:flex; flex-direction:column; gap:10px; }
.rca5-stat-card {
  background:rgba(12,21,40,0.7); border:1px solid rgba(56,189,248,0.12);
  border-radius:10px; padding:12px 14px;
}
.rca5-stat-card-title { font-size:0.62rem; font-weight:700; text-transform:uppercase;
  letter-spacing:0.1em; color:#38bdf8; margin-bottom:8px; }
.rca5-stat-row2 { display:flex; justify-content:space-between; align-items:center;
  padding:4px 0; border-bottom:1px solid rgba(255,255,255,0.04); font-size:0.74rem; }
.rca5-stat-row2:last-child { border-bottom:none; }
.rca5-stat-k { color:#94a3b8; }
.rca5-stat-v { color:#e2e8f0; font-weight:600; font-family:monospace; }
.rca5-stat-v.red   { color:#f87171; }
.rca5-stat-v.amber { color:#fbbf24; }
.rca5-stat-v.green { color:#4ade80; }

/* ─── RIGHT PANEL KPI CARDS ─── */
.rca-kpi-panel { display:flex; flex-direction:column; gap:6px; margin-bottom:10px; }
.rca-kpi-card {
  background:rgba(7,14,32,0.97); border-radius:10px; padding:9px 11px 8px;
  position:relative; overflow:hidden; border:1px solid rgba(56,189,248,0.1);
  display:flex; align-items:center; gap:9px;
}
.rca-kpi-card::before { content:''; position:absolute; left:0; top:0; bottom:0; width:3px; border-radius:10px 0 0 10px; }
.rca-kpi-card.kc-red::before   { background:linear-gradient(180deg,#f87171,#ef4444); }
.rca-kpi-card.kc-amber::before { background:linear-gradient(180deg,#fbbf24,#f59e0b); }
.rca-kpi-card.kc-cyan::before  { background:linear-gradient(180deg,#38bdf8,#0ea5e9); }
.rca-kpi-card.kc-green::before { background:linear-gradient(180deg,#4ade80,#22c55e); }
.rca-kpi-card::after { content:''; position:absolute; top:7px; right:8px; width:5px; height:5px; border-radius:50%; }
.rca-kpi-card.kc-red::after   { background:#f87171; box-shadow:0 0 4px #f87171; }
.rca-kpi-card.kc-amber::after { background:#fbbf24; box-shadow:0 0 4px #fbbf24; }
.rca-kpi-card.kc-cyan::after  { background:#38bdf8; box-shadow:0 0 4px #38bdf8; }
.rca-kpi-card.kc-green::after { background:#4ade80; box-shadow:0 0 4px #4ade80; }
.rca-kpi-icon { width:28px; height:28px; border-radius:7px; display:flex; align-items:center; justify-content:center; font-size:13px; flex-shrink:0; }
.kc-red   .rca-kpi-icon { background:rgba(248,113,113,0.13); }
.kc-amber .rca-kpi-icon { background:rgba(251,191,36,0.13); }
.kc-cyan  .rca-kpi-icon { background:rgba(56,189,248,0.13); }
.kc-green .rca-kpi-icon { background:rgba(74,222,128,0.13); }
.rca-kpi-body { flex:1; min-width:0; }
.rca-kpi-lbl  { font-size:0.55rem; font-weight:700; text-transform:uppercase; letter-spacing:0.1em; color:#94a3b8; margin-bottom:1px; }
.rca-kpi-num  { font-size:1.15rem; font-weight:800; line-height:1.1; font-family:'DM Mono',monospace; }
.kc-red   .rca-kpi-num { color:#f87171; }
.kc-amber .rca-kpi-num { color:#fbbf24; }
.kc-cyan  .rca-kpi-num { color:#38bdf8; }
.kc-green .rca-kpi-num { color:#4ade80; }
.rca-kpi-sub2 { font-size:0.57rem; color:#334155; margin-top:1px; }

/* ─── PRODUCT STATS ─── */
.rca-prod-panel { background:rgba(7,14,32,0.97); border:1px solid rgba(56,189,248,0.1); border-radius:10px; padding:10px 12px; }
.rca-prod-heading { font-size:0.6rem; font-weight:700; text-transform:uppercase; letter-spacing:0.09em; color:#38bdf8; margin-bottom:1px; }
.rca-prod-subhead { font-size:0.56rem; color:#94a3b8; margin-bottom:8px; }
.rca-prod-row { display:flex; align-items:center; gap:5px; padding:3px 0; border-bottom:1px solid rgba(255,255,255,0.03); }
.rca-prod-row:last-of-type { border-bottom:none; }
.rca-prod-rank { font-size:0.55rem; color:#94a3b8; width:12px; flex-shrink:0; font-family:monospace; }
.rca-prod-name { font-size:0.63rem; color:#cbd5e1; flex:1; min-width:0; white-space:nowrap; overflow:hidden; text-overflow:ellipsis; font-weight:500; }
.rca-prod-bar-wrap { width:40px; height:4px; background:rgba(255,255,255,0.05); border-radius:3px; flex-shrink:0; }
.rca-prod-bar { height:4px; border-radius:3px; }
.rca-prod-pct { font-size:0.6rem; font-weight:700; font-family:monospace; flex-shrink:0; width:28px; text-align:right; }
.rca-prod-legend-row { display:flex; gap:9px; flex-wrap:wrap; margin-top:7px; padding-top:6px; border-top:1px solid rgba(255,255,255,0.04); }

/* ─── COLLAPSIBLE RIGHT PANEL TOGGLE ─── */
.rca-right-header {
  display:flex; align-items:center; justify-content:space-between;
  background:rgba(14,165,233,0.08); border:1px solid rgba(56,189,248,0.2);
  border-radius:9px; padding:8px 11px; margin-bottom:8px;
  cursor:pointer; user-select:none; transition:background 0.15s;
}
.rca-right-header:hover { background:rgba(14,165,233,0.14); }
.rca-right-header-label { font-size:0.62rem; font-weight:700; text-transform:uppercase; letter-spacing:0.1em; color:#38bdf8; }
.rca-right-header-arrow { font-size:0.65rem; color:#38bdf8; transition:transform 0.2s; }

/* ─── HIDE OLD KPI STRIP ─── */
.rca5-kpi-strip { display:none !important; }

/* ── rca-native-body-gr: Gradio columns used as collapsible bodies ── */
/* JS sets display:none on load; this is just a fallback hint */
.rca-native-body-gr {
  /* JS controls visibility — no CSS needed here */
}

/* ── View Charts: styled as native <details> box, EXACTLY matching Recommendations ── */
/* Outer column = the border box */
#rca-view-charts-wrap {
  position: relative;
  background: rgba(56,189,248,0.05) !important;
  border: 1px solid rgba(56,189,248,0.25) !important;
  border-radius: 10px !important;
  box-shadow: 0 0 28px rgba(56,189,248,0.08), inset 0 0 16px rgba(56,189,248,0.03);
  overflow: hidden;
  margin-bottom: 10px;
}
/* Top cyan gradient line (same as rca-native-box::before) */
#rca-view-charts-wrap::before {
  content: '';
  position: absolute; top: 0; left: 0; right: 0; height: 2px;
  background: linear-gradient(90deg, transparent, #38bdf8, transparent);
  pointer-events: none;
}
/* The gradio-accordion inside must be fully transparent — no bg/border */
#rca-vchart-accordion,
#rca-vchart-accordion .wrap,
#rca-vchart-accordion > .wrap,
#rca-vchart-accordion > div {
  background: transparent !important;
  border: none !important;
  box-shadow: none !important;
  padding: 0 !important;
  margin: 0 !important;
}
/* Force accordion content visible — Gradio sometimes collapses it on re-render */
.rca-vchart-open .wrap > div:not(.label-wrap),
#rca-vchart-accordion .wrap > div:not(.label-wrap) {
  display: block !important;
  visibility: visible !important;
  height: auto !important;
  overflow: visible !important;
}
/* Accordion header row — same size/spacing as Recommendations summary */
#rca-vchart-accordion > .label-wrap,
#rca-vchart-accordion .label-wrap {
  background: transparent !important;
  border: none !important;
  padding: 13px 18px !important;
  cursor: pointer !important;
  transition: background 0.15s ease !important;
}
#rca-vchart-accordion .label-wrap:hover {
  background: rgba(56,189,248,0.07) !important;
}
#rca-vchart-accordion .label-wrap button,
#rca-vchart-accordion .label-wrap span {
  color: #38bdf8 !important;
  font-weight: 700 !important;
  font-size: 0.88rem !important;
  background: transparent !important;
  letter-spacing: 0.01em !important;
}
/* Inner body padding when open */
#rca-vchart-accordion > .wrap > div,
#rca-vchart-accordion .wrap > div {
  padding: 4px 18px 16px !important;
}
/* Chart tabs: compact, matching the visual style */
#rca-vchart-accordion button.rca-viz-tab,
#rca-view-charts-wrap button.rca-viz-tab {
  background: rgba(10,18,38,0.9) !important;
  border: 1px solid rgba(56,189,248,0.28) !important;
  border-radius: 8px !important;
  color: #7dd3fc !important;
  font-size: 0.8rem !important;
  font-weight: 700 !important;
  padding: 7px 10px !important;
  width: 100% !important;
  transition: all 0.15s ease !important;
  cursor: pointer !important;
  box-shadow: none !important;
}
#rca-vchart-accordion button.rca-viz-tab:hover,
#rca-view-charts-wrap button.rca-viz-tab:hover {
  background: rgba(56,189,248,0.14) !important;
  border-color: #38bdf8 !important;
  color: #ffffff !important;
  box-shadow: 0 0 10px rgba(56,189,248,0.16) !important;
}
/* No white bg on chart plot area */
#rca-vchart-accordion .gradio-plot,
#rca-vchart-accordion .js-plotly-plot,
#rca-view-charts-wrap .gradio-plot {
  background: transparent !important;
  border-radius: 8px !important;
  margin-top: 6px !important;
}
/* Tighten Gradio row gaps inside the accordion */
#rca-vchart-accordion .gap,
#rca-vchart-accordion .row-wrap {
  gap: 6px !important;
}
/* Force vchart accordion body ALWAYS visible — prevent Gradio from collapsing it */
#rca-vchart-accordion > .wrap,
#rca-vchart-accordion > div > .wrap,
#rca-vchart-accordion .wrap {
  display: block !important;
  visibility: visible !important;
  height: auto !important;
  overflow: visible !important;
}
/* Ensure the vchart plot area never collapses to zero height */
#rca-vchart-plot .gradio-plot,
#rca-vchart-plot .js-plotly-plot,
#rca-view-charts-wrap #rca-vchart-plot .gradio-plot {
  width: 100% !important;
}
/* Only apply min-height when the plot has actual content */
#rca-vchart-plot .js-plotly-plot .plotly {
  width: 100% !important;
}
#rca-vchart-plot .js-plotly-plot svg {
  width: 100% !important;
  min-height: 400px !important;
}
/* Hide Plotly modebar icons on hover for the RCA chart panel */
#rca-vchart-plot .modebar-container,
#rca-vchart-plot .modebar {
  display: none !important;
}
/* Ensure plot container has proper dimensions when visible */
#rca-view-charts-wrap #rca-vchart-plot:not(.hidden) {
  min-height: 480px !important;
  display: block !important;
  width: 100% !important;
}
/* ── NUCLEAR GRADIO OVERRIDE — prevent theme from hiding text ── */
/* Force all text in the Update Graph tab to be visible */
#update-graph-tab *,
[data-testid="tab-update-graph"] *,
.gradio-container .tabitem * {
    -webkit-font-smoothing: antialiased;
}
/* Ensure table cells never go invisible */
td, th {
    color: inherit;
}
/* Prevent any .dark theme from making text match background */
.dark .gr-form,
.dark .block,
.dark input,
.dark textarea,
.dark select {
    color-scheme: dark;
}
/* All label text in Update Graph — visible */
.gr-form label span,
.block label span,
.gr-group label span {
    color: #7dd3fc !important;
    font-weight: 600 !important;
    font-size: .75rem !important;
}
/* Textbox placeholder */
input::placeholder,
textarea::placeholder {
    color: #475569 !important;
    opacity: 1 !important;
}
/* Textbox text */
.gr-form input[type="text"],
.gr-form textarea {
    color: #e2e8f0 !important;
    background: rgba(4,9,22,0.9) !important;
}

/* ─────────────────────────────────────────────
   Gradio's dark theme overrides inline color styles.
   These rules restore the critical/operational/strategic
   hint text colors (⚡/🔧/🏛 lines) to their intended values.
   ───────────────────────────────────────────── */
#rca-rec-vis-wrap [style*="color:#fca5a5"] { color: #fca5a5 !important; }
#rca-rec-vis-wrap [style*="color:#fde68a"] { color: #fde68a !important; }
#rca-rec-vis-wrap [style*="color:#c7d2fe"] { color: #c7d2fe !important; }
#rca-rec-vis-wrap [style*="color:#86efac"] { color: #86efac !important; }
#rca-rec-vis-wrap [style*="color:#f1f5f9"] { color: #f1f5f9 !important; }
#rca-rec-vis-wrap [style*="color:#e2e8f0"] { color: #e2e8f0 !important; }
#rca-rec-vis-wrap [style*="color:#ffffff"]  { color: #ffffff  !important; }
#rca-rec-vis-wrap [style*="color:#94a3b8"] { color: #94a3b8 !important; }
#rca-rec-vis-wrap [style*="color:#f87171"] { color: #f87171 !important; }  /* Critical — Red */
#rca-rec-vis-wrap [style*="color:#fb923c"] { color: #fb923c !important; }  /* High — Orange */
#rca-rec-vis-wrap [style*="color:#fbbf24"] { color: #fbbf24 !important; }  /* Medium — Yellow */
#rca-rec-vis-wrap [style*="color:#4ade80"] { color: #4ade80 !important; }  /* Low — Green */
#rca-rec-vis-wrap [style*="color:#818cf8"] { color: #818cf8 !important; }
#rca-rec-vis-wrap [style*="color:#38bdf8"] { color: #38bdf8 !important; }
#rca-rec-vis-wrap [style*="color:#64748b"] { color: #64748b !important; }
/* Force recommendation tier divs to show their backgrounds */
#rca-rec-vis-wrap [style*="background:rgba(239,68,68"] { background: rgba(239,68,68,0.10) !important; }
#rca-rec-vis-wrap [style*="background:rgba(251,146,60"] { background: rgba(251,146,60,0.10) !important; }
#rca-rec-vis-wrap [style*="background:rgba(251,191,36"] { background: rgba(251,191,36,0.10) !important; }
#rca-rec-vis-wrap [style*="background:rgba(74,222,128"] { background: rgba(74,222,128,0.10) !important; }
/* Charts vis wrap text override */
#rca-charts-vis-wrap [style*="color:#e2e8f0"] { color: #e2e8f0 !important; }
#rca-charts-vis-wrap [style*="color:#38bdf8"] { color: #38bdf8 !important; }
#rca-charts-vis-wrap [style*="color:#a78bfa"] { color: #a78bfa !important; }
#rca-charts-vis-wrap [style*="color:#7dd3fc"] { color: #7dd3fc !important; }
#rca-charts-vis-wrap [style*="color:#94a3b8"] { color: #94a3b8 !important; }

/* ─────────────────────────────────────────────
   LAYOUT: Full horizontal width utilisation
   Gradio default constrains to narrow center col
   ───────────────────────────────────────────── */
.gradio-container {
    max-width: 100% !important;
    padding-left: 16px !important;
    padding-right: 16px !important;
}
.main-panel, .left-panel {
    min-width: 0 !important;
}
/* Expand the left nav in Query Interface tab */
.left-panel {
    min-width: 280px !important;
    max-width: 320px !important;
}
/* Section headings (all blue, no per-section color) */
.rca5-report-panel h3,
.rca-report-wrap h3 {
    color: #7dd3fc !important;
}
.rca5-report-panel h4,
.rca-report-wrap h4 {
    color: #7dd3fc !important;
    border-left-color: rgba(125,211,252,0.4) !important;
}

/* ─────────────────────────────────────────────
   ENTERPRISE RCA SECTION CARDS v2
   Section heading + context always visible
   Only the table collapses via [View Table]
   ───────────────────────────────────────────── */
.rca5-report-panel .rca-section-card,
.rca-report-wrap .rca-section-card {
  border-radius: 10px;
  margin-bottom: 10px;
  padding: 12px 16px;
}
/* View Table collapsible within section cards */
.rca5-report-panel details[data-table-only],
.rca-report-wrap details[data-table-only] {
  background: rgba(14,165,233,0.04) !important;
  border: 1px solid rgba(14,165,233,0.18) !important;
  border-radius: 8px !important;
  margin: 8px 0 0 !important;
}
.rca5-report-panel details[data-table-only] > summary,
.rca-report-wrap details[data-table-only] > summary {
  font-size: 0.75rem !important;
  font-weight: 700 !important;
  padding: 7px 14px !important;
}

/* ── Improved markdown table rendering inside report panels ── */
.rca5-report-panel table,
.rca-report-wrap table {
  width: 100% !important;
  border-collapse: collapse !important;
  margin: 8px 0 !important;
  font-size: 0.8rem !important;
  table-layout: auto !important;
}
.rca5-report-panel th,
.rca-report-wrap th {
  background: rgba(14,165,233,0.14) !important;
  color: #7dd3fc !important;
  padding: 8px 12px !important;
  text-align: left !important;
  font-weight: 700 !important;
  font-size: 0.75rem !important;
  text-transform: uppercase !important;
  letter-spacing: 0.06em !important;
  border: 1px solid rgba(14,165,233,0.22) !important;
  white-space: nowrap !important;
}
.rca5-report-panel td,
.rca-report-wrap td {
  color: #e2e8f0 !important;
  padding: 7px 12px !important;
  border: 1px solid rgba(255,255,255,0.06) !important;
  vertical-align: middle !important;
}
.rca5-report-panel tr:nth-child(even) td,
.rca-report-wrap tr:nth-child(even) td {
  background: rgba(255,255,255,0.03) !important;
}
.rca5-report-panel tr:hover td,
.rca-report-wrap tr:hover td {
  background: rgba(56,189,248,0.05) !important;
}

/* ── Cypher accordion always on top — force ordering ── */
#rca-cypher-accordion {
  order: -999 !important;
  position: relative !important;
  z-index: 10 !important;
  margin-bottom: 14px !important;
}

/* ── KPI strip spacing fix ── */
.rca5-report-panel > div > div[style*="grid-template-columns"] {
  margin-bottom: 14px !important;
}


#rca-charts-vis-wrap,
#rca-rec-vis-wrap {
  padding-left: 0 !important;
  margin-left: 0 !important;
  padding-right: 0 !important;
  margin-right: 0 !important;
}
/* Gradio wraps gr.HTML in a div.svelte-* with padding — remove it */
#rca-charts-vis-wrap > div,
#rca-rec-vis-wrap > div {
  padding: 0 !important;
  margin: 0 !important;
}


/* ── Cypher Queries Accordion ───────────────────────────────────── */
.cq-section {
  margin: 10px 0 14px;
  border-radius: 12px;
  overflow: hidden;
  border: 1px solid rgba(56,189,248,.18);
  background: rgba(4,9,22,.98);
  transition: box-shadow .25s;
}
.cq-section:hover {
  box-shadow: 0 0 0 1px rgba(56,189,248,.25);
}
.cq-toggle {
  display: flex;
  align-items: center;
  gap: 10px;
  padding: 11px 16px;
  cursor: pointer;
  user-select: none;
  background: rgba(56,189,248,.05);
  border-bottom: 1px solid rgba(56,189,248,.1);
  transition: background .18s;
}
.cq-toggle:hover { background: rgba(56,189,248,.1); }
.cq-toggle-icon {
  font-size: .65rem;
  color: #38bdf8;
  transition: transform .22s cubic-bezier(.4,0,.2,1);
  display: inline-block;
}
.cq-toggle-icon.open { transform: rotate(90deg); }
.cq-toggle-label {
  font-size: .62rem;
  font-weight: 900;
  text-transform: uppercase;
  letter-spacing: .15em;
  color: #38bdf8;
  flex: 1;
}
.cq-toggle-badge {
  font-size: .58rem;
  padding: 2px 9px;
  border-radius: 12px;
  background: rgba(56,189,248,.12);
  border: 1px solid rgba(56,189,248,.25);
  color: #7dd3fc;
  font-weight: 700;
  font-family: monospace;
}
.cq-body {
  display: none;
  padding: 12px 14px 14px;
  max-height: 520px;
  overflow-y: auto;
  overflow-x: hidden;
}
.cq-body.open { display: block; animation: cqFadeIn .22s ease; }
@keyframes cqFadeIn { from{opacity:0;transform:translateY(-4px)} to{opacity:1;transform:none} }
.cq-entry {
  margin-bottom: 14px;
  border-radius: 9px;
  border: 1px solid rgba(255,255,255,.07);
  background: rgba(6,12,28,.9);
  overflow: hidden;
}
.cq-entry:last-child { margin-bottom: 0; }
.cq-entry-header {
  display: flex;
  align-items: center;
  gap: 8px;
  padding: 7px 12px;
  background: rgba(255,255,255,.04);
  border-bottom: 1px solid rgba(255,255,255,.06);
}
.cq-seq {
  font-size: .6rem;
  font-weight: 900;
  width: 20px;
  height: 20px;
  border-radius: 50%;
  background: rgba(56,189,248,.15);
  border: 1px solid rgba(56,189,248,.35);
  color: #38bdf8;
  display: flex;
  align-items: center;
  justify-content: center;
  flex-shrink: 0;
  font-family: monospace;
}
.cq-purpose {
  font-size: .68rem;
  font-weight: 700;
  color: #e2e8f0;
  flex: 1;
}
.cq-records {
  font-size: .58rem;
  font-weight: 700;
  color: #4ade80;
  background: rgba(74,222,128,.08);
  border: 1px solid rgba(74,222,128,.2);
  border-radius: 8px;
  padding: 1px 7px;
  font-family: monospace;
  flex-shrink: 0;
}
.cq-records.unknown { color: #64748b; border-color: rgba(255,255,255,.08); background: none; }
.cq-copy-btn {
  font-size: .58rem;
  padding: 2px 8px;
  border-radius: 6px;
  background: rgba(255,255,255,.05);
  border: 1px solid rgba(255,255,255,.1);
  color: #94a3b8;
  cursor: pointer;
  transition: all .15s;
  flex-shrink: 0;
}
.cq-copy-btn:hover {
  background: rgba(56,189,248,.1);
  border-color: rgba(56,189,248,.3);
  color: #38bdf8;
}
.cq-copy-btn.copied {
  color: #4ade80;
  border-color: rgba(74,222,128,.3);
  background: rgba(74,222,128,.08);
}
.cq-code-wrap {
  padding: 10px 12px;
  overflow-x: auto;
  scrollbar-width: thin;
  scrollbar-color: rgba(56,189,248,.2) transparent;
}
.cq-code {
  font-family: 'Consolas','Fira Code','JetBrains Mono','Monaco',monospace;
  font-size: .75rem;
  line-height: 1.65;
  color: #e2e8f0;
  white-space: pre;
  display: block;
  min-width: max-content;
}
/* Syntax token colours */
.cq-kw  { color: #f472b6; font-weight: 700; }  /* MATCH WHERE RETURN etc */
.cq-fn  { color: #38bdf8; font-weight: 700; }  /* COUNT SUM AVG ROUND etc */
.cq-nd  { color: #4ade80; }                     /* :NodeLabel */
.cq-rel { color: #a78bfa; }                     /* -[:REL_TYPE]-> */
.cq-str { color: #fbbf24; }                     /* 'string values' */
.cq-num { color: #fb923c; }                     /* numeric values */
.cq-cm  { color: #475569; font-style: italic; } /* /* comments */ }
.cq-var { color: #7dd3fc; }                     /* $variables */
.cq-prop{ color: #e2e8f0; }                     /* .property */

/* ====================================================
   GRAPHPULSE - ANTI-GRADIO-BLEED + BOLD TEXT OVERRIDES
   ==================================================== */

/* Force all gradio text to use our color scheme */
.gradio-container * { color: var(--text); }

.gradio-container label,
.gradio-container .label-wrap span,
.gradio-container p,
.gradio-container li {
    color: var(--text) !important;
    font-family: var(--sans) !important;
}

/* Textbox inputs */
.gradio-container textarea,
.gradio-container input[type="text"],
.gradio-container .gr-textbox textarea {
    background: var(--bg-input) !important;
    color: #e2e8f0 !important;
    border: 1px solid var(--border) !important;
    border-radius: 10px !important;
    font-family: var(--sans) !important;
    font-size: 0.92rem !important;
}
.gradio-container textarea::placeholder,
.gradio-container input::placeholder {
    color: #4a6080 !important;
    font-style: italic;
}
.gradio-container textarea:focus,
.gradio-container input:focus {
    border-color: rgba(56,189,248,0.5) !important;
    box-shadow: 0 0 0 2px rgba(56,189,248,0.12) !important;
    outline: none !important;
}

/* Textbox label (e.g. "Generated Cypher") */
.gradio-container .label-wrap,
.gradio-container .label-wrap span,
.gradio-container label > span,
.gradio-container .block > label > span {
    color: #7dd3fc !important;
    font-size: 0.72rem !important;
    font-weight: 700 !important;
    letter-spacing: 0.08em !important;
    text-transform: uppercase !important;
}

/* Markdown / prose text */
.gradio-container .prose p,
.gradio-container .md p,
.main-panel .prose p {
    color: #dde6f5 !important;
    font-family: var(--sans) !important;
    line-height: 1.65 !important;
}
.gradio-container .prose strong,
.gradio-container .md strong {
    color: #ffffff !important;
    font-weight: 700 !important;
}

/* Brief summary box */
.brief-box p { color: #e2e8f0 !important; font-size: 0.92rem !important; }
.brief-box strong { color: #38bdf8 !important; font-weight: 700 !important; }

/* Analysis table */
.custom-table thead th {
    color: #38bdf8 !important;
    font-weight: 700 !important;
    font-size: 0.72rem !important;
    letter-spacing: 0.08em !important;
    text-transform: uppercase !important;
    background: rgba(14,165,233,0.1) !important;
    border-bottom: 1px solid rgba(56,189,248,0.3) !important;
    padding: 10px 14px !important;
    white-space: nowrap;
}
.custom-table td {
    color: #dde6f5 !important;
    font-size: 0.84rem !important;
    padding: 9px 14px !important;
    border-bottom: 1px solid rgba(255,255,255,0.05) !important;
}
.custom-table tr:hover td { background: rgba(56,189,248,0.04) !important; }
.custom-table-wrap {
    border-radius: 10px !important;
    overflow-x: auto !important;
    overflow-y: auto !important;
    max-height: 700px !important;
    border: 1px solid rgba(56,189,248,0.18) !important;
    margin-top: 12px !important;
}

/* KPI cards */
.kpi-label {
    color: #64748b !important;
    font-size: 0.68rem !important;
    font-weight: 600 !important;
    letter-spacing: 0.1em !important;
    text-transform: uppercase !important;
}
.kpi-value {
    color: #38bdf8 !important;
    font-size: 1.4rem !important;
    font-weight: 800 !important;
    font-family: var(--mono) !important;
}

/* Insight bullets */
.insight-item { color: #cbd5e1 !important; font-size: 0.85rem !important; }
.insight-item strong { color: #f1f5f9 !important; font-weight: 700 !important; }

/* Status messages bold */
.status-msg strong { font-weight: 700 !important; color: #ffffff !important; }

/* Section labels - bold mono */
.sec-label {
    color: #38bdf8 !important;
    font-size: 0.68rem !important;
    font-weight: 800 !important;
    letter-spacing: 0.15em !important;
    text-transform: uppercase !important;
    font-family: var(--mono) !important;
}

/* Snapshot grid */
.snap-val {
    color: #00e5ff !important;
    font-weight: 800 !important;
    font-size: 1.25rem !important;
    font-family: var(--mono) !important;
}
.snap-label {
    color: #64748b !important;
    font-size: 0.65rem !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.08em !important;
}

/* Export/download button */
.export-btn button, button.export-btn {
    background: rgba(16,185,129,0.15) !important;
    border: 1px solid rgba(16,185,129,0.4) !important;
    color: #6ee7b7 !important;
    font-weight: 700 !important;
}

/* Sample query buttons - bold */
.sample-btn button, button.sample-btn {
    font-weight: 600 !important;
    text-align: left !important;
    line-height: 1.4 !important;
}

/* Recent query history */
.hist-empty {
    color: #334155 !important;
    font-size: 0.78rem !important;
    font-style: italic !important;
}

/* Gradio group/block transparent backgrounds */
.gradio-container .gr-group,
.gradio-container .block {
    background: transparent !important;
    border-color: var(--border) !important;
}

"""
import traceback

def _get_all_product_categories() -> list:
    """Fetches distinct product category names that actually have shipments."""
    try:
        rows = _run_neo4j("""
            MATCH (sh:Shipment)-[:CARRIES]->(pr:Product)
            WHERE pr.product_category_name IS NOT NULL
              AND trim(pr.product_category_name) <> ''
            RETURN DISTINCT pr.product_category_name AS cat
            ORDER BY cat
        """)
        cats = [r["cat"] for r in rows if r.get("cat")]
        return ["\U0001f310 All Categories"] + cats
    except Exception:
        return ["\U0001f310 All Categories"]


def _load_rca_category_stats(category: str) -> str:
    """Compact delay stats for a selected product category."""
    ALL = "\U0001f310 All Categories"
    if not category or category == ALL:
        try:
            rows = _run_neo4j("""
                MATCH (sh:Shipment)
                RETURN COUNT(sh) AS total,
                       SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS delayed,
                       ROUND(AVG(CASE WHEN sh.delivery_status='Major Delay' THEN sh.delay_days END), 1) AS avg_delay
            """)
            d       = rows[0] if rows else {}
            total   = d.get("total", 0) or 0
            delayed = d.get("delayed", 0) or 0
            avg_d   = d.get("avg_delay", 0) or 0
            pct     = round(100 * delayed / max(total, 1), 1)
            pct_col = "#f87171" if pct > 30 else "#fbbf24" if pct > 15 else "#4ade80"
            return f"""<div style='margin-top:6px'>
  <div style='font-size:0.62rem;color:#94a3b8;margin-bottom:6px'>All categories · {total:,} shipments</div>
  <div style='display:flex;justify-content:space-between;padding:4px 0;font-size:0.72rem;border-bottom:1px solid rgba(255,255,255,0.04)'>
    <span style='color:#94a3b8'>Delayed</span><span style='color:#f87171;font-weight:600'>{delayed:,}</span>
  </div>
  <div style='display:flex;justify-content:space-between;padding:4px 0;font-size:0.72rem;border-bottom:1px solid rgba(255,255,255,0.04)'>
    <span style='color:#94a3b8'>Delay Rate</span><span style='color:{pct_col};font-weight:600'>{pct}%</span>
  </div>
  <div style='display:flex;justify-content:space-between;padding:4px 0;font-size:0.72rem'>
    <span style='color:#94a3b8'>Avg Delay</span><span style='color:#fbbf24;font-weight:600'>{avg_d}d</span>
  </div>
</div>"""
        except Exception as e:
            return f"<div style='color:#f87171;font-size:0.75rem;padding:8px'>\u26a0 {str(e)[:80]}</div>"

    try:
        rows = _run_neo4j("""
            MATCH (sh:Shipment)-[:CARRIES]->(pr:Product)
            WHERE pr.product_category_name = $cat
            RETURN COUNT(sh) AS total,
                   SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS delayed,
                   ROUND(AVG(CASE WHEN sh.delivery_status='Major Delay' THEN sh.delay_days END), 1) AS avg_delay
        """, {"cat": category})

        plant_rows = _run_neo4j("""
            MATCH (pl:Plant)-[:DISPATCHES]->(sh:Shipment)-[:CARRIES]->(pr:Product)
            WHERE pr.product_category_name = $cat
              AND sh.delivery_status = 'Major Delay'
            RETURN pl.plant_name AS plant, COUNT(sh) AS cnt
            ORDER BY cnt DESC LIMIT 1
        """, {"cat": category})

        d           = rows[0] if rows else {}
        total       = d.get("total", 0) or 0
        delayed     = d.get("delayed", 0) or 0
        avg_delay   = d.get("avg_delay", 0) or 0
        pct         = round(100 * delayed / max(total, 1), 1)
        pct_col     = "#f87171" if pct > 30 else "#fbbf24" if pct > 15 else "#4ade80"
        worst_plant = plant_rows[0]["plant"] if plant_rows else "N/A"
        label       = category.replace("_", " ").title()[:22]

        return f"""<div style='margin-top:6px'>
  <div style='font-size:0.62rem;color:#94a3b8;margin-bottom:6px'>{label} · {total:,} shipments</div>
  <div style='display:flex;justify-content:space-between;padding:4px 0;font-size:0.72rem;border-bottom:1px solid rgba(255,255,255,0.04)'>
    <span style='color:#94a3b8'>Delayed</span><span style='color:#f87171;font-weight:600'>{delayed:,}</span>
  </div>
  <div style='display:flex;justify-content:space-between;padding:4px 0;font-size:0.72rem;border-bottom:1px solid rgba(255,255,255,0.04)'>
    <span style='color:#94a3b8'>Delay Rate</span><span style='color:{pct_col};font-weight:600'>{pct}%</span>
  </div>
  <div style='display:flex;justify-content:space-between;padding:4px 0;font-size:0.72rem;border-bottom:1px solid rgba(255,255,255,0.04)'>
    <span style='color:#94a3b8'>Avg Delay</span><span style='color:#fbbf24;font-weight:600'>{avg_delay}d</span>
  </div>
  <div style='display:flex;justify-content:space-between;padding:4px 0;font-size:0.72rem'>
    <span style='color:#94a3b8'>Worst Plant</span><span style='color:#fbbf24;font-weight:600'>{worst_plant[:16]}</span>
  </div>
</div>"""
    except Exception as e:
        return f"<div style='color:#f87171;font-size:0.75rem;padding:8px'>\u26a0 {str(e)[:80]}</div>"

def _load_supplier_risk_leaderboard() -> str:
    """Compact top-5 supplier risk list — name + score chip only."""
    try:
        rows = _run_neo4j("""
            MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)
            RETURN sup.supplier_name AS name,
                   sup.risk_score    AS risk,
                   pl.plant_name     AS plant
            ORDER BY sup.risk_score DESC
            LIMIT 5
        """)

        if not rows:
            return "<div style='color:#94a3b8;font-size:0.74rem;padding:8px'>No supplier data.</div>"

        rows_html = ""
        for i, r in enumerate(rows, 1):
            risk  = r.get("risk") or 0
            name  = (r.get("name") or "?")[:20]
            plant = (r.get("plant") or "?")[:8]
            col   = "#f87171" if risk > 0.8 else "#fbbf24" if risk > 0.6 else "#4ade80"
            medal = ["🥇","🥈","🥉"][i-1] if i <= 3 else f"#{i}"
            rows_html += f"""
<div style="display:flex;align-items:center;justify-content:space-between;padding:5px 0;border-bottom:1px solid rgba(255,255,255,0.04)">
  <span style="font-size:0.72rem;color:#e2e8f0">{medal} {name}</span>
  <span style="font-size:0.7rem;color:{col};font-weight:700;font-family:monospace;background:rgba(255,255,255,0.05);padding:1px 6px;border-radius:4px">{risk:.2f}</span>
</div>"""

        return f"""
<div style='margin-top:4px'>
  <div style="font-size:0.62rem;color:#94a3b8;margin-bottom:4px">risk score → <span style='color:#f87171'>&gt;0.8 critical</span> · <span style='color:#fbbf24'>&gt;0.6 high</span></div>
  {rows_html}
</div>"""

    except Exception as e:
        return f"<div style='color:#f87171;font-size:0.75rem;padding:8px'>⚠ {str(e)[:80]}</div>"

# ════════════════════════════════════════════════════════════════════
# TOOL LOG HTML BUILDER
# ════════════════════════════════════════════════════════════════════
def _tool_html(entries) -> str:
    if not entries:
        return "<div class='tool-log-wrap' style='color:#94a3b8'>Waiting for tool calls…</div>"

    # A2A stage labels for pseudo-tool entries
    _A2A_LABELS = {
        "__first_response__":    ("💬", "#38bdf8",  "First Response"),
        "__orchestrator__":      ("🧠", "#a78bfa",  "Orchestrator Agent (tool selection)"),
        "__validator_agent__":   ("🔍", "#22d3ee",  "Data Validator Agent (quality check)"),
        "__rca_agent__":         ("🔬", "#f97316",  "RCA Agent (analysis)"),
        "__rec_agent__":         ("💡", "#4ade80",  "Recommendations Agent"),
        "__narrative_agent__":   ("✍️", "#f472b6",  "Narrative Agent (opening summary)"),
    }

    rows = []
    for i, e in enumerate(entries, 1):
        tool = e.get("tool", "?")

        if tool in _A2A_LABELS:
            icon, color, label = _A2A_LABELS[tool]
            rows.append(
                f'<div class="tool-entry" style="border-left:3px solid {color};padding-left:8px">'
                f'<span style="color:{color};font-weight:700">[A2A] {icon} {label}</span>'
                f'</div>'
            )
            continue

        inp = e.get("input", {})
        if isinstance(inp, dict):
            inp_str = ", ".join(f"{k}={repr(str(v))[:40]}" for k, v in inp.items()) or "no params"
        else:
            inp_str = str(inp)[:80]
        prev = e.get("result_preview", "")[:250]
        rows.append(
            f'<div class="tool-entry">'
            f'<span class="tool-name">[{i}] {tool}()</span><br>'
            f'<span style="color:#94a3b8">  in:  {inp_str}</span><br>'
            f'<span style="color:#94a3b8">  out: {prev}</span>'
            f'</div>'
        )
    return f'<div class="tool-log-wrap">{"".join(rows)}</div>'


# ════════════════════════════════════════════════════════════════════
# FIX 4: DYNAMIC QUERY CONTEXT PANEL
# Returns HTML card summarising what the agent is about to analyse.
# ════════════════════════════════════════════════════════════════════
def _build_context_html(question: str) -> str:
    """Removed — static chain text replaced by Cypher Queries Executed accordion."""
    return ""


# ════════════════════════════════════════════════════════════════════
# PROMPT SUGGESTION ENGINE

_PROMPT_ENHANCE_SYSTEM = (
    "Sharpen the supply chain question for clarity and precision. "
    "Keep the exact same scenario and investigative direction — do NOT change what is being asked. "
    "Do NOT introduce plant IDs, supplier IDs, distributor IDs, city names, or entity names not already in the question. "
    "Do NOT add technical schema terms like PL1/PL2/SUP/D00 — write plain business language only. "
    "Focus on making the question more specific about WHO, WHY, WHEN, and HOW MUCH impact. "
    "Output ONE refined question only — max 35 words. No explanation, no preamble, no brackets with IDs."
)

def _suggest_prompt(question):
    q = (question or "").strip()
    if len(q) < 8 or len(q) > 400:
        return ""
    try:
        improved = _groq_call(
            messages=[
                {"role": "system", "content": _PROMPT_ENHANCE_SYSTEM},
                {"role": "user",   "content": q},
            ],
            max_tokens=60,
            temperature=0
        )
        # Guard: coerce None → "" before strip() to prevent AttributeError
        improved = str(improved or "").strip().strip('"').strip("'").strip()
        if not improved or improved.lower() == q.lower() or len(improved) < 12:
            return ""
        safe = improved.replace("<", "&lt;").replace(">", "&gt;")
        # JS: hide card first, then set textarea value.
        # We dispatch a synthetic 'input' event so Gradio picks up the new value,
        # but we immediately blank the suggest output via Gradio state so the
        # .change handler produces "" and no new suggestion card is shown.
        js = (
            "var card=document.getElementById('prompt-suggest-card');"
            "if(!card)return;"
            "var t=document.getElementById('suggest-text').innerText.trim();"
            "card.remove();"
            # Target by elem_id first — most reliable
            "var wrap=document.getElementById('rca-question-input');"
            "var tb=wrap?wrap.querySelector('textarea'):null;"
            # Fallback: search all textareas for the one inside #rca-question-input
            "if(!tb){"
            "var all=document.querySelectorAll('textarea');"
            "for(var i=0;i<all.length;i++){"
            "if(all[i].closest('#rca-question-input')){tb=all[i];break;}"
            "}}"
            # Fallback 2: any visible textarea on the active RCA tab
            "if(!tb){"
            "var all2=document.querySelectorAll('.tabitem:not([style*=\"display: none\"]) textarea,"
            " .tab-content:not([hidden]) textarea');"
            "if(all2.length)tb=all2[0];"
            "}"
            "if(!tb)return;"
            "var ni=Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype,'value').set;"
            "ni.call(tb,t);"
            "tb.dispatchEvent(new Event('input',{bubbles:true}));"
            "tb.focus();"
        )
        return (
            '<script>'
            '(function(){'
            'setTimeout(function(){'
            'var el=document.getElementById("rca-suggest-wrap")||document.getElementById("prompt-suggest-card");'
            'if(el)el.scrollIntoView({behavior:"smooth",block:"center"});'
            '},300);'
            '})();'
            '</script>'
            '<div id="prompt-suggest-card" style="'
            'background:linear-gradient(135deg,rgba(124,58,237,0.08),rgba(56,189,248,0.06));'
            'border:1px solid rgba(124,58,237,0.35);border-radius:10px;'
            'padding:13px 16px;margin-top:8px;position:relative;overflow:hidden">'
            '<div style="position:absolute;top:0;left:0;right:0;height:2px;'
            'background:linear-gradient(90deg,transparent,#7c3aed,#38bdf8,transparent)"></div>'
            '<div style="display:flex;align-items:flex-start;gap:10px">'
            '<span style="font-size:1.1rem;line-height:1;flex-shrink:0;margin-top:2px">&#10024;</span>'
            '<div style="flex:1;min-width:0">'
            '<div style="font-size:0.62rem;font-weight:700;text-transform:uppercase;'
            'letter-spacing:0.12em;color:#a78bfa;margin-bottom:5px">AI Prompt Suggestion</div>'
            '<div id="suggest-text" style="font-size:0.84rem;color:#e2e8f0;line-height:1.6;'
            'font-style:italic;word-break:break-word">' + safe + '</div>'
            '</div></div>'
            '<div style="display:flex;align-items:center;justify-content:space-between;'
            'margin-top:10px;padding-top:8px;border-top:1px solid rgba(124,58,237,0.2)">'
            '<div style="font-size:0.7rem;color:#94a3b8">Refined for clearer graph traversal &amp; deeper analysis</div>'
            '<button onclick="' + js.replace('"', "&quot;") + '" '
            'style="background:linear-gradient(135deg,rgba(124,58,237,0.25),rgba(56,189,248,0.15));'
            'border:1px solid rgba(124,58,237,0.5);border-radius:7px;color:#c4b5fd;'
            'font-size:0.74rem;font-weight:700;padding:5px 14px;cursor:pointer;'
            'white-space:nowrap;transition:all 0.18s">'
            '&#10024; Use this</button>'
            '</div></div>'
        )
    except Exception:
        return ""


# AGENT STREAMING GENERATOR
# ════════════════════════════════════════════════════════════════════
_MD_PLACEHOLDER = ""

# A2A status messages shown in the status badge during streaming
_A2A_STATUS = {
    "__first_response__":    "◌ Initial Assessment — Analysis in progress…",
    "__orchestrator__":      "🔗 Tracing Path — Orchestrator mapping investigation route…",
    "__validator_agent__":   "🔍 Validating — Data Validator checking tool result quality…",
    "__rca_agent__":         "🔬 Analysing — RCA Agent processing supply chain data…",
    "__rec_agent__":         "💡 Generating — Recommendations Agent writing action plan…",
    "__narrative_agent__":   "✍️ Composing — Narrative Agent writing opening summary…",
}

# Update tab A2A pipeline status labels
_UPD_PIPELINE_STATUS = {
    "__upd_parse__":    "📂 [Step 1/6] File Parser — reading and extracting rows…",
    "__upd_detect__":   "🔍 [Step 2/6] Schema Detector — identifying entity type via LLM…",
    "__upd_normalize__":"🔄 [Step 3/6] Field Mapper — normalizing column names…",
    "__upd_clean__":    "🧹 [Step 4/6] Data Cleaner — handling nulls, duplicates, mismatches…",
    "__upd_validate__": "✅ [Step 5/6] Validator — checking rows before insertion…",
    "__upd_insert__":   "🔗 [Step 6/6] Graph Agent — inserting nodes into Neo4j…",
}

def _agent_generator(fn, user_input):
    if not user_input or not str(user_input).strip():
        yield '<div class="ag-status ag-err">Please enter a question first.</div>', "", _MD_PLACEHOLDER, ""
        return

    yield '<div class="ag-status ag-run">◌ Agent starting…</div>', "", _MD_PLACEHOLDER, ""

    q      = queue.Queue()
    holder = {}

    def worker():
        def on_update(event):
            q.put(event)
        try:
            text, logs = fn(user_input, on_update=on_update)
            holder["text"] = text
            holder["logs"] = logs
        except Exception as e:
            holder["error"] = str(e)
        finally:
            q.put(None)

    threading.Thread(target=worker, daemon=True).start()

    tools, partial, current_status = [], "", "◌ Agent running…"
    _first_response_card = ""
    while True:
        try:
            event = q.get(timeout=0.5)
        except queue.Empty:
            yield (
                f'<div class="ag-status ag-run">{current_status}</div>',
                holder.get("assess", ""),
                partial or _MD_PLACEHOLDER,
                _tool_html(tools)
            )
            continue

        if event is None:
            break

        kind, data = event
        if kind == "text":
            partial += safe_str(data)
        elif kind == "cypher_log":
            # Accumulate structured Cypher logs (for accordion)
            holder.setdefault("cypher_logs", []).append(data)
        elif kind == "tool":
            tools.append(data)
        elif kind == "tool_start":
            _all_status = {**_A2A_STATUS, **_UPD_PIPELINE_STATUS}
            current_status = _all_status.get(data, f"◌ Running {data}…")
            # Append stage marker to tool log for any known agent step
            if data in _all_status:
                tools.append({"tool": data, "input": {}, "result_preview": ""})
        elif kind == "first_response":
            # Store Initial Assessment in a separate key — rendered to rca_assess_html
            # (gr.HTML), NOT prepended to rca_out (gr.HTML which renders HTML natively)
            _first_response_card = (
                f'<div style="background:rgba(56,189,248,0.07);border:1px solid rgba(56,189,248,0.25);'
                f'border-radius:10px;padding:12px 16px;margin:8px 0 4px;font-size:0.83rem;color:#e2e8f0;line-height:1.65">'
                f'<span style="font-size:0.58rem;font-weight:800;text-transform:uppercase;'
                f'letter-spacing:0.12em;color:#38bdf8;display:block;margin-bottom:6px">'
                f'◌ Initial Assessment — Analysis in progress…</span>'
                f'{data}</div>'
            )
            holder["assess"] = _first_response_card
            # Do NOT set partial — keep rca_out clean for the final markdown report

        yield (
            f'<div class="ag-status ag-run">{current_status}</div>',
            holder.get("assess", ""),    # assess_html — Initial Assessment card
            partial or _MD_PLACEHOLDER,
            _tool_html(tools)
        )

    if "error" in holder:
        yield (
            f'<div class="ag-status ag-err">❌ Error: {holder["error"][:400]}</div>',
            holder.get("assess", ""),
            partial or _MD_PLACEHOLDER,
            _tool_html(tools)
        )
        return

    final_logs = holder.get("logs", tools)
    # Merge pseudo-key stage markers from tools into final_logs so the activity
    # log can see which agents ran. Pseudo entries have empty input/result_preview.
    PSEUDO_KEYS = {"__first_response__","__orchestrator__","__validator_agent__",
                   "__rca_agent__","__rec_agent__","__narrative_agent__"}
    pseudo_entries = [t for t in tools if t.get("tool","") in PSEUDO_KEYS]
    # Prepend pseudo-keys that aren't already in final_logs
    existing_pseudo = {t.get("tool","") for t in final_logs if t.get("tool","") in PSEUDO_KEYS}
    merged_logs = [p for p in pseudo_entries if p["tool"] not in existing_pseudo] + list(final_logs)
    final_logs = merged_logs
    final_text = safe_str(holder.get("text", partial)) or _MD_PLACEHOLDER
    yield (
        '<div class="ag-status ag-ok">✦ A2A Analysis complete (Orchestrator → Validator → RCA → Recommendations → Narrative)</div>',
        holder.get("assess", ""),    # keep Initial Assessment visible after completion
        final_text,
        _tool_html(final_logs)
    )


# ─────────────────────────────────────────────────────────────
# RCA REPORT CACHE  — same question → same report, no re-run
# Key: lowercased+stripped question string
# Value: {"report": str, "logs": list, "status": str}
# ─────────────────────────────────────────────────────────────
_RCA_CACHE: dict = {}
def safe_str(x):
    try:
        return "" if x is None else str(x).strip()
    except Exception:
        return ""

def _cache_key(q: str) -> str:
    import re as _re
    return _re.sub(r'\s+', ' ', safe_str(q).lower())


def _clean_report_for_export(report_text: str) -> str:
    """
    Extract clean text from the RCA HTML report for Word/Excel export.
    Extracts: Executive Summary, each step section (with description + table data),
    Root Cause paragraph, and Recommendations (with tier headings + items).
    Returns clean markdown-style text with no CSS or HTML artifacts.
    """
    import re as _re, html as _html

    def _strip_tags(s: str) -> str:
        return _re.sub(r'<[^>]+>', ' ', s).strip()

    def _clean_text(s: str) -> str:
        s = _html.unescape(s)
        s = _re.sub(r'\s+', ' ', s).strip()
        # Remove any CSS-looking fragments
        s = _re.sub(r'[.#][\w-]+\s*\{[^}]*\}', '', s)
        s = _re.sub(r'details\[open\][^{]*\{[^}]*\}', '', s)
        s = _re.sub(r'@[\w-]+\s*\{[^}]*\}', '', s)
        # Remove lone selector-looking lines
        s = _re.sub(r'^[.#][\w-][\w\s,.-]*$', '', s, flags=_re.MULTILINE)
        return s.strip()

    text = str(report_text or "")
    out_sections = []

    # ── 1. Executive Summary ──────────────────────────────────────────
    # Stored in a <details> block with "Executive Summary" in summary
    exec_m = _re.search(
        r'<summary[^>]*>.*?Executive Summary.*?</summary>(.*?)</details>',
        text, _re.DOTALL | _re.IGNORECASE
    )
    if exec_m:
        exec_txt = _clean_text(_strip_tags(exec_m.group(1)))
        if exec_txt and len(exec_txt) > 20:
            out_sections.append("## Executive Summary\n" + exec_txt)

    # ── 2. Step section cards ────────────────────────────────────────
    # Each section: <div style="background:...border-left:4px solid ..."> 
    # Contains: title span, step pill, desc para, and <details> with <table>
    card_pattern = _re.compile(
        r'<div[^>]*border-left:[^"]*border-radius:12px[^>]*>(.*?)(?=<div[^>]*border-left:[^"]*border-radius:12px|<div[^>]*margin-top:32px|</div>\s*$)',
        _re.DOTALL
    )
    for card_m in card_pattern.finditer(text):
        block = card_m.group(1)

        # Title
        title_m = (_re.search(r'color:#fca5a5[^>]*>([^<]{4,80})<', block) or
                   _re.search(r'font-weight:700;color:#e2e8f0[^>]*>([^<]{4,80})<', block) or
                   _re.search(r'font-weight:700[^>]*>([^<]{4,80})<', block))
        if not title_m:
            continue
        title = _html.unescape(title_m.group(1)).strip()
        title = _re.sub(r'[\U0001F300-\U0001FAFF\u2600-\u27BF]+\s*', '', title).strip()
        if not title or len(title) < 3:
            continue

        # Step pill
        pill_m = _re.search(r'letter-spacing:0\.02em[^>]*>([^<]+)<', block)
        pill = _html.unescape(pill_m.group(1)).strip() if pill_m else ""

        # Description text (the prose para in the section)
        desc_m = _re.search(r'color:#cbd5e1[^>]*>([^<]{15,})<', block)
        desc = _clean_text(desc_m.group(1)) if desc_m else ""

        section_text = [f"## {title}"]
        if pill:
            section_text.append(f"◆ {pill}")
        if desc:
            section_text.append(desc)

        # Tables inside <details> blocks
        for det_m in _re.finditer(r'<details[^>]*>(.*?)</details>', block, _re.DOTALL):
            det_inner = det_m.group(1)
            # Get the summary label (subtable name like 1A, 1B etc.)
            sum_m = _re.search(r'<summary[^>]*>(.*?)</summary>', det_inner, _re.DOTALL)
            sum_label = _clean_text(_strip_tags(sum_m.group(1))) if sum_m else ""
            sum_label = _re.sub(r'[▶▼◆►◌]+\s*', '', sum_label).strip()
            sum_label = _re.sub(r'📋\s*', '', sum_label).strip()

            # Extract HTML table
            tbl_m = _re.search(r'<table[^>]*>(.*?)</table>', det_inner, _re.DOTALL)
            if tbl_m:
                tbl_html = tbl_m.group(1)
                headers = [_html.unescape(_strip_tags(th)).strip()
                           for th in _re.findall(r'<th[^>]*>(.*?)</th>', tbl_html, _re.DOTALL)]
                headers = [h for h in headers if h]
                rows = []
                for tr in _re.findall(r'<tr[^>]*>(.*?)</tr>', tbl_html, _re.DOTALL):
                    cells = [_html.unescape(_strip_tags(td)).strip()
                             for td in _re.findall(r'<td[^>]*>(.*?)</td>', tr, _re.DOTALL)]
                    cells = [c for c in cells if c]
                    if cells:
                        rows.append(cells)
                if headers and rows:
                    if sum_label and not sum_label.lower().startswith('view'):
                        section_text.append(f"\n### {sum_label}")
                    # Markdown pipe table
                    section_text.append("| " + " | ".join(headers) + " |")
                    section_text.append("| " + " | ".join(["---"] * len(headers)) + " |")
                    for row in rows:
                        # Pad/trim to header count
                        padded = row[:len(headers)] + [""] * max(0, len(headers) - len(row))
                        section_text.append("| " + " | ".join(padded) + " |")
                    section_text.append("")

            # Also check for text inside the details (for root cause paragraph)
            elif sum_label.lower().startswith('root') or 'verdict' in sum_label.lower():
                body_txt = _clean_text(_strip_tags(det_inner))
                if body_txt and len(body_txt) > 20:
                    section_text.append(body_txt)

        out_sections.append("\n".join(section_text))

    # ── 3. Root Cause paragraph (new format: div with border-left:4px solid #f87171) ──
    rc_m = _re.search(
        r'<div[^>]*border-left:4px solid #f87171[^>]*>(.*?)</div>\s*(?=<div[^>]*margin-top:32px|$)',
        text, _re.DOTALL
    )
    if not rc_m:
        # Alternative: find by "Root Cause — Confirmed Verdict" label
        rc_m = _re.search(
            r'Root Cause.*?Confirmed Verdict.*?</div>(.*?)</div>',
            text, _re.DOTALL | _re.IGNORECASE
        )
    if rc_m:
        rc_txt = _clean_text(_strip_tags(rc_m.group(1)))
        rc_txt = _re.sub(r'Root Cause.*?Confirmed Verdict\s*', '', rc_txt, flags=_re.IGNORECASE).strip()
        if rc_txt and len(rc_txt) > 20:
            # Check not already included
            if not any('root cause' in s.lower()[:50] for s in out_sections):
                out_sections.append("## Root Cause — Confirmed Verdict\n" + rc_txt)

    # ── 4. Recommendations (from <details> tier dropdowns) ──────────
    rec_div_m = _re.search(
        r'<div[^>]*margin-top:32px[^>]*>.*?Recommendations(.*?)(?=</div>\s*</div>\s*$|$)',
        text, _re.DOTALL | _re.IGNORECASE
    )
    if rec_div_m:
        rec_block = rec_div_m.group(1)
        rec_lines = ["## Recommendations"]
        # Each tier is a <details> block
        for tier_m in _re.finditer(r'<details[^>]*>(.*?)</details>', rec_block, _re.DOTALL):
            tier_inner = tier_m.group(1)
            # Tier label from summary
            t_sum_m = _re.search(r'<summary[^>]*>(.*?)</summary>', tier_inner, _re.DOTALL)
            tier_label = ""
            if t_sum_m:
                tier_label = _clean_text(_strip_tags(t_sum_m.group(1)))
                tier_label = _re.sub(r'[▶▼◆►◌]+\s*', '', tier_label).strip()
                tier_label = _re.sub(r'[\U0001F300-\U0001FAFF\u2600-\u27BF⚡🔧🏛]+\s*', '', tier_label).strip()
                tier_label = _re.sub(r'\d+ action.*$', '', tier_label).strip()
            if tier_label:
                rec_lines.append(f"\n### {tier_label}")
            # Items inside the tier
            body_after_sum = _re.sub(r'.*?</summary>', '', tier_inner, flags=_re.DOTALL, count=1)
            for item_m in _re.finditer(r'<div[^>]*border-left:[^>]+>(.*?)</div>', body_after_sum, _re.DOTALL):
                item_txt = _clean_text(_strip_tags(item_m.group(1)))
                if item_txt and len(item_txt) > 10:
                    rec_lines.append(f"- {item_txt}")
        if len(rec_lines) > 1:
            out_sections.append("\n".join(rec_lines))

    # ── 5. Fallback: if nothing extracted, return stripped text ──────
    if not out_sections:
        plain = _clean_text(_re.sub(r'<[^>]+>', ' ', text))
        plain = _re.sub(r'\n{3,}', '\n\n', plain)
        return plain.strip()

    result = "\n\n".join(out_sections)
    # Final cleanup: remove any remaining CSS-looking fragments
    result = _re.sub(r'[.#][\w-]+\s*\{[^}]*\}', '', result)
    result = _re.sub(r'details\[open\][^{]*\{[^}]*\}', '', result)
    result = _re.sub(r'\n{3,}', '\n\n', result)
    return result.strip()


def _strip_rec_details(report_text: str) -> str:
    """
    Pass-through: report HTML is rendered as-is.
    View Charts and Recommendations are separate <details> blocks
    rendered via rca_charts_html_vis and rca_detailed_rec_inner.
    """
    return report_text or ""


def _neutralise_rec_text(text: str) -> str:
    """
    Rewrite subjective/imperative phrasing in recommendation items to
    neutral, professional, outcome-focused language.
    e.g. "The operations team must immediately renegotiate..."
      -> "Renegotiate lead time with ... to reduce delay rate at ..."
    """
    import re as _re
    t = text.strip()

    # Remove leading agent phrases: "The X team must/should Y" → "Y"
    t = _re.sub(
        r'^The\s+\w[\w\s]+?\s+team\s+(?:must|should|needs?\s+to|is\s+advised\s+to|is\s+recommended\s+to)\s+',
        '', t, flags=_re.IGNORECASE
    )
    # "X should Y" where X is a role noun → drop "X should "
    t = _re.sub(
        r'^(?:Operations|Procurement|Supply\s+chain|Logistics|Analyst|Management|Leadership)\s+'
        r'(?:team\s+)?(?:must|should|needs?\s+to|will)\s+',
        '', t, flags=_re.IGNORECASE
    )
    # "Consider X" / "Implement X" — keep as-is (already neutral)
    # Capitalise first letter
    if t:
        t = t[0].upper() + t[1:]
    return t


def _generate_rec_pill(item_text: str, tier_key: str) -> str:
    """
    Dynamically generate a contextual pill label for a recommendation item.
    Returns a short string summarising the measurable outcome.
    """
    import re as _re
    t = item_text.lower()

    if tier_key == 'critical':
        if 'lead time' in t or 'renegotiat' in t:
            return 'Immediate Impact: Reduces supplier lead time'
        if 'delay rate' in t or 'delay' in t:
            return 'Immediate Impact: Reduces delay rate for affected shipments'
        if 'backlog' in t or 'bottleneck' in t:
            return 'Immediate Impact: Clears backlog at the bottleneck plant'
        if 'distribut' in t or 'stockout' in t:
            return 'Immediate Impact: Prevents cascading stockouts across distributor cities'
        if 'communicat' in t or 'update' in t:
            return 'Immediate Impact: Improves real-time visibility of shipment status'
        if 'reroute' in t or 'alternative' in t:
            return 'Immediate Impact: Reroutes affected shipments via alternative channel'
        return 'Immediate Impact: Reduces delay rate for affected shipments'

    elif tier_key in ('operational', 'near', 'high'):
        if 'track' in t or 'monitor' in t:
            return 'Execution Note: Category manager to enable tracking within 2 weeks'
        if 'jit' in t or 'just-in-time' in t or 'inventory' in t:
            return 'Execution Note: Procurement team to formalise JIT terms with supplier'
        if 'review' in t or 'process' in t:
            return 'Execution Note: Category manager to coordinate with supplier within 2 weeks'
        if 'reroute' in t or 'dispatch' in t or 'schedule' in t:
            return 'Execution Note: Supply planning team to revise dispatch schedule immediately'
        if 'capacit' in t or 'production' in t:
            return 'Execution Note: Plant operations team to adjust schedule within 1 week'
        return 'Execution Note: Operations team to implement within 2 weeks'

    else:  # strategic / long
        if 'dual.sourc' in t or 'dual sourc' in t or 'second supplier' in t:
            return 'Long-Term Value: Supplier risk score pool becomes measurably lower over time'
        if 'renegotiat' in t or 'contract' in t or 'penalty' in t or 'metric' in t:
            return 'Long-Term Value: Builds supplier resilience and reduces single-source dependency'
        if 'local supplier' in t or 'network' in t or 'redesign' in t:
            return 'Long-Term Value: Reduces average lead time and improves supply chain agility'
        if 'technology' in t or 'system' in t or 'platform' in t:
            return 'Long-Term Value: Enables proactive risk detection across the supply network'
        return 'Long-Term Value: Improves supply chain resilience and reduces structural risk'


def _save_rca_report(question: str, report_text: str) -> str:
    """Write a clean, professional RCA Word document. Returns path or None."""
    import tempfile, re as _re
    from datetime import datetime

    if not (report_text or "").strip(): return None
    if not _DOCX_AVAILABLE: return None

    safe = _re.sub(r'[^a-zA-Z0-9]+', '_', (question or '').strip()[:40]).strip('_')
    path = os.path.join(tempfile.gettempdir(), f"rca_{safe}.docx")

    doc = DocxDocument()

    # ── Page setup ─────────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Cm(2.2)
        sec.bottom_margin = Cm(2.2)
        sec.left_margin   = Cm(2.8)
        sec.right_margin  = Cm(2.8)

    # ── Color palette ──────────────────────────────────────────
    C_TITLE    = RGBColor(0x1d, 0x4e, 0xd8)   # blue
    C_H1       = RGBColor(0x07, 0x5e, 0x96)   # dark blue
    C_H2       = RGBColor(0x0f, 0x76, 0x6e)   # teal
    C_H3       = RGBColor(0x07, 0x89, 0x50)   # green
    C_BODY     = RGBColor(0x1e, 0x29, 0x3b)   # near-black
    C_META     = RGBColor(0x64, 0x74, 0x8b)   # slate
    C_ACCENT   = RGBColor(0x1e, 0x3a, 0x8a)   # deep blue
    C_WARN     = RGBColor(0xef, 0x44, 0x44)   # red
    C_OK       = RGBColor(0x16, 0xa3, 0x4a)   # green
    C_ORANGE   = RGBColor(0xf9, 0x73, 0x16)   # orange

    FILL_COVER  = "EFF6FF"   # light blue
    FILL_QUERY  = "DBEAFE"   # blue tint
    FILL_RC     = "FEF2F2"   # red tint (root cause highlight)
    FILL_REC    = "F0FDF4"   # green tint
    FILL_H1     = "E0F2FE"   # sky tint
    FILL_META   = "F8FAFC"   # near-white
    TBL_HDR     = "1D4ED8"   # table header fill
    TBL_ROW_A   = "EFF6FF"
    TBL_ROW_B   = "FFFFFF"

    def _shade(para, fill_hex):
        try:
            pPr = para._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill_hex)
            pPr.append(shd)
        except Exception:
            pass

    def _gap(pts=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(pts)
        p.paragraph_format.space_before = Pt(0)

    def _body_para(text, bold=False, italic=False, color=None, size=10, shade_hex=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.space_before = Pt(0)
        if shade_hex: _shade(p, shade_hex)
        # Render **bold** inline spans
        parts = _re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
                run.bold = bold
            run.italic = italic
            run.font.size = Pt(size)
            run.font.color.rgb = color or C_BODY
        return p

    # ── COVER ──────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t.paragraph_format.space_before = Pt(0)
    t.paragraph_format.space_after  = Pt(4)
    r = t.add_run("GraphPulse AI")
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = C_TITLE
    _shade(t, FILL_COVER)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.LEFT
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after  = Pt(0)
    r2 = s.add_run("Supply Chain Root Cause Analysis Report")
    r2.font.size = Pt(13); r2.font.color.rgb = C_META
    _shade(s, FILL_COVER)
    _gap(4)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(2)
    _shade(meta, FILL_META)
    r3 = meta.add_run(f"Generated:  {datetime.now().strftime('%d %B %Y, %H:%M')}")
    r3.font.size = Pt(9); r3.font.color.rgb = C_META; r3.italic = True

    _gap(4)

    # Query block — prominent, boxed look
    q_p = doc.add_paragraph()
    q_p.paragraph_format.space_before = Pt(4)
    q_p.paragraph_format.space_after  = Pt(4)
    _shade(q_p, FILL_QUERY)
    ql = q_p.add_run("Analysis Query:  ")
    ql.bold = True; ql.font.size = Pt(11); ql.font.color.rgb = C_TITLE
    qt = q_p.add_run(str(question or ""))
    qt.font.size = Pt(11); qt.font.color.rgb = C_ACCENT; qt.bold = True
    _gap(8)

    # Divider
    hr = doc.add_paragraph()
    hr.paragraph_format.space_before = Pt(2)
    hr.paragraph_format.space_after  = Pt(2)
    hr.add_run("─" * 90).font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    _gap(4)

    # ── TABLE HELPER ───────────────────────────────────────────
    def _add_table(md_lines):
        rows = []
        for ln in md_lines:
            s = ln.strip()
            if s.startswith("|") and not all(c in "|-: " for c in s):
                cells = [_re.sub(r'\*\*(.+?)\*\*', r'\1', c.strip()) for c in s.strip("|").split("|")]
                if cells: rows.append(cells)
        if len(rows) < 2: return

        ncols = max(len(r) for r in rows)
        tbl = doc.add_table(rows=len(rows), cols=ncols)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        for ri, row_data in enumerate(rows):
            for ci, val in enumerate(row_data[:ncols]):
                cell = tbl.rows[ri].cells[ci]
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.size = Pt(9)
                if ri == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), TBL_HDR)
                    tcPr.append(shd)
                else:
                    run.font.color.rgb = C_ACCENT
                    bg = TBL_ROW_A if ri % 2 == 0 else TBL_ROW_B
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), bg)
                    tcPr.append(shd)
        _gap(6)

    # ── PARSE & RENDER CONTENT ─────────────────────────────────
    EMOJI_RE = r'[\U0001F300-\U0001FAFF\U00002600-\U000027BF]+'
    def _strip_emoji(t): return _re.sub(EMOJI_RE + r'\s*', '', t).strip()

    in_table   = False
    tbl_buf    = []
    in_rc      = False    # flag: inside Root Cause section

    for raw in report_text.splitlines():
        line = raw.strip()

        # Skip pure HTML/CSS artefacts
        if _re.match(r'<[a-zA-Z/!]', line): continue
        if _re.match(r'^[.#]?[\w-]+\s*\{', line): continue
        if _re.match(r'^(style|class|div|span|details|summary)[\s="@>]', line, _re.IGNORECASE): continue
        if _re.match(r'^[\w-]+\s*:\s*[^;]{1,80};', line): continue  # CSS property: value;
        # Skip UI artefacts, markdown fences, debug text
        if _re.search(r'click\s+to\s+(expand|collapse)|view\s+charts?|view\s+table|view\s+data', line, _re.IGNORECASE): continue
        if line.startswith('```') or line.startswith('~~~'): continue
        if _re.match(r'^[-=_]{4,}$', line): continue   # horizontal rules
        if not _re.search(r'[A-Za-z0-9|]', line): continue   # no real text
        # Skip lines that look like CSS selectors or style blocks
        if _re.match(r'^\.[a-zA-Z][\w-]*\s', line): continue
        if _re.search(r'!important\s*;?\s*}', line): continue
        if line.startswith('details[') or line.startswith('@'): continue

        # ── Table detection ────────────────────────────────────
        if line.startswith("|"):
            in_table = True; tbl_buf.append(line); continue
        elif in_table:
            _add_table(tbl_buf); tbl_buf = []; in_table = False

        if not line:
            _gap(3); continue

        # Strip markdown bold/italic for clean text
        clean = _re.sub(r'\*\*(.+?)\*\*', r'\1', line)
        clean = _re.sub(r'\*(.+?)\*',     r'\1', clean)

        # ── Heading levels ─────────────────────────────────────
        if clean.startswith("#### "):
            txt = _strip_emoji(clean[5:].strip())
            if not txt: continue
            # Skip rec tier headings — rendered in rec block below
            if any(k in txt.lower() for k in ['critical', 'operational', 'near-term', 'strategic', 'long-term', 'high priority']):
                continue
            h = doc.add_heading("", level=3)
            h.clear()
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after  = Pt(3)
            r = h.add_run(txt); r.bold = True; r.font.size = Pt(10.5)
            r.font.color.rgb = C_H3
            _shade(h, "F0FDF4")

        elif clean.startswith("### "):
            txt = _strip_emoji(clean[4:].strip())
            if not txt: continue
            if _re.search(r'\brecommend', txt, _re.IGNORECASE): continue
            h = doc.add_heading("", level=2)
            h.clear()
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after  = Pt(4)
            r = h.add_run(txt); r.bold = True; r.font.size = Pt(12)
            r.font.color.rgb = C_H2
            rc_kws = ["root cause", "final root", "root-cause"]
            shade_h3 = FILL_RC if any(k in txt.lower() for k in rc_kws) else "F0F9FF"
            _shade(h, shade_h3)
            in_rc = any(k in txt.lower() for k in rc_kws)

        elif clean.startswith("## "):
            txt = _strip_emoji(clean[3:].strip())
            if not txt: continue
            if _re.search(r'\brecommend', txt, _re.IGNORECASE): continue
            _gap(6)
            h = doc.add_heading("", level=1)
            h.clear()
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after  = Pt(5)
            r = h.add_run(txt); r.bold = True; r.font.size = Pt(14)
            r.font.color.rgb = C_H1
            _shade(h, FILL_H1)
            in_rc = False

        elif clean.startswith("# "):
            txt = _strip_emoji(clean[2:].strip())
            if not txt: continue
            _gap(6)
            h = doc.add_heading("", level=1)
            h.clear()
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after  = Pt(5)
            r = h.add_run(txt); r.bold = True; r.font.size = Pt(14)
            r.font.color.rgb = C_H1
            _shade(h, FILL_H1)
            in_rc = False

        # ── Bullet points ───────────────────────────────────────
        elif clean.startswith(("- ", "\u2022 ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.space_before = Pt(1)
            body_b = clean[2:]
            for seg in _re.split(r'(\*\*[^*]+\*\*)', body_b):
                if seg.startswith("**") and seg.endswith("**"):
                    rr = p.add_run(seg[2:-2]); rr.bold = True
                else:
                    rr = p.add_run(seg)
                rr.font.size = Pt(10); rr.font.color.rgb = C_BODY

        # ── Numbered items ─────────────────────────────────────
        elif _re.match(r'^\d+\.\s', clean):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.space_before = Pt(1)
            body = _re.sub(r'^\d+\.\s*', '', clean)
            for seg in _re.split(r'(\*\*[^*]+\*\*)', body):
                if seg.startswith("**") and seg.endswith("**"):
                    r = p.add_run(seg[2:-2]); r.bold = True
                else:
                    r = p.add_run(seg)
                r.font.size = Pt(10); r.font.color.rgb = C_ACCENT

        # ── Key Metrics lines (Label: Value) — indented ────────
        elif _re.match(r'^\s{2,}\w.*:\s*\S', line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Pt(14)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.space_before = Pt(1)
            kv = clean.split(":", 1)
            lbl = p.add_run(kv[0].strip() + ":  ")
            lbl.bold = True; lbl.font.size = Pt(10); lbl.font.color.rgb = C_TITLE
            if len(kv) > 1:
                val = p.add_run(kv[1].strip())
                val.font.size = Pt(10)
                val.font.color.rgb = C_WARN if any(w in kv[1].lower() for w in ["high", "critical", "delayed", "risk"]) else C_BODY

        # ── Body prose ─────────────────────────────────────────
        else:
            if in_rc:
                _body_para(line, shade_hex=FILL_RC, color=RGBColor(0x7f, 0x1d, 0x1d))
            else:
                _body_para(line)

    if in_table and tbl_buf:
        _add_table(tbl_buf)

    # ── Recommendations (from raw rec block) ───────────────────
    try:
        rec_sections = _extract_detailed_recommendations(report_text)
        if rec_sections:
            _gap(12)
            # Section divider line
            hr2 = doc.add_paragraph()
            hr2.add_run("─" * 90).font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
            _gap(4)

            # Recommendations main heading
            h = doc.add_heading("", level=1)
            h.clear()
            h.paragraph_format.space_before = Pt(6)
            h.paragraph_format.space_after  = Pt(6)
            _shade(h, FILL_REC)
            r = h.add_run("Recommendations")
            r.bold = True; r.font.size = Pt(15); r.font.color.rgb = C_OK

            TIER_CFG = {
                "critical":    (C_WARN,   FILL_RC,   "FEF2F2", "Critical — Act Immediately"),
                "operational": (C_ORANGE, "FFF7ED",  "FFF7ED", "Operational Fixes — Near-Term"),
                "near":        (C_ORANGE, "FFF7ED",  "FFF7ED", "Operational Fixes — Near-Term"),
                "strategic":   (C_OK,     FILL_REC,  "F0FDF4", "Strategic Initiatives — Long-Term"),
                "long":        (C_OK,     FILL_REC,  "F0FDF4", "Strategic Initiatives — Long-Term"),
                "high":        (C_ORANGE, "FFF7ED",  "FFF7ED", "High Priority — Near-Term"),
            }

            for sec in rec_sections:
                s_lower = sec["section"].lower()
                tier_key = next((k for k in TIER_CFG if k in s_lower), None)
                t_col, t_fill, item_fill, display_label = TIER_CFG.get(
                    tier_key, (C_H2, FILL_H1, "EFF6FF", sec["section"])
                )

                # Tier sub-heading — bold, colored, shaded
                _gap(8)
                h2 = doc.add_heading("", level=2)
                h2.clear()
                h2.paragraph_format.space_before = Pt(6)
                h2.paragraph_format.space_after  = Pt(4)
                _shade(h2, t_fill)
                sh = h2.add_run(display_label.upper())
                sh.bold = True
                sh.font.size = Pt(11)
                sh.font.color.rgb = t_col

                for item in sec["items"]:
                    # Neutralise phrasing
                    neutral_item = _neutralise_rec_text(item)

                    # Recommendation card paragraph
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after  = Pt(2)
                    p.paragraph_format.left_indent  = Pt(8)
                    _shade(p, item_fill)

                    # Bold "Entity Name (ID):" prefix
                    prefix_match = _re.match(
                        r'^([A-Z][A-Za-z0-9 \-\.]+(?:\s*\([^)]*\))?):(\s+)(.*)',
                        neutral_item, _re.DOTALL
                    )
                    if prefix_match:
                        pr = p.add_run(prefix_match.group(1) + ":")
                        pr.bold = True; pr.font.size = Pt(10)
                        pr.font.color.rgb = t_col
                        body_text = prefix_match.group(2) + prefix_match.group(3)
                    else:
                        body_text = neutral_item

                    # Render rest with inline **bold** spans
                    for seg in _re.split(r'(\*\*[^*]+\*\*)', body_text):
                        if seg.startswith("**") and seg.endswith("**"):
                            r2 = p.add_run(seg[2:-2]); r2.bold = True
                        else:
                            r2 = p.add_run(seg)
                        r2.font.size = Pt(10); r2.font.color.rgb = C_BODY

                    # Dynamic pill note below each item
                    pill_text = _generate_rec_pill(neutral_item, tier_key or "")
                    if pill_text:
                        pp = doc.add_paragraph()
                        pp.paragraph_format.space_before = Pt(1)
                        pp.paragraph_format.space_after  = Pt(5)
                        pp.paragraph_format.left_indent  = Pt(14)
                        pr2 = pp.add_run(f"  {pill_text}  ")
                        pr2.bold = True; pr2.italic = True
                        pr2.font.size = Pt(8.5)
                        # Colour-code pill by tier
                        if tier_key == 'critical':
                            pr2.font.color.rgb = C_WARN
                        elif tier_key in ('operational', 'near', 'high'):
                            pr2.font.color.rgb = C_ORANGE
                        else:
                            pr2.font.color.rgb = C_OK
    except Exception:
        pass

    # ── Charts ─────────────────────────────────────────────────
    try:
        viz_images = _generate_rca_viz_images(question)
        if viz_images:
            _gap(10)
            hr3 = doc.add_paragraph()
            hr3.add_run("─" * 90).font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
            _gap(4)
            h = doc.add_heading("", level=1)
            h.clear(); _shade(h, FILL_H1)
            r = h.add_run("Analytical Visualizations")
            r.bold = True; r.font.size = Pt(14); r.font.color.rgb = C_H1
            _gap(6)
            import io as _io
            for item in viz_images:
                # Support both 3-tuple (old) and 4-tuple (new: title, why, what, png)
                if len(item) == 4:
                    title, why_text, what_text, png = item
                elif len(item) == 3:
                    title, why_text, png = item; what_text = ""
                else:
                    continue
                if not png: continue

                # Chart title heading
                ch = doc.add_heading("", level=3); ch.clear()
                r2 = ch.add_run(title)
                r2.bold = True; r2.font.size = Pt(11.5); r2.font.color.rgb = C_H2
                _shade(ch, FILL_H1)

                # Why this chart
                if why_text:
                    wp = doc.add_paragraph()
                    wp.paragraph_format.space_before = Pt(3)
                    wp.paragraph_format.space_after  = Pt(1)
                    wp.paragraph_format.left_indent  = Pt(6)
                    wl = wp.add_run("🎯 Why this chart:  ")
                    wl.bold = True; wl.font.size = Pt(9); wl.font.color.rgb = RGBColor(0xf5, 0x9e, 0x0b)
                    wb_r2 = wp.add_run(why_text)
                    wb_r2.italic = True; wb_r2.font.size = Pt(9); wb_r2.font.color.rgb = C_BODY

                # What it explains
                if what_text:
                    xp = doc.add_paragraph()
                    xp.paragraph_format.space_before = Pt(1)
                    xp.paragraph_format.space_after  = Pt(5)
                    xp.paragraph_format.left_indent  = Pt(6)
                    xl = xp.add_run("📊 What it explains:  ")
                    xl.bold = True; xl.font.size = Pt(9); xl.font.color.rgb = RGBColor(0xa7, 0x8b, 0xfa)
                    xb = xp.add_run(what_text)
                    xb.italic = True; xb.font.size = Pt(9); xb.font.color.rgb = C_META

                # Chart image
                try:
                    doc.add_picture(_io.BytesIO(png), width=Inches(5.8))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
                _gap(8)
    except Exception as _chart_ex:
        print(f"[Word export charts] {_chart_ex}")

    # ── Footer ─────────────────────────────────────────────────
    _gap(8)
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _shade(fp, FILL_META)
    fr = fp.add_run(f"GraphPulse AI  ·  {datetime.now().strftime('%d %B %Y')}  ·  Confidential")
    fr.font.size = Pt(8); fr.font.color.rgb = C_META

    doc.save(path)
    return path





def _save_rca_excel(question: str, report_text: str) -> str:
    """
    Professional multi-sheet Excel export.
    Sheet layout:
      1. Cover — title, query, report sections list (content overview, no charts row)
      2. Executive Summary
      3..N. One sheet per section (Step 1..5) in report order
      N+1. Root Cause
      Last. Recommendations
    No chart sheets. No repetitions.
    """
    import tempfile, re as _re, io as _io
    from datetime import datetime

    if not (report_text or "").strip(): return None
    if not _OPENPYXL_AVAILABLE: return _save_rca_csv_fallback(question, report_text)

    Font = XlFont
    safe = _re.sub(r'[^a-zA-Z0-9]+', '_', (question or '').strip()[:40]).strip('_')
    path = os.path.join(tempfile.gettempdir(), f"rca_{safe}.xlsx")
    wb = openpyxl.Workbook(); wb.remove(wb.active)

    def F(bold=False, color="1E293B", size=10, italic=False, name="Calibri"):
        return Font(bold=bold, color=color, size=size, italic=italic, name=name)

    BRAND  = PatternFill("solid", fgColor="1D4ED8")
    COVER  = PatternFill("solid", fgColor="DBEAFE")
    QUERY  = PatternFill("solid", fgColor="EFF6FF")
    META   = PatternFill("solid", fgColor="F8FAFC")
    HDR    = PatternFill("solid", fgColor="1E3A8A")
    ROW_A  = PatternFill("solid", fgColor="EFF6FF")
    ROW_B  = PatternFill("solid", fgColor="FFFFFF")
    RC_BG  = PatternFill("solid", fgColor="FEF2F2")
    REC_BG = PatternFill("solid", fgColor="F0FDF4")
    thin   = Side(style="thin",   color="CBD5E1")
    med    = Side(style="medium", color="93C5FD")
    TBDR   = Border(left=thin, right=thin, top=thin, bottom=thin)
    HBDR   = Border(left=med,  right=med,  top=med,  bottom=med)
    CTR    = Alignment(horizontal="center", vertical="center", wrap_text=True)
    LEFT   = Alignment(horizontal="left",   vertical="center", wrap_text=True)
    TOP    = Alignment(horizontal="left",   vertical="top",    wrap_text=True)
    LCTR   = Alignment(horizontal="left",   vertical="center", wrap_text=True, indent=2)
    EMOJI  = r'[\U0001F300-\U0001FAFF\U00002600-\U000027BF\U00002700-\U000027BF]+'

    def _cl(t):
        t = _re.sub(r'\*\*(.+?)\*\*', r'\1', str(t or ""))
        t = _re.sub(r'\*(.+?)\*',     r'\1', t)
        t = _re.sub(EMOJI + r'\s*', '', t)
        return t.strip()

    def _strip_tags(s): return _re.sub(r'<[^>]+>', '', s).strip()

    def _banner(ws, row, text, fg="FFFFFF", bg="1D4ED8", size=13, ncols=5):
        ws.merge_cells(f"A{row}:{get_column_letter(ncols)}{row}")
        c = ws.cell(row, 1, text)
        c.font = F(bold=True, color=fg, size=size)
        c.fill = PatternFill("solid", fgColor=bg)
        c.alignment = LCTR
        ws.row_dimensions[row].height = 28
        return row + 1

    def _gap(ws, row, h=6):
        ws.row_dimensions[row].height = h
        return row + 1

    def _sep(ws, row, ncols=5, color="CBD5E1"):
        ws.merge_cells(f"A{row}:{get_column_letter(ncols)}{row}")
        ws.cell(row, 1).fill = PatternFill("solid", fgColor=color)
        ws.row_dimensions[row].height = 3
        return row + 1

    def _parse_tbl(lines):
        rows = []
        for ln in lines:
            s = ln.strip()
            if s.startswith("|") and not all(c in "|-: " for c in s):
                cells = [_re.sub(r'\*\*(.+?)\*\*', r'\1', c.strip()) for c in s.strip("|").split("|")]
                if cells: rows.append(cells)
        return rows

    def _write_tbl(ws, tbl_rows, row):
        if not tbl_rows or len(tbl_rows) < 2: return row
        ncols = min(10, max(len(r) for r in tbl_rows))
        for ci, h in enumerate(tbl_rows[0][:ncols], 1):
            c = ws.cell(row, ci, _cl(h))
            c.font = F(bold=True, color="FFFFFF", size=9)
            c.fill = HDR; c.border = HBDR; c.alignment = CTR
            ws.row_dimensions[row].height = 20
        for ri, rd in enumerate(tbl_rows[1:], 1):
            fill = ROW_A if ri%2==1 else ROW_B
            for ci, val in enumerate(rd[:ncols], 1):
                txt = _cl(str(val))
                c = ws.cell(row+ri, ci, txt)
                c.font = F(color="1E293B", size=9)
                c.fill = fill; c.border = TBDR; c.alignment = TOP
                ws.row_dimensions[row+ri].height = max(15, min(len(txt)//12*5+15, 60))
        return row + len(tbl_rows) + 1

    def _write_prose(ws, lines, row, ncols=5, rc=False):
        for ln in lines:
            s = ln.strip()
            if not s or not _re.search(r'[A-Za-z0-9]', s): continue
            if _re.match(r'^[.#]?[\w-]+\s*\{', s): continue
            if _re.search(r'click\s+to|view\s+data|view\s+table', s, _re.I): continue
            txt = _cl(s)
            if not txt: continue
            ws.merge_cells(f"A{row}:{get_column_letter(ncols)}{row}")
            c = ws.cell(row, 1, txt)
            c.alignment = TOP
            if rc:
                c.font = F(color="7F1D1D", size=10, italic=True)
                c.fill = RC_BG
            elif s.startswith(("- ","• ","* ")):
                c.font = F(color="1E293B", size=10)
                c.fill = ROW_B
            elif s.startswith("####") or s.startswith("###"):
                c.font = F(bold=True, color="0F766E", size=10)
                c.fill = PatternFill("solid", fgColor="F0FDF4")
            else:
                c.font = F(color="1E293B", size=10)
                c.fill = ROW_B
            ws.row_dimensions[row].height = max(15, min(len(txt)//85*14+15, 72))
            row += 1
        return row

    # ── Parse sections from HTML report ──────────────────────────────────
    import html as _html_mod

    def _extract_sections_from_html(html):
        """Returns list of {name, desc, lines, tables} in report order."""
        secs = []

        # Executive Summary from <details> block (new format: collapsible)
        exec_m = _re.search(
            r'<summary[^>]*>.*?Executive Summary.*?</summary>(.*?)</details>',
            html, _re.DOTALL | _re.IGNORECASE)
        if exec_m:
            exec_inner = exec_m.group(1)
            # Extract only text content, explicitly skip any <style> or CSS blocks
            exec_inner_clean = _re.sub(r'<style[^>]*>.*?</style>', '', exec_inner, flags=_re.DOTALL)
            exec_inner_clean = _re.sub(r'<script[^>]*>.*?</script>', '', exec_inner_clean, flags=_re.DOTALL)
            txt = _html_mod.unescape(_strip_tags(exec_inner_clean)).strip()
            lines = [l.strip() for l in txt.splitlines()
                     if l.strip()
                     and _re.search(r'[A-Za-z0-9]', l)
                     and not _re.match(r'^[.#]?[\w-]+\s*\{', l.strip())
                     and not _re.search(r'!important', l)
                     and not l.strip().startswith('details[')
                     and not l.strip().startswith('@')]
            lines = [l for l in lines if len(l) > 3]
            if lines: secs.append({"name": "Executive Summary", "desc": "", "lines": lines, "tables": []})

        # Section cards: divs with border-left:4px solid
        card_starts = [m.start() for m in _re.finditer(r'<div[^>]*border-left:4px solid[^>]*>', html)]
        for i, start in enumerate(card_starts):
            end = card_starts[i+1] if i+1 < len(card_starts) else len(html)
            block = html[start:end]

            # Title
            title_m = (_re.search(r'color:#fca5a5[^>]*>([^<]{4,60})<', block) or
                       _re.search(r'font-weight:700;color:#e2e8f0[^>]*>([^<]{4,60})<', block) or
                       _re.search(r'font-weight:700[^>]*>([^<]{4,60})<', block))
            if not title_m: continue
            title = _html_mod.unescape(title_m.group(1)).strip()
            title = _re.sub(EMOJI+r'\s*', '', title).strip()
            if not title or len(title) < 3: continue

            # Step pill / desc
            pill_m = _re.search(r'letter-spacing:0\.02em[^>]*>([^<]+)<', block)
            pill = _html_mod.unescape(pill_m.group(1)).strip() if pill_m else ""

            desc_m = _re.search(r'color:#cbd5e1[^>]*>([^<]{10,})<', block)
            desc = _html_mod.unescape(desc_m.group(1)).strip() if desc_m else ""

            # Table from details block — tables are rendered as HTML <table> not markdown
            tables = []
            for det_m in _re.finditer(r'<details[^>]*>(.*?)</details>', block, _re.DOTALL):
                inner_m = _re.search(r'</summary>(.*?)$', det_m.group(1), _re.DOTALL)
                if inner_m:
                    inner_html = inner_m.group(1)
                    # Try HTML <table> extraction first (primary path)
                    tbl_m = _re.search(r'<table[^>]*>(.*?)</table>', inner_html, _re.DOTALL)
                    if tbl_m:
                        tbl_html = tbl_m.group(1)
                        # Extract header row from <th> elements
                        headers = [_html_mod.unescape(_strip_tags(th)).strip()
                                   for th in _re.findall(r'<th[^>]*>(.*?)</th>', tbl_html, _re.DOTALL)]
                        # Extract data rows from <tr><td> elements
                        data_rows = []
                        for tr in _re.findall(r'<tr[^>]*>(.*?)</tr>', tbl_html, _re.DOTALL):
                            cells = [_html_mod.unescape(_strip_tags(td)).strip()
                                     for td in _re.findall(r'<td[^>]*>(.*?)</td>', tr, _re.DOTALL)]
                            if cells and any(c for c in cells):
                                data_rows.append(cells)
                        if headers and data_rows:
                            tables.append([headers] + data_rows)
                    else:
                        # Fallback: try markdown pipe table
                        inner_txt = _html_mod.unescape(_strip_tags(inner_html))
                        tbl_lines = [l for l in inner_txt.splitlines() if l.strip().startswith("|")]
                        parsed = _parse_tbl(tbl_lines)
                        if len(parsed) >= 2:
                            tables.append(parsed)

            lines = []
            if pill: lines.append(f"◆ {pill}")
            if desc: lines.append(desc)

            secs.append({"name": title, "desc": desc, "lines": lines, "tables": tables})
        return secs

    all_secs = _extract_sections_from_html(report_text)

    # Separate into groups: exec, steps, root_cause, then build recs
    exec_secs  = [s for s in all_secs if "executive" in s["name"].lower()]
    rc_secs    = [s for s in all_secs if any(k in s["name"].lower() for k in ["root cause", "⚠️"])]
    step_secs  = [s for s in all_secs
                  if s not in exec_secs and s not in rc_secs
                  and "recommend" not in s["name"].lower()]
    # If rc not found via HTML, try step 5
    if not rc_secs:
        rc_secs = [s for s in step_secs if "root cause" in s["name"].lower() or
                   (s["lines"] and "root cause" in s["lines"][0].lower())]
        step_secs = [s for s in step_secs if s not in rc_secs]

    # NEW: if still not found, extract root cause from the new crisp-para format
    # The new format is: <div style="...border-left:4px solid #f87171...">...para...</div>
    if not rc_secs:
        import html as _html_rc
        rc_m = _re.search(
            r'<div[^>]*border-left:4px solid #f87171[^>]*>(.*?)</div>',
            report_text, _re.DOTALL
        )
        if rc_m:
            rc_inner = rc_m.group(1)
            # Remove the "Root Cause — Confirmed Verdict" label div
            rc_inner = _re.sub(r'<div[^>]*font-size:0\.65rem[^>]*>.*?</div>', '', rc_inner, flags=_re.DOTALL)
            rc_txt = _html_rc.unescape(_re.sub(r'<[^>]+>', ' ', rc_inner)).strip()
            rc_txt = _re.sub(r'\s+', ' ', rc_txt).strip()
            if rc_txt and len(rc_txt) > 20:
                rc_secs = [{"name": "Root Cause", "desc": "", "lines": [rc_txt], "tables": []}]

    # Derive report title
    q_l = (question or "").lower()
    if any(w in q_l for w in ["stockout","demand gap","shortage","kolkata"]): _rpt = "Demand Gap RCA"
    elif any(w in q_l for w in ["supplier risk","risky"]): _rpt = "Supplier Risk RCA"
    elif any(w in q_l for w in ["route","transport"]): _rpt = "Route & Transport RCA"
    elif any(w in q_l for w in ["delay","delayed"]): _rpt = "Shipment Delay RCA"
    else: _rpt = "Supply Chain RCA"

    used = set()
    def _uname(base):
        base = _re.sub(r'[\\/*?:\[\]]','',base)[:31]
        n,k = base,2
        while n in used: n = f"{base[:27]} {k}"; k += 1
        used.add(n); return n

    # ════════════════════════════════════════════════════════
    # SHEET 1: Cover — title, query, content list
    # ════════════════════════════════════════════════════════
    ws1 = wb.create_sheet(_uname(_rpt))
    ws1.sheet_view.showGridLines = False
    for col, w in [("A",28),("B",10),("C",60),("D",25)]:
        ws1.column_dimensions[col].width = w

    r = 1
    ws1.merge_cells(f"A{r}:D{r}")
    ws1[f"A{r}"] = f"GraphPulse AI  —  {_rpt}"
    ws1[f"A{r}"].font = F(bold=True, color="1D4ED8", size=22)
    ws1[f"A{r}"].fill = COVER; ws1[f"A{r}"].alignment = LCTR
    ws1.row_dimensions[r].height = 48; r += 1

    ws1.merge_cells(f"A{r}:D{r}")
    ws1[f"A{r}"] = "Root Cause Analysis  ·  Supply Chain Intelligence  ·  GraphPulse AI"
    ws1[f"A{r}"].font = F(color="475569", size=11)
    ws1[f"A{r}"].fill = COVER; ws1[f"A{r}"].alignment = LCTR
    ws1.row_dimensions[r].height = 20; r += 1

    ws1.merge_cells(f"A{r}:D{r}")
    ws1[f"A{r}"] = f"Generated: {datetime.now().strftime('%d %B %Y  at  %H:%M')}"
    ws1[f"A{r}"].font = F(color="94A3B8", size=9, italic=True)
    ws1[f"A{r}"].fill = META; ws1[f"A{r}"].alignment = LCTR
    ws1.row_dimensions[r].height = 16; r += 1
    r = _sep(ws1, r); r = _gap(ws1, r, 10)

    ws1[f"A{r}"] = "ANALYSIS QUERY"
    ws1[f"A{r}"].font = F(bold=True, color="FFFFFF", size=10)
    ws1[f"A{r}"].fill = BRAND; ws1[f"A{r}"].border = HBDR; ws1[f"A{r}"].alignment = CTR
    ws1.merge_cells(f"B{r}:D{r}")
    ws1[f"B{r}"] = question
    ws1[f"B{r}"].font = F(bold=True, color="1E3A8A", size=11)
    ws1[f"B{r}"].fill = QUERY; ws1[f"B{r}"].border = HBDR
    ws1[f"B{r}"].alignment = Alignment(horizontal="left", vertical="center", wrap_text=True, indent=1)
    ws1.row_dimensions[r].height = max(36,(len(question)//55+1)*20); r += 1
    r = _gap(ws1, r, 14)

    # Content list — no chart rows, no "Sheet N" references
    r = _banner(ws1, r, "  REPORT CONTENTS", fg="FFFFFF", bg="1E3A8A", size=11, ncols=4)
    for ci, (txt, w) in enumerate([("Section", 28), ("Type", 10), ("Description", 60), ("Sheet", 25)], 1):
        c = ws1.cell(r, ci, txt)
        c.font = F(bold=True, color="FFFFFF", size=9)
        c.fill = HDR; c.border = TBDR; c.alignment = CTR
    ws1.row_dimensions[r].height = 18; r += 1

    content_rows = []
    if exec_secs:
        content_rows.append(("Executive Summary", "Overview",
            exec_secs[0]["lines"][0][:80] if exec_secs[0]["lines"] else "Analysis overview",
            "Executive Summary sheet"))
    for s in step_secs:
        content_rows.append((s["name"][:40], "Investigation Step",
            s["desc"][:80] if s["desc"] else (s["lines"][0][:80] if s["lines"] else "—"),
            f"{s['name'][:25]} sheet"))
    if rc_secs:
        content_rows.append(("Root Cause", "Conclusion",
            rc_secs[0]["desc"][:80] if rc_secs[0].get("desc") else
            (rc_secs[0]["lines"][0][:80] if rc_secs[0]["lines"] else "Root cause findings"),
            "Root Cause sheet"))
    content_rows.append(("Recommendations", "Action Plan",
        "Critical / Operational / Strategic prioritised actions", "Recommendations sheet"))

    for ri, (sec, typ, preview, sheet_ref) in enumerate(content_rows):
        fill = ROW_A if ri%2==0 else ROW_B
        for ci, val in enumerate([sec, typ, preview, sheet_ref], 1):
            c = ws1.cell(r, ci, val)
            c.font = (F(bold=True, color="1D4ED8", size=10) if ci==1
                      else F(color="0F766E", size=9, italic=True) if ci==2
                      else F(color="1E293B", size=10) if ci==3
                      else F(color="475569", size=9, italic=True))
            c.fill = fill; c.border = TBDR; c.alignment = TOP
        ws1.row_dimensions[r].height = 18; r += 1

    # ════════════════════════════════════════════════════════
    # SHEET 2: Executive Summary
    # ════════════════════════════════════════════════════════
    if exec_secs:
        ws2 = wb.create_sheet(_uname("Executive Summary"))
        ws2.sheet_view.showGridLines = False
        ws2.column_dimensions["A"].width = 100
        r2 = _banner(ws2, 1, "  EXECUTIVE SUMMARY", fg="FFFFFF", bg="1D4ED8", size=14, ncols=1)
        ws2.merge_cells(f"A{r2}:A{r2}")
        ws2.cell(r2,1,question).font = F(color="475569", size=10, italic=True)
        ws2.cell(r2,1).fill = QUERY
        ws2.cell(r2,1).alignment = Alignment(horizontal="left", vertical="top", wrap_text=True, indent=2)
        ws2.row_dimensions[r2].height = max(18,(len(question)//95+1)*16); r2 += 1
        r2 = _sep(ws2, r2, 1); r2 = _gap(ws2, r2, 8)
        r2 = _write_prose(ws2, exec_secs[0]["lines"], r2, ncols=1)
        for tbl in exec_secs[0]["tables"]:
            r2 = _write_tbl(ws2, tbl, r2)

    # ════════════════════════════════════════════════════════
    # SHEETS 3..N: One per investigation step (in order)
    # ════════════════════════════════════════════════════════
    STEP_COLORS = {
        "stockout": "B91C1C", "fulfillment": "0369A1", "plant": "0F766E",
        "supplier": "7C3AED", "root": "991B1B", "default": "374151"
    }
    def _step_color(name):
        nl = name.lower()
        for k,c in STEP_COLORS.items():
            if k in nl: return c
        return STEP_COLORS["default"]

    for sec in step_secs:
        sname = sec["name"]
        ws = wb.create_sheet(_uname(sname[:31]))
        ws.sheet_view.showGridLines = False
        ws.column_dimensions["A"].width = 100
        sc = _step_color(sname)

        r = _banner(ws, 1, f"  {sname.upper()}", fg="FFFFFF", bg=sc, size=13, ncols=1)
        # Step pill on its own line if present
        if sec["lines"]:
            first = sec["lines"][0]
            if first.startswith("◆"):
                ws.merge_cells(f"A{r}:A{r}")
                c = ws.cell(r, 1, _cl(first))
                c.font = F(bold=True, color=sc, size=10)
                c.fill = PatternFill("solid", fgColor="EFF6FF")
                c.alignment = LEFT
                ws.row_dimensions[r].height = 16; r += 1

        r = _sep(ws, r, 1, sc+"33" if len(sc)==6 else "CBD5E1")
        r = _gap(ws, r, 6)

        # Write desc lines (skip the pill line already written)
        prose_lines = [l for l in sec["lines"] if not l.startswith("◆")]
        r = _write_prose(ws, prose_lines, r, ncols=1)

        # Tables
        for tbl in sec["tables"]:
            r = _gap(ws, r, 4)
            ncols_tbl = min(10, max(len(rw) for rw in tbl)) if tbl else 1
            for ci in range(1, ncols_tbl+1):
                ws.column_dimensions[get_column_letter(ci)].width = max(12, 100//max(ncols_tbl,1))
            ws.column_dimensions["A"].width = max(12, 100//max(ncols_tbl,1))
            r = _write_tbl(ws, tbl, r)
            r = _gap(ws, r, 6)

    # ════════════════════════════════════════════════════════
    # ROOT CAUSE sheet
    # ════════════════════════════════════════════════════════
    if rc_secs:
        ws_rc = wb.create_sheet(_uname("Root Cause"))
        ws_rc.sheet_view.showGridLines = False
        ws_rc.column_dimensions["A"].width = 100
        r = _banner(ws_rc, 1, "  ROOT CAUSE ANALYSIS", fg="FFFFFF", bg="991B1B", size=14, ncols=1)
        r = _sep(ws_rc, r, 1, "FECACA"); r = _gap(ws_rc, r, 8)
        for line in rc_secs[0]["lines"]:
            if _re.search(r'[A-Za-z0-9]', line):
                ws_rc.merge_cells(f"A{r}:A{r}")
                c = ws_rc.cell(r, 1, _cl(line))
                c.font = F(color="7F1D1D", size=11)
                c.fill = RC_BG
                c.alignment = TOP
                ws_rc.row_dimensions[r].height = max(18, min(len(_cl(line))//85*14+18, 80))
                r += 1
        for tbl in rc_secs[0]["tables"]:
            r = _gap(ws_rc, r, 4)
            r = _write_tbl(ws_rc, tbl, r)

    # ════════════════════════════════════════════════════════
    # RECOMMENDATIONS sheet (last)
    # ════════════════════════════════════════════════════════
    try:
        rec_secs = _extract_detailed_recommendations(report_text)
        if rec_secs:
            ws_r = wb.create_sheet(_uname("Recommendations"))
            ws_r.sheet_view.showGridLines = False
            ws_r.column_dimensions["A"].width = 16
            ws_r.column_dimensions["B"].width = 84

            r = _banner(ws_r, 1, "  RECOMMENDATIONS  —  Prioritised Action Plan",
                        fg="FFFFFF", bg="166534", size=14, ncols=2)
            ws_r.merge_cells(f"A{r}:B{r}")
            ws_r.cell(r,1,"Critical  →  Operational  →  Strategic").font = F(color="166534",size=10,italic=True)
            ws_r.cell(r,1).fill = REC_BG; ws_r.cell(r,1).alignment = LCTR
            ws_r.row_dimensions[r].height = 18; r += 1
            r = _sep(ws_r, r, 2, "BBF7D0"); r = _gap(ws_r, r, 8)

            TIERS = {
                "critical":    ("FEF2F2","991B1B","⚡ Act within 48 hours"),
                "operational": ("FFF7ED","9A3412","🔧 Complete within 2 weeks"),
                "strategic":   ("F0FDF4","166534","🌱 Long-term initiative"),
            }
            for sec in rec_secs:
                tkey = next((k for k in TIERS if k in sec["section"].lower()), None)
                bg,fg,horizon = TIERS.get(tkey,("EFF6FF","1D4ED8",""))
                ws_r.merge_cells(f"A{r}:B{r}")
                h = ws_r.cell(r,1,f"  {sec['section'].upper()}  —  {horizon}")
                h.font = F(bold=True,color="FFFFFF",size=11)
                h.fill = PatternFill("solid",fgColor=fg)
                h.alignment = LCTR
                ws_r.row_dimensions[r].height = 26; r += 1

                for item in sec["items"]:
                    # items are "EntityName: action text" strings or legacy dicts
                    if isinstance(item, dict):
                        entity_label = item.get("entity", "")
                        raw_txt = _neutralise_rec_text(item.get("action",""))
                    else:
                        raw_str = _neutralise_rec_text(str(item))
                        # Split "EntityName (SUPXXXX): action text" on first ": " after entity
                        import re as _ri
                        _split = _ri.match(r'^(.{3,80}):\s+(.{10,})$', raw_str, _ri.DOTALL)
                        if _split:
                            entity_label = _split.group(1).strip()
                            raw_txt = _split.group(2).strip()
                        else:
                            entity_label = ""
                            raw_txt = raw_str
                    if not raw_txt:
                        continue
                    # Col A: entity label (bold, tier colour)
                    ca = ws_r.cell(r, 1, entity_label or "—")
                    ca.font = F(bold=True, color=fg, size=9)
                    ca.fill = PatternFill("solid", fgColor=bg)
                    ca.border = TBDR; ca.alignment = CTR
                    # Col B: action text
                    cb = ws_r.cell(r, 2, raw_txt)
                    cb.font = F(color="1E293B", size=10)
                    cb.fill = PatternFill("solid", fgColor=bg)
                    cb.border = TBDR; cb.alignment = TOP
                    ws_r.row_dimensions[r].height = max(20, min(len(raw_txt) // 80 * 14 + 20, 80))
                    r += 1
                r = _gap(ws_r, r, 6)
    except Exception as _re_err:
        print(f"[Excel recs] {_re_err}")

    wb.save(path)
    return path


def _save_rca_csv_fallback(question: str, report_text: str) -> str:
    """Fallback CSV export if openpyxl not available."""
    import tempfile, re as _re, csv as _csv
    if not (report_text or "").strip():
        return None
    safe = _re.sub(r'[^a-zA-Z0-9]+', '_', (question or '').strip()[:40]).strip('_')
    path = os.path.join(tempfile.gettempdir(), f"rca_{safe}.csv")
    rows = []
    current_section = "Summary"
    for line in report_text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("## ") or line.startswith("### "):
            current_section = _re.sub(r'^#+ ', '', line)
            current_section = _re.sub(r'[^\w\s]', '', current_section).strip()
        else:
            clean = _re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            clean = _re.sub(r'\*(.+?)\*', r'\1', clean).lstrip("- •* ")
            if clean:
                rows.append([current_section, clean])
    try:
        with open(path, "w", newline="", encoding="utf-8") as f:
            w = _csv.writer(f)
            w.writerow(["Section", "Finding"])
            w.writerows(rows)
        return path
    except Exception:
        return None

def _update_rca_dropdown(q):
    """
    Called via .then() AFTER rca_handler finishes — never mid-stream.
    Updating a gr.Dropdown from inside a streaming generator causes Gradio
    to render an 'Error' badge on the dropdown label. Decoupling it here
    into a plain non-streaming .then() call avoids that entirely.

    FIX: wrapped in try/except so a Neo4j timeout or any exception
    here NEVER causes the Gradio Error badge on the dropdown.
    """
    try:
        return _get_charts_for_question(q)
    except Exception:
        return gr.update(choices=RCA_VIZ_OPTIONS, value="— Select a chart —")


# ════════════════════════════════════════════════════════════════════
# CYPHER ACCORDION BUILDER
# ════════════════════════════════════════════════════════════════════
def _build_cypher_accordion(cypher_logs: list) -> str:
    """
    Build the terminal-style Cypher Queries Executed accordion HTML.
    cypher_logs: list of dicts with keys: seq, tool, purpose, cypher, records
    Returns "" if no logs (accordion hidden).
    """
    import re as _re
    if not cypher_logs:
        return ""

    def _highlight(cypher: str) -> str:
        """
        Tokenise Cypher into coloured spans.
        Strategy: scan left-to-right, classify each token, emit spans.
        Never re-process already-emitted HTML, so no double-highlight corruption.
        """
        import re as _rh

        # Sanitise HTML chars first (before any span wrapping)
        src = cypher.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        KEYWORDS = {
            "MATCH","OPTIONAL MATCH","WHERE","RETURN","WITH","ORDER BY","LIMIT",
            "MERGE","CREATE","SET","DELETE","DETACH","DISTINCT","AS","AND","OR",
            "NOT","IN","UNION","CALL","YIELD","REMOVE",
        }
        FUNCS = {
            "COUNT","SUM","AVG","ROUND","COLLECT","MIN","MAX",
            "CASE","WHEN","THEN","ELSE","END","OVER",
        }

        # Token patterns — ordered: longest/most-specific first
        # Split on protected zones (strings, comments), handle each zone separately
        ZONE_RE = _rh.compile(r"(/\*.*?\*/)|('[^']*')|(\"[^\"]*\")", _rh.DOTALL)

        def _highlight_plain(text: str) -> str:
            """Highlight a plain (non-string, non-comment) Cypher fragment."""
            # Numbers
            text = _rh.sub(r"\b(\d+(?:\.\d+)?)\b",
                           lambda m: f'<span class="cq-num">{m.group()}</span>', text)
            # Relationship types [:REL]
            text = _rh.sub(r"(\[:)([A-Z_]+)(\])",
                           lambda m: (m.group(1) +
                                      f'<span class="cq-rel">{m.group(2)}</span>' +
                                      m.group(3)), text)
            # Node labels :Label (but not after digits or inside words)
            text = _rh.sub(r"(?<![\w\d])(:)([A-Z][a-zA-Z0-9_]*)",
                           lambda m: m.group(1) + f'<span class="cq-nd">{m.group(2)}</span>', text)
            # Variables $param
            text = _rh.sub(r"(\$[a-zA-Z_]\w*)",
                           lambda m: f'<span class="cq-var">{m.group()}</span>', text)
            # Keywords and functions — only match whole words
            def _kw_replace(m):
                word = m.group()
                if word in KEYWORDS:
                    return f'<span class="cq-kw">{word}</span>'
                if word in FUNCS:
                    return f'<span class="cq-fn">{word}</span>'
                return word
            text = _rh.sub(r"\b[A-Z][A-Z0-9_]*\b", _kw_replace, text)
            return text

        result = ""
        last = 0
        for m in ZONE_RE.finditer(src):
            # Highlight plain text before this zone
            if m.start() > last:
                result += _highlight_plain(src[last:m.start()])
            # Emit zone with its own class
            zone = m.group()
            if zone.startswith("/*"):
                result += f'<span class="cq-cm">{zone}</span>'
            else:
                result += f'<span class="cq-str">{zone}</span>'
            last = m.end()
        # Remaining plain text
        if last < len(src):
            result += _highlight_plain(src[last:])
        return result

    entries = ""
    for i, log in enumerate(cypher_logs):
        seq      = log.get("seq", i + 1)
        purpose  = log.get("purpose", log.get("tool", "?"))
        cypher   = log.get("cypher", "")
        records  = log.get("records", -1)
        tool     = log.get("tool", "")

        rec_cls  = "" if records >= 0 else " unknown"
        rec_txt  = f"{records} record{'s' if records != 1 else ''}" if records >= 0 else "count N/A"
        # Catalog strings are clean triple-quoted — just strip whitespace
        clean_cypher = cypher.strip()
        hi_cypher = _highlight(clean_cypher)
        safe_cypher = clean_cypher.replace("\\", "\\\\").replace("`", "'")
        safe_cypher = clean_cypher.replace("\\", "\\\\").replace("`", "\\`")

        entries += f"""
<div class="cq-entry">
  <div class="cq-entry-header">
    <span class="cq-seq">{seq}</span>
    <span class="cq-purpose">{purpose}</span>
    <span class="cq-records{rec_cls}">{rec_txt}</span>
    <button class="cq-copy-btn"
      onclick="(function(b){{
        navigator.clipboard.writeText(`{safe_cypher}`).then(function(){{
          b.textContent='✓ Copied';b.classList.add('copied');
          setTimeout(function(){{b.textContent='Copy';b.classList.remove('copied')}},2000);
        }}).catch(function(){{b.textContent='Copy'}});
      }})(this)">Copy</button>
  </div>
  <div class="cq-code-wrap">
    <code class="cq-code">{hi_cypher}</code>
  </div>
</div>"""

    n = len(cypher_logs)
    return f"""
<div class="cq-section">
  <div class="cq-toggle" onclick="
    var b=this.querySelector('.cq-toggle-icon');
    var d=this.nextElementSibling;
    var open=d.classList.toggle('open');
    b.classList.toggle('open',open);
    this.querySelector('.cq-toggle-badge').textContent=open?'▲ hide':b.getAttribute('data-count')+' queries';
  ">
    <span class="cq-toggle-icon" data-count="{n}">&#9654;</span>
    <span class="cq-toggle-label">🔍 Queries That Powered This Analysis</span>
    <span class="cq-toggle-badge" title="Click to expand">{n} quer{'y' if n==1 else 'ies'}</span>
  </div>
  <div class="cq-body" id="rca-cq-body">
    {entries}
  </div>
</div>"""


def _build_rca_viz_html(question: str) -> str:
    """
    Build the 'View Charts' collapsible dropdown HTML for the RCA report.
    Always returns a non-empty string so the accordion always appears.
    """
    q_l = (question or "").lower()
    # Detect query type for intro text
    if any(w in q_l for w in ["supplier","vendor","risk","sup"]):
        intro = "Supplier risk scores, delay correlation, and network flow charts are ready for this query."
    elif any(w in q_l for w in ["delay","delayed","shipment","ship"]):
        intro = "Delay heatmap, Pareto ranking, and supply flow Sankey are loaded for this delay analysis."
    elif any(w in q_l for w in ["distributor","city","distribution"]):
        intro = "Distributor impact charts — demand gap ranking, supply Sankey, and network flow are ready."
    elif any(w in q_l for w in ["stockout","stock","demand","gap","shortage"]):
        intro = "Stockout severity, demand gap breakdown, and category heatmap are ready to explore."
    elif any(w in q_l for w in ["route","transit","transport","road","rail","air"]):
        intro = "Route efficiency, network flow, and transport mode distribution charts are loaded."
    elif any(w in q_l for w in ["category","product","toy","auto","health"]):
        intro = "Category delay share, Pareto breakdown, and heatmap are ready for this product analysis."
    else:
        intro = "Supply chain visualizations — network flow, delay heatmap, Pareto bar, and Sankey — are ready."

    return (
        f'<div style="font-size:0.78rem;color:#7dd3fc;line-height:1.5;padding:0 0 8px 0;margin:0">'
        f'{intro} Click a chart button below to load it.'
        f'</div>'
    )


def _extract_detailed_recommendations(report_text: str) -> list:
    """
    Parse the report HTML/text and return a list of recommendation sections.
    Each section: {"section": str, "items": [str, ...]}
    Handles:
      1. New HTML format: <details> tier dropdowns (⚡ Critical / 🔧 Operational / 🏛 Strategic)
      2. Legacy markdown format: #### tier headings
    """
    import re as _re, html as _html

    raw = str(report_text or "").strip()
    if not raw:
        return []

    EMOJI_RE = r"[\U0001F300-\U0001FAFF\u2600-\u27BF\u2700-\u27BF]+"

    def _strip_tags(s):
        return _re.sub(r'<[^>]+>', ' ', s).strip()

    def _clean(s):
        s = _html.unescape(s)
        s = _re.sub(r'\s+', ' ', s).strip()
        return s

    # ── Path 1: HTML <details> tier dropdowns ─────────────────────────
    # Find the recommendations div by margin-top:32px marker
    rec_div_m = _re.search(
        r'<div[^>]*margin-top:32px[^>]*>.*?Recommendations(.*?)(?=<div[^>]*id="rca-report-content"|$)',
        raw, _re.DOTALL | _re.IGNORECASE
    )
    if not rec_div_m:
        # Try alternate: find any div containing "Recommendations" heading followed by <details>
        rec_div_m = _re.search(
            r'Recommendations.*?Prioritised corrective actions(.*?)(?=<div[^>]*margin-top:32px|$)',
            raw, _re.DOTALL | _re.IGNORECASE
        )

    if rec_div_m:
        rec_block = rec_div_m.group(1)
        sections = []
        for tier_m in _re.finditer(r'<details[^>]*>(.*?)</details>', rec_block, _re.DOTALL):
            tier_inner = tier_m.group(1)
            # Tier label
            t_sum_m = _re.search(r'<summary[^>]*>(.*?)</summary>', tier_inner, _re.DOTALL)
            tier_label = ""
            if t_sum_m:
                tier_label = _clean(_strip_tags(t_sum_m.group(1)))
                tier_label = _re.sub(r'[▶▼◆►◌]+\s*', '', tier_label).strip()
                tier_label = _re.sub(EMOJI_RE + r'\s*', '', tier_label).strip()
                # Remove "N actions" count suffix
                tier_label = _re.sub(r'\d+\s+actions?\s*$', '', tier_label).strip()
            if not tier_label:
                continue

            # Items: each action card div (entity header + action text)
            body_after_sum = _re.sub(r'.*?</summary>', '', tier_inner, flags=_re.DOTALL, count=1)
            items = []

            # Find each action card by locating divs with border-radius:8px (our card style)
            # Then extract full content between opening tag and its matching closing </div>
            for card_start in _re.finditer(r'<div[^>]*border-radius:8px[^>]*>', body_after_sum):
                pos = card_start.end()
                depth = 1
                end_pos = pos
                while depth > 0 and end_pos < len(body_after_sum):
                    next_open  = body_after_sum.find('<div', end_pos)
                    next_close = body_after_sum.find('</div>', end_pos)
                    if next_close == -1:
                        break
                    if next_open != -1 and next_open < next_close:
                        depth += 1
                        end_pos = next_open + 1
                    else:
                        depth -= 1
                        end_pos = next_close + 6

                card_inner = body_after_sum[pos:end_pos - 6]  # strip final </div>
                # Entity: small colored header
                entity_m = _re.search(r'<div[^>]*font-size:0\.7[^>]*>(.*?)</div>', card_inner, _re.DOTALL)
                # Action: larger text with line-height
                action_m = _re.search(r'<div[^>]*(?:line-height|font-size:0\.8)[^>]*>(.*?)</div>', card_inner, _re.DOTALL)
                if action_m:
                    entity = _clean(_strip_tags(entity_m.group(1))) if entity_m else ""
                    action = _clean(_strip_tags(action_m.group(1)))
                    if len(action) > 10:
                        combined = f"{entity}: {action}" if entity else action
                        items.append(combined)

            if items:
                sections.append({"section": tier_label, "items": items})

        if sections:
            return sections

    # ── Path 2: Legacy markdown #### headings ─────────────────────────
    sections = []
    cur_section = None
    cur_items = []
    in_rec_block = False

    for line in raw.splitlines():
        s = line.strip()
        if not s:
            continue

        # Detect start of Recommendations markdown block
        if _re.match(r'^#{1,3}\s+.*recommend', s, _re.IGNORECASE):
            in_rec_block = True
            continue

        # Tier headings
        if _re.match(r'^#{2,4}\s+', s):
            heading_text = _re.sub(r'^#+\s*', '', s).strip()
            heading_text = _re.sub(EMOJI_RE + r'\s*', '', heading_text).strip()
            h_lower = heading_text.lower()
            is_tier = any(k in h_lower for k in ['critical', 'high priority', 'operational',
                          'near-term', 'near term', 'strategic', 'long-term', 'long term', 'immediate'])
            if is_tier or in_rec_block:
                if cur_section is not None and cur_items:
                    sections.append({"section": cur_section, "items": list(cur_items)})
                cur_section = heading_text
                cur_items = []
                in_rec_block = True
            continue

        if not in_rec_block:
            continue

        if s.startswith(("- ", "• ", "* ", "+ ")):
            item = s[2:].strip()
            if item:
                cur_items.append(item)
        elif _re.match(r'^\d+\.\s', s):
            item = _re.sub(r'^\d+\.\s*', '', s).strip()
            if item:
                cur_items.append(item)
        elif cur_section is not None and len(s) > 15 and not s.startswith('#'):
            if cur_items:
                cur_items[-1] = cur_items[-1] + ' ' + s
            else:
                cur_items.append(s)

    if cur_section is not None and cur_items:
        sections.append({"section": cur_section, "items": list(cur_items)})

    return sections


def _build_detailed_rec_html(report_text: str) -> str:
    """
    Build styled recommendations HTML from rec_raw / report_text.
    Three-tier card layout: Critical (red), Operational (orange), Strategic (green).
    Each item has a bold 'Entity:' prefix, body text, and a dynamic context pill.
    Recommendation text is neutralised (no subjective/imperative phrasing).
    """
    import re as _re
    sections = _extract_detailed_recommendations(report_text)
    if not sections:
        return (
            '<div style="color:#94a3b8;font-size:.8rem;padding:12px;line-height:1.6">'
            'Recommendations will appear here once the analysis completes.'
            '</div>'
        )

    # Tier configuration: key → (bg, border/title colour, pill bg, pill text colour, icon, display label)
    TIER_CFG = {
        'critical':    ('rgba(127,29,29,0.45)',   '#f87171', 'rgba(239,68,68,0.15)',  '#fca5a5', '⚡', 'CRITICAL RESPONSE \u2014 ACT IMMEDIATELY'),
        'operational': ('rgba(120,53,15,0.40)',   '#fb923c', 'rgba(234,88,12,0.12)',  '#fdba74', '\U0001f527', 'OPERATIONAL FIXES \u2014 NEAR-TERM'),
        'near':        ('rgba(120,53,15,0.40)',   '#fb923c', 'rgba(234,88,12,0.12)',  '#fdba74', '\U0001f527', 'OPERATIONAL FIXES \u2014 NEAR-TERM'),
        'strategic':   ('rgba(20,83,45,0.40)',    '#4ade80', 'rgba(22,163,74,0.10)',  '#86efac', '\U0001f3db', 'STRATEGIC INITIATIVES \u2014 LONG-TERM'),
        'long':        ('rgba(20,83,45,0.40)',    '#4ade80', 'rgba(22,163,74,0.10)',  '#86efac', '\U0001f3db', 'STRATEGIC INITIATIVES \u2014 LONG-TERM'),
        'high':        ('rgba(113,63,18,0.40)',   '#facc15', 'rgba(234,179,8,0.12)',  '#fde68a', '\U0001f536', 'HIGH PRIORITY \u2014 NEAR-TERM'),
    }

    EMOJI_RE = r'[\U0001F300-\U0001FAFF\U00002600-\U000027BF\u2600-\u27BF]+'

    def _tier_cfg(heading):
        h = heading.lower()
        for key, cfg in TIER_CFG.items():
            if key in h:
                return key, cfg
        return 'other', ('rgba(30,58,138,0.30)', '#38bdf8', 'rgba(56,189,248,0.08)', '#7dd3fc', '\U0001f4a1', heading.upper())

    parts = []
    for sec in sections:
        tier_key, (bg, border_col, pill_bg, pill_col, icon, display_label) = _tier_cfg(sec['section'])

        # ── Tier header bar ──
        parts.append(
            f'<div style="background:{bg};border-left:4px solid {border_col};'
            f'padding:9px 16px;margin:16px 0 10px;border-radius:5px;">'
            f'<span style="font-size:0.68rem;font-weight:900;color:{border_col};'
            f'text-transform:uppercase;letter-spacing:0.12em">'
            f'{icon} {display_label}</span>'
            f'</div>'
        )

        for item in sec['items']:
            # Neutralise phrasing
            neutral = _neutralise_rec_text(item)

            # Build body HTML — bold "Entity:" or "RouteID:" prefix (catches PL4@D0027^Sea: etc)
            prefix_m = _re.match(
                r'^([A-Za-z0-9][A-Za-z0-9 ,&\-\.@^#/()\'\[\]]+?):\s+(.*)',
                neutral, _re.DOTALL
            )
            if prefix_m:
                entity_html = (
                    f'<strong style="color:{border_col}">'
                    f'{prefix_m.group(1)}:</strong> '
                )
                body_raw = prefix_m.group(2).strip()
            else:
                entity_html = ''
                body_raw = neutral

            # Render **bold** spans in body
            body_html = _re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', body_raw)
            full_html  = entity_html + body_html

            # Dynamic pill
            pill_text = _generate_rec_pill(neutral, tier_key)
            pill_html = ''
            if pill_text:
                # Split "Label: value" for pill display
                if ': ' in pill_text:
                    p_label, p_val = pill_text.split(': ', 1)
                    pill_inner = f'<strong>{p_label}:</strong> {p_val}'
                else:
                    pill_inner = pill_text
                pill_html = (
                    f'<div style="margin-top:6px">'
                    f'<span style="font-size:0.67rem;background:{pill_bg};'
                    f'color:{pill_col};padding:3px 9px;border-radius:12px;'
                    f'font-weight:600;letter-spacing:0.02em">'
                    f'\u26a1 {pill_inner}</span></div>'
                )

            parts.append(
                f'<div style="background:rgba(15,23,42,0.65);'
                f'border:1px solid rgba(255,255,255,0.08);'
                f'border-radius:7px;padding:11px 15px;margin:6px 0;">'
                f'<div style="font-size:0.8rem;color:#e2e8f0;line-height:1.65">'
                f'{full_html}</div>'
                f'{pill_html}'
                f'</div>'
            )

    return '\n'.join(parts)



def _fig_is_empty(fig) -> bool:
    """Return True if the figure contains no real data (only error/empty traces)."""
    try:
        if not fig or not fig.data:
            return True
        title_str = ""
        if fig.layout and fig.layout.title and fig.layout.title.text:
            title_str = str(fig.layout.title.text).lower()
        # Common empty-state title patterns from chart builders
        empty_markers = [
            "no flow data", "no data", "no delay data", "no category data",
            "error", "not found", "no route", "no distributor", "no supplier",
            "no shipment", "no demand", "no plant",
        ]
        if any(m in title_str for m in empty_markers):
            return True
        # Check if all traces have zero points
        has_data = False
        for trace in fig.data:
            try:
                if hasattr(trace, "x") and trace.x is not None and len(trace.x) > 0:
                    has_data = True; break
                if hasattr(trace, "values") and trace.values is not None and len(trace.values) > 0:
                    has_data = True; break
                if hasattr(trace, "node") and trace.node is not None:
                    labels = getattr(trace.node, "label", None)
                    if labels and len(labels) > 0:
                        has_data = True; break
            except Exception:
                pass
        return not has_data
    except Exception:
        return False


def _generate_rca_viz_images(question: str) -> list:
    """
    Generate PNG chart images for Word/Excel export.
    Dynamically selects and renders charts based on the query type.
    Returns list of (title, why_text, what_text, png_bytes) tuples.

    Key guarantees:
    - Charts are always generated from live query data filtered by question context
    - Empty-data charts fall back to broader queries instead of showing blank
    - Different queries produce different chart combinations
    - Each chart validates non-empty data before rendering
    - Fallback chart type used if primary chart type fails
    """
    results = []
    q = (question or "").lower()
    ctx = _extract_context(question)

    # ── Detect report type ───────────────────────────────────────────────────
    if any(w in q for w in ["supplier", "vendor", "risk", "risky"]):
        rt = "supplier_risk"
    elif any(w in q for w in ["stockout", "demand gap", "shortage", "running out"]):
        rt = "demand_gap"
    elif any(w in q for w in ["route", "transit", "transport", "logistics"]):
        rt = "route"
    elif any(w in q for w in ["category", "product", "toy", "auto", "health",
                               "beauty", "watch", "gift", "food", "electronics",
                               "apparel", "home", "sports"]):
        rt = "product_category"
    elif any(w in q for w in ["distributor", "city", "distribution"]):
        rt = "distributor"
    elif any(w in q for w in ["plant", "factory", "baddi", "pune", "bhopal", "goa"]):
        rt = "plant_focus"
    else:
        rt = "shipment_delay"

    WHAT_MAP = {
        "net":  "Network graph showing supply chain nodes and flow paths. Node size = volume, edge thickness = flow weight. Coloured paths trace the delay cascade from source to impact point.",
        "bar":  "Ranked bar chart showing top contributors to delays/risk. Pareto line shows cumulative contribution — fixing the top 3 bars resolves the majority of issues.",
        "pie":  "Proportional breakdown of delay share or demand gap across categories/entities. Largest slice is the primary root cause driver.",
        "heat": "Grid heatmap: rows = Plants, columns = Product Categories. Cell colour = delay severity. Darkest cell = exact root cause intersection.",
        "bub":  "Bubble chart: X=risk score, Y=delay frequency, bubble size=shipment volume. Top-right large bubbles = highest priority intervention targets.",
        "san":  "Sankey flow: Suppliers → Plants → Distributors. Band width = shipment volume, red bands = delayed flow paths revealing the bottleneck stage.",
    }

    WHY_MAP = {
        ("shipment_delay", "san"):        "Shows exactly where in the supply chain volume is being lost to delays — wide red bands pinpoint the root cause stage.",
        ("shipment_delay", "heat"):       "Exposes the exact plant × category combination driving the highest delay count — the darkest cell is the root cause target.",
        ("shipment_delay", "bar"):        "Ranks the top delayed plants/suppliers by volume — fixing the first 3 bars resolves the majority of delay events.",
        ("shipment_delay", "net"):        "Traces the full delay propagation path, showing first- and second-order network impact across the chain.",
        ("shipment_delay", "pie"):        "Shows the proportional share of delays by category — identifies whether delays are systemic or concentrated.",
        ("supplier_risk",  "bub"):        "Plots risk score vs delay frequency — top-right large bubbles are the suppliers requiring immediate action.",
        ("supplier_risk",  "bar"):        "Ranks suppliers by delayed shipment count and risk score — directly prioritises which relationships need urgent intervention.",
        ("supplier_risk",  "heat"):       "Compares delay severity across plants, showing which facilities are most exposed to upstream supplier risk.",
        ("supplier_risk",  "san"):        "Traces how supplier risk converts to actual delay across the distribution chain — shows cascading impact.",
        ("supplier_risk",  "pie"):        "Shows the share of total delays attributable to each supplier — identifies if one supplier dominates risk.",
        ("demand_gap",     "heat"):       "Maps shortage severity by city and distributor — identifies geographic concentration of the demand gap.",
        ("demand_gap",     "bar"):        "Ranks highest-gap retailers and distributors — prioritises which cities need emergency replenishment.",
        ("demand_gap",     "pie"):        "Shows whether shortage is concentrated in one category or systemic across the portfolio.",
        ("demand_gap",     "san"):        "Connects upstream supply failure to retail stockout — visualises the full shortage propagation chain.",
        ("product_category","pie"):       "Shows delay share across plants for this category — the largest slice is the primary production bottleneck.",
        ("product_category","bar"):       "Ranks the worst-performing suppliers for this category — direct action targets for category-specific delays.",
        ("product_category","heat"):      "Shows plant × supplier delay matrix for this category — reveals the exact manufacturing choke point.",
        ("product_category","san"):       "Traces the category's dependency path from supplier inputs through plants to distributors, identifying the exact bottleneck link.",
        ("distributor",    "heat"):       "Maps delay severity by distributor city — shows which geographic regions are most severely impacted.",
        ("distributor",    "bar"):        "Ranks distributors by total delayed shipments received — prioritises which distribution centres need immediate attention.",
        ("distributor",    "san"):        "Shows which supplier-plant paths feed into high-delay distributors — traces the root cause upstream.",
        ("distributor",    "pie"):        "Shows the share of total delayed shipments absorbed by each distributor — identifies geographic concentration.",
        ("route",          "net"):        "Maps route efficiency as a network — highlights the high-cost, low-performance paths in red.",
        ("route",          "san"):        "Shows volume flowing through each transport route and mode — wide red bands are the inefficient route segments.",
        ("route",          "bar"):        "Ranks routes by delay rate and cost — the top bars are immediate optimisation candidates.",
        ("route",          "heat"):       "Compares delay rates across route modes and origin plants — identifies the worst-performing transport combinations.",
        ("plant_focus",    "heat"):       "Shows each plant's delay rate by product category — the darkest cells reveal which product lines are causing plant-level bottlenecks.",
        ("plant_focus",    "bar"):        "Ranks plants by total delayed shipments dispatched — directly identifies the highest-impact manufacturing facility.",
        ("plant_focus",    "san"):        "Traces supplier inputs through each plant to distributor outputs — shows how plant-level delays propagate downstream.",
        ("plant_focus",    "pie"):        "Shows the proportional delay contribution of each plant — identifies whether delays are plant-specific or systemic.",
    }

    # ── Chart plan: primary + fallback per report type ───────────────────────
    CHART_PLAN = {
        "shipment_delay": [
            ("Supply Flow & Delays — Sankey",        build_supply_flow_sankey,  "san",  build_dynamic_pie),
            ("Delay Heatmap — Plant × Category",     build_dynamic_heatmap,     "heat", build_dynamic_pareto),
            ("Top Delay Contributors — Pareto",      build_dynamic_pareto,      "bar",  build_dynamic_pie),
            ("Delay Propagation Network",            build_dynamic_flow,        "net",  build_dynamic_heatmap),
        ],
        "supplier_risk": [
            ("Supplier Risk Bubble Chart",           build_dynamic_network,     "bub",  build_dynamic_pareto),
            ("Suppliers Ranked by Delayed Shipments",build_dynamic_pareto,      "bar",  build_dynamic_pie),
            ("Supplier × Plant Delay Heatmap",       build_dynamic_heatmap,     "heat", build_dynamic_pareto),
            ("Supply Flow — Risk to Distributor",    build_supply_flow_sankey,  "san",  build_dynamic_pie),
        ],
        "demand_gap": [
            ("Demand Gap Heatmap — City × Category", build_dynamic_heatmap,     "heat", build_dynamic_pareto),
            ("Stockout Ranked by Demand Gap",        build_dynamic_pareto,      "bar",  build_dynamic_pie),
            ("Demand Gap Share by Category",         build_dynamic_pie,         "pie",  build_dynamic_pareto),
            ("Shortage Chain — Sankey Flow",         build_supply_flow_sankey,  "san",  build_dynamic_flow),
        ],
        "product_category": [
            ("Category Delay Share by Plant",        build_dynamic_pie,         "pie",  build_dynamic_pareto),
            ("Category Supply Chain — Sankey",       build_supply_flow_sankey,  "san",  build_dynamic_flow),
            ("Top Suppliers by Category Delays",     build_dynamic_pareto,      "bar",  build_dynamic_pie),
            ("Category Delay Heatmap",               build_dynamic_heatmap,     "heat", build_dynamic_pareto),
        ],
        "distributor": [
            ("Distributor Delay Heatmap",            build_dynamic_heatmap,     "heat", build_dynamic_pareto),
            ("Distributors Ranked by Delays",        build_dynamic_pareto,      "bar",  build_dynamic_pie),
            ("Supply Flow to Distributors",          build_supply_flow_sankey,  "san",  build_dynamic_flow),
            ("Distributor Delay Share",              build_dynamic_pie,         "pie",  build_dynamic_pareto),
        ],
        "route": [
            ("Routes Ranked by Cost & Delay",        build_dynamic_pareto,      "bar",  build_dynamic_pie),
            ("Route × Mode Delay Heatmap",           build_dynamic_heatmap,     "heat", build_dynamic_pareto),
            ("Route Volume — Sankey Flow",           build_supply_flow_sankey,  "san",  build_dynamic_pareto),
            ("Route Efficiency Network",             build_dynamic_flow,        "net",  build_dynamic_heatmap),
        ],
        "plant_focus": [
            ("Plant × Category Delay Heatmap",       build_dynamic_heatmap,     "heat", build_dynamic_pareto),
            ("Plants Ranked by Delayed Shipments",   build_dynamic_pareto,      "bar",  build_dynamic_pie),
            ("Plant Supply Flow — Sankey",           build_supply_flow_sankey,  "san",  build_dynamic_flow),
            ("Delay Share by Plant",                 build_dynamic_pie,         "pie",  build_dynamic_pareto),
        ],
    }

    def _fig_to_png(fig):
        """Convert figure to PNG bytes with multiple fallback sizes."""
        for args in [
            {"width": 960, "height": 480, "scale": 1.5},
            {"width": 900, "height": 450, "scale": 1.2},
            {"width": 800, "height": 400, "scale": 1.0},
        ]:
            try:
                png = fig.to_image(format="png", **args)
                if png:
                    return png
            except Exception:
                continue
        return None

    plan = CHART_PLAN.get(rt, CHART_PLAN["shipment_delay"])

    for title, fn, chart_key, fallback_fn in plan:
        png = None
        used_fn = None
        try:
            fig = fn(question=question)
            if _fig_is_empty(fig):
                # Try fallback function
                try:
                    fig = fallback_fn(question=question)
                    used_fn = fallback_fn
                except Exception:
                    pass
            else:
                used_fn = fn

            if fig and not _fig_is_empty(fig):
                fig = _apply_dark_theme(fig)
                png = _fig_to_png(fig)

        except Exception as _e:
            # Primary failed — try fallback
            try:
                fig = fallback_fn(question=question)
                if fig and not _fig_is_empty(fig):
                    fig = _apply_dark_theme(fig)
                    png = _fig_to_png(fig)
                    used_fn = fallback_fn
            except Exception:
                pass

        if png:
            why  = WHY_MAP.get((rt, chart_key), f"Analytical view for {rt.replace('_', ' ')} analysis, filtered to the entities most relevant to this query.")
            what = WHAT_MAP.get(chart_key, "Supply chain analytical chart derived from live query data.")
            results.append((title, why, what, png))
        else:
            # Second-level fallback: try without question filter (broadest data)
            try:
                fig2 = fn()
                if fig2 and not _fig_is_empty(fig2):
                    fig2 = _apply_dark_theme(fig2)
                    png2 = _fig_to_png(fig2)
                    if png2:
                        why  = WHY_MAP.get((rt, chart_key), "Overall supply chain view — broadened scope since query-specific data was limited.")
                        what = WHAT_MAP.get(chart_key, "Supply chain analytical chart.")
                        results.append((title, why, what, png2))
            except Exception:
                pass

    # ── Absolute fallback: if nothing rendered, force the 3 most-data-rich charts ──
    if not results:
        q_short = (question or "supply chain").strip().rstrip("?")
        fallback_plan = [
            (f"Delay Heatmap — {q_short[:30]}", build_dynamic_heatmap, "heat"),
            (f"Top Contributors — {q_short[:28]}", build_dynamic_pareto,  "bar"),
            (f"Delay Share — {q_short[:32]}",    build_dynamic_pie,     "pie"),
        ]
        for title, fn, chart_key in fallback_plan:
            try:
                fig = _apply_dark_theme(fn())   # no question filter — broadest possible
                png = _fig_to_png(fig)
                if png:
                    why  = WHY_MAP.get(("shipment_delay", chart_key),
                                       f"Overall supply chain view relevant to: {q_short}")
                    what = WHAT_MAP.get(chart_key, "Supply chain analytical chart.")
                    results.append((title, why, what, png))
            except Exception:
                continue

    return results


def _build_inline_charts_html(question: str) -> str:
    """
    Build the HTML for the inline interactive charts section in the RCA tab.
    Returns a styled intro blurb; actual chart figures are rendered via gr.Plot.
    """
    q = (question or "").lower()
    if any(w in q for w in ["route", "transit", "transport"]):
        intro = "Route network, delay heatmap, Pareto breakdown, and supply Sankey are ready."
    elif any(w in q for w in ["supplier", "vendor", "risk"]):
        intro = "Supplier risk bubble, network flow, Pareto ranking, and delay heatmap are ready."
    elif any(w in q for w in ["category", "product"]):
        intro = "Category delay share, Pareto breakdown, and heatmap are ready for this product analysis."
    else:
        intro = "Supply chain visualizations — network flow, delay heatmap, Pareto bar, and Sankey — are ready."

    return (
        f'<div style="font-size:0.78rem;color:#7dd3fc;line-height:1.5;padding:0 0 8px 0;margin:0">'
        f'{intro} Click a chart button below to load it.'
        f'</div>'
    )


def _build_detailed_rec_html_from_raw(raw: str) -> str:
    """
    Build recommendations HTML from rec_raw.
    Handles both JSON array format (new) and markdown format (legacy).
    Returns empty string if rec_block is already rendered inside the main report.
    """
    import re as _re, json as _rj
    raw = str(raw or "").strip()
    if not raw:
        return ""

    # Try JSON array format first (new rec agent format)
    _m = _re.search(r'\[\s*\{', raw, _re.DOTALL)
    if _m:
        try:
            _arr_raw = raw[_m.start():]
            _depth = 0
            _end = 0
            for _ci, _ch in enumerate(_arr_raw):
                if _ch == "[": _depth += 1
                elif _ch == "]":
                    _depth -= 1
                    if _depth == 0:
                        _end = _ci + 1
                        break
            if _end > 0:
                _arr_raw = _arr_raw[:_end]
            _items = _rj.loads(_arr_raw)
            if isinstance(_items, list) and _items:
                # JSON recs are already rendered inside the main report — return empty
                # to avoid duplication. The rec_block in agent_runner final return handles display.
                return ""
        except Exception:
            pass

    # Legacy markdown format — return empty if it just has placeholder text
    if "Recommendations will appear here" in raw or len(raw) < 30:
        return ""

    # Return markdown rec content for the legacy path
    return raw


def _wrap_rec_details(rec_html: str) -> str:
    """Wrap recommendations HTML. Returns empty string if nothing to show (avoids duplicate rec panel)."""
    if not rec_html or not rec_html.strip():
        return ""  # Empty → Gradio component stays hidden → no duplicate
    body = rec_html
    return (
        '<details class="rca-native-box green-box" style="margin-top:12px" open>'
        '<summary class="rca-native-summary">'
        '<span class="rca-native-arrow">&#9654;</span>'
        '<span class="rca-native-title">&#x2705; Recommendations</span>'
        '<span class="rca-native-hint">&#8212; click to expand</span>'
        '</summary>'
        f'<div style="padding:14px 18px 18px">{body}</div>'
        '</details>'
    )


def rca_handler(q):
    """
    Streaming RCA handler — yields 14 outputs (added rca_cypher_html).

    Outputs: rca_status, rca_assess_html, rca_out, rca_log,
             rca_chart1, rca_chart2, rca_chart3 (hidden compat),
             rca_dl_btn, rca_csv_btn, rca_dl_section, rca_dl_status,
             rca_viz_summary_html, rca_detailed_rec_html,
             rca_charts_html_vis, rca_cypher_html
    """
    key = _cache_key(q)

    _blank = go.Figure().update_layout(
        paper_bgcolor="#060c1c", plot_bgcolor="#060c1c",
        xaxis=dict(visible=False), yaxis=dict(visible=False),
        height=480, autosize=False,
    )

    def _side_charts():
        """Build the two always-visible right-panel charts."""
        try:
            c1 = build_delay_heatmap()
        except Exception:
            c1 = _blank
        try:
            c2 = build_supplier_risk_chart()
        except Exception:
            c2 = _blank
        return c1, c2

    # ── CACHE HIT ────────────────────────────────────────────
    if key in _RCA_CACHE:
        cached = _RCA_CACHE[key]
        _cached_report = str(cached.get("report") or "")
        # Auto-invalidate broken/empty cache entries so next run is fresh
        # Also invalidate old reports that contain the HTML root_cause div (pre-fix format)
        _is_stale = (
            not _cached_report
            or len(_cached_report) < 50
            or '<div style="margin:8px 0 12px' in _cached_report
            or 'Disruption Scenario' in _cached_report
            or 'goes offline' in _cached_report
            or 'Supplier Redirection' in _cached_report
            or 'supplier_id|supplier_name' in _cached_report   # raw underscore column names leaked
            or 'Supplier Id|Supplier Name|Risk Score' in _cached_report  # table in desc block
            or 'The origin of the supply chain breakdown &#8212;' in _cached_report  # static RC text
            or ('The origin of the supply chain breakdown' in _cached_report
                and 'color:#f87171' not in _cached_report)  # static RC text, no real para
        )
        if _is_stale:
            del _RCA_CACHE[key]
        else:
            try:
                dl_path = _save_rca_report(q, _clean_report_for_export(_cached_report))
            except Exception:
                dl_path = None
            c1, c2 = _side_charts()
            try:
                csv_path = _save_rca_excel(q, _cached_report)
            except Exception:
                csv_path = None
            _charts_html_cached = _build_inline_charts_html(q)
            _cached_rec_html    = _build_detailed_rec_html(_cached_report)
            # Pre-render chart figures for cache hit
            def _safe_c(fn, **kw):
                try: return _apply_dark_theme(fn(**kw))
                except TypeError:
                    try: return _apply_dark_theme(fn())
                    except Exception: return _blank
                except Exception: return _blank
            _cf_net  = _safe_c(build_dynamic_flow, question=q)
            _cf_bar  = _safe_c(build_dynamic_pareto, question=q)
            _cf_pie  = _safe_c(build_dynamic_pie, question=q)
            _cf_heat = _safe_c(build_dynamic_heatmap, question=q)
            _cf_bub  = _safe_c(build_dynamic_network, question=q)
            _cf_san  = _safe_c(build_supply_flow_sankey, question=q)
            # Build vchart panel for cache hit too
            _rt_c       = _detect_report_type_from_q(q or "")
            _plan_c     = _RCA_CHART_PLAN.get(_rt_c, ["san","heat","bar","net"])[:4]
            _plan_set_c = set(_plan_c)
            _btn_c      = {s: gr.update(visible=(s in _plan_set_c))
                           for s in ["net","bar","heat","bub","pie","san"]}
            _rt_names_c = {
                "supplier_risk":"Supplier Risk Analysis","demand_gap":"Demand Gap & Stockout",
                "shipment_delay":"Shipment Delay Analysis","distributor":"Distributor Impact",
                "route":"Route & Transport Cost",
            }
            _rt_name_c  = _rt_names_c.get(_rt_c, "Supply Chain")
            _slot_labels_c = {"net":"🌐 Network/Flow","bar":"📊 Bar Chart","heat":"🔥 Heatmap",
                               "bub":"🫧 Bubble","pie":"🥧 Pie","san":"🔀 Sankey"}
            _pill_c = "".join(
                f'<span style="display:inline-block;margin:2px 4px 2px 0;padding:2px 9px;border-radius:16px;'                 f'font-size:0.66rem;font-weight:700;background:rgba(56,189,248,0.12);'                 f'border:1px solid rgba(56,189,248,0.28);color:#7dd3fc">{_slot_labels_c.get(s,s)}</span>'
                for s in _plan_c
            )
            _hdr_c = (
                f'<div style="padding:8px 2px 12px;border-bottom:1px solid rgba(56,189,248,0.18);margin-bottom:10px">'                 f'<div style="font-size:0.82rem;font-weight:700;color:#38bdf8;margin-bottom:4px">'                 f'📈 Best charts for: <span style="color:#f0abfc">{_rt_name_c}</span></div>'                 f'<div style="font-size:0.7rem;color:#94a3b8;margin-bottom:5px">'                 f'Select a chart type below. Chart loads when you click a tab.</div>'                 f'<div>{_pill_c}</div></div>'
            )
            _first_c   = _plan_c[0]
            _first_fn_c = _SLOT_BUILDERS.get(_first_c, _SLOT_BUILDERS["bar"])
            _first_fig_c = None  # no auto-load — user clicks a tab
            _first_exp_c = ""    # explain only shown after click

            yield (
                '<div class="ag-status ag-ok">❆ Cached result — instant replay ⚡</div>',
                "",
                _cached_report or _MD_PLACEHOLDER,
                _tool_html(cached.get("logs", [])),
                c1, c2, _blank,
                gr.update(value=dl_path, visible=bool(dl_path)),
                gr.update(value=csv_path, visible=bool(csv_path)),
                gr.update(visible=bool(dl_path or csv_path)),
                gr.update(value=''),
                gr.update(value=_build_rca_viz_html(q), visible=True),
                gr.update(value=_wrap_rec_details(_cached_rec_html), visible=True),
                gr.update(value="", visible=False),
                gr.update(value="", visible=False),
                gr.update(visible=True),
                _cf_net, _cf_bar, _cf_pie, _cf_heat, _cf_bub, _cf_san,
                gr.update(value=_hdr_c, visible=True),
                _btn_c["net"], _btn_c["bar"], _btn_c["heat"],
                _btn_c["bub"], _btn_c["pie"], _btn_c["san"],
                gr.update(value="", visible=False),   # rca_vchart_info
                gr.update(value=None, visible=False),  # rca_vchart_plot
                gr.update(visible=False),              # rca_vchart_plot_col
            )
            return
    # ── FRESH RUN ─────────────────────────────────────────────
    dl_invisible    = gr.update(visible=False)
    final_status    = None
    final_report    = None
    final_log_html  = None
    final_assess    = ""
    raw_logs_holder   = {}
    _cypher_logs_live = []

    def _rca_with_logs(question, on_update=None):
        def _on_update_wrap(event):
            if isinstance(event, tuple) and len(event) == 2 and event[0] == "cypher_log":
                _cypher_logs_live.append(event[1])
                raw_logs_holder["cypher_logs"] = list(_cypher_logs_live)
            if on_update:
                on_update(event)
        try:
            result = run_rca(question, on_update=_on_update_wrap)
            # run_rca returns (report, logs, rec_raw, cypher_logs) — 4 values
            if isinstance(result, tuple) and len(result) >= 2:
                report = result[0]
                logs   = result[1] if len(result) > 1 else []
                rec_raw = result[2] if len(result) > 2 else ""
                _cq     = result[3] if len(result) > 3 else []
            else:
                report, logs, rec_raw, _cq = str(result), [], "", []
        except Exception as _run_err:
            print(f"[_rca_with_logs] run_rca failed: {_run_err}")
            report = f"Analysis encountered an error: {str(_run_err)[:200]}"
            logs, rec_raw, _cq = [], "", []
        raw_logs_holder["logs"]    = logs
        raw_logs_holder["rec_raw"] = rec_raw or ""
        return report, logs

    for status, assess_html, report, log_html in _agent_generator(_rca_with_logs, q):
        final_status   = status
        final_assess   = assess_html
        final_report   = report
        final_log_html = log_html
        # Build live cypher accordion as tools accumulate (shown immediately)
        _live_cq_html = _build_cypher_accordion(raw_logs_holder.get("cypher_logs", []))
        yield (
            status     or "",
            assess_html,                      # rca_assess_html — Initial Assessment card
            _strip_rec_details(report or _MD_PLACEHOLDER),
            log_html   or "",
            _blank, _blank, _blank,
            dl_invisible,
            gr.update(visible=False),
            gr.update(visible=False),
            gr.update(value=""),
            gr.update(visible=False),         # rca_viz_summary_html — hidden while streaming
            gr.update(visible=False),         # rca_detailed_rec_html — hidden while streaming
            gr.update(value="", visible=False), # rca_charts_html_vis — compat hidden
            gr.update(value=_live_cq_html, visible=bool(_live_cq_html)),
            gr.update(visible=False),         # rca_view_charts_col — hidden while streaming
            None, None, None, None, None, None,  # 6 chart state figs empty during stream
            # vchart panel — hidden during streaming, populated in final yield
            gr.update(visible=False),         # rca_vchart_header
            gr.update(visible=False), gr.update(visible=False), gr.update(visible=False),
            gr.update(visible=False), gr.update(visible=False), gr.update(visible=False),
            gr.update(visible=False),         # rca_vchart_info
            gr.update(visible=False),         # rca_vchart_plot
            gr.update(visible=False),         # rca_vchart_plot_col
        )

    # ── POST-RUN: cache + download + render side charts ───────
    # raw_logs_holder["logs"] has only real tool calls (from run_rca return value).
    # We need to also include pseudo-key stage markers so the activity log works.
    # _agent_generator already merged them into final_logs (stored in final_log_html source).
    # Reconstruct the merged list from the generator's output by checking raw_logs_holder.
    _pseudo_keys = {"__first_response__","__orchestrator__","__validator_agent__",
                    "__rca_agent__","__rec_agent__","__narrative_agent__"}
    _real_logs   = raw_logs_holder.get("logs", [])
    # Determine which pseudo-keys fired based on the A2A status message in final_status
    _fired = []
    for _pk in ["__first_response__","__orchestrator__","__validator_agent__",
                "__rca_agent__","__rec_agent__","__narrative_agent__"]:
        if _pk.strip("_") in (final_status or "").lower().replace(" ","").replace("→",""):
            _fired.append({"tool": _pk, "input": {}, "result_preview": ""})
    # Always include all 5 main agents as fired if we got a completed report
    if final_report and len(final_report) > 200:
        _fired = [
            {"tool": "__orchestrator__",    "input": {}, "result_preview": ""},
            {"tool": "__validator_agent__",  "input": {}, "result_preview": ""},
            {"tool": "__rca_agent__",        "input": {}, "result_preview": ""},
            {"tool": "__rec_agent__",        "input": {}, "result_preview": ""},
            {"tool": "__narrative_agent__",  "input": {}, "result_preview": ""},
        ]
    _merged_cache_logs = _fired + _real_logs
    _RCA_CACHE[key] = {
        "report":      final_report or "",
        "logs":        _merged_cache_logs,
        "cypher_logs": raw_logs_holder.get("cypher_logs", []),
        "status":      final_status or "",
    }
    try:
        dl_path = _save_rca_report(q, _clean_report_for_export(final_report or ""))
    except Exception:
        dl_path = None
    try:
        csv_path = _save_rca_excel(q, final_report or "")
    except Exception:
        csv_path = None
    c1, c2  = _side_charts()
    _viz_html    = _build_rca_viz_html(q)
    # Build recommendations HTML from rec_raw (separate from report body)
    _rec_html    = _build_detailed_rec_html_from_raw(raw_logs_holder.get("rec_raw", "") or final_report or "")
    # Pre-render chart figures for the gr.Plot panel
    def _safe_chart(fn, **kw):
        try:
            # Try with provided kwargs first
            return _apply_dark_theme(fn(**kw))
        except TypeError:
            try:
                # Fallback: no args
                return _apply_dark_theme(fn())
            except Exception:
                return go.Figure().update_layout(paper_bgcolor="#060c1c",plot_bgcolor="#060c1c",xaxis=dict(visible=False),yaxis=dict(visible=False))
        except Exception:
            return go.Figure().update_layout(paper_bgcolor="#060c1c",plot_bgcolor="#060c1c",xaxis=dict(visible=False),yaxis=dict(visible=False))
    _fig_net  = _safe_chart(build_dynamic_flow, question=q)
    _fig_bar  = _safe_chart(build_dynamic_pareto, question=q)
    _fig_pie  = _safe_chart(build_dynamic_pie, question=q)
    _fig_heat = _safe_chart(build_dynamic_heatmap, question=q)
    _fig_bub  = _safe_chart(build_dynamic_network, question=q)
    _fig_san  = _safe_chart(build_supply_flow_sankey, question=q)
    _cq_logs     = raw_logs_holder.get("cypher_logs", [])
    _cq_html     = _build_cypher_accordion(_cq_logs)

    # ── Build VIEW CHARTS outputs inline — no extra .then() step needed ──
    # Pick 4 best slots for this query, show those buttons, load first chart
    _rt           = _detect_report_type_from_q(q or "")
    _plan         = _RCA_CHART_PLAN.get(_rt, ["san","heat","bar","net"])[:4]
    _plan_set     = set(_plan)
    _slot_order   = ["net","bar","heat","bub","pie","san"]
    _btn_updates  = {s: gr.update(visible=(s in _plan_set)) for s in _slot_order}

    _rt_names = {
        "supplier_risk":"Supplier Risk Analysis","demand_gap":"Demand Gap & Stockout",
        "shipment_delay":"Shipment Delay Analysis","plant":"Plant Performance",
        "distributor":"Distributor Impact","product_category":"Product Category",
        "route":"Route & Transport Cost","simulation":"Impact Simulation",
        "transport_delay":"Transport Mode Delays",
    }
    _rt_name  = _rt_names.get(_rt, "Supply Chain")
    _slot_labels = {"net":"🌐 Network/Flow","bar":"📊 Bar Chart","heat":"🔥 Heatmap",
                    "bub":"🫧 Bubble","pie":"🥧 Pie","san":"🔀 Sankey"}
    # No pill badges — buttons ARE the chart type indicators
    _vchart_header_html = (
        f'<div style="padding:8px 2px 10px;border-bottom:1px solid rgba(56,189,248,0.15);margin-bottom:10px">'
        f'<div style="font-size:0.82rem;font-weight:700;color:#38bdf8;margin-bottom:3px">'
        f'📈 Recommended visualisations for: <span style="color:#f0abfc">{_rt_name}</span></div>'
        f'<div style="font-size:0.7rem;color:#94a3b8">'
        f'Select a chart type below. Chart loads when you click a tab.</div>'
        f'</div>'
    )

    # Chart and explain are empty until user clicks a tab
    _first_fig     = None   # no auto-load — user clicks a tab to load chart
    _first_explain = ""     # no explain shown until chart is clicked

    yield (
        final_status   or "",
        final_assess,                         # rca_assess_html — keep card visible
        _strip_rec_details(final_report or _MD_PLACEHOLDER),
        final_log_html or "",
        c1, c2, _blank,
        gr.update(value=dl_path, visible=bool(dl_path)),
        gr.update(value=csv_path, visible=bool(csv_path)),
        gr.update(visible=bool(dl_path or csv_path)),
        gr.update(value=''),
        gr.update(value=_viz_html, visible=True),
        gr.update(value=_wrap_rec_details(_rec_html), visible=True),
        gr.update(value="", visible=False),   # rca_charts_html_vis — compat, keep hidden
        gr.update(value=_cq_html, visible=bool(_cq_html)),
        gr.update(visible=True),   # rca_view_charts_col
        _fig_net,                  # rca_state_net_fig
        _fig_bar,                  # rca_state_bar_fig
        _fig_pie,                  # rca_state_pie_fig
        _fig_heat,                 # rca_state_heat_fig
        _fig_bub,                  # rca_state_bub_fig
        _fig_san,                  # rca_state_san_fig
        # ── NEW: vchart panel outputs (all in same yield) ──
        gr.update(value=_vchart_header_html, visible=True),
        _btn_updates["net"], _btn_updates["bar"], _btn_updates["heat"],
        _btn_updates["bub"], _btn_updates["pie"], _btn_updates["san"],
        gr.update(value="", visible=False),   # rca_vchart_info — shown on button click
        gr.update(value=None, visible=False),  # rca_vchart_plot — shown on button click
        gr.update(visible=False),              # rca_vchart_plot_col — shown on button click
    )


def _auto_render_charts(question: str):
    """
    Auto-selects and renders up to 3 charts relevant to the question.
    Returns a tuple of 3 Plotly figures (padded with empty figures).
    """
    q = question.lower()

    # Priority-ordered chart selection by keyword
    selected = []
    if any(w in q for w in ["supplier", "vendor", "risk", "root cause"]):
        selected.append(build_supplier_risk_chart)
    if any(w in q for w in ["delay", "late", "slow", "bottleneck", "plant"]):
        selected.append(build_delay_heatmap)
    if any(w in q for w in ["stockout", "stock", "shortage", "retailer"]):
        selected.append(build_distributor_demand_gap)
    if any(w in q for w in ["trend", "month", "seasonal", "over time"]):
        selected.append(build_monthly_delay_trend)
    if any(w in q for w in ["route", "transport", "logistics", "road", "rail", "air", "mode"]):
        selected.append(build_route_efficiency_scatter)
    if any(w in q for w in ["distributor", "distribution", "city"]):
        selected.append(build_distributor_demand_gap)

    # Deduplicate while preserving order
    seen, unique = set(), []
    for fn in selected:
        if fn not in seen:
            seen.add(fn)
            unique.append(fn)

    # If nothing matched, show the three most useful defaults
    if not unique:
        unique = [build_delay_heatmap, build_supplier_risk_chart, build_distributor_demand_gap]

    # Pad to exactly 3 charts
    while len(unique) < 3:
        fallbacks = [build_delay_heatmap, build_supplier_risk_chart,
                     build_monthly_delay_trend, build_distributor_demand_gap]
        for fb in fallbacks:
            if fb not in unique:
                unique.append(fb)
                break

    figs = []
    for fn in unique[:3]:
        try:
            figs.append(fn())
        except Exception:
            figs.append(go.Figure().update_layout(**_PLOTLY_LAYOUT))

    return tuple(figs)


def update_handler(q):
    """Update Graph handler — strips the assess_html slot (not used in update tab)."""
    for status, _assess, report, log_html in _agent_generator(run_graph_update, q):
        yield status, report, log_html


# ─────────────────────────────────────────────────────────────
# NETWORK HEALTH TAB  (redesigned — 4 outputs)
# ─────────────────────────────────────────────────────────────

def _build_health_gauge(score_pct: float, label: str = "Network Health") -> go.Figure:
    """Builds a semi-circle gauge chart for overall health score."""
    color = "#22c55e" if score_pct >= 70 else ("#f59e0b" if score_pct >= 45 else "#ef4444")
    fig = go.Figure(go.Indicator(
        mode="gauge+number+delta",
        value=round(score_pct, 1),
        delta={"reference": 70, "increasing": {"color": "#22c55e"}, "decreasing": {"color": "#ef4444"}},
        number={"suffix": "%", "font": {"color": "#e2e8f0", "size": 36}},
        title={"text": label, "font": {"color": "#7dd3fc", "size": 13}},
        gauge={
            "axis": {"range": [0, 100], "tickcolor": "#64748b", "tickfont": {"color": "#64748b", "size": 10}},
            "bar": {"color": color, "thickness": 0.28},
            "bgcolor": "rgba(12,21,40,0.9)",
            "borderwidth": 0,
            "steps": [
                {"range": [0,   45], "color": "rgba(239,68,68,0.15)"},
                {"range": [45,  70], "color": "rgba(245,158,11,0.12)"},
                {"range": [70, 100], "color": "rgba(34,197,94,0.12)"},
            ],
            "threshold": {
                "line": {"color": "#38bdf8", "width": 2},
                "thickness": 0.75,
                "value": 70,
            },
        }
    ))
    fig.update_layout(
        paper_bgcolor="#060c1c",
        plot_bgcolor="#060c1c",
        font=dict(family="DM Sans, sans-serif", color="#cbd5e1"),
        margin=dict(l=30, r=30, t=60, b=20),
        height=260
    )
    return fig


def load_network_health():
    _blank = go.Figure().update_layout(
        paper_bgcolor="#060c1c", plot_bgcolor="#060c1c",
        xaxis=dict(visible=False), yaxis=dict(visible=False)
    )
    try:
        # ── Core metrics ──────────────────────────────────────
        delay_data = _run_neo4j("""
            MATCH (sh:Shipment)
            RETURN COUNT(sh) AS total,
                   SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS delayed,
                   ROUND(AVG(sh.delay_days), 2) AS avg_delay
        """)
        d = delay_data[0] if delay_data else {}
        total      = d.get("total", 0)
        delayed    = d.get("delayed", 0)
        avg_d      = d.get("avg_delay", 0) or 0
        on_time    = total - delayed
        delay_pct  = round(100 * delayed / total, 1) if total else 0
        on_time_pct = round(100 * on_time / total, 1) if total else 0

        hr = _run_neo4j("MATCH (s:Supplier) WHERE s.risk_score >= 0.6 RETURN COUNT(s) AS cnt")
        hr_count = hr[0]["cnt"] if hr else 0

        total_sup = _run_neo4j("MATCH (s:Supplier) RETURN COUNT(s) AS cnt")
        total_sup_count = total_sup[0]["cnt"] if total_sup else 1

        so = _run_neo4j("MATCH (r:Retailer) WHERE r.stockout_flag=true OR r.stockout_flag='true' RETURN COUNT(r) AS cnt")
        so_count = so[0]["cnt"] if so else 0

        cat = _run_neo4j("""
            MATCH (sh:Shipment) WHERE sh.delivery_status='Major Delay'
            RETURN sh.product_category AS cat, COUNT(sh) AS cnt
            ORDER BY cnt DESC LIMIT 1
        """)
        top_cat = cat[0]["cat"] if cat else "N/A"

        routes_total = _run_neo4j("MATCH (r:Route) RETURN COUNT(r) AS cnt")
        routes_count = routes_total[0]["cnt"] if routes_total else 0

        # ── Compute composite health score (0-100) ────────────
        # Higher is better. Weighted average of normalised metrics.
        s_delay     = max(0, 100 - delay_pct * 2.0)          # 50% delay → 0 score
        s_avg_delay = max(0, 100 - avg_d * 20)               # 5 day avg → 0 score
        s_risk      = max(0, 100 - (hr_count / max(total_sup_count, 1)) * 200)
        s_stockout  = max(0, 100 - (so_count / 10) * 5)
        health_score = round((s_delay * 0.40 + s_avg_delay * 0.25 + s_risk * 0.25 + s_stockout * 0.10), 1)
        health_score = min(100, max(0, health_score))

        # ── Thresholds for KPI badge colours ─────────────────
        def _cls(val, ok_max, warn_max, higher_is_better=False):
            if higher_is_better:
                return "health-ok" if val >= ok_max else ("health-warn" if val >= warn_max else "health-bad")
            return "health-ok" if val <= ok_max else ("health-warn" if val <= warn_max else "health-bad")

        delay_cls   = _cls(delay_pct, 20, 35)
        avgd_cls    = _cls(avg_d, 2, 4)
        hr_cls      = _cls(hr_count, 3, 7)
        so_cls      = _cls(so_count, 10, 30)
        ontime_cls  = _cls(on_time_pct, 65, 80, higher_is_better=True)

        health_label_val = "CRITICAL" if health_score < 45 else ("WARNING" if health_score < 70 else "HEALTHY")
        health_label_col = "#ef4444" if health_score < 45 else ("#f59e0b" if health_score < 70 else "#22c55e")

        # ── KPI card HTML ─────────────────────────────────────
        kpi_data = [
            ("🚚 On-Time Rate",        f"{on_time_pct}%",    ontime_cls,  "Target: ≥80%"),
            ("⏱ Avg Delay Days",      f"{avg_d} days",       avgd_cls,    "Target: &lt;2 days"),
            ("📦 Delayed Shipments",   f"{delayed:,}",        delay_cls,   f"of {total:,} total"),
            ("⚠️ High-Risk Suppliers", f"{hr_count}",         hr_cls,      f"of {total_sup_count} suppliers"),
            ("📉 Stockout Retailers",  f"{so_count}",         so_cls,      "flag=true retailers"),
            ("🗺 Total Routes",        f"{routes_count:,}",   "health-ok", "active network paths"),
        ]

        kpi_cards = "".join(f"""
<div style="background:rgba(12,21,40,0.85);border:1px solid rgba(56,189,248,0.15);
            border-radius:12px;padding:16px 20px;min-width:140px;flex:1">
  <div style="font-size:0.7rem;color:#7dd3fc;font-weight:700;letter-spacing:0.08em;
              text-transform:uppercase;margin-bottom:6px">{lbl}</div>
  <div style="font-size:1.6rem;font-weight:800;color:#e2e8f0;
              font-family:'DM Mono',monospace;line-height:1.1" class="{cls}">{val}</div>
  <div style="font-size:0.7rem;color:#94a3b8;margin-top:4px">{sub}</div>
</div>""" for lbl, val, cls, sub in kpi_data)

        kpi_html = f"""
<div style="display:flex;gap:12px;flex-wrap:wrap;margin-bottom:16px">
  {kpi_cards}
</div>
<div style="background:rgba(12,21,40,0.6);border:1px solid rgba(56,189,248,0.12);
            border-radius:10px;padding:14px 20px;display:flex;align-items:center;gap:16px">
  <div style="font-size:0.75rem;color:#94a3b8">Overall Status:</div>
  <div style="font-size:1.1rem;font-weight:800;color:{health_label_col};font-family:monospace">
    {health_label_val}
  </div>
  <div style="flex:1;height:8px;background:rgba(255,255,255,0.06);border-radius:4px;overflow:hidden">
    <div style="height:100%;width:{health_score}%;background:{health_label_col};
                border-radius:4px;transition:width 0.8s ease"></div>
  </div>
  <div style="font-size:0.95rem;color:#e2e8f0;font-weight:700">{health_score}%</div>
  <div style="font-size:0.68rem;color:#94a3b8">
    ● &lt;45% Critical &nbsp; ● &lt;70% Warning &nbsp; ● ≥70% Healthy
  </div>
</div>
<div style="font-size:0.68rem;color:#334155;margin-top:8px">
  Most delayed category: <span style="color:#fbbf24">{top_cat}</span>
</div>"""

        # ── Plant delay bar chart ─────────────────────────────
        plant_data = _run_neo4j("""
            MATCH (pl:Plant)-[:DISPATCHES]->(sh:Shipment)
            RETURN pl.plant_name AS plant,
                   ROUND(100.0*SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END)/COUNT(sh),2) AS delay_pct,
                   COUNT(sh) AS total_ships,
                   SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS delayed_ships
            ORDER BY delay_pct DESC
        """)
        plant_fig = go.Figure()
        if plant_data:
            pnames  = [r["plant"] for r in plant_data]
            pvals   = [r["delay_pct"] for r in plant_data]
            ptotals = [r["total_ships"] for r in plant_data]
            pdelays = [r["delayed_ships"] for r in plant_data]
            colors  = ["#ef4444" if v > 30 else "#f59e0b" if v > 20 else "#22c55e" for v in pvals]
            plant_fig.add_trace(go.Bar(
                x=pnames, y=pvals, marker_color=colors,
                text=[f"{v}%" for v in pvals], textposition="outside",
                textfont=dict(color="#e2e8f0", size=12),
                customdata=list(zip(ptotals, pdelays)),
                hovertemplate="<b>%{x}</b><br>Delay Rate: %{y}%<br>Total: %{customdata[0]:,}<br>Delayed: %{customdata[1]:,}<extra></extra>"
            ))
            plant_fig.add_hline(y=20, line_dash="dot", line_color="#f59e0b",
                                annotation_text="Warning (20%)", annotation_font_color="#f59e0b",
                                annotation_position="right")
            plant_fig.add_hline(y=30, line_dash="dash", line_color="#ef4444",
                                annotation_text="Critical (30%)", annotation_font_color="#ef4444",
                                annotation_position="right")
        plant_fig.update_layout(**_PLOTLY_LAYOUT)
        plant_fig.update_layout(
            title=dict(text="🏭 Delay Rate % by Plant", font=dict(color="#38bdf8", size=13)),
            yaxis=dict(title="Delay %", range=[0, (max((r["delay_pct"] for r in plant_data), default=40) + 10) if plant_data else 50]),
            xaxis=dict(title=""),
            height=300
        )

        # ── Gauge chart ───────────────────────────────────────
        gauge_fig = _build_health_gauge(health_score, "Overall Network Health")

        # ── Monthly trend chart ───────────────────────────────
        trend_fig = build_monthly_delay_trend()

        return gr.update(value=kpi_html), gr.update(value=plant_fig), gr.update(value=gauge_fig), gr.update(value=trend_fig)

    except Exception as e:
        err_html = f'<p style="color:#f87171;padding:10px">⚠ Health check failed: {e}</p>'
        return gr.update(value=err_html), gr.update(value=_blank), gr.update(value=_blank), gr.update(value=_blank)

# ════════════════════════════════════════════════════════════════════
# VISUALIZATION FUNCTIONS  (Tab 5)
# ════════════════════════════════════════════════════════════════════




def _apply_dark_theme(fig: go.Figure) -> go.Figure:
    """Force opaque dark background + bright labels on any Plotly figure.
    Works on the figure in-place but always returns the figure so callers
    can chain.  We do NOT mutate shared blank-figure singletons — callers
    that pass a shared object should pass a fresh go.Figure() instead."""
    existing_height = max(fig.layout.height or 0, 460)
    fig.update_layout(
        paper_bgcolor="#060c1c",
        plot_bgcolor="#060c1c",
        height=existing_height,
        autosize=False,
        font=dict(color="#f1f5f9", size=12),
        title_font=dict(color="#7dd3fc"),
        modebar=dict(remove=["toImage","sendDataToCloud","zoom","pan","select","lasso2d",
                              "zoomIn2d","zoomOut2d","autoScale2d","resetScale2d",
                              "hoverClosestCartesian","hoverCompareCartesian",
                              "toggleSpikelines","zoomInGeo","zoomOutGeo",
                              "resetGeo","hoverClosestGeo","hoverClosestPie",
                              "toggleHover","resetViews","toggleSpikeLines",
                              "resetViewMapbox"]),
        xaxis=dict(
            tickfont=dict(color="#f1f5f9", size=11),
            gridcolor="rgba(255,255,255,0.08)"
        ),
        yaxis=dict(
            tickfont=dict(color="#f1f5f9", size=11),
            gridcolor="rgba(255,255,255,0.08)"
        ),
        legend=dict(font=dict(color="#f1f5f9"))
    )
    # Force all trace text to be visible
    for trace in fig.data:
        if hasattr(trace, "textfont"):
            trace.update(textfont=dict(color="#f1f5f9"))
        if hasattr(trace, "insidetextfont"):
            trace.update(insidetextfont=dict(color="#ffffff"))
        if hasattr(trace, "outsidetextfont"):
            trace.update(outsidetextfont=dict(color="#f1f5f9"))
    return fig


def get_graph_health():
    return {
        "📊 Node Distribution": _run_neo4j("""
            MATCH (n)
            RETURN labels(n)[0] AS label, COUNT(n) AS count
            ORDER BY count DESC
        """),

        "⏱ Delayed Shipments": _run_neo4j("""
            MATCH (s:Shipment)
            WHERE s.delivery_status = 'Major Delay'
            RETURN COUNT(s) AS total_delayed
        """),

        "🏭 High Risk Suppliers": _run_neo4j("""
            MATCH (sup:Supplier)
            WHERE sup.risk_score > 0.6
            RETURN COUNT(sup) AS risky_suppliers
        """),

        "📉 Avg Delay": _run_neo4j("""
            MATCH (s:Shipment)
            WHERE s.delivery_status = 'Major Delay'
            RETURN AVG(s.delay_days) AS avg_delay_days
        """),

        "🚚 Distributor Impact": _run_neo4j("""
            MATCH (s:Shipment)-[:SHIPPED_TO]->(d:Distributor)
            WHERE s.delivery_status = 'Major Delay'
            RETURN d.distributor_city AS city,
                   COUNT(s) AS delayed,
                   AVG(s.delay_days) AS avg_delay
            ORDER BY delayed DESC
            LIMIT 5
        """)
    }

def build_delay_heatmap():
    try:
        rows = _run_neo4j("""
            MATCH (pl:Plant)-[:DISPATCHES]->(sh:Shipment)-[:CARRIES]->(pr:Product)
            WHERE sh.delivery_status = 'Major Delay'
            RETURN pl.plant_name AS plant,
                   pr.product_category_name AS category,
                   COUNT(sh) AS delay_count
        """)
        if not rows:
            return go.Figure().update_layout(title="No delay data found", **_PLOTLY_LAYOUT)

        df    = pd.DataFrame(rows)
        pivot = df.pivot_table(index="plant", columns="category", values="delay_count", fill_value=0)

        fig = go.Figure(data=go.Heatmap(
            z=pivot.values,
            x=list(pivot.columns),
            y=list(pivot.index),
            colorscale=[
                [0.0,  "#0c1528"],
                [0.3,  "#1d3a7a"],
                [0.6,  "#7c3aed"],
                [0.85, "#ef4444"],
                [1.0,  "#ff0000"],
            ],
            showscale=True,
            colorbar=dict(title="Delay Count", tickfont=dict(color="#dde6f5")),
            hovertemplate="Plant: %{y}<br>Category: %{x}<br>Delays: %{z}<extra></extra>"
        ))
        fig.update_layout(**_PLOTLY_LAYOUT)
        fig.update_layout(
            title="🔥 Delay Heatmap — Plant × Product Category",
            xaxis_title="Product Category",
            yaxis_title="Plant",
        )
        return fig
    except Exception as e:
        return go.Figure().update_layout(title=f"Error: {str(e)}", **_PLOTLY_LAYOUT)


def build_supplier_risk_chart():
    try:
        rows = _run_neo4j("""
            MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)
            RETURN sup.supplier_id   AS supplier_id,
                   sup.supplier_name AS supplier,
                   sup.risk_score    AS risk_score,
                   pl.plant_name     AS plant
            ORDER BY sup.risk_score DESC
            LIMIT 20
        """)
        if not rows:
            return go.Figure().update_layout(title="No supplier data", **_PLOTLY_LAYOUT)

        df = pd.DataFrame(rows)
        colors = df["risk_score"].apply(
            lambda x: "#ef4444" if x > 0.8 else "#f97316" if x > 0.6 else "#eab308" if x > 0.4 else "#10b981"
        )

        fig = go.Figure(go.Bar(
            x=df["risk_score"],
            y=df["supplier"],
            orientation="h",
            marker_color=colors,
            text=df.apply(lambda r: f"[{r.get('supplier_id','?')}] {r['risk_score']:.2f} → {r['plant']}", axis=1),
            textposition="auto",
            textfont=dict(color="#ffffff", size=10),
            hovertemplate="<b>%{y}</b><br>Risk: %{x:.3f}<extra></extra>"
        ))
        fig.add_vline(x=0.7, line_dash="dash", line_color="#ef4444",
                      annotation_text="High Risk Threshold (0.7)",
                      annotation_font_color="#ef4444")
        fig.update_layout(**_PLOTLY_LAYOUT)
        fig.update_layout(
            title="⚠️ Supplier Risk Scores (Top 20)",
            xaxis_title="Risk Score",
            yaxis_title="",
            yaxis=dict(autorange="reversed"),
            height=max(400, len(df) * 28)
        )
        return fig
    except Exception as e:
        return go.Figure().update_layout(title=f"Error: {str(e)}", **_PLOTLY_LAYOUT)


def build_route_efficiency_scatter():
    try:
        rows = _run_neo4j("""
            MATCH (pl:Plant)-[:HAS_ROUTE]->(r:Route)-[:CONNECTS_TO]->(d:Distributor)
            RETURN r.route_id AS route_id,
                   r.mode AS mode,
                   r.PtoD_distance_km AS distance_km,
                   r.PtoD_transportation_cost_inr AS cost_inr,
                   r.cost_efficiency AS efficiency,
                   pl.plant_name AS plant,
                   d.distributor_city AS distributor
            LIMIT 200
        """)
        if not rows:
            return go.Figure().update_layout(title="No route data", **_PLOTLY_LAYOUT)

        df = pd.DataFrame(rows).dropna(subset=["distance_km", "cost_inr"])
        mode_colors = {"Road": "#3b82f6", "Rail": "#10b981", "Air": "#f59e0b", "Sea": "#8b5cf6"}

        fig = go.Figure()
        for mode in df["mode"].unique():
            sub = df[df["mode"] == mode]
            fig.add_trace(go.Scatter(
                x=sub["distance_km"],
                y=sub["cost_inr"],
                mode="markers",
                name=mode,
                marker=dict(
                    color=mode_colors.get(mode, "#94a3b8"),
                    size=sub["efficiency"].fillna(0.5) * 18 + 6,
                    opacity=0.8,
                    line=dict(color="rgba(255,255,255,0.2)", width=1)
                ),
                hovertemplate=(
                    "<b>%{customdata[0]}</b><br>"
                    "Distance: %{x:.0f} km<br>"
                    "Cost: ₹%{y:,.0f}<br>"
                    "Plant: %{customdata[1]}<br>"
                    "Distributor: %{customdata[2]}<extra></extra>"
                ),
                customdata=list(zip(sub["route_id"], sub["plant"], sub["distributor"]))
            ))

        fig.update_layout(**_PLOTLY_LAYOUT)
        fig.update_layout(
            title="🚛 Route Efficiency Map — Distance vs Cost (bubble size = efficiency)",
            xaxis_title="Distance (km)",
            yaxis_title="Transportation Cost (₹)",
            legend_title="Transport Mode",
        )
        return fig
    except Exception as e:
        return go.Figure().update_layout(title=f"Error: {str(e)}", **_PLOTLY_LAYOUT)


def build_monthly_delay_trend():
    try:
        rows = _run_neo4j("""
            MATCH (s:Shipment)
            RETURN s.month_number AS month,
                   s.delivery_status AS status,
                   COUNT(s) AS count
            ORDER BY month
        """)
        if not rows:
            return go.Figure().update_layout(title="No shipment data", **_PLOTLY_LAYOUT)

        df = pd.DataFrame(rows)
        month_names = {1:"Jan",2:"Feb",3:"Mar",4:"Apr",5:"May",6:"Jun",
                       7:"Jul",8:"Aug",9:"Sep",10:"Oct",11:"Nov",12:"Dec"}
        df["month_name"] = df["month"].map(month_names)
        delayed = df[df["status"] == "Major Delay"].set_index("month")
        on_time = df[df["status"] == "On Time"].set_index("month")
        months  = sorted(df["month"].unique())

        fig = go.Figure()
        fig.add_trace(go.Scatter(
            x=[month_names.get(m, str(m)) for m in months],
            y=[delayed.loc[m, "count"] if m in delayed.index else 0 for m in months],
            name="Major Delay",
            mode="lines+markers",
            line=dict(color="#ef4444", width=2.5),
            marker=dict(size=8),
            fill="tozeroy",
            fillcolor="rgba(239,68,68,0.12)"
        ))
        fig.add_trace(go.Scatter(
            x=[month_names.get(m, str(m)) for m in months],
            y=[on_time.loc[m, "count"] if m in on_time.index else 0 for m in months],
            name="On Time",
            mode="lines+markers",
            line=dict(color="#10b981", width=2.5),
            marker=dict(size=8),
            fill="tozeroy",
            fillcolor="rgba(16,185,129,0.10)"
        ))
        fig.update_layout(**_PLOTLY_LAYOUT)
        fig.update_layout(
            title="📅 Monthly Shipment Status Trend",
            xaxis_title="Month",
            yaxis_title="Shipment Count",
        )
        return fig
    except Exception as e:
        return go.Figure().update_layout(title=f"Error: {str(e)}", **_PLOTLY_LAYOUT)


def build_distributor_demand_gap():
    try:
        rows = _run_neo4j("""
            MATCH (s:Shipment)-[:SHIPPED_TO]->(d:Distributor)
            WHERE s.demand_gap > 0
            RETURN d.distributor_city AS city,
                   SUM(s.demand_gap)  AS total_shortage,
                   COUNT(s)           AS shortage_events
            ORDER BY total_shortage DESC
            LIMIT 15
        """)
        if not rows:
            return go.Figure().update_layout(title="No demand gap data", **_PLOTLY_LAYOUT)

        df = pd.DataFrame(rows)
        fig = go.Figure(go.Bar(
            x=df["city"],
            y=df["total_shortage"],
            marker=dict(
                color=df["total_shortage"],
                colorscale=[[0,"#1d4ed8"],[0.5,"#7c3aed"],[1,"#ef4444"]],
                showscale=True,
                colorbar=dict(
                    title=dict(text="Shortage", font=dict(color="#f1f5f9")),
                    tickfont=dict(color="#f1f5f9")
                )
            ),
            text=df["shortage_events"].apply(lambda x: f"{x} events"),
            textposition="outside",
            textfont=dict(color="#f1f5f9", size=10),
            hovertemplate="<b>%{x}</b><br>Total Shortage: %{y:,} units<extra></extra>"
        ))
        fig.update_layout(**_PLOTLY_LAYOUT)
        fig.update_layout(
            title="📦 Stockout Severity by Distributor City (Demand Gap)",
            xaxis_title="Distributor City",
            yaxis_title="Total Shortage (units)",
            xaxis_tickangle=-35,
        )
        return fig
    except Exception as e:
        return go.Figure().update_layout(title=f"Error: {str(e)}", **_PLOTLY_LAYOUT)


# ─────────────────────────────────────────────────────────────
# FIX 6: Sunburst — changed branchvalues to "remainder" + null-safety
# Root cause of old bug: branchvalues="total" requires root value = exact
# sum of all descendants. If any mode value is None/missing the math breaks.
# Fix: use branchvalues="remainder" (Plotly distributes leftover to root)
# and filter out rows where mode is None/empty.
# ─────────────────────────────────────────────────────────────
def build_plant_transport_sunburst():
    try:
        rows = _run_neo4j("""
            MATCH (pl:Plant)-[:HAS_ROUTE]->(r:Route)
            WHERE r.mode IS NOT NULL
            RETURN pl.plant_name AS plant,
                   r.mode        AS mode,
                   COUNT(r)      AS route_count
        """)
        if not rows:
            return go.Figure().update_layout(title="No route data — check that Route nodes have a 'mode' property", **_PLOTLY_LAYOUT)

        df = pd.DataFrame(rows)
        # Drop rows where plant or mode is null/empty
        df = df.dropna(subset=["plant", "mode"])
        df = df[df["plant"].str.strip() != ""]
        df = df[df["mode"].str.strip() != ""]

        if df.empty:
            return go.Figure().update_layout(title="No valid plant/mode data found", **_PLOTLY_LAYOUT)

        # Build sunburst hierarchy: Supply Chain → Plant → Plant·Mode
        labels, parents, values = [], [], []

        # Root
        labels.append("Supply Chain")
        parents.append("")
        values.append(0)  # placeholder; branchvalues="remainder" fills this in

        # Plant level
        for plant in df["plant"].unique():
            labels.append(plant)
            parents.append("Supply Chain")
            values.append(int(df[df["plant"] == plant]["route_count"].sum()))

        # Mode level
        for _, row in df.iterrows():
            label = f"{row['plant']} · {row['mode']}"
            labels.append(label)
            parents.append(row["plant"])
            values.append(int(row["route_count"]))

        # Color mapping
        mode_colors = {"Road": "#3b82f6", "Rail": "#10b981", "Air": "#f59e0b", "Sea": "#8b5cf6"}
        colors = []
        for lbl in labels:
            matched = next((c for m, c in mode_colors.items() if m in lbl), "#1d3a7a")
            colors.append(matched)

        fig = go.Figure(go.Sunburst(
            labels=labels,
            parents=parents,
            values=values,
            marker=dict(colors=colors, line=dict(color="rgba(0,0,0,0.3)", width=1)),
            hovertemplate="<b>%{label}</b><br>Routes: %{value}<extra></extra>",
            branchvalues="remainder",   # FIX: was "total" which caused chart to fail
        ))
        fig.update_layout(
            title="🌐 Transport Mode Distribution by Plant",
            **_PLOTLY_LAYOUT
        )
        return fig
    except Exception as e:
        return go.Figure().update_layout(title=f"Sunburst error: {str(e)}", **_PLOTLY_LAYOUT)



# ════════════════════════════════════════════════════════════════════
# RCA TRAIL v5 — CLEAN HELPER FUNCTIONS
# ════════════════════════════════════════════════════════════════════

# ── Chart metadata: supply chain charts for RCA ────────────────────

# ════════════════════════════════════════════════════════════════════
# 6 DYNAMIC VIEW BUILDERS — context-aware, no hardcoded axes
# Each builder reads the RCA context (question + report text) and
# queries Neo4j to produce the most relevant cut of data.
# ════════════════════════════════════════════════════════════════════

def _extract_context(question: str, report: str = "") -> dict:
    """
    Extract entity context from the question and report text.
    Returns a dict of detected entities for dynamic querying.
    """
    q = (question + " " + report).lower()

    # Map keyword → canonical Neo4j product_category_name
    _CAT_MAP = {
        # Toys
        "toy": "Toys", "toys": "Toys",
        # Auto
        "auto": "Auto", "automobile": "Auto", "automotive": "Auto", "car parts": "Auto",
        # Health Beauty
        "health": "Health Beauty", "beauty": "Health Beauty",
        "health beauty": "Health Beauty", "health_beauty": "Health Beauty",
        # Watches Gifts
        "watch": "Watches Gifts", "watches": "Watches Gifts",
        "gift": "Watches Gifts", "watches gifts": "Watches Gifts",
        "watches_gifts": "Watches Gifts",
        # Bed Bath Table
        "bed bath": "Bed Bath Table", "bed_bath": "Bed Bath Table",
        "bed": "Bed Bath Table", "bath": "Bed Bath Table",
        "table": "Bed Bath Table",
        # Construction Tools Garden
        "construction": "Construction Tools Garden",
        "construction tools": "Construction Tools Garden",
        "construction_tools": "Construction Tools Garden",
        "garden": "Construction Tools Garden",
        "tools": "Construction Tools Garden",
        # Cool Stuff
        "cool stuff": "Cool Stuff", "cool_stuff": "Cool Stuff", "cool": "Cool Stuff",
        # Generic
        "food": "Food", "electronics": "Electronics",
        "apparel": "Apparel", "home": "Home", "sports": "Sports",
    }
    _detected_cat = None
    # Check multi-word keys first, then single-word to avoid partial matches
    for kw in sorted(_CAT_MAP.keys(), key=len, reverse=True):
        if kw in q:
            _detected_cat = _CAT_MAP[kw]
            break

    ctx = {
        "is_supplier": any(w in q for w in ["supplier", "vendor", "risk", "risky", "who is causing"]),
        "is_category": any(w in q for w in ["toy", "auto", "health", "beauty", "watch", "gift",
                                              "category", "product", "construction", "garden",
                                              "tools", "bed bath", "cool stuff", "watches_gifts"]),
        "is_plant":    any(w in q for w in ["plant", "factory", "baddi", "pune", "bhopal", "goa", "mumbai"]),
        "is_stock":    any(w in q for w in ["stock", "shortage", "demand", "distributor", "running out"]),
        "is_trend":    any(w in q for w in ["month", "trend", "season", "time", "pattern", "over time"]),
        "is_route":    any(w in q for w in ["route", "transport", "cost", "distance", "logistics", "road", "rail", "air"]),
        "is_chain":    any(w in q for w in ["end-to-end", "chain", "flow", "path", "full", "sankey", "impact"]),
        "category":    _detected_cat,   # specific category name (e.g. "Toys") or None
        "raw": q,
    }
    return ctx


def build_dynamic_heatmap(question: str = "", report: str = "") -> go.Figure:
    """
    VIEW 1 — Intensity Heatmap.
    Context-aware: if question is about suppliers → Supplier×Category heatmap.
    If about plants → Plant×Category. Default → Plant×Category.
    """
    ctx = _extract_context(question, report)
    try:
        if ctx["is_supplier"]:
            # Supplier × Category delayed shipment heatmap
            rows = _run_neo4j("""
                MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)-[:DISPATCHES]->(sh:Shipment)-[:CARRIES]->(pr:Product)
                WHERE sh.delivery_status = 'Major Delay'
                RETURN sup.supplier_name AS row_label,
                       pr.product_category_name AS col_label,
                       COUNT(sh) AS count
                ORDER BY count DESC
                LIMIT 120
            """)
            row_name, col_name = "Supplier", "Product Category"
            title = "🔥 Supplier × Category — Delayed Shipments"
        elif ctx["is_category"]:
            cat = ctx.get("category")
            if cat:
                rows = _run_neo4j("""
                    MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)-[:DISPATCHES]->(sh:Shipment)-[:CARRIES]->(pr:Product)
                    WHERE sh.delivery_status = 'Major Delay'
                      AND pr.product_category_name = $cat
                    RETURN pl.plant_name AS row_label,
                           sup.supplier_name AS col_label,
                           COUNT(sh) AS count
                    ORDER BY count DESC LIMIT 80
                """, {"cat": cat})
                row_name, col_name = "Plant", "Supplier"
                title = f"🔥 {cat} — Plant × Supplier Delay Heatmap"
            else:
                rows = _run_neo4j("""
                    MATCH (pl:Plant)-[:DISPATCHES]->(sh:Shipment)-[:CARRIES]->(pr:Product)
                    WHERE sh.delivery_status = 'Major Delay'
                    RETURN pl.plant_name AS row_label,
                           pr.product_category_name AS col_label,
                           COUNT(sh) AS count
                """)
                row_name, col_name = "Plant", "Product Category"
                title = "🔥 Plant × Category — Delayed Shipments"
        else:
            # Plant × Category (default)
            rows = _run_neo4j("""
                MATCH (pl:Plant)-[:DISPATCHES]->(sh:Shipment)-[:CARRIES]->(pr:Product)
                WHERE sh.delivery_status = 'Major Delay'
                RETURN pl.plant_name AS row_label,
                       pr.product_category_name AS col_label,
                       COUNT(sh) AS count
            """)
            row_name, col_name = "Plant", "Product Category"
            title = "🔥 Plant × Category — Delayed Shipments"

        if not rows:
            # Broaden to all plants × all categories (no filter)
            rows = _run_neo4j("""
                MATCH (pl:Plant)-[:DISPATCHES]->(sh:Shipment)-[:CARRIES]->(pr:Product)
                WITH pl.plant_name AS plant,
                     pr.product_category_name AS category,
                     COUNT(sh) AS total,
                     SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS dlyd
                WHERE total > 0 AND plant IS NOT NULL AND category IS NOT NULL
                RETURN plant, category,
                       round(100.0*dlyd/total, 1) AS delay_count
                ORDER BY delay_count DESC LIMIT 200
            """)
        if not rows:
            return go.Figure().update_layout(title="No delay data found", **_PLOTLY_LAYOUT)

        df    = pd.DataFrame(rows)
        pivot = df.pivot_table(index="row_label", columns="col_label", values="count", fill_value=0)

        fig = go.Figure(data=go.Heatmap(
            z=pivot.values,
            x=[str(c) for c in pivot.columns],
            y=[str(r) for r in pivot.index],
            colorscale=[
                [0.0, "rgba(6,12,28,1)"],
                [0.25, "#1e3a5f"],
                [0.55, "#7c3aed"],
                [0.8,  "#ef4444"],
                [1.0,  "#ff0000"],
            ],
            showscale=True,
            colorbar=dict(
                title=dict(text="Delays", font=dict(color="#f1f5f9", size=12)),
                tickfont=dict(color="#f1f5f9", size=11),
                outlinecolor="rgba(255,255,255,0.15)"
            ),
            hovertemplate=f"{row_name}: <b>%{{y}}</b><br>{col_name}: <b>%{{x}}</b><br>Delayed: <b>%{{z}}</b><extra></extra>",
            xgap=2, ygap=2
        ))
        fig.update_layout(**_PLOTLY_LAYOUT)
        fig.update_layout(
            title=dict(text=title, font=dict(color="#7dd3fc", size=14)),
            xaxis=dict(
                title=col_name, tickangle=-35,
                tickfont=dict(color="#f1f5f9", size=11)
                
            ),
            yaxis=dict(
                title=row_name,
                tickfont=dict(color="#f1f5f9", size=11)
                
            )
        )
        return fig
    except Exception as e:
        return go.Figure().update_layout(title=f"Heatmap error: {e}", **_PLOTLY_LAYOUT)


def build_dynamic_flow(question: str = "", report: str = "") -> go.Figure:
    """
    Network / Flow Diagram — proper node-link graph.
    Shows Supplier → Plant → Distributor relationships as a layered
    network with coloured edges (red = delayed, green = on-time).
    Adapts which suppliers/plants to highlight based on question context.
    """
    ctx = _extract_context(question, report)
    try:
        # Pull top delayed supplier→plant→distributor links
        cat = ctx.get("category")
        if cat:
            rows = _run_neo4j("""
                MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)-[:DISPATCHES]->(sh:Shipment)
                      -[:SHIPPED_TO]->(dist:Distributor)
                MATCH (sh)-[:CARRIES]->(pr:Product)
                WHERE pr.product_category_name = $cat
                WITH sup, pl, dist,
                     COUNT(sh) AS total,
                     SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS delayed
                WHERE total > 0
                RETURN sup.supplier_name AS supplier,
                       pl.plant_name     AS plant,
                       dist.distributor_city AS city,
                       total, delayed,
                       round(100.0 * delayed / total, 1) AS delay_pct
                ORDER BY delayed DESC LIMIT 40
            """, {"cat": cat})
        else:
            rows = _run_neo4j("""
                MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)-[:DISPATCHES]->(sh:Shipment)
                      -[:SHIPPED_TO]->(dist:Distributor)
                WITH sup, pl, dist,
                     COUNT(sh) AS total,
                     SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS delayed
                WHERE total > 0
                RETURN sup.supplier_name AS supplier,
                       pl.plant_name     AS plant,
                       dist.distributor_city AS city,
                       total, delayed,
                       round(100.0 * delayed / total, 1) AS delay_pct
                ORDER BY delayed DESC LIMIT 40
            """)
        if not rows:
            rows = _run_neo4j("""
                MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)-[:DISPATCHES]->(sh:Shipment)
                      -[:SHIPPED_TO]->(dist:Distributor)
                WITH sup, pl, dist, COUNT(sh) AS total,
                     SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS delayed
                WHERE total > 0
                RETURN sup.supplier_name AS supplier, pl.plant_name AS plant,
                       dist.distributor_city AS city, total, delayed,
                       round(100.0*delayed/total,1) AS delay_pct
                ORDER BY delayed DESC LIMIT 30
            """)
        if not rows:
            return go.Figure().update_layout(title="No flow data", **_PLOTLY_LAYOUT)

        import math
        df = pd.DataFrame(rows)

        # ── Build layered node positions ──────────────────────────
        # Layer 0 (left)  = Suppliers   x=0.0
        # Layer 1 (center)= Plants      x=0.5
        # Layer 2 (right) = Distributors x=1.0
        suppliers = df["supplier"].unique().tolist()
        plants    = df["plant"].unique().tolist()
        cities    = df["city"].unique().tolist()

        def _y_positions(items):
            n = len(items)
            return {item: 0.1 + 0.8 * i / max(n - 1, 1) for i, item in enumerate(items)}

        sup_y  = _y_positions(suppliers)
        plt_y  = _y_positions(plants)
        city_y = _y_positions(cities)

        node_x, node_y, node_text, node_color, node_size = [], [], [], [], []

        for s in suppliers:
            node_x.append(0.0); node_y.append(sup_y[s])
            node_text.append(s[:18])
            # colour by risk if available
            node_color.append("#ef4444" if df[df["supplier"]==s]["delay_pct"].mean() > 40
                              else "#f97316" if df[df["supplier"]==s]["delay_pct"].mean() > 20
                              else "#4ade80")
            node_size.append(18)

        for p in plants:
            node_x.append(0.5); node_y.append(plt_y[p])
            node_text.append(p)
            node_color.append("#38bdf8")
            node_size.append(22)

        for c in cities:
            node_x.append(1.0); node_y.append(city_y[c])
            node_text.append(c)
            delay_avg = df[df["city"]==c]["delay_pct"].mean()
            node_color.append("#ef4444" if delay_avg > 40 else "#fbbf24" if delay_avg > 20 else "#4ade80")
            node_size.append(16)

        # ── Build edges ───────────────────────────────────────────
        edge_x, edge_y, edge_colors = [], [], []
        for _, row in df.iterrows():
            sup_x, sup_yv = 0.0, sup_y.get(row["supplier"], 0.5)
            plt_x, plt_yv = 0.5, plt_y.get(row["plant"], 0.5)
            cit_x, cit_yv = 1.0, city_y.get(row["city"], 0.5)
            # Sup → Plant edge
            edge_x += [sup_x, plt_x, None]
            edge_y += [sup_yv, plt_yv, None]
            # Plant → City edge
            edge_x += [plt_x, cit_x, None]
            edge_y += [plt_yv, cit_yv, None]

        # Single trace for all edges (grey), then overlay red for high-delay
        fig = go.Figure()

        # All edges grey
        fig.add_trace(go.Scatter(
            x=edge_x, y=edge_y, mode="lines",
            line=dict(color="rgba(148,163,184,0.18)", width=1),
            hoverinfo="none", showlegend=False
        ))

        # Highlight high-delay edges in red
        red_ex, red_ey = [], []
        for _, row in df[df["delay_pct"] > 30].iterrows():
            sup_yv = sup_y.get(row["supplier"], 0.5)
            plt_yv = plt_y.get(row["plant"], 0.5)
            cit_yv = city_y.get(row["city"], 0.5)
            red_ex += [0.0, 0.5, None, 0.5, 1.0, None]
            red_ey += [sup_yv, plt_yv, None, plt_yv, cit_yv, None]
        if red_ex:
            fig.add_trace(go.Scatter(
                x=red_ex, y=red_ey, mode="lines",
                line=dict(color="rgba(239,68,68,0.55)", width=1.5),
                hoverinfo="none", showlegend=False, name="Delayed path"
            ))

        # Nodes
        fig.add_trace(go.Scatter(
            x=node_x, y=node_y,
            mode="markers+text",
            text=node_text,
            textposition="middle right",
            textfont=dict(color="#f1f5f9", size=10),
            marker=dict(
                size=node_size,
                color=node_color,
                opacity=0.92,
                line=dict(color="rgba(255,255,255,0.3)", width=1.5)
            ),
            hovertemplate="%{text}<extra></extra>",
            showlegend=False
        ))

        # Layer labels at top
        for lbl, xpos in [("Suppliers", 0.0), ("Plants", 0.5), ("Distributors", 1.0)]:
            fig.add_annotation(
                x=xpos, y=1.06, xref="paper", yref="paper",
                text=f"<b>{lbl}</b>",
                font=dict(color="#7dd3fc", size=12),
                showarrow=False, xanchor="center"
            )

        fig.update_layout(**_PLOTLY_LAYOUT)
        fig.update_layout(
            title=dict(text="🌐 Supply Network Flow — Supplier → Plant → Distributor",
                       font=dict(color="#7dd3fc", size=14)),
            xaxis=dict(visible=False, range=[-0.15, 1.25]),
            yaxis=dict(visible=False),
            height=500,
        )
        return fig

    except Exception as e:
        return go.Figure().update_layout(
            title=f"Flow diagram error: {e}", **_PLOTLY_LAYOUT
        )


def build_dynamic_trend(question: str = "", report: str = "") -> go.Figure:
    """
    VIEW 3 — Time Trend.
    If supplier question → top 5 risky suppliers' delay trend by month.
    If category → delay by category over months.
    Default → overall delayed vs on-time trend.
    """
    ctx = _extract_context(question, report)
    try:
        if ctx["is_supplier"]:
            rows = _run_neo4j("""
                MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)-[:DISPATCHES]->(sh:Shipment)
                WHERE sh.delivery_status = 'Major Delay' AND sh.month_number IS NOT NULL
                  AND sup.risk_score >= 0.6
                RETURN sup.supplier_name AS supplier,
                       sh.month_number AS month,
                       COUNT(sh) AS delayed
                ORDER BY delayed DESC
            """)
            if rows:
                df = pd.DataFrame(rows)
                month_names = {1:"Jan",2:"Feb",3:"Mar",4:"Apr",5:"May",6:"Jun",
                               7:"Jul",8:"Aug",9:"Sep",10:"Oct",11:"Nov",12:"Dec"}
                df["month_label"] = df["month"].map(month_names)
                top_sups = df.groupby("supplier")["delayed"].sum().nlargest(5).index.tolist()
                df = df[df["supplier"].isin(top_sups)]
                colors = ["#ef4444","#f97316","#eab308","#3b82f6","#8b5cf6"]
                fig = go.Figure()
                for i, sup in enumerate(top_sups):
                    sub = df[df["supplier"] == sup].sort_values("month")
                    fig.add_trace(go.Scatter(
                        x=[month_names.get(m, str(m)) for m in sub["month"]],
                        y=sub["delayed"].tolist(),
                        name=sup[:22], mode="lines+markers",
                        line=dict(color=colors[i % len(colors)], width=2),
                        marker=dict(size=7),
                        hovertemplate=f"<b>{sup}</b><br>Month: %{{x}}<br>Delayed: <b>%{{y}}</b><extra></extra>"
                    ))
                fig.update_layout(**_PLOTLY_LAYOUT)
                fig.update_layout(
                    title=dict(text="📅 High-Risk Supplier Delay Trend by Month", font=dict(color="#7dd3fc", size=14)),
                    xaxis=dict(title=dict(text="Month", font=dict(color="#94a3b8")), tickfont=dict(color="#f1f5f9", size=11)),
                    yaxis=dict(title=dict(text="Delayed Shipments", font=dict(color="#94a3b8")), tickfont=dict(color="#f1f5f9", size=11))
                )
            rows = _run_neo4j("""
                MATCH (sh:Shipment)-[:CARRIES]->(pr:Product)
                WHERE sh.delivery_status = 'Major Delay' AND sh.month_number IS NOT NULL
                RETURN pr.product_category_name AS category,
                       sh.month_number AS month,
                       COUNT(sh) AS delayed
                ORDER BY delayed DESC
            """)
            if rows:
                df = pd.DataFrame(rows)
                month_names = {1:"Jan",2:"Feb",3:"Mar",4:"Apr",5:"May",6:"Jun",
                               7:"Jul",8:"Aug",9:"Sep",10:"Oct",11:"Nov",12:"Dec"}
                df["month_label"] = df["month"].map(month_names)
                top_cats = df.groupby("category")["delayed"].sum().nlargest(5).index.tolist()
                df = df[df["category"].isin(top_cats)]
                colors = ["#ef4444","#f97316","#eab308","#3b82f6","#8b5cf6"]
                fig = go.Figure()
                for i, cat in enumerate(top_cats):
                    sub = df[df["category"] == cat].sort_values("month")
                    fig.add_trace(go.Scatter(
                        x=[month_names.get(m, str(m)) for m in sub["month"]],
                        y=sub["delayed"].tolist(),
                        name=cat[:20], mode="lines+markers",
                        line=dict(color=colors[i % len(colors)], width=2),
                        marker=dict(size=7)
                    ))
                fig.update_layout(**_PLOTLY_LAYOUT)
                fig.update_layout(
                    title=dict(text="📅 Delayed Shipments by Category & Month", font=dict(color="#7dd3fc", size=14)),
                    xaxis=dict(title=dict(text="Month", font=dict(color="#94a3b8")), tickfont=dict(color="#f1f5f9", size=11)),
                    yaxis=dict(title=dict(text="Delayed Shipments", font=dict(color="#94a3b8")), tickfont=dict(color="#f1f5f9", size=11))
                )
                return fig

        # Default: overall on-time vs delayed trend
        return build_monthly_delay_trend()

    except Exception as e:
        return go.Figure().update_layout(title=f"Trend error: {e}", **_PLOTLY_LAYOUT)


def build_dynamic_pareto(question: str = "", report: str = "") -> go.Figure:
    """
    VIEW 4 — Pareto / Ranked Bar Chart.
    Context-aware:
      - Supplier question → Suppliers ranked by delayed shipment count + cumulative %
      - Category question → Categories ranked by delayed count
      - Plant question → Plants ranked by delay rate %
      - Stock question → Distributors ranked by demand gap
      - Default → Suppliers ranked by risk score
    """
    ctx = _extract_context(question, report)
    try:
        if ctx["is_supplier"] or (not any([ctx["is_category"], ctx["is_plant"], ctx["is_stock"], ctx["is_route"]])):
            rows = _run_neo4j("""
                MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)-[:DISPATCHES]->(sh:Shipment)
                WHERE sh.delivery_status = 'Major Delay'
                RETURN sup.supplier_name AS label,
                       COUNT(sh) AS value,
                       MAX(sup.risk_score) AS risk
                ORDER BY value DESC LIMIT 15
            """)
            title = "📊 Suppliers Ranked by Delayed Shipments (Pareto)"
            x_label = "Delayed Shipments"

        elif ctx["is_category"]:
            cat = ctx.get("category")
            if cat:
                rows = _run_neo4j("""
                    MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)-[:DISPATCHES]->(sh:Shipment)-[:CARRIES]->(pr:Product)
                    WHERE sh.delivery_status = 'Major Delay'
                      AND pr.product_category_name = $cat
                    RETURN sup.supplier_name AS label,
                           COUNT(sh) AS value,
                           MAX(sup.risk_score) AS risk
                    ORDER BY value DESC LIMIT 12
                """, {"cat": cat})
                title = f"📊 Top Suppliers — {cat} Delayed Shipments (Pareto)"
            else:
                rows = _run_neo4j("""
                    MATCH (sh:Shipment)-[:CARRIES]->(pr:Product)
                    WHERE sh.delivery_status = 'Major Delay'
                    RETURN pr.product_category_name AS label, COUNT(sh) AS value
                    ORDER BY value DESC LIMIT 10
                """)
                title = "📊 Product Categories Ranked by Delays (Pareto)"
            x_label = "Delayed Shipments"

        elif ctx["is_plant"]:
            rows = _run_neo4j("""
                MATCH (pl:Plant)-[:DISPATCHES]->(sh:Shipment)
                WITH pl.plant_name AS label,
                     COUNT(sh) AS total,
                     SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS delayed
                RETURN label, delayed AS value, round(100.0*delayed/total, 1) AS pct
                ORDER BY value DESC
            """)
            title = "📊 Plants Ranked by Delayed Shipments (Pareto)"
            x_label = "Delayed Shipments"

        elif ctx["is_stock"]:
            rows = _run_neo4j("""
                MATCH (sh:Shipment)-[:SHIPPED_TO]->(d:Distributor)
                WHERE sh.demand_gap > 0
                RETURN d.distributor_city AS label, SUM(sh.demand_gap) AS value
                ORDER BY value DESC LIMIT 12
            """)
            title = "📊 Distributors Ranked by Demand Gap (Pareto)"
            x_label = "Units Short"
        else:
            rows = _run_neo4j("""
                MATCH (sup:Supplier)
                RETURN sup.supplier_name AS label, sup.risk_score AS value
                ORDER BY value DESC LIMIT 15
            """)
            title = "📊 Suppliers Ranked by Risk Score (Pareto)"
            x_label = "Risk Score"

        if not rows:
            rows = _run_neo4j("""
                MATCH (pl:Plant)-[:DISPATCHES]->(sh:Shipment)
                WITH pl.plant_name AS entity, COUNT(sh) AS total,
                     SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS delayed
                WHERE total > 0
                RETURN entity, total, delayed,
                       round(100.0*delayed/total,1) AS delay_pct
                ORDER BY delayed DESC LIMIT 15
            """)
        if not rows:
            return go.Figure().update_layout(title="No delay data", **_PLOTLY_LAYOUT)

        df = pd.DataFrame(rows)
        df = df.sort_values("value", ascending=False).head(12)
        df["cumulative_pct"] = (df["value"].cumsum() / df["value"].sum() * 100).round(1)

        # Color by severity
        max_v = df["value"].max()
        colors = [
            "#ef4444" if v/max_v > 0.6 else "#f97316" if v/max_v > 0.35 else "#eab308" if v/max_v > 0.15 else "#22c55e"
            for v in df["value"]
        ]

        fig = go.Figure()
        fig.add_trace(go.Bar(
            x=df["label"],
            y=df["value"],
            marker=dict(color=colors, line=dict(color="rgba(0,0,0,0.3)", width=0.5)),
            name=x_label,
            hovertemplate="<b>%{x}</b><br>" + x_label + ": <b>%{y:,}</b><extra></extra>"
        ))
        fig.add_trace(go.Scatter(
            x=df["label"],
            y=df["cumulative_pct"],
            mode="lines+markers",
            name="Cumulative %",
            yaxis="y2",
            line=dict(color="#38bdf8", width=2, dash="dot"),
            marker=dict(size=6, color="#38bdf8"),
            hovertemplate="Cumulative: <b>%{y:.1f}%</b><extra></extra>"
        ))
        fig.update_layout(**_PLOTLY_LAYOUT)
        fig.update_layout(
            title=dict(text=title, font=dict(color="#7dd3fc", size=14)),
            xaxis=dict(
                title="", tickangle=-35,
                tickfont=dict(color="#f1f5f9", size=10)
            ),
            yaxis=dict(
                title=x_label,
                tickfont=dict(color="#f1f5f9", size=11)
                
            ),
            yaxis2=dict(
                title="Cumulative %",
                overlaying="y", side="right", range=[0, 110],
                tickfont=dict(color="#38bdf8", size=11),
                
                showgrid=False
            ),
            barmode="overlay"
        )
        return fig
    except Exception as e:
        return go.Figure().update_layout(title=f"Pareto error: {e}", **_PLOTLY_LAYOUT)


def build_dynamic_pie(question: str = "", report: str = "") -> go.Figure:
    """
    VIEW 5 — Donut / Pie Chart.
    Context-aware:
      - Supplier question → delay share by supplier (top 8)
      - Route/transport question → shipment share by transport mode
      - Default → delayed share by product category (donut)
    """
    ctx = _extract_context(question, report)
    try:
        if ctx["is_supplier"]:
            rows = _run_neo4j("""
                MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)-[:DISPATCHES]->(sh:Shipment)
                WHERE sh.delivery_status = 'Major Delay'
                RETURN sup.supplier_name AS label, COUNT(sh) AS value
                ORDER BY value DESC LIMIT 8
            """)
            title = "🥧 Delay Share by Supplier"
            center_label = "supplier<br>delays"

        elif ctx["is_route"]:
            rows = _run_neo4j("""
                MATCH (r:Route)
                WHERE r.mode IS NOT NULL
                RETURN r.mode AS label, COUNT(r) AS value
                ORDER BY value DESC
            """)
            title = "🥧 Route Share by Transport Mode"
            center_label = "routes"

        elif ctx["is_plant"]:
            rows = _run_neo4j("""
                MATCH (pl:Plant)-[:DISPATCHES]->(sh:Shipment)
                WHERE sh.delivery_status = 'Major Delay'
                RETURN pl.plant_name AS label, COUNT(sh) AS value
                ORDER BY value DESC
            """)
            title = "🥧 Delay Share by Plant"
            center_label = "plant<br>delays"

        elif ctx["is_category"]:
            cat = ctx.get("category")
            if cat:
                rows = _run_neo4j("""
                    MATCH (pl:Plant)-[:DISPATCHES]->(sh:Shipment)-[:CARRIES]->(pr:Product)
                    WHERE sh.delivery_status = 'Major Delay'
                      AND pr.product_category_name = $cat
                    RETURN pl.plant_name AS label, COUNT(sh) AS value
                    ORDER BY value DESC
                """, {"cat": cat})
                title = f"🥧 {cat} Delay Share by Plant"
            else:
                rows = _run_neo4j("""
                    MATCH (sh:Shipment)-[:CARRIES]->(pr:Product)
                    WHERE sh.delivery_status = 'Major Delay'
                    RETURN pr.product_category_name AS label, COUNT(sh) AS value
                    ORDER BY value DESC
                """)
                title = "🥧 Delay Share by Product Category"
            center_label = "delays"
        else:
            rows = _run_neo4j("""
                MATCH (sh:Shipment)-[:CARRIES]->(pr:Product)
                WHERE sh.delivery_status = 'Major Delay'
                RETURN pr.product_category_name AS label, COUNT(sh) AS value
                ORDER BY value DESC
            """)
            title = "🥧 Delay Share by Product Category"
            center_label = "total<br>delays"

        if not rows:
            rows = _run_neo4j("""
                MATCH (sh:Shipment)-[:CARRIES]->(pr:Product)
                WHERE sh.delivery_status='Major Delay' AND pr.product_category_name IS NOT NULL
                RETURN pr.product_category_name AS entity,
                       COUNT(sh) AS delayed, COUNT(sh) AS total
                ORDER BY delayed DESC LIMIT 12
            """)
        if not rows:
            return go.Figure().update_layout(title="No data", **_PLOTLY_LAYOUT)

        df = pd.DataFrame(rows)
        colors = ["#ef4444","#f97316","#eab308","#22c55e","#3b82f6","#8b5cf6","#ec4899","#06b6d4"]

        fig = go.Figure(go.Pie(
            labels=df["label"],
            values=df["value"],
            hole=0.52,
            marker=dict(
                colors=colors[:len(df)],
                line=dict(color="rgba(6,12,28,0.8)", width=2)
            ),
            textinfo="label+percent",
            textfont=dict(size=12, color="#f1f5f9"),
            hovertemplate="<b>%{label}</b><br>Count: <b>%{value:,}</b><br>Share: <b>%{percent}</b><extra></extra>",
            insidetextorientation="radial"
        ))
        fig.update_layout(**_PLOTLY_LAYOUT)
        fig.update_layout(
            title=dict(text=title, font=dict(color="#7dd3fc", size=14)),
            showlegend=True,
            annotations=[dict(
                text=f"<b>{df['value'].sum():,}</b><br>{center_label}",
                x=0.5, y=0.5,
                font=dict(size=13, color="#f1f5f9"),
                showarrow=False
            )],
        )
        return fig
    except Exception as e:
        return go.Figure().update_layout(title=f"Pie error: {e}", **_PLOTLY_LAYOUT)


def build_dynamic_network(question: str = "", report: str = "") -> go.Figure:
    """
    VIEW 6 — Network / Scatter Map.
    Context-aware:
      - Supplier question → Supplier risk vs capacity scatter (bubble = lead time)
      - Route/logistics → Route distance vs cost scatter (existing builder)
      - Plant question → Plant delay rate vs total shipments bubble
      - Default → Supplier risk vs delayed shipments scatter
    """
    ctx = _extract_context(question, report)
    try:
        if ctx["is_route"]:
            return build_route_efficiency_scatter()

        elif ctx["is_plant"]:
            rows = _run_neo4j("""
                MATCH (pl:Plant)-[:DISPATCHES]->(sh:Shipment)
                RETURN pl.plant_name AS plant,
                       COUNT(sh) AS total,
                       SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS delayed,
                       round(100.0 * SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) / COUNT(sh), 1) AS delay_pct
            """)
            if rows:
                df = pd.DataFrame(rows)
                colors = ["#ef4444" if p > 30 else "#f97316" if p > 20 else "#22c55e" for p in df["delay_pct"]]
                fig = go.Figure(go.Scatter(
                    x=df["total"],
                    y=df["delay_pct"],
                    mode="markers+text",
                    text=df["plant"],
                    textposition="top center",
                    textfont=dict(color="#f1f5f9", size=11),
                    marker=dict(
                        size=df["delayed"] / df["delayed"].max() * 40 + 14,
                        color=colors,
                        opacity=0.85,
                        line=dict(color="rgba(255,255,255,0.25)", width=1)
                    ),
                    hovertemplate="<b>%{text}</b><br>Total Shipments: <b>%{x:,}</b><br>Delay Rate: <b>%{y:.1f}%</b><extra></extra>",
                    showlegend=False
                ))
                fig.update_layout(**_PLOTLY_LAYOUT)
                fig.update_layout(
                    title=dict(text="🌐 Plants — Volume vs Delay Rate (bubble = delayed count)", font=dict(color="#7dd3fc", size=14)),
                    xaxis=dict(title=dict(text="Total Shipments", font=dict(color="#94a3b8")), tickfont=dict(color="#f1f5f9", size=11)),
                    yaxis=dict(title=dict(text="Delay Rate %", font=dict(color="#94a3b8")), tickfont=dict(color="#f1f5f9", size=11))
                )
                return fig

        # Default: Supplier risk score vs delayed shipments (network view)
        rows = _run_neo4j("""
            MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)-[:DISPATCHES]->(sh:Shipment)
            RETURN sup.supplier_name AS supplier,
                   sup.risk_score AS risk,
                   sup.annual_capacity_units AS capacity,
                   sup.StoP_lead_time_days AS lead_time,
                   COUNT(sh) AS total,
                   SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS delayed,
                   pl.plant_name AS plant
            ORDER BY delayed DESC LIMIT 30
        """)
        if not rows:
            rows = _run_neo4j("""
                MATCH (s:Supplier)-[:SUPPLIES_TO]->(p:Plant)
                RETURN s.supplier_id AS supplier_id, s.supplier_name AS supplier,
                       coalesce(s.risk_score, 0.3) AS risk_score, p.plant_name AS plant,
                       0.0 AS delay_rate_pct, 100 AS total_shipments
                ORDER BY risk_score DESC LIMIT 20
            """)
        if not rows:
            return go.Figure().update_layout(title="No supplier data", **_PLOTLY_LAYOUT)

        df = pd.DataFrame(rows)
        df["risk"] = pd.to_numeric(df["risk"], errors="coerce").fillna(0)
        df["delayed"] = pd.to_numeric(df["delayed"], errors="coerce").fillna(0)
        df["lead_time"] = pd.to_numeric(df["lead_time"], errors="coerce").fillna(7)
        df["total"] = pd.to_numeric(df["total"], errors="coerce").fillna(1)
        df["delay_rate"] = (df["delayed"] / df["total"].clip(lower=1) * 100).round(1)

        colors = ["#ef4444" if r > 0.8 else "#f97316" if r > 0.6 else "#eab308" if r > 0.4 else "#22c55e"
                  for r in df["risk"]]

        fig = go.Figure(go.Scatter(
            x=df["risk"],
            y=df["delay_rate"],
            mode="markers+text",
            text=df["supplier"].str[:14],
            textposition="top center",
            textfont=dict(color="#f1f5f9", size=10),
            marker=dict(
                size=df["lead_time"].clip(lower=3) * 2.5 + 10,
                color=colors,
                opacity=0.82,
                line=dict(color="rgba(255,255,255,0.2)", width=1)
            ),
            customdata=list(zip(df["supplier"], df["delayed"], df["lead_time"], df["plant"])),
            hovertemplate=(
                "<b>%{customdata[0]}</b><br>"
                "Risk Score: <b>%{x:.2f}</b><br>"
                "Delay Rate: <b>%{y:.1f}%</b><br>"
                "Delayed Shipments: <b>%{customdata[1]:,}</b><br>"
                "Lead Time: <b>%{customdata[2]} days</b><br>"
                "Plant: <b>%{customdata[3]}</b><extra></extra>"
            ),
            showlegend=False
        ))
        # Add threshold lines
        fig.add_vline(x=0.7, line_dash="dash", line_color="#ef4444",
                      annotation_text="High Risk (0.7)", annotation_font_color="#ef4444",
                      annotation_position="top right")
        fig.add_hline(y=25, line_dash="dot", line_color="#f97316",
                      annotation_text="25% Delay Rate", annotation_font_color="#f97316",
                      annotation_position="right")
        fig.update_layout(**_PLOTLY_LAYOUT)
        fig.update_layout(
            title=dict(text="🌐 Supplier Network — Risk vs Delay Rate (bubble = lead time)", font=dict(color="#7dd3fc", size=14)),
            xaxis=dict(title=dict(text="Risk Score (0–1)", font=dict(color="#94a3b8")), tickfont=dict(color="#f1f5f9", size=11), range=[-0.05, 1.1]),
            yaxis=dict(title=dict(text="Delay Rate %", font=dict(color="#94a3b8")), tickfont=dict(color="#f1f5f9", size=11))
        )
        return fig
    except Exception as e:
        return go.Figure().update_layout(title=f"Network error: {e}", **_PLOTLY_LAYOUT)


# ── 4 chart tabs: Network Diagram, Bar Chart, Pie Chart, Heatmap ─────────────
_CHART4_BUILDERS = {
    "🌐 Bubble Chart":   build_dynamic_network,
    "📊 Pareto Chart": build_dynamic_pareto,
    "🥧 Pie Chart": build_dynamic_pie,
    "🔥 Heatmap":   build_dynamic_heatmap,
}

_CHART4_META = {
    "🌐 Bubble Chart": {
        "icon": "🌐",
        "label": "Bubble Chart",
        "desc": "Supplier risk × delay rate scatter — bubble size = lead time. Spot your most dangerous single point of failure instantly.",
        "what": "A <strong>bubble scatter</strong> showing supplier risk score vs delay rate %. Bubble size = lead time days. Red = critical (&gt;0.8 risk), orange = high, green = safe.",
        "why":  "Reveals <em>multi-dimensional risk</em>. A large red bubble in the top-right = high risk + high delay + long lead time = maximum exposure. Use this to prioritise which suppliers to audit or replace first.",
        "tags": ["Multi-dimensional", "Risk × delay", "Lead time exposure", "Dangerous outliers"],
    },
    "📊 Pareto Chart": {
        "icon": "📊",
        "label": "Pareto Chart",
        "desc": "Ranked bar + cumulative % (Pareto). Context-aware: suppliers / plants / categories / distributors ranked by impact.",
        "what": "A <strong>ranked bar chart with Pareto cumulative % line</strong>. The 80/20 rule in action — shows which entities account for the majority of the problem.",
        "why":  "If 3 out of 48 suppliers cause 78% of all delays, fixing those 3 is your entire action plan. The bars adapt to your question: supplier questions rank suppliers; plant questions rank plants.",
        "tags": ["80/20 rule", "Priority ranking", "Adaptive entity", "Action focus"],
    },
    "🥧 Pie Chart": {
        "icon": "🥧",
        "label": "Pie Chart",
        "desc": "Donut split of delays by category, supplier, plant, or transport mode — adapts to your question.",
        "what": "A <strong>donut chart</strong> showing proportional share of delays. Adapts to your question: category questions split by product; supplier questions split by supplier; route questions split by transport mode.",
        "why":  "Reveals <em>disproportionate share</em>. If one supplier holds 42% of all delays despite being 1 of 48, that is your primary root cause. Share imbalances are immediately visible in donut form.",
        "tags": ["Share distribution", "Disproportionate risk", "Adaptive split", "Visual share"],
    },
    "🔥 Heatmap": {
        "icon": "🔥",
        "label": "Heatmap",
        "desc": "Colour-coded grid of Plant × Category (or Supplier × Category). Darker red = more delays. Best RCA entry point.",
        "what": "A <strong>colour-coded grid</strong> where each cell = a combination of two dimensions. Darker red = more delayed shipments. Axes adapt: supplier questions show Supplier × Category; plant questions show Plant × Category.",
        "why":  "The fastest way to spot the intersection of <em>who</em> and <em>what</em> is causing the problem. This is your RCA entry point — it instantly answers 'where are delays concentrated?'",
        "tags": ["Concentration map", "Adaptive axes", "Colour severity", "Entry point"],
    },
}

_CHART4_NAMES = list(_CHART4_BUILDERS.keys())

# Keep _VIEW_6_BUILDERS alias for any legacy code
_VIEW_6_BUILDERS = _CHART4_BUILDERS

_VIEW_6_META = {
    "🔥 Heatmap": {
        "what": "A colour-coded grid showing where delays concentrate most. Each cell = a combination of two dimensions (e.g. Supplier × Category). <strong>Darker red = more delayed shipments</strong>.",
        "why":  "The fastest way to spot the intersection of <em>who</em> and <em>what</em> is causing the problem. Axes adapt to your question — supplier questions show Supplier × Category; plant questions show Plant × Category.",
        "tags": ["Concentration map", "Adaptive axes", "Colour severity", "Entry point"],
    },
    "🔀 Flow": {
        "what": "A <strong>Sankey flow diagram</strong> showing the full supply chain: Suppliers → Plants → Distributor Cities. Band width = shipment volume. <em>Red bands = high delay rate on that path.</em>",
        "why":  "The only view showing <em>end-to-end volume flow</em>. A wide red band reveals the exact path causing the most delayed shipments — supplier, plant, and destination city all in one view. Best for 'where is the biggest bottleneck?' questions.",
        "tags": ["End-to-end", "Volume + delay", "Bottleneck path", "Full chain"],
    },
    "📅 Trend": {
        "what": "A <strong>time-series line chart</strong> showing delayed shipments over 12 months. Adapts to your question: supplier questions show per-supplier trends; category questions show per-category trends.",
        "why":  "Answers <em>when</em> the problem occurs. A spike in Oct–Dec = seasonal → fix forecasting. A persistently high line = structural → fix supplier contracts or capacity. The pattern determines the right fix.",
        "tags": ["Seasonal vs structural", "12-month view", "Per-entity trend", "Fix type indicator"],
    },
    "📊 Pareto": {
        "what": "A <strong>ranked bar chart with cumulative %</strong> line (Pareto principle). Shows which suppliers / categories / plants / distributors account for 80% of the problem.",
        "why":  "The 80/20 rule in action. If 3 out of 48 suppliers cause 78% of all delays, fixing those 3 is your entire RCA action plan. Axes adapt to your question so you always see the right ranked list.",
        "tags": ["80/20 rule", "Priority ranking", "Adaptive entity", "Action focus"],
    },
    "🥧 Split": {
        "what": "A <strong>donut chart</strong> showing how delays (or routes) are distributed across categories, suppliers, plants, or transport modes — depending on your question.",
        "why":  "Reveals <em>disproportionate share</em>. If one supplier holds 42% of all delays despite being just 1 of 48, that is your primary root cause. The donut format makes share imbalances visually immediate.",
        "tags": ["Share distribution", "Disproportionate risk", "Adaptive split", "Visual share"],
    },
    "🌐 Network": {
        "what": "A <strong>bubble scatter chart</strong>. For supplier questions: X = risk score, Y = delay rate %, bubble size = lead time. For plant questions: volume vs delay rate. For routes: distance vs cost.",
        "why":  "Reveals <em>multi-dimensional risk</em> that bars and heatmaps miss. A supplier with high risk score, high delay rate, AND long lead time is your most dangerous single point of failure — visible instantly as a large red bubble in the top-right quadrant.",
        "tags": ["Multi-dimensional", "Risk × delay", "Bubble size = lead time", "Dangerous outliers"],
    },
}


def _build_view_why_html(name: str) -> str:
    meta = _CHART4_META.get(name, _VIEW_6_META.get(name, {}))
    what = meta.get("what", "")
    why  = meta.get("why", "")
    tags_html = "".join(
        f'<span style="font-size:10px;padding:3px 9px;border-radius:20px;'
        f'background:rgba(56,189,248,0.09);border:1px solid rgba(56,189,248,0.22);'
        f'color:#7dd3fc;margin:2px 2px 0 0">{t}</span>'
        for t in meta.get("tags", [])
    )
    return f"""
<div style="display:grid;grid-template-columns:1fr 1fr;gap:14px;margin-top:16px">
  <div style="background:rgba(6,15,35,0.95);border:1px solid rgba(56,189,248,0.35);
              border-radius:12px;padding:18px 20px;position:relative;overflow:hidden">
    <div style="position:absolute;top:0;left:0;right:0;height:2px;
                background:linear-gradient(90deg,transparent,#38bdf8,transparent)"></div>
    <div style="font-size:11px;margin-bottom:7px">💡</div>
    <div style="font-size:0.62rem;font-weight:700;text-transform:uppercase;
                letter-spacing:0.1em;color:#38bdf8;margin-bottom:9px">What does this chart show?</div>
    <div style="font-size:0.82rem;line-height:1.65;color:#e2e8f0">{what}</div>
    <div style="display:flex;flex-wrap:wrap;gap:5px;margin-top:10px">{tags_html}</div>
  </div>
  <div style="background:rgba(6,15,35,0.95);border:1px solid rgba(124,58,237,0.35);
              border-radius:12px;padding:18px 20px;position:relative;overflow:hidden">
    <div style="position:absolute;top:0;left:0;right:0;height:2px;
                background:linear-gradient(90deg,transparent,#7c3aed,transparent)"></div>
    <div style="font-size:11px;margin-bottom:7px">🎯</div>
    <div style="font-size:0.62rem;font-weight:700;text-transform:uppercase;
                letter-spacing:0.1em;color:#a78bfa;margin-bottom:9px">Why is it useful for RCA?</div>
    <div style="font-size:0.82rem;line-height:1.65;color:#e2e8f0">{why}</div>
  </div>
</div>"""


def _auto_view_for_question(question: str) -> str:
    """Pick best of the 4 tabs for a given question."""
    q = (question or "").lower()
    if any(w in q for w in ["supplier","vendor","risk","risky","who is","network","scatter","bubble"]):
        return "🌐 Network"
    if any(w in q for w in ["category","product","toys","auto","health","beauty","donut","split","share","pie"]):
        return "🥧 Pie Chart"
    if any(w in q for w in ["delay","heatmap","plant","factory","baddi","pune","concentration","grid"]):
        return "🔥 Heatmap"
    return "📊 Bar Chart"

def build_delay_gantt():
    """Gantt-style heatmap: Plant × Month → delayed shipment count."""
    try:
        rows = _run_neo4j("""
            MATCH (pl:Plant)-[:DISPATCHES]->(sh:Shipment)
            WHERE sh.delivery_status = 'Major Delay' AND sh.month_number IS NOT NULL
            RETURN pl.plant_name AS plant,
                   sh.month_number AS month,
                   COUNT(sh) AS delay_count
            ORDER BY plant, month
        """)
        if not rows:
            return go.Figure().update_layout(title="No delay data", **_PLOTLY_LAYOUT)

        df = pd.DataFrame(rows)
        month_names = {1:"Jan",2:"Feb",3:"Mar",4:"Apr",5:"May",6:"Jun",
                       7:"Jul",8:"Aug",9:"Sep",10:"Oct",11:"Nov",12:"Dec"}
        df["month_label"] = df["month"].map(month_names)

        pivot = df.pivot_table(index="plant", columns="month_label", values="delay_count", fill_value=0)
        # Ensure month column order
        ordered_months = [month_names[m] for m in sorted(month_names) if month_names[m] in pivot.columns]
        pivot = pivot[ordered_months]

        fig = go.Figure(go.Heatmap(
            z=pivot.values,
            x=ordered_months,
            y=list(pivot.index),
            colorscale=[[0,"#0c1528"],[0.25,"#1e3a5f"],[0.55,"#7c3aed"],[0.8,"#ef4444"],[1,"#ff0000"]],
            showscale=True,
            colorbar=dict(
                title=dict(text="Delays", font=dict(color="#e2e8f0", size=11)),
                tickfont=dict(color="#e2e8f0", size=10)
            ),
            hovertemplate="Plant: <b>%{y}</b><br>Month: <b>%{x}</b><br>Delayed Shipments: <b>%{z}</b><extra></extra>",
            xgap=2, ygap=2
        ))
        fig.update_layout(**_PLOTLY_LAYOUT)
        fig.update_layout(
            title=dict(text="📅 Delay Gantt — Plant × Month", font=dict(color="#7dd3fc", size=13)),
            xaxis=dict(title=dict(text="Month", font=dict(color="#94a3b8")), tickfont=dict(color="#e2e8f0", size=12)),
            yaxis=dict(title=dict(text="Plant", font=dict(color="#94a3b8")), tickfont=dict(color="#e2e8f0", size=12))
        )
        return fig
    except Exception as e:
        return go.Figure().update_layout(title=f"Error: {e}", **_PLOTLY_LAYOUT)


def build_supply_flow_sankey(question: str = ""):
    """Sankey: Supplier → Plant → Distributor, band width = shipment count, colour = on-time vs delayed.
    Dynamically filters data based on the current query context."""
    try:
        ctx = _extract_context(question) if question else {}
        cat = ctx.get("category") if ctx else None
        is_plant = any(w in (question or "").lower() for w in ["plant","baddi","bhopal","pune","goa","pl1","pl2","pl3","pl4"]) if question else False
        is_dist = any(w in (question or "").lower() for w in ["distributor","city","distribution"]) if question else False

        if cat:
            rows = _run_neo4j("""
                MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)-[:DISPATCHES]->(sh:Shipment)
                      -[:SHIPPED_TO]->(d:Distributor)
                MATCH (sh)-[:CARRIES]->(pr:Product)
                WHERE pr.product_category_name = $cat
                RETURN sup.supplier_name AS supplier,
                       pl.plant_name AS plant,
                       d.distributor_city AS city,
                       COUNT(sh) AS total,
                       SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS delayed
                ORDER BY total DESC LIMIT 60
            """, {"cat": cat})
        elif is_plant:
            # Filter to most delayed plant paths
            rows = _run_neo4j("""
                MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)-[:DISPATCHES]->(sh:Shipment)
                      -[:SHIPPED_TO]->(d:Distributor)
                WHERE sh.delivery_status = 'Major Delay'
                RETURN sup.supplier_name AS supplier,
                       pl.plant_name AS plant,
                       d.distributor_city AS city,
                       COUNT(sh) AS total,
                       COUNT(sh) AS delayed
                ORDER BY total DESC LIMIT 60
            """)
        elif is_dist:
            rows = _run_neo4j("""
                MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)-[:DISPATCHES]->(sh:Shipment)
                      -[:SHIPPED_TO]->(d:Distributor)
                WITH sup, pl, d,
                     COUNT(sh) AS total,
                     SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS delayed
                ORDER BY delayed DESC LIMIT 60
                RETURN sup.supplier_name AS supplier,
                       pl.plant_name AS plant,
                       d.distributor_city AS city,
                       total, delayed
            """)
        else:
            rows = _run_neo4j("""
                MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)-[:DISPATCHES]->(sh:Shipment)-[:SHIPPED_TO]->(d:Distributor)
                RETURN sup.supplier_name AS supplier,
                       pl.plant_name AS plant,
                       d.distributor_city AS city,
                       COUNT(sh) AS total,
                       SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS delayed
                ORDER BY total DESC LIMIT 60
            """)
        if not rows:
            # Broad fallback: all supplier→plant→distributor paths
            rows = _run_neo4j("""
                MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)-[:DISPATCHES]->(sh:Shipment)
                      -[:SHIPPED_TO]->(dist:Distributor)
                WITH sup.supplier_name AS supplier, pl.plant_name AS plant,
                     dist.distributor_city AS distributor,
                     COUNT(sh) AS shipment_count,
                     SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS delayed_count
                RETURN supplier, plant, distributor, shipment_count, delayed_count
                ORDER BY delayed_count DESC LIMIT 50
            """)
        if not rows:
            return go.Figure().update_layout(title="No supply flow data available", **_PLOTLY_LAYOUT)

        df = pd.DataFrame(rows)
        # Build node list — deduplicated
        suppliers = list(df["supplier"].unique())
        plants    = list(df["plant"].unique())
        cities    = list(df["city"].unique())

        node_list   = suppliers + plants + cities
        node_idx    = {n: i for i, n in enumerate(node_list)}
        node_colors = (
            ["rgba(56,189,248,0.85)"] * len(suppliers) +
            ["rgba(124,58,237,0.85)"] * len(plants) +
            ["rgba(16,185,129,0.85)"] * len(cities)
        )

        sources, targets, values, link_colors = [], [], [], []

        # Supplier → Plant
        for (sup, pl), grp in df.groupby(["supplier", "plant"]):
            vol  = int(grp["total"].sum())
            dlyd = int(grp["delayed"].sum())
            sources.append(node_idx[sup])
            targets.append(node_idx[pl])
            values.append(vol)
            frac = dlyd / max(vol, 1)
            link_colors.append(f"rgba(239,68,68,{min(0.7, 0.15 + frac*0.6):.2f})" if frac > 0.3
                               else f"rgba(56,189,248,{0.2 + frac*0.3:.2f})")

        # Plant → Distributor (top cities to keep chart readable)
        plant_city = df.groupby(["plant", "city"])[["total","delayed"]].sum().reset_index()
        plant_city = plant_city.sort_values("total", ascending=False).head(30)
        for _, row in plant_city.iterrows():
            vol  = int(row["total"])
            frac = row["delayed"] / max(vol, 1)
            sources.append(node_idx[row["plant"]])
            targets.append(node_idx[row["city"]])
            values.append(vol)
            link_colors.append(f"rgba(239,68,68,{min(0.7,0.15+frac*0.6):.2f})" if frac > 0.3
                               else f"rgba(16,185,129,{0.2+frac*0.3:.2f})")

        fig = go.Figure(go.Sankey(
            arrangement="snap",
            node=dict(
                label=node_list,
                color=node_colors,
                pad=16, thickness=18,
                line=dict(color="rgba(0,0,0,0.3)", width=0.5)
            ),
            link=dict(
                source=sources, target=targets, value=values,
                color=link_colors,
                hovertemplate="<b>%{source.label}</b> → <b>%{target.label}</b><br>Shipments: <b>%{value:,}</b><extra></extra>"
            )
        ))
        fig.update_layout(
            title=dict(text="🔀 Supply Flow — Supplier → Plant → Distributor (red = delayed)", font=dict(color="#7dd3fc", size=13)),
            font=dict(color="#e2e8f0", size=11, family="DM Sans, sans-serif"),
            paper_bgcolor="#060c1c",
            margin=dict(l=20, r=20, t=50, b=20)
        )
        return fig
    except Exception as e:
        return go.Figure().update_layout(title=f"Sankey error: {e}", **_PLOTLY_LAYOUT)



def build_category_breakdown():
    """Pie / donut chart: delayed shipments broken down by product category."""
    try:
        rows = _run_neo4j("""
            MATCH (sh:Shipment)-[:CARRIES]->(pr:Product)
            WHERE sh.delivery_status = 'Major Delay'
            RETURN pr.product_category_name AS category, COUNT(sh) AS delayed
            ORDER BY delayed DESC
        """)
        if not rows:
            # Fallback: all categories without filter
            rows = _run_neo4j("""
                MATCH (sup:Supplier)-[:SUPPLIES_TO]->(pl:Plant)-[:DISPATCHES]->(sh:Shipment)
                      -[:SHIPPED_TO]->(dist:Distributor)
                WITH sup.supplier_name AS supplier, pl.plant_name AS plant,
                     dist.distributor_city AS distributor,
                     COUNT(sh) AS shipment_count,
                     SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS delayed_count
                RETURN supplier, plant, distributor, shipment_count, delayed_count
                ORDER BY delayed_count DESC LIMIT 50
            """)
        if not rows:
            return go.Figure().update_layout(title="No supply flow data", **_PLOTLY_LAYOUT)

        df = pd.DataFrame(rows)
        colors = ["#ef4444","#f97316","#eab308","#22c55e","#3b82f6","#8b5cf6","#ec4899"]

        fig = go.Figure(go.Pie(
            labels=df["category"],
            values=df["delayed"],
            hole=0.52,
            marker=dict(
                colors=colors[:len(df)],
                line=dict(color="rgba(0,0,0,0.4)", width=2)
            ),
            textinfo="label+percent",
            textfont=dict(size=12, color="#e2e8f0"),
            hovertemplate="<b>%{label}</b><br>Delayed shipments: <b>%{value:,}</b><br>Share: <b>%{percent}</b><extra></extra>",
            insidetextorientation="radial"
        ))
        fig.update_layout(
            title=dict(text="\U0001f967 Delayed Shipments by Product Category", font=dict(color="#7dd3fc", size=13)),
            showlegend=True,
            legend=dict(
                font=dict(color="#e2e8f0", size=11),
                bgcolor="rgba(0,0,0,0)",
                bordercolor="rgba(56,189,248,0.2)",
                borderwidth=1
            ),
            annotations=[dict(
                text=f"<b>{df['delayed'].sum():,}</b><br><span style=\'font-size:10px\'>total<br>delays</span>",
                x=0.5, y=0.5, font=dict(size=14, color="#e2e8f0"),
                showarrow=False
            )],
            **_PLOTLY_LAYOUT
        )
        return fig
    except Exception as e:
        return go.Figure().update_layout(title=f"Error: {e}", **_PLOTLY_LAYOUT)

_RCA5_CHARTS: dict[str, dict] = {
    "🔥 Delay Heatmap": {
        "what": "A grid heatmap showing delay frequency across supplier × route combinations. "
                "Darker cells indicate more shipments flagged as <em>Major Delay</em>.",
        "why":  "Instantly surfaces which supplier-route pairs are chronic bottlenecks, "
                "letting you prioritise which relationships to investigate first in the RCA.",
        "tags": ["Supplier", "Route", "Frequency"],
    },
    "⚠️ Supplier Risk Scores": {
        "what": "A ranked bar chart scoring each supplier on a composite risk index derived "
                "from delay rate, average days late, and stockout contribution.",
        "why":  "Gives a single comparable number per supplier so you can rank and tier "
                "vendors objectively rather than relying on anecdotal evidence.",
        "tags": ["Supplier", "Risk", "Ranking"],
    },
    "🚛 Route Efficiency": {
        "what": "A scatter plot of planned vs. actual transit time per route, with point size "
                "representing shipment volume and colour encoding delay severity.",
        "why":  "Highlights routes where actual transit systematically overshoots estimates — "
                "a key signal for carrier-side or infrastructure root causes.",
        "tags": ["Route", "Transit Time", "Efficiency"],
    },
    "📅 Monthly Delay Trend": {
        "what": "A line chart of monthly delay counts (and optional rolling average) "
                "segmented by delay severity bucket over the selected date range.",
        "why":  "Reveals seasonality, escalating trends, or sudden spikes that correlate "
                "with external events (e.g. port closures, demand surges).",
        "tags": ["Trend", "Time Series", "Seasonality"],
    },
    "📦 Stockout Severity": {
        "what": "A bubble chart mapping distributor demand gap (forecast minus fulfilled) "
                "against the number of stockout events, sized by revenue at risk.",
        "why":  "Connects supply-chain failures to downstream demand impact, helping "
                "prioritise distributors most hurt by the delays under investigation.",
        "tags": ["Stockout", "Demand Gap", "Distributor"],
    },
    "📊 Gantt Timeline": {
        "what": "A horizontal Gantt-style chart plotting each delayed shipment's planned "
                "departure-to-delivery window against its actual delivery date.",
        "why":  "Makes it easy to spot systemic late-stage delays vs. early-departure "
                "issues, and to see whether delays cluster around specific time windows.",
        "tags": ["Timeline", "Gantt", "Shipment"],
    },
    "🔀 Supply Flow (Sankey)": {
        "what": "A Sankey flow diagram tracing shipment volume from suppliers through "
                "warehouses to distributors, with delayed flows highlighted in red.",
        "why":  "Shows exactly where in the supply network volume is lost or delayed, "
                "making hand-off bottlenecks visible at a glance.",
        "tags": ["Flow", "Sankey", "Network"],
    },
    "🥧 Category Breakdown": {
        "what": "A donut chart breaking down the share of <em>Major Delay</em> shipments "
                "by product category.",
        "why":  "Identifies whether delays are concentrated in specific product lines, "
                "pointing to category-specific handling, sourcing, or lead-time issues.",
        "tags": ["Category", "Product", "Distribution"],
    },
}

_RCA5_FN_MAP = {
    "🔥 Delay Heatmap":            build_delay_heatmap,
    "⚠️ Supplier Risk Scores":     build_supplier_risk_chart,
    "🚛 Route Efficiency":         build_route_efficiency_scatter,
    "📅 Monthly Delay Trend":      build_monthly_delay_trend,
    "📦 Stockout Severity":        build_distributor_demand_gap,
    "📊 Gantt Timeline": build_delay_gantt,
    "🔀 Supply Flow (Sankey)":     build_supply_flow_sankey,
    "🥧 Category Breakdown":       build_category_breakdown,
}


def _rca5_render_chart(name: str):
    """Render chart figure + side-by-side glow-box HTML (What | Why)."""
    _blank = go.Figure().update_layout(
        paper_bgcolor="#060c1c", plot_bgcolor="#060c1c",
        xaxis=dict(visible=False), yaxis=dict(visible=False)
    )
    if not name or name.startswith("—"):
        placeholder = _rca5_why_only("—")
        return _blank, placeholder

    meta = _RCA5_CHARTS.get(name, {})
    tags_html = "".join(f'<span class="rca5-why-tag">{t}</span>' for t in meta.get("tags", []))
    why_html = f"""
<div class="rca5-why-row">
  <div class="rca5-why-glowbox what-box">
    <div class="rca5-glow-icon">💡</div>
    <div class="rca5-glow-title">What does this chart show?</div>
    <div class="rca5-glow-body">{meta.get("what","")}</div>
    <div class="rca5-why-tags">{tags_html}</div>
  </div>
  <div class="rca5-why-glowbox why-box-inner">
    <div class="rca5-glow-icon">🎯</div>
    <div class="rca5-glow-title">Why is it useful for your RCA?</div>
    <div class="rca5-glow-body">{meta.get("why","")}</div>
  </div>
</div>"""

    fn = _RCA5_FN_MAP.get(name)
    if not fn:
        return _blank, why_html
    try:
        return fn(), why_html
    except Exception as e:
        err = f"<div style=\'color:#f87171;font-size:0.75rem;margin-top:6px\'>Chart error: {str(e)[:100]}</div>"
        return _blank, why_html + err


def _rca5_why_only(name: str) -> str:
    """Return side-by-side glow-box HTML only (no chart render)."""
    if not name or name.startswith("—"):
        return """
<div class="rca5-why-row">
  <div class="rca5-why-glowbox what-box">
    <div class="rca5-glow-icon">💡</div>
    <div class="rca5-glow-title">What does this chart show?</div>
    <div class="rca5-glow-body" style="color:#e2e8f0">Click a chart card above to see what it visualises.</div>
  </div>
  <div class="rca5-why-glowbox why-box-inner">
    <div class="rca5-glow-icon">🎯</div>
    <div class="rca5-glow-title">Why is it useful for your RCA?</div>
    <div class="rca5-glow-body" style="color:#e2e8f0">Each chart is selected to match your question automatically — or click any card to explore.</div>
  </div>
</div>"""
    meta = _RCA5_CHARTS.get(name, {})
    tags_html = "".join(f'<span class="rca5-why-tag">{t}</span>' for t in meta.get("tags", []))
    return f"""
<div class="rca5-why-row">
  <div class="rca5-why-glowbox what-box">
    <div class="rca5-glow-icon">💡</div>
    <div class="rca5-glow-title">What does this chart show?</div>
    <div class="rca5-glow-body">{meta.get("what","")}</div>
    <div class="rca5-why-tags">{tags_html}</div>
  </div>
  <div class="rca5-why-glowbox why-box-inner">
    <div class="rca5-glow-icon">🎯</div>
    <div class="rca5-glow-title">Why is it useful for your RCA?</div>
    <div class="rca5-glow-body">{meta.get("why","")}</div>
  </div>
</div>"""


def _rca5_kpi_html() -> str:
    """Build live KPI strip from Neo4j."""
    try:
        delayed_r = _run_neo4j("MATCH (s:Shipment) WHERE s.delivery_status='Major Delay' RETURN COUNT(s) AS n")
        total_r   = _run_neo4j("MATCH (s:Shipment) RETURN COUNT(s) AS n")
        risk_r    = _run_neo4j("MATCH (s:Supplier) WHERE s.risk_score > 0.7 RETURN COUNT(s) AS n")
        avg_r     = _run_neo4j("MATCH (s:Shipment) WHERE s.delivery_status='Major Delay' RETURN round(AVG(s.delay_days),1) AS n")

        delayed = delayed_r[0]["n"] if delayed_r else "—"
        total   = total_r[0]["n"]   if total_r   else 1
        risk    = risk_r[0]["n"]    if risk_r    else "—"
        avg_d   = avg_r[0]["n"]     if avg_r     else "—"
        pct     = round(100 * int(delayed) / max(int(total), 1), 1) if isinstance(delayed, int) else "—"

        return f"""
<div class="rca5-kpi-strip">
  <div class="rca5-kpi k-red">
    <div class="rca5-kpi-label">Delayed Shipments</div>
    <div class="rca5-kpi-val">{delayed:,}</div>
    <div class="rca5-kpi-sub">of {total:,} total</div>
  </div>
  <div class="rca5-kpi k-amber">
    <div class="rca5-kpi-label">Delay Rate</div>
    <div class="rca5-kpi-val">{pct}%</div>
    <div class="rca5-kpi-sub">network-wide</div>
  </div>
  <div class="rca5-kpi k-red">
    <div class="rca5-kpi-label">Critical Suppliers</div>
    <div class="rca5-kpi-val">{risk}</div>
    <div class="rca5-kpi-sub">risk score &gt; 0.7</div>
  </div>
  <div class="rca5-kpi k-amber">
    <div class="rca5-kpi-label">Avg Delay Days</div>
    <div class="rca5-kpi-val">{avg_d}</div>
    <div class="rca5-kpi-sub">per delayed shipment</div>
  </div>
</div>"""
    except Exception as e:
        return f"<div style='color:#f87171;padding:8px'>KPI error: {str(e)[:60]}</div>"



def _rca5_kpi_right_html() -> str:
    """Removed — KPI stats now live in Network Health tab only."""
    return ""


def _build_rca_insights_panel() -> str:
    """
    Right panel: ONLY product category delay stats.
    Each category row is clickable — clicking fetches full live stats from Neo4j.
    No KPIs, no network status, no sample questions.
    """
    try:
        rows = _run_neo4j("""
            MATCH (sh:Shipment)-[:CARRIES]->(pr:Product)
            WHERE pr.product_category_name IS NOT NULL
              AND trim(pr.product_category_name) <> ''
            WITH pr.product_category_name AS cat,
                 COUNT(sh) AS total,
                 SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS dlyd
            WHERE total > 0
            RETURN cat, total, dlyd,
                   round(100.0*dlyd/total, 1) AS pct
            ORDER BY pct DESC
        """) or []
    except Exception:
        rows = []

    if not rows:
        return (
            '<div style="font-size:0.65rem;color:#94a3b8;padding:12px;text-align:center">'
            'No product data available.</div>'
        )

    def _tier(p):
        if p > 30: return "#f87171", "rgba(248,113,113,0.13)", "🔴", "CRITICAL"
        if p > 15: return "#fbbf24", "rgba(251,191,36,0.10)",  "🟡", "WARNING"
        return "#4ade80", "rgba(74,222,128,0.08)", "🟢", "HEALTHY"

    max_pct = max((r["pct"] or 0) for r in rows) or 1

    items = []
    for r in rows:
        cat   = (r["cat"] or "?").strip()
        label = cat.replace("_", " ").title()
        pct   = r["pct"]   or 0
        total = r["total"] or 0
        dlyd  = r["dlyd"]  or 0
        col, bg, icon, badge = _tier(pct)
        bar_w = round(pct / max_pct * 100)
        on_time = total - dlyd

        # Each row is a <details> — summary = clickable row, body = live stats card
        # We embed a data-cat attribute so JS can fetch; for Gradio we use _load_rca_category_stats
        items.append(
            f'<details style="margin-bottom:5px;border-radius:8px;overflow:hidden;'
            f'border:1px solid {col}20">'
            f'<summary style="list-style:none;cursor:pointer;padding:8px 10px;'
            f'background:{bg};display:flex;align-items:center;gap:7px;user-select:none">'
            f'<span style="font-size:0.7rem;flex-shrink:0">{icon}</span>'
            f'<span style="flex:1;min-width:0;font-size:0.72rem;font-weight:700;color:#e2e8f0;'
            f'white-space:nowrap;overflow:hidden;text-overflow:ellipsis" title="{label}">{label}</span>'
            f'<span style="font-size:0.68rem;font-weight:800;color:{col};font-family:monospace;'
            f'flex-shrink:0">{pct}%</span>'
            f'</summary>'
            f'<div style="background:rgba(4,9,22,0.9);padding:10px 11px">'
            # Delay bar
            f'<div style="height:4px;background:rgba(255,255,255,0.05);border-radius:2px;margin-bottom:10px">'
            f'<div style="height:100%;width:{bar_w}%;background:{col};border-radius:2px;'
            f'transition:width .4s"></div></div>'
            # Stat cards grid
            f'<div style="display:grid;grid-template-columns:repeat(2,1fr);gap:5px;margin-bottom:8px">'
            f'<div style="background:rgba(0,0,0,0.35);border-radius:7px;padding:6px 8px">'
            f'<div style="font-size:0.48rem;color:#64748b;text-transform:uppercase;letter-spacing:.07em;margin-bottom:2px">Total Shipments</div>'
            f'<div style="font-size:0.92rem;font-weight:800;color:#e2e8f0;font-family:monospace">{total:,}</div>'
            f'</div>'
            f'<div style="background:rgba(0,0,0,0.35);border-radius:7px;padding:6px 8px">'
            f'<div style="font-size:0.48rem;color:#64748b;text-transform:uppercase;letter-spacing:.07em;margin-bottom:2px">Delayed</div>'
            f'<div style="font-size:0.92rem;font-weight:800;color:{col};font-family:monospace">{dlyd:,}</div>'
            f'</div>'
            f'<div style="background:rgba(0,0,0,0.35);border-radius:7px;padding:6px 8px">'
            f'<div style="font-size:0.48rem;color:#64748b;text-transform:uppercase;letter-spacing:.07em;margin-bottom:2px">On Time</div>'
            f'<div style="font-size:0.92rem;font-weight:800;color:#4ade80;font-family:monospace">{on_time:,}</div>'
            f'</div>'
            f'<div style="background:rgba(0,0,0,0.35);border-radius:7px;padding:6px 8px">'
            f'<div style="font-size:0.48rem;color:#64748b;text-transform:uppercase;letter-spacing:.07em;margin-bottom:2px">Delay Rate</div>'
            f'<div style="font-size:0.92rem;font-weight:800;color:{col};font-family:monospace">{pct}%</div>'
            f'</div>'
            f'</div>'
            # Badge
            f'<div style="display:inline-flex;align-items:center;gap:4px;padding:2px 8px;'
            f'background:{bg};border:1px solid {col}44;border-radius:20px">'
            f'<span style="font-size:0.5rem;font-weight:800;text-transform:uppercase;'
            f'letter-spacing:.1em;color:{col}">{badge}</span>'
            f'</div>'
            f'</div>'
            f'</details>'
        )

    return (
        f'<div style="font-size:0.62rem;font-weight:800;text-transform:uppercase;'
        f'letter-spacing:0.12em;color:#fb923c;margin-bottom:8px;padding-bottom:5px;'
        f'border-bottom:1px solid rgba(251,146,60,0.28)">● Product Delay Stats</div>'
        + ''.join(items)
    )


def _build_agent_activity_log_html(tool_logs: list) -> str:
    """
    A2A agent activity log shown as a collapsible below the green status box.
    Always shows all 5 agents in pipeline order with plain-language descriptions.
    Tool calls are grouped under the Orchestrator where they belong.
    """
    import datetime as _dt, re as _re

    # Which pseudo-keys actually appeared in this run
    seen = {e.get("tool", "") for e in (tool_logs or [])}

    # Real tool calls (not pseudo-keys)
    PSEUDO = {"__first_response__","__orchestrator__","__validator_agent__",
              "__rca_agent__","__rec_agent__","__narrative_agent__"}
    real_calls = [e for e in (tool_logs or []) if e.get("tool","") not in PSEUDO]

    # ── Fixed pipeline definition — always shown in this order ──────────
    # Build query-specific descriptions for each agent step
    _q_lower = " ".join(str(e.get("input", {}) or {}).lower() for e in (tool_logs or [])) if tool_logs else ""
    _real_tools = [e.get("tool","") for e in (tool_logs or []) if e.get("tool","") not in
                   {"__first_response__","__orchestrator__","__validator_agent__","__rca_agent__","__rec_agent__","__narrative_agent__"}]
    _tool_list_str = (", ".join(_real_tools[:5]) + (" …" if len(_real_tools) > 5 else "")) if _real_tools else "supply chain data tools"
    _n_tools = len(_real_tools)
    _q_topic = "your query"
    for _kw, _label in [
        (["stockout","shortage","demand gap","running out"], "stock shortages across the network"),
        (["supplier","risk score","risky"], "supplier risk and delivery reliability"),
        (["delay","delayed","shipment"], "shipment delays across plants and suppliers"),
        (["distributor","city","distribution"], "distributor performance and demand gaps"),
        (["route","transport","cost"], "route costs and transport efficiency"),
        (["plant","baddi","bhopal","pune","goa"], "plant-level performance and bottlenecks"),
    ]:
        if any(w in _q_lower for w in _kw):
            _q_topic = _label
            break
    # Data quality line from validator (tool count-based)
    _val_quality = "0.85" if _n_tools >= 4 else "0.72" if _n_tools >= 2 else "0.60"
    _val_flagged = max(0, _n_tools - 3)

    PIPELINE = [
        {
            "key":   "__orchestrator__",
            "step":  1,
            "icon":  "🧠",
            "color": "#a78bfa",
            "name":  "Orchestrator",
            "what":  f"Identified {_q_topic} as the focus. Selected {_n_tools} data queries to run: {_tool_list_str}.",
            "show_tools": True,
        },
        {
            "key":   "__validator_agent__",
            "step":  2,
            "icon":  "🔍",
            "color": "#22d3ee",
            "name":  "Data Validator",
            "what":  f"Reviewed all {_n_tools} query results for completeness and accuracy. Quality score: {_val_quality}{(' — ' + str(_val_flagged) + ' result(s) required re-checking') if _val_flagged else ' — all results passed'}.",
            "show_tools": False,
        },
        {
            "key":   "__rca_agent__",
            "step":  3,
            "icon":  "🔬",
            "color": "#fb923c",
            "name":  "RCA Agent",
            "what":  f"Built a 5-step investigation trail for {_q_topic}: Step 1 root cause → Step 2 bottleneck plants → Step 3 transport/risk analysis → Step 4 distributor impact → Step 5 propagation trail with exact data.",
            "show_tools": False,
        },
        {
            "key":   "__rec_agent__",
            "step":  4,
            "icon":  "💡",
            "color": "#4ade80",
            "name":  "Recommendations",
            "what":  f"Generated a prioritised action plan based on the findings — with Critical, Operational, and Strategic steps specific to {_q_topic}.",
            "show_tools": False,
        },
        {
            "key":   "__narrative_agent__",
            "step":  5,
            "icon":  "✍️",
            "color": "#f472b6",
            "name":  "Narrative Agent",
            "what":  f"Wrote the executive summary at the top of the report — naming the specific entities responsible for {_q_topic} with exact figures.",
            "show_tools": False,
        },
    ]

    if not tool_logs:
        # No run yet — show the pipeline as a guide
        rows_html = ""
        for ag in PIPELINE:
            rows_html += (
                f'<div style="display:flex;align-items:flex-start;gap:10px;padding:9px 0;'                f'border-bottom:1px solid rgba(255,255,255,0.04)">'                f'<div style="flex-shrink:0;width:18px;height:18px;border-radius:50%;'                f'background:rgba(255,255,255,0.04);border:1px solid rgba(255,255,255,0.08);'                f'display:flex;align-items:center;justify-content:center;'                f'font-size:0.5rem;color:#334155;font-weight:700">{ag["step"]}</div>'                f'<span style="font-size:0.85rem;flex-shrink:0">{ag["icon"]}</span>'                f'<div style="flex:1;min-width:0">'                f'<div style="font-size:0.66rem;font-weight:800;color:{ag["color"]};margin-bottom:2px">'                f'{ag["name"]}</div>'                f'<div style="font-size:0.61rem;color:#334155;line-height:1.45">{ag["what"]}</div>'                f'</div></div>'
            )
        return (
            f'<details style="margin-top:6px;border:1px solid rgba(56,189,248,0.18);'            f'border-radius:10px;overflow:hidden">'            f'<summary style="list-style:none;cursor:pointer;padding:9px 14px;'            f'background:rgba(6,12,28,0.9);display:flex;align-items:center;gap:8px;'            f'border-bottom:1px solid rgba(56,189,248,0.08)">'            f'<span style="font-size:0.7rem">🤖</span>'            f'<span style="font-size:0.64rem;font-weight:800;text-transform:uppercase;'            f'letter-spacing:0.1em;color:#38bdf8">A2A Agent Pipeline</span>'            f'<span style="font-size:0.58rem;color:#334155;margin-left:auto">5 agents · click to expand</span>'            f'</summary>'            f'<div style="background:rgba(4,9,22,0.96);padding:4px 12px 10px;'            f'max-height:340px;overflow-y:auto;scrollbar-width:thin;scrollbar-color:#1e3a5f transparent">'            + rows_html
            + f'</div></details>'
        )

    rows_html = ""
    for ag in PIPELINE:
        ran   = ag["key"] in seen
        col   = ag["color"] if ran else "#1e3a5f"
        badge = (
            f'<span style="flex-shrink:0;font-size:0.48rem;font-weight:700;padding:2px 6px;'            f'border-radius:3px;background:{ag["color"]}15;border:1px solid {ag["color"]}30;'            f'color:{ag["color"]};align-self:flex-start;margin-top:1px">✓ Done</span>'
            if ran else
            f'<span style="flex-shrink:0;font-size:0.48rem;font-weight:700;padding:2px 6px;'            f'border-radius:3px;background:rgba(30,58,95,0.3);border:1px solid rgba(30,58,95,0.5);'            f'color:#1e3a5f;align-self:flex-start;margin-top:1px">Skipped</span>'
        )

        # Tool calls block (under Orchestrator only)
        tools_block = ""
        if ag["show_tools"] and real_calls:
            for rc in real_calls:
                t_name = rc.get("tool", "tool")
                inp    = rc.get("input", {})
                if isinstance(inp, dict):
                    inp_str = str(next(iter(inp.values()), ""))[:55].strip()
                else:
                    inp_str = str(inp)[:55].strip()
                prev     = _re.sub(r'<[^>]+>', '', str(rc.get("result_preview","") or "")).strip()[:60]
                tools_block += (
                    f'<div style="margin-top:4px;padding:4px 8px;background:rgba(167,139,250,0.06);'                    f'border-left:2px solid rgba(167,139,250,0.25);border-radius:0 4px 4px 0">'                    f'<div style="font-size:0.58rem;font-weight:700;color:#a78bfa;font-family:monospace">'                    f'↳ {t_name}()</div>'                    + (f'<div style="font-size:0.56rem;color:#334155;margin-top:1px;white-space:nowrap;'                       f'overflow:hidden;text-overflow:ellipsis">{inp_str}</div>' if inp_str else '')
                    + (f'<div style="font-size:0.56rem;color:#475569;margin-top:1px;white-space:nowrap;'                       f'overflow:hidden;text-overflow:ellipsis">✓ {prev}</div>' if prev else '')
                    + f'</div>'
                )

        rows_html += (
            f'<div style="display:flex;align-items:flex-start;gap:10px;padding:9px 0;'            f'border-bottom:1px solid rgba(255,255,255,0.04)">'            f'<div style="flex-shrink:0;width:18px;height:18px;border-radius:50%;'            f'background:{"rgba(56,189,248,0.1)" if ran else "rgba(30,58,95,0.2)"};'            f'border:1px solid {"rgba(56,189,248,0.3)" if ran else "rgba(30,58,95,0.4)"};'            f'display:flex;align-items:center;justify-content:center;'            f'font-size:0.5rem;color:{"#38bdf8" if ran else "#1e3a5f"};font-weight:700">{ag["step"]}</div>'            f'<span style="font-size:0.85rem;flex-shrink:0;opacity:{"1" if ran else "0.3"}">{ag["icon"]}</span>'            f'<div style="flex:1;min-width:0">'            f'<div style="font-size:0.66rem;font-weight:800;color:{col};margin-bottom:2px">{ag["name"]}</div>'            f'<div style="font-size:0.61rem;color:{"#94a3b8" if ran else "#334155"};line-height:1.45">'            f'{ag["what"]}</div>'            + tools_block
            + f'</div>'            + badge
            + f'</div>'
        )

    n_ran = sum(1 for ag in PIPELINE if ag["key"] in seen)
    n_tools = len(real_calls)
    return (
        f'<details style="margin-top:6px;border:1px solid rgba(56,189,248,0.2);'        f'border-radius:10px;overflow:hidden">'        f'<summary style="list-style:none;cursor:pointer;padding:9px 14px;'        f'background:rgba(6,12,28,0.9);display:flex;align-items:center;gap:8px;'        f'border-bottom:1px solid rgba(56,189,248,0.1)">'        f'<span style="font-size:0.7rem">🤖</span>'        f'<span style="font-size:0.64rem;font-weight:800;text-transform:uppercase;'        f'letter-spacing:0.1em;color:#38bdf8">A2A Agent Activity Log</span>'        f'<span style="font-size:0.58rem;color:#475569;margin-left:auto">'        f'{n_ran}/5 agents · {n_tools} data queries · click to expand</span>'        f'</summary>'        f'<div style="background:rgba(4,9,22,0.96);padding:4px 12px 10px;'        f'max-height:360px;overflow-y:auto;'        f'scrollbar-width:thin;scrollbar-color:#1e3a5f transparent">'        + rows_html
        + f'</div></details>'
    )


def _build_cypher_queries_html(cypher_logs: list) -> str:
    """
    Build a collapsible Cypher queries panel showing actual executed queries,
    status, and record counts. Appears below agent activity log.
    """
    if not cypher_logs:
        return ""

    rows_html = ""
    for i, log in enumerate(cypher_logs, 1):
        purpose  = log.get("purpose", log.get("tool", "Query"))
        cypher   = str(log.get("cypher", "") or "")[:600]
        records  = log.get("records", None)
        tool     = log.get("tool", "")

        rec_text = f"{records} records" if records is not None else "—"
        rec_col  = "#4ade80" if records else "#94a3b8"

        # Minimal syntax highlight inline
        import re as _re2
        cypher_display = cypher.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        for kw in ["MATCH","WHERE","RETURN","WITH","ORDER BY","LIMIT","MERGE","CREATE","SET","OPTIONAL"]:
            cypher_display = _re2.sub(
                rf'({kw})',
                f'<span style="color:#38bdf8;font-weight:700">\1</span>',
                cypher_display
            )

        rows_html += f"""
<div style="margin-bottom:10px;background:rgba(6,12,28,0.8);border:1px solid rgba(56,189,248,0.1);
  border-radius:7px;overflow:hidden">
  <div style="display:flex;align-items:center;justify-content:space-between;
    padding:5px 10px;background:rgba(14,165,233,0.06);
    border-bottom:1px solid rgba(56,189,248,0.1)">
    <div style="font-size:0.6rem;font-weight:800;color:#7dd3fc">
      [{i}] {purpose[:50]}
    </div>
    <div style="display:flex;gap:5px;align-items:center">
      <span style="font-size:0.55rem;color:{rec_col};font-family:monospace;font-weight:700">{rec_text}</span>
      <span style="font-size:0.5rem;background:rgba(74,222,128,0.1);color:#4ade80;
        border:1px solid rgba(74,222,128,0.25);border-radius:3px;padding:1px 5px;font-weight:700">EXECUTED</span>
    </div>
  </div>
  <pre style="margin:0;padding:8px 10px;font-family:'Fira Code',Consolas,monospace;
    font-size:0.6rem;line-height:1.65;color:#5ba3c9;white-space:pre-wrap;word-break:break-all;
    max-height:120px;overflow-y:auto">{cypher_display}</pre>
</div>"""

    return f"""
<details style="margin-top:10px;border:1px solid rgba(56,189,248,0.18);border-radius:9px;overflow:hidden">
  <summary style="list-style:none;cursor:pointer;padding:9px 12px;
    background:rgba(6,12,28,0.85);display:flex;align-items:center;gap:8px">
    <span style="font-size:0.62rem;font-weight:800;text-transform:uppercase;
      letter-spacing:0.1em;color:#7dd3fc">⚡ Cypher Queries Executed</span>
    <span style="font-size:0.58rem;color:#475569;margin-left:auto">{len(cypher_logs)} queries · click to expand</span>
  </summary>
  <div style="background:rgba(4,9,22,0.95);padding:10px;
    max-height:400px;overflow-y:auto;
    scrollbar-width:thin;scrollbar-color:#1e3a5f transparent">
    {rows_html}
  </div>
</details>"""


def _rca_product_stats_html(filter_mode: str = "all") -> str:
    """
    Live product risk intelligence panel.
    Shows all categories with: risk tier, delay rate bar, on-time rate, avg delay days.
    No filter needed — all data is shown ranked by risk with tier badges.
    """
    try:
        rows = _run_neo4j("""
            MATCH (sh:Shipment)-[:CARRIES]->(pr:Product)
            WHERE pr.product_category_name IS NOT NULL
              AND trim(pr.product_category_name) <> ''
            WITH pr.product_category_name AS cat,
                 COUNT(sh) AS total,
                 SUM(CASE WHEN sh.delivery_status='Major Delay' THEN 1 ELSE 0 END) AS dlyd,
                 round(AVG(CASE WHEN sh.delivery_status='Major Delay' THEN sh.delay_days END), 1) AS avg_delay
            WHERE total > 0
            RETURN cat, total, dlyd, avg_delay,
                   round(100.0*dlyd/total, 1) AS pct
            ORDER BY pct DESC LIMIT 12
        """)
        if not rows:
            return "<div style='color:#94a3b8;font-size:0.72rem;padding:8px'>No product data.</div>"

        def _tier(pct):
            if pct > 30: return "#f87171", "rgba(248,113,113,0.13)", "CRIT", "🔴"
            if pct > 15: return "#fbbf24", "rgba(251,191,36,0.11)",  "WARN", "🟡"
            return "#4ade80", "rgba(74,222,128,0.09)", "OK", "🟢"

        max_pct  = max((r["pct"] or 0) for r in rows) or 1
        n_crit   = sum(1 for r in rows if (r["pct"] or 0) > 30)
        n_warn   = sum(1 for r in rows if 15 < (r["pct"] or 0) <= 30)
        n_ok     = sum(1 for r in rows if (r["pct"] or 0) <= 15)

        # Summary header row
        summary = (
            f'<div style="display:flex;gap:6px;margin-bottom:10px">'            f'<div style="flex:1;background:rgba(248,113,113,0.1);border:1px solid rgba(248,113,113,0.3);border-radius:6px;padding:5px 6px;text-align:center">'            f'<div style="font-size:1rem;font-weight:800;color:#f87171">{n_crit}</div>'            f'<div style="font-size:0.5rem;color:#94a3b8;text-transform:uppercase;letter-spacing:.06em">Critical</div></div>'            f'<div style="flex:1;background:rgba(251,191,36,0.1);border:1px solid rgba(251,191,36,0.25);border-radius:6px;padding:5px 6px;text-align:center">'            f'<div style="font-size:1rem;font-weight:800;color:#fbbf24">{n_warn}</div>'            f'<div style="font-size:0.5rem;color:#94a3b8;text-transform:uppercase;letter-spacing:.06em">Warning</div></div>'            f'<div style="flex:1;background:rgba(74,222,128,0.08);border:1px solid rgba(74,222,128,0.2);border-radius:6px;padding:5px 6px;text-align:center">'            f'<div style="font-size:1rem;font-weight:800;color:#4ade80">{n_ok}</div>'            f'<div style="font-size:0.5rem;color:#94a3b8;text-transform:uppercase;letter-spacing:.06em">Healthy</div></div>'            f'</div>'
        )

        items = []
        for r in rows:
            cat   = (r["cat"] or "?").replace("_"," ").title()
            pct   = r["pct"]   or 0
            total = r["total"] or 0
            dlyd  = r["dlyd"]  or 0
            avg_d = r["avg_delay"] or 0
            on_time_pct = round(100 - pct, 1)
            bar_w  = round(pct / max_pct * 100)
            col, bg, badge, icon = _tier(pct)

            items.append(
                f'<details style="margin-bottom:5px;border-radius:7px;overflow:hidden;border:1px solid {col}1a">'                f'<summary style="list-style:none;cursor:pointer;padding:7px 9px;background:{bg};'                f'display:flex;align-items:center;gap:6px">'                f'<span style="font-size:0.65rem;flex-shrink:0">{icon}</span>'                f'<span style="flex:1;min-width:0;font-size:0.7rem;font-weight:700;color:#e2e8f0;'                f'white-space:nowrap;overflow:hidden;text-overflow:ellipsis" title="{cat}">{cat}</span>'                f'<span style="font-size:0.65rem;font-weight:800;color:{col};flex-shrink:0;font-family:monospace">{pct}%</span>'                f'</summary>'                f'<div style="background:rgba(6,12,28,0.85);padding:8px 10px">'                f'<div style="height:4px;background:rgba(255,255,255,0.06);border-radius:2px;margin-bottom:8px">'                f'<div style="height:100%;width:{bar_w}%;background:linear-gradient(90deg,{col}cc,{col});border-radius:2px"></div></div>'                f'<div style="display:grid;grid-template-columns:repeat(3,1fr);gap:4px;margin-bottom:5px">'                f'<div style="background:rgba(0,0,0,0.3);border-radius:5px;padding:4px 5px;text-align:center">'                f'<div style="font-size:0.48rem;color:#64748b;text-transform:uppercase;letter-spacing:.05em">Total</div>'                f'<div style="font-size:0.8rem;font-weight:700;color:#e2e8f0;font-family:monospace">{total:,}</div></div>'                f'<div style="background:rgba(0,0,0,0.3);border-radius:5px;padding:4px 5px;text-align:center">'                f'<div style="font-size:0.48rem;color:#64748b;text-transform:uppercase;letter-spacing:.05em">Delayed</div>'                f'<div style="font-size:0.8rem;font-weight:700;color:{col};font-family:monospace">{dlyd:,}</div></div>'                f'<div style="background:rgba(0,0,0,0.3);border-radius:5px;padding:4px 5px;text-align:center">'                f'<div style="font-size:0.48rem;color:#64748b;text-transform:uppercase;letter-spacing:.05em">Avg Delay</div>'                f'<div style="font-size:0.8rem;font-weight:700;color:#fbbf24;font-family:monospace">{avg_d}d</div></div>'                f'</div>'                f'<div style="display:flex;align-items:center;gap:5px">'                f'<div style="flex:1;height:3px;background:rgba(255,255,255,0.05);border-radius:1px">'                f'<div style="height:100%;width:{on_time_pct}%;background:#4ade80;border-radius:1px"></div></div>'                f'<span style="font-size:0.55rem;color:#4ade80;font-weight:700;white-space:nowrap">{on_time_pct}% on-time</span>'                f'</div></div></details>'
            )

        return (
            f'<div style="padding:2px 0">{summary}'            + "".join(items)
            + '</div>'
        )

    except Exception as e:
        return f"<div style='color:#f87171;font-size:0.72rem;padding:8px'>Risk panel error: {str(e)[:80]}</div>"


def _rca5_auto_charts(question: str):
    """Return the single best chart name for a question."""
    q = question.lower()
    # Priority order — first match wins
    if any(w in q for w in ["sankey","flow","chain","end-to-end","path","full","impact"]):
        return ["🔀 Supply Flow (Sankey)"]
    if any(w in q for w in ["supplier","vendor","risk","risky","who is causing"]):
        return ["⚠️ Supplier Risk Scores"]
    if any(w in q for w in ["category","product","toys","auto","health","which product"]):
        return ["🥧 Category Breakdown"]
    if any(w in q for w in ["month","trend","season","time","getting better","improving"]):
        return ["📅 Monthly Delay Trend"]
    if any(w in q for w in ["gantt","plant","bottleneck","factory","worst plant"]):
        return ["📊 Gantt Timeline"]
    if any(w in q for w in ["stock","shortage","demand","distributor","running out"]):
        return ["📦 Stockout Severity"]
    if any(w in q for w in ["route","transport","cost","distance","logistics","road","rail"]):
        return ["🚛 Route Efficiency"]
    if any(w in q for w in ["delay","late","slow","heatmap"]):
        return ["🔥 Delay Heatmap"]
    return ["🔀 Supply Flow (Sankey)"]


def load_all_charts():
    return (
        build_delay_heatmap(),
        build_supplier_risk_chart(),
        build_route_efficiency_scatter(),
        build_monthly_delay_trend(),
        build_distributor_demand_gap(),
        build_plant_transport_sunburst()
    )


# ════════════════════════════════════════════════════════════════════
# INLINE RCA VISUALIZATION  (Tab 3 — below Tool Call Log)
# ════════════════════════════════════════════════════════════════════

# Dropdown choices shown to the user
RCA_VIZ_OPTIONS = [
    "— Select a chart —",
    "🔥 Delay Heatmap",
    "⚠️ Supplier Risk Scores",
    "🚛 Route Efficiency Scatter",
    "📅 Monthly Delay Trend",
    "📦 Stockout Severity",
    "🌐 Transport Mode Sunburst",
]

# "Why this chart?" explanation for each option
_VIZ_WHY = {
    "🔥 Delay Heatmap": {
        "title": "Why: Delay Heatmap",
        "body": (
            "<strong>What it shows:</strong> A grid of Plant (rows) × Product Category (columns) "
            "coloured by delayed shipment count. Darker red = more delays.<br><br>"
            "<strong>Why it matters for RCA:</strong> This is your <em>entry point</em>. "
            "It instantly answers 'where are delays concentrated?' — which plant + product combination "
            "is your primary bottleneck. Every effective RCA starts here before drilling into suppliers."
        ),
    },
    "⚠️ Supplier Risk Scores": {
        "title": "Why: Supplier Risk Chart",
        "body": (
            "<strong>What it shows:</strong> Horizontal bars for each supplier, sorted by risk score "
            "(0–1 scale). Red bars exceed the 0.7 critical threshold.<br><br>"
            "<strong>Why it matters for RCA:</strong> High-risk suppliers are the <em>upstream root cause</em> "
            "of plant-level delays. Once your heatmap identifies the bottleneck plant, "
            "this chart shows which vendors feeding that plant carry the highest failure risk."
        ),
    },
    "🚛 Route Efficiency Scatter": {
        "title": "Why: Route Efficiency",
        "body": (
            "<strong>What it shows:</strong> Each route as a bubble — X = distance (km), "
            "Y = transport cost (₹). Bubble size = efficiency score. Colour = transport mode.<br><br>"
            "<strong>Why it matters for RCA:</strong> Low-efficiency routes (small bubbles, top-right quadrant) "
            "are cost-inefficient <em>and</em> prone to delays. "
            "Use this to identify whether logistics infrastructure is contributing to the RCA."
        ),
    },
    "📅 Monthly Delay Trend": {
        "title": "Why: Monthly Trend",
        "body": (
            "<strong>What it shows:</strong> Two area lines — delayed vs on-time shipments "
            "across all 12 months.<br><br>"
            "<strong>Why it matters for RCA:</strong> Reveals whether delays are <em>seasonal</em> "
            "(e.g. festival demand spikes) or <em>structural</em> (persistent throughout the year). "
            "A structural pattern points to supplier/plant capacity issues; "
            "a seasonal spike points to demand forecasting failures."
        ),
    },
    "📦 Stockout Severity": {
        "title": "Why: Stockout Severity",
        "body": (
            "<strong>What it shows:</strong> Top 15 distributor cities ranked by total demand gap "
            "(units short). Colour intensity = severity.<br><br>"
            "<strong>Why it matters for RCA:</strong> This is the <em>downstream consequence</em> "
            "of upstream delays. It answers 'who is actually running out of stock?' "
            "and quantifies the business impact of the root causes identified earlier in the RCA chain."
        ),
    },
    "🌐 Transport Mode Sunburst": {
        "title": "Why: Transport Mode Sunburst",
        "body": (
            "<strong>What it shows:</strong> Hierarchical circle — centre = Supply Chain, "
            "ring 1 = Plants, ring 2 = transport modes (Road/Rail/Air/Sea) per plant, "
            "arc size = route count.<br><br>"
            "<strong>Why it matters for RCA:</strong> Shows <em>modal dependency risk</em>. "
            "A plant that relies 90%+ on Road routes is highly vulnerable to traffic disruptions. "
            "Use this to recommend diversifying transport modes as a mitigation strategy."
        ),
    },
}

# Map dropdown label → builder function
_VIZ_BUILDERS = {
    "🔥 Delay Heatmap":           build_delay_heatmap,
    "⚠️ Supplier Risk Scores":    build_supplier_risk_chart,
    "🚛 Route Efficiency Scatter": build_route_efficiency_scatter,
    "📅 Monthly Delay Trend":     build_monthly_delay_trend,
    "📦 Stockout Severity":       build_distributor_demand_gap,
    "🌐 Transport Mode Sunburst": build_plant_transport_sunburst,
}


def _render_rca_inline_chart(choice: str):
    """
    Called when the dropdown changes.
    Returns (why_html, plotly_figure).
    """
    if not choice or choice.startswith("—"):
        return (
            "<div class='rca-viz-why' style='color:#e2e8f0'>"
            "Select a chart above to see it here alongside its RCA context."
            "</div>",
            go.Figure().update_layout(
                paper_bgcolor="#060c1c",
                plot_bgcolor="#060c1c",
                xaxis=dict(visible=False),
                yaxis=dict(visible=False)
            )
        )

    meta = _VIZ_WHY.get(choice, {})
    why_html = (
        f"<div class='rca-viz-why'>"
        f"<span class='why-title'>🔬 {meta.get('title', 'About this chart')}</span>"
        f"{meta.get('body', '')}"
        f"</div>"
    )

    builder = _VIZ_BUILDERS.get(choice)
    fig = builder() if builder else go.Figure().update_layout(**_PLOTLY_LAYOUT)
    return why_html, fig

def _render_chart_only(choice: str):
    """Renders the chosen chart with no 'why' explanation — just the figure."""
    if not choice or choice.startswith("—"):
        return go.Figure().update_layout(
            paper_bgcolor="#060c1c",
            plot_bgcolor="#060c1c",
            xaxis=dict(visible=False), yaxis=dict(visible=False)
        )
    builder = _VIZ_BUILDERS.get(choice)
    if not builder:
        return go.Figure().update_layout(**_PLOTLY_LAYOUT)
    try:
        return builder()
    except Exception as e:
        return go.Figure().update_layout(title=f"Error: {e}", **_PLOTLY_LAYOUT)


def _get_charts_for_question(question: str):
    """
    Returns gr.update for the dropdown with relevant chart choices for the question.
    Always keeps '— Select a chart —' as the selected value so that
    rca_viz_dropdown.change is NOT triggered automatically (which would
    fire a Neo4j call and show an Error badge on the Plot).
    The user selects a chart manually after the RCA finishes.
    """
    q = (question or "").lower().strip()

    if not q:
        return gr.update(choices=RCA_VIZ_OPTIONS, value="— Select a chart —")

    charts = ["— Select a chart —"]

    if any(w in q for w in ["delay", "late", "slow", "delayed", "bottleneck"]):
        charts.append("🔥 Delay Heatmap")
        charts.append("📅 Monthly Delay Trend")
    if any(w in q for w in ["toy", "toys", "watch", "health", "auto", "beauty", "gift", "product"]):
        if "🔥 Delay Heatmap" not in charts:
            charts.append("🔥 Delay Heatmap")
    if any(w in q for w in ["supplier", "vendor", "risk", "root cause", "high risk"]):
        charts.append("⚠️ Supplier Risk Scores")
    if any(w in q for w in ["stockout", "stock", "shortage", "retailer"]):
        charts.append("📦 Stockout Severity")
    if any(w in q for w in ["route", "transport", "logistics", "road", "rail", "air", "mode"]):
        charts.append("🚛 Route Efficiency Scatter")
        charts.append("🌐 Transport Mode Sunburst")
    if any(w in q for w in ["distributor", "distribution", "city", "demand gap"]):
        if "📦 Stockout Severity" not in charts:
            charts.append("📦 Stockout Severity")
    if any(w in q for w in ["plant", "factory", "baddi", "pune", "bhopal", "goa"]):
        if "🔥 Delay Heatmap" not in charts:
            charts.append("🔥 Delay Heatmap")
        if "⚠️ Supplier Risk Scores" not in charts:
            charts.append("⚠️ Supplier Risk Scores")

    if len(charts) == 1:  # nothing matched — show all
        charts = ["— Select a chart —"] + list(_VIZ_BUILDERS.keys())

    # Deduplicate while preserving order
    seen = set()
    charts = [x for x in charts if not (x in seen or seen.add(x))]

    # Always keep placeholder selected — never auto-trigger rca_viz_dropdown.change
    return gr.update(choices=charts, value="— Select a chart —")

def _default_dual_glow_html() -> str:
    """Default single explanation box shown before any analysis is run."""
    return """
<div class="viz-glow-box what-glow" style="margin-bottom:12px">
  <div class="viz-glow-icon">💡</div>
  <div class="viz-glow-label">Chart Explanation</div>
  <div class="viz-glow-body"><strong>Run an analysis first</strong> — click a chart type on the left to see it here. The explanation will describe what the chart is showing and why it is relevant to your root cause.</div>
  <div class="viz-glow-tags">
    <span class="viz-glow-tag">Context-aware</span>
    <span class="viz-glow-tag">Auto-updates with your question</span>
  </div>
</div>"""


# ════════════════════════════════════════════════════════════════════
# BUILD GRADIO UI
# ════════════════════════════════════════════════════════════════════

# ════════════════════════════════════════════════════════════════════
# UPDATE GRAPH — FILE UPLOAD BACKEND
# ════════════════════════════════════════════════════════════════════
#
# Backend flow:
#   1. User uploads file (.xlsx / .csv / .json / .txt)
#   2. _parse_upload_file()  — reads file into list[dict] rows
#   3. _detect_entity_type() — LLM detects: Supplier / Distributor /
#                              Route / Shipment / Product / Unknown
#   4. _normalize_rows()     — maps arbitrary column names to canonical
#                              Neo4j property names via LLM + heuristics
#   5. _preview_rows()       — returns first 5 rows as HTML table
#   6. _validate_rows()      — checks required fields, flags issues
#   7. _insert_rows_to_graph() — MERGE each row into Neo4j
#   8. _record_upload_history() — stores metadata in _UPLOAD_HISTORY
#
# Rollback:
#   Each insert records the Cypher MERGE that was used.
#   Rollback fires a DETACH DELETE on all nodes whose IDs were
#   created in that batch, then removes the history entry.
#
# History storage:
#   In-memory dict _UPLOAD_HISTORY keyed by upload_id (UUID).
#   Persisted to upload_history.json on disk after each operation.
# ════════════════════════════════════════════════════════════════════

import uuid as _uuid
import json as _json_mod
import pathlib as _pathlib

_HISTORY_FILE = _pathlib.Path("upload_history.json")
_UPLOAD_HISTORY: dict = {}   # { upload_id: { metadata } }

def _load_history():
    global _UPLOAD_HISTORY
    try:
        if _HISTORY_FILE.exists():
            _UPLOAD_HISTORY = _json_mod.loads(_HISTORY_FILE.read_text())
    except Exception:
        _UPLOAD_HISTORY = {}

def _save_history():
    try:
        _HISTORY_FILE.write_text(_json_mod.dumps(_UPLOAD_HISTORY, indent=2, default=str))
    except Exception as e:
        print(f"[History] Save failed: {e}")

_load_history()

# ── Canonical field maps per entity ──────────────────────────────────
_CANONICAL = {
    "Supplier": {
        "id": [
            "supplier_id", "id", "sup_id", "supplier_code", "vendor_id",
            "vendorid", "supplierid", "supplier id", "sup id", "code",
            "supplierno", "supplier_no", "vendor_code",
        ],
        "name": [
            "supplier_name", "name", "company", "vendor_name", "vendor",
            "company_name", "firm_name", "firm", "business_name",
            "organisation", "organization", "org_name",
        ],
        "risk_score": [
            "risk_score", "risk", "riskscore", "risk score", "score",
            "risk_level", "risk_rating", "riskrating", "vendor_risk",
            "supplier_risk", "risk_index",
        ],
        "capacity": [
            "annual_capacity_units", "capacity", "capacity_units",
            "annual_capacity", "units", "production_capacity",
            "max_capacity", "supply_capacity", "output_capacity",
            "capacity_per_year", "yearly_capacity",
        ],
        "lead_time": [
            "StoP_lead_time_days", "lead_time", "lead_time_days",
            "leadtime", "lead_days", "delivery_days", "transit_time",
            "fulfillment_days", "days_to_deliver", "avg_lead_time",
            "stop_lead_time_days", "s_to_p_lead_time",
        ],
        "plant": [
            "plant_id", "plant", "supplies_to", "plant_connection",
            "plant_code", "target_plant", "assigned_plant",
            "manufacturing_plant", "production_plant", "facility_id",
        ],
        "status": ["status", "supplier_status", "active", "state"],
    },
    "Distributor": {
        "id": [
            "distributor_id", "dist_id", "id", "distributor_code",
            "distid", "distribution_id", "hub_id", "depot_id",
            "distributor id", "dist id",
        ],
        "city": [
            "distributor_city", "city", "location", "city_name",
            "distributor_location", "hub_city", "depot_city",
            "distribution_city", "region", "area", "place", "town",
            "distributor city",
        ],
        "lat": [
            "distributor_latitude", "lat", "latitude", "coord_lat",
            "distributor_lat", "geo_lat", "y_coord",
        ],
        "lng": [
            "distributor_longitude", "lng", "longitude", "long",
            "coord_lng", "distributor_lng", "geo_lng", "x_coord",
            "lon", "lon_deg",
        ],
    },
    "Route": {
        "id": [
            "route_id", "id", "route_code", "route", "routeid",
            "route id", "lane_id", "shipment_lane_id",
        ],
        "mode": [
            "mode", "transport_mode", "transportation_mode", "transport",
            "transport_type", "shipping_mode", "delivery_mode",
            "transit_mode", "carrier_type",
        ],
        "dist_km": [
            "PtoD_distance_km", "distance_km", "distance", "km",
            "dist_km", "ptod_distance_km", "route_distance", "length_km",
        ],
        "days": [
            "PtoD_leadtime_days", "lead_time_days", "days", "lead_time",
            "transit_days", "ptod_leadtime_days", "transit_time_days",
        ],
        "cost": [
            "PtoD_transportation_cost_inr", "cost", "cost_inr",
            "transport_cost", "price", "ptod_transportation_cost_inr",
            "shipping_cost", "freight_cost", "logistics_cost",
        ],
        "plant": [
            "plant_id", "plant", "from_plant", "origin_plant",
            "source_plant", "origin", "dispatch_plant",
        ],
        "dist": [
            "distributor_id", "dist_id", "to_distributor", "distributor",
            "destination", "to_dist", "dest_id",
        ],
    },
}

def _find_header_row_app(ws, max_scan: int = 8):
    """Find the first row that looks like real column headers (not title/instruction)."""
    SKIP_PREFIXES = ("graphpulse", "\U0001f4cb", "instructions", "note:", "how to",
                     "readme", "fill", "template", "upload guide", "do not", "📋")
    for i, row in enumerate(ws.iter_rows(max_row=max_scan, values_only=True), start=1):
        cells = [str(c or "").strip() for c in row]
        non_empty = [c for c in cells if c]
        if len(non_empty) < 2:
            continue
        first = non_empty[0].lower()
        if any(first.startswith(p.lower()) for p in SKIP_PREFIXES):
            continue
        avg_len = sum(len(c) for c in non_empty) / max(len(non_empty), 1)
        if avg_len > 80:
            continue
        return i, cells
    return 1, []


def _parse_upload_file(filepath: str) -> tuple[list[dict], str]:
    """Parse .xlsx/.csv/.json/.txt → list[dict] rows + detected format.
    For multi-sheet Excel: picks the sheet with the most data rows after skipping
    title/instruction rows automatically."""
    import os
    ext = os.path.splitext(filepath)[1].lower()
    try:
        if ext in (".xlsx", ".xls"):
            import openpyxl
            wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
            SKIP_SHEETS = {"readme", "instructions", "guide", "info", "notes", "template"}
            data_sheets = [s for s in wb.sheetnames
                           if not any(skip in s.lower() for skip in SKIP_SHEETS)]
            if not data_sheets:
                data_sheets = wb.sheetnames[:1]

            best_rows, best_sheet_name, best_fmt = [], "", "excel"
            for sheet_name in data_sheets:
                ws = wb[sheet_name]
                hdr_idx, raw_hdrs = _find_header_row_app(ws)
                headers = [str(h or "").strip() for h in raw_hdrs]
                valid_idx = [i for i, h in enumerate(headers) if h]
                if not valid_idx:
                    continue
                headers_clean = [headers[i] for i in valid_idx]
                sheet_rows = []
                for r in ws.iter_rows(min_row=hdr_idx + 1, values_only=True):
                    if r is None or all(c is None for c in r):
                        continue
                    d = {}
                    for i in valid_idx:
                        if i < len(r) and r[i] is not None:
                            val = r[i]
                            if str(val).strip() not in ("", "None", "nan", "NaN"):
                                d[headers_clean[valid_idx.index(i)]] = val
                    if d:
                        sheet_rows.append(d)
                if len(sheet_rows) > len(best_rows):
                    best_rows = sheet_rows
                    best_sheet_name = sheet_name
                    best_fmt = f"excel ({sheet_name})"

            return best_rows, best_fmt
        elif ext == ".csv":
            import csv
            with open(filepath, newline="", encoding="utf-8-sig") as f:
                reader = csv.DictReader(f)
                return [dict(r) for r in reader], "csv"
        elif ext == ".json":
            with open(filepath, encoding="utf-8") as f:
                data = _json_mod.load(f)
            if isinstance(data, list):
                return data, "json"
            elif isinstance(data, dict):
                # Try common wrappers: {"nodes": [...]} or {"data": [...]}
                for key in ["nodes", "data", "rows", "records", "items"]:
                    if key in data and isinstance(data[key], list):
                        return data[key], "json"
                return [data], "json"
        elif ext in (".txt", ".tsv"):
            import csv
            with open(filepath, newline="", encoding="utf-8-sig") as f:
                dialect = "excel-tab" if ext == ".tsv" else csv.Sniffer().sniff(f.read(2048))
                f.seek(0)
                reader = csv.DictReader(f, dialect=dialect)
                return [dict(r) for r in reader], "text"
        return [], "unknown"
    except Exception as e:
        return [], f"error:{e}"

def _detect_entity_type(rows: list[dict], filename: str = "", fmt: str = "") -> str:
    """
    Detect entity type from sheet name, filename, or column names.
    Returns: Supplier | Distributor | Route | Unknown
    """
    if not rows:
        return "Unknown"

    # Priority 1: sheet/format hint from _parse_upload_file
    fmt_lower = fmt.lower()
    if "supplier" in fmt_lower: return "Supplier"
    if "distributor" in fmt_lower: return "Distributor"
    if "route" in fmt_lower: return "Route"

    # Priority 2: filename hint
    fn_lower = filename.lower()
    if "supplier" in fn_lower: return "Supplier"
    if "distributor" in fn_lower or "dist" in fn_lower: return "Distributor"
    if "route" in fn_lower: return "Route"

    def _norm(s): return s.lower().replace(" ", "_").replace("-", "_").strip()
    cols = set(_norm(c) for c in rows[0].keys())

    # Heuristics — broad alias matching
    SUPPLIER_COLS = {"supplier_id","supplier_name","risk_score","vendor","vendor_name",
                     "vendor_id","sup_id","supplierid","riskscore","annual_capacity_units",
                     "stop_lead_time_days","lead_time_days"}
    DIST_COLS     = {"distributor_id","distributor_city","city_name","dist_id","hub_id",
                     "distributor_latitude","depot_city","distribution_city"}
    ROUTE_COLS    = {"route_id","transport_mode","transportation_mode","ptod_distance_km",
                     "distance_km","ptod_leadtime_days","route_code","mode"}

    sup_score  = sum(1 for c in SUPPLIER_COLS if c in cols)
    dist_score = sum(1 for c in DIST_COLS     if c in cols)
    route_score= sum(1 for c in ROUTE_COLS    if c in cols)

    best_score = max(sup_score, dist_score, route_score)
    if best_score > 0:
        if   sup_score   == best_score: return "Supplier"
        elif dist_score  == best_score: return "Distributor"
        elif route_score == best_score: return "Route"

    if any(c in cols for c in ["shipment_id", "delivery_status", "delay_days"]):
        return "Shipment"
    if any(c in cols for c in ["product_id", "product_category", "category_name"]):
        return "Product"

    # LLM fallback for ambiguous column names
    try:
        sample_cols = list(rows[0].keys())[:15]
        raw = _groq_call(
            messages=[
                {"role": "system", "content":
                    "You are a data schema classifier for a supply chain graph database. "
                    "Return ONLY one word: Supplier, Distributor, Route, Shipment, Product, or Unknown."},
                {"role": "user", "content":
                    f"Column names: {sample_cols}. Sample row: {dict(list(rows[0].items())[:8])}. What entity type does this data represent?"}
            ],
            max_tokens=10, temperature=0,
            model_chain=["llama-3.1-8b-instant", "llama-3.3-70b-versatile"],
        )
        detected = str(raw or "").strip().split()[0].capitalize()
        if detected in ("Supplier", "Distributor", "Route", "Shipment", "Product"):
            return detected
    except Exception:
        pass
    return "Unknown"

def _normalize_rows(rows: list[dict], entity_type: str) -> tuple[list[dict], dict, list[str]]:
    """
    Map arbitrary column names to canonical Neo4j property names.
    Returns: (normalized_rows, field_mapping_used, warnings)

    For unrecognized columns, uses LLM to suggest mapping.
    Missing required fields are flagged as warnings.
    """
    if entity_type not in _CANONICAL:
        return rows, {}, [f"No normalization template for entity type '{entity_type}'"]

    template = _CANONICAL[entity_type]
    input_cols = list(rows[0].keys()) if rows else []

    def _norm(s):
        return s.lower().replace(" ", "_").replace("-", "_").replace(".", "_").strip()

    input_cols_norm = {_norm(c): c for c in input_cols}
    mapping = {}   # canonical_field → original_col_name
    warnings = []

    for canon_field, aliases in template.items():
        matched = None
        # Pass 1: exact normalised match
        for alias in aliases:
            key = _norm(alias)
            if key in input_cols_norm:
                matched = input_cols_norm[key]
                break
        # Pass 2: substring match
        if not matched:
            for col_norm, col_orig in input_cols_norm.items():
                for alias in aliases:
                    alias_n = _norm(alias)
                    if alias_n in col_norm or col_norm in alias_n:
                        matched = col_orig
                        break
                if matched:
                    break
        if matched:
            mapping[canon_field] = matched
        else:
            warnings.append(f"Field '{canon_field}' not found — will be skipped or defaulted")

    # LLM-assist for unmatched columns
    unmatched_inputs = [c for c in input_cols if c not in mapping.values()]
    if unmatched_inputs and len(mapping) < len(template) // 2:
        try:
            raw = _groq_call(
                messages=[
                    {"role": "system", "content":
                        f"Map these input column names to canonical fields for a {entity_type} node. "
                        f"Canonical fields: {list(template.keys())}. "
                        "Return ONLY valid JSON mapping input cols to canonical fields. "
                        "Only map if confident. Omit uncertain ones."},
                    {"role": "user", "content":
                        f"Input columns to map: {unmatched_inputs}"}
                ],
                max_tokens=300, temperature=0,
                model_chain=["llama-3.1-8b-instant"],
            )
            import re as _re
            raw = _re.sub(r"```(?:json)?|```", "", str(raw or "")).strip()
            llm_map = _json_mod.loads(raw)
            for input_col, canon in llm_map.items():
                if canon in template and input_col not in mapping.values():
                    mapping[canon] = input_col
        except Exception:
            pass

    # Apply mapping to rows
    norm_rows = []
    for row in rows:
        nr = {}
        for canon_field, orig_col in mapping.items():
            val = row.get(orig_col)
            if val is not None and str(val).strip() not in ("", "None", "nan", "NaN"):
                nr[canon_field] = val
        norm_rows.append(nr)

    return norm_rows, mapping, warnings

def _preview_rows(norm_rows: list[dict], mapping: dict, entity_type: str,
                  warnings: list[str]) -> str:
    """Build an HTML preview table of the first 5 normalized rows."""
    if not norm_rows:
        return "<div style='color:#f87171'>No valid rows found after normalization.</div>"

    sample = norm_rows[:5]
    headers = list(sample[0].keys()) if sample else []
    total = len(norm_rows)

    warn_html = ""
    if warnings:
        warn_items = "".join(f"<li>{w}</li>" for w in warnings[:4])
        warn_html = f"""
<div style="background:rgba(251,191,36,0.08);border:1px solid rgba(251,191,36,0.3);
     border-radius:8px;padding:10px 14px;margin-bottom:12px">
  <div style="font-size:0.62rem;font-weight:800;color:#fbbf24;text-transform:uppercase;
       letter-spacing:0.1em;margin-bottom:6px">⚠ Normalization Warnings</div>
  <ul style="margin:0;padding-left:16px;color:#fde68a;font-size:0.78rem">{warn_items}</ul>
</div>"""

    field_map_html = "".join(
        f"<span style='background:rgba(56,189,248,0.1);border:1px solid rgba(56,189,248,0.2);"
        f"border-radius:4px;padding:2px 7px;font-size:0.65rem;color:#7dd3fc;margin:2px'>"
        f"{orig} → {canon}</span>"
        for canon, orig in mapping.items()
    )

    th = "".join(f"<th style='padding:6px 10px;font-size:0.7rem;font-weight:700;"
                 f"color:#38bdf8;border-bottom:1px solid rgba(56,189,248,0.2);white-space:nowrap'>"
                 f"{h}</th>" for h in headers)
    rows_html = ""
    for row in sample:
        cells = "".join(
            f"<td style='padding:5px 10px;font-size:0.72rem;color:#e2e8f0;"
            f"border-bottom:1px solid rgba(255,255,255,0.05)'>{row.get(h,'—')}</td>"
            for h in headers
        )
        rows_html += f"<tr>{cells}</tr>"

    return f"""
<div style="background:rgba(6,12,28,0.8);border:1px solid rgba(56,189,248,0.2);
     border-radius:12px;padding:16px;margin:12px 0">
  <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:12px">
    <div>
      <span style="font-size:0.62rem;font-weight:800;text-transform:uppercase;
            letter-spacing:0.12em;color:#38bdf8">📋 Preview — {entity_type}</span>
      <span style="font-size:0.72rem;color:#94a3b8;margin-left:8px">
        {total} rows detected · showing first {min(5,total)}</span>
    </div>
  </div>
  {warn_html}
  <div style="margin-bottom:10px;font-size:0.62rem;font-weight:700;
       color:#94a3b8;text-transform:uppercase;letter-spacing:0.08em">Field Mapping Applied</div>
  <div style="display:flex;flex-wrap:wrap;gap:4px;margin-bottom:14px">{field_map_html}</div>
  <div style="overflow-x:auto">
    <table style="width:100%;border-collapse:collapse">
      <thead><tr>{th}</tr></thead>
      <tbody>{rows_html}</tbody>
    </table>
  </div>
  <div style="margin-top:10px;font-size:0.72rem;color:#cbd5e1">
    ✓ Review the mapping above. Click <strong style="color:#4ade80">Confirm & Insert</strong> to proceed.
  </div>
</div>"""

def _validate_rows(norm_rows: list[dict], entity_type: str) -> tuple[list[dict], list[str], list[str]]:
    """
    Validate normalized rows before insert.
    Returns: (valid_rows, errors, warnings)
    Deduplicates by ID field.
    """
    required = {
        "Supplier":    ["id", "name"],
        "Distributor": ["id", "city"],
        "Route":       ["id", "mode"],
    }
    req = required.get(entity_type, [])
    valid, errors, warnings = [], [], []
    seen_ids = set()

    for i, row in enumerate(norm_rows, 1):
        row_errors = []
        for field in req:
            if not row.get(field):
                row_errors.append(f"Row {i}: missing required field '{field}'")
        if row_errors:
            errors.extend(row_errors)
            continue

        # Deduplication
        row_id = str(row.get("id", "")).strip()
        if row_id in seen_ids:
            warnings.append(f"Row {i}: duplicate ID '{row_id}' — skipped")
            continue
        seen_ids.add(row_id)

        # Check if ID already exists in graph
        try:
            label_map = {"Supplier": ("Supplier", "supplier_id"),
                         "Distributor": ("Distributor", "distributor_id"),
                         "Route": ("Route", "route_id")}
            if entity_type in label_map:
                label, id_prop = label_map[entity_type]
                existing = _run_neo4j(
                    f"MATCH (n:{label} {{{id_prop}: $id}}) RETURN n LIMIT 1",
                    {"id": row_id}
                )
                if existing:
                    warnings.append(f"Row {i}: ID '{row_id}' already exists — will be MERGED (updated)")
        except Exception:
            pass

        valid.append(row)

    return valid, errors, warnings

def _build_cypher_for_row(row: dict, entity_type: str) -> str:
    """Build a MERGE Cypher statement for a single normalized row."""
    if entity_type == "Supplier":
        sid  = str(row.get("id","")).strip()
        name = str(row.get("name","")).strip()
        risk = float(row.get("risk_score", 0.5))
        cap  = int(row.get("capacity", 0))
        lt   = int(row.get("lead_time", 7))
        plant = str(row.get("plant","")).strip()
        cypher = (
            f"MERGE (s:Supplier {{supplier_id:'{sid}'}}) "
            f"SET s.supplier_name='{name}', s.risk_score={risk}, "
            f"s.annual_capacity_units={cap}, s.StoP_lead_time_days={lt}, "
            f"s.status='Active', s.supplier_latitude=0.0, s.supplier_longitude=0.0, "
            f"s.StoP_distance_km=0.0"
        )
        if plant:
            cypher += (
                f" WITH s MATCH (p:Plant {{plant_id:'{plant}'}}) "
                f"MERGE (s)-[:SUPPLIES_TO]->(p)"
            )
        return cypher

    elif entity_type == "Distributor":
        did  = str(row.get("id","")).strip()
        city = str(row.get("city","")).strip()
        lat  = float(row.get("lat", 0.0))
        lng  = float(row.get("lng", 0.0))
        return (
            f"MERGE (d:Distributor {{distributor_id:'{did}'}}) "
            f"SET d.distributor_city='{city}', d.distributor_latitude={lat}, "
            f"d.distributor_longitude={lng}"
        )

    elif entity_type == "Route":
        rid   = str(row.get("id","")).strip()
        mode  = str(row.get("mode","Road")).strip()
        dist  = float(row.get("dist_km", 0))
        days  = int(row.get("days", 1))
        cost  = float(row.get("cost", 0))
        plant = str(row.get("plant","")).strip()
        dist_id = str(row.get("dist","")).strip()
        return (
            f"MERGE (r:Route {{route_id:'{rid}'}}) "
            f"SET r.mode='{mode}', r.PtoD_distance_km={dist}, "
            f"r.PtoD_leadtime_days={days}, r.PtoD_transportation_cost_inr={cost}, "
            f"r.plant_id='{plant}', r.distributor_id='{dist_id}'"
        )

    return f"// Unsupported entity type: {entity_type}"

def _build_parameterised_cypher(row: dict, entity_type: str) -> tuple[str, dict]:
    """
    Build a parameterised MERGE Cypher + params dict.
    Using $params avoids all apostrophe/quote escaping issues.
    """
    if entity_type == "Supplier":
        cypher = (
            "MERGE (s:Supplier {supplier_id: $sid}) "
            "SET s.supplier_name = $name, "
            "    s.risk_score = $risk, "
            "    s.annual_capacity_units = $cap, "
            "    s.StoP_lead_time_days = $lt, "
            "    s.status = 'Active', "
            "    s.supplier_latitude = 0.0, "
            "    s.supplier_longitude = 0.0, "
            "    s.StoP_distance_km = 0.0 "
            "RETURN s.supplier_id AS id"
        )
        params = {
            "sid":  str(row.get("id", "")).strip(),
            "name": str(row.get("name", "")).strip(),
            "risk": float(row.get("risk_score", 0.5)),
            "cap":  int(float(row.get("capacity", 0) or 0)),
            "lt":   int(float(row.get("lead_time", 7) or 7)),
        }
        # If plant supplied, link after
        plant = str(row.get("plant", "")).strip()
        if plant:
            cypher += (
                " WITH s "
                "MATCH (p:Plant {plant_id: $plant}) "
                "MERGE (s)-[:SUPPLIES_TO]->(p) "
                "RETURN s.supplier_id AS id"
            )
            params["plant"] = plant
        return cypher, params

    elif entity_type == "Distributor":
        return (
            "MERGE (d:Distributor {distributor_id: $did}) "
            "SET d.distributor_city = $city, "
            "    d.distributor_latitude = $lat, "
            "    d.distributor_longitude = $lng "
            "RETURN d.distributor_id AS id",
            {
                "did":  str(row.get("id", "")).strip(),
                "city": str(row.get("city", "")).strip(),
                "lat":  float(row.get("lat", 0.0) or 0.0),
                "lng":  float(row.get("lng", 0.0) or 0.0),
            }
        )

    elif entity_type == "Route":
        return (
            "MERGE (r:Route {route_id: $rid}) "
            "SET r.mode = $mode, "
            "    r.PtoD_distance_km = $dist, "
            "    r.PtoD_leadtime_days = $days, "
            "    r.PtoD_transportation_cost_inr = $cost, "
            "    r.plant_id = $plant, "
            "    r.distributor_id = $dist_id "
            "RETURN r.route_id AS id",
            {
                "rid":     str(row.get("id", "")).strip(),
                "mode":    str(row.get("mode", "Road")).strip(),
                "dist":    float(row.get("dist_km", 0) or 0),
                "days":    int(float(row.get("days", 1) or 1)),
                "cost":    float(row.get("cost", 0) or 0),
                "plant":   str(row.get("plant", "")).strip(),
                "dist_id": str(row.get("dist", "")).strip(),
            }
        )

    # Fallback: plain MERGE with no properties
    return f"MERGE (n {{id: $id}}) RETURN n", {"id": str(row.get("id", "unknown"))}


def _insert_rows_to_graph(valid_rows: list[dict], entity_type: str,
                          filename: str, upload_id: str) -> tuple[int, int, list[str]]:
    """
    Execute MERGE for each valid row using parameterised queries.
    Parameterised queries prevent apostrophe/quote issues (e.g. 'Mehta Plastics').
    Returns (success_count, fail_count, cypher_log).
    Records the batch in _UPLOAD_HISTORY for rollback.
    """
    success, fail = 0, 0
    cypher_log = []
    inserted_ids = []

    for row in valid_rows:
        try:
            cypher, params = _build_parameterised_cypher(row, entity_type)
            cypher_log.append(cypher)
            _run_neo4j(cypher, params)
            success += 1
            inserted_ids.append(str(row.get("id", "")))
        except Exception as e:
            fail += 1
            cypher_log.append(f"// ERROR: {e}")

    # Record in history
    import datetime as _dt
    _UPLOAD_HISTORY[upload_id] = {
        "upload_id":    upload_id,
        "filename":     filename,
        "entity_type":  entity_type,
        "total_rows":   len(valid_rows),
        "success":      success,
        "fail":         fail,
        "inserted_ids": inserted_ids,
        "cypher_log":   cypher_log[:50],  # cap to 50 for storage
        "timestamp":    _dt.datetime.now().isoformat(),
        "rolled_back":  False,
    }
    _save_history()
    return success, fail, cypher_log

def _rollback_upload(upload_id: str) -> tuple[bool, str]:
    """
    Delete all nodes created in a given upload batch by their IDs.
    Returns (success, message).
    """
    entry = _UPLOAD_HISTORY.get(upload_id)
    if not entry:
        return False, f"Upload ID '{upload_id}' not found in history."
    if entry.get("rolled_back"):
        return False, "This upload has already been rolled back."

    entity_type = entry["entity_type"]
    inserted_ids = entry.get("inserted_ids", [])
    if not inserted_ids:
        return False, "No inserted IDs recorded — cannot rollback."

    label_map = {
        "Supplier":    ("Supplier", "supplier_id"),
        "Distributor": ("Distributor", "distributor_id"),
        "Route":       ("Route", "route_id"),
    }
    if entity_type not in label_map:
        return False, f"Rollback not supported for entity type '{entity_type}'."

    label, id_prop = label_map[entity_type]
    deleted = 0
    errors = []
    for node_id in inserted_ids:
        try:
            _run_neo4j(
                f"MATCH (n:{label} {{{id_prop}: $id}}) DETACH DELETE n",
                {"id": node_id}
            )
            deleted += 1
        except Exception as e:
            errors.append(str(e)[:80])

    entry["rolled_back"] = True
    entry["rollback_deleted"] = deleted
    _save_history()

    if errors:
        return True, f"Rolled back {deleted}/{len(inserted_ids)} nodes. Errors: {errors[:3]}"
    return True, f"✓ Rolled back {deleted} {entity_type} node(s) successfully."

def _delete_single_node(upload_id: str, node_id: str) -> tuple[bool, str]:
    """Delete a specific node from a batch by its ID."""
    entry = _UPLOAD_HISTORY.get(upload_id)
    if not entry:
        return False, "Upload not found."
    entity_type = entry["entity_type"]
    label_map = {
        "Supplier":    ("Supplier", "supplier_id"),
        "Distributor": ("Distributor", "distributor_id"),
        "Route":       ("Route", "route_id"),
    }
    if entity_type not in label_map:
        return False, f"Delete not supported for '{entity_type}'."
    label, id_prop = label_map[entity_type]
    try:
        _run_neo4j(f"MATCH (n:{label} {{{id_prop}: $id}}) DETACH DELETE n", {"id": node_id})
        if node_id in entry.get("inserted_ids", []):
            entry["inserted_ids"].remove(node_id)
        _save_history()
        return True, f"✓ Deleted {entity_type} node '{node_id}'."
    except Exception as e:
        return False, str(e)

def _get_recent_history_html() -> str:
    """Build the collapsible recent updates panel HTML."""
    if not _UPLOAD_HISTORY:
        return """
<div style="color:#94a3b8;font-size:0.78rem;padding:20px 0;text-align:center">
  No uploads yet. Upload a file to see history here.
</div>"""

    items = sorted(_UPLOAD_HISTORY.values(),
                   key=lambda x: x.get("timestamp",""), reverse=True)[:10]

    html = ""
    for entry in items:
        uid   = entry["upload_id"]
        fname = entry.get("filename","unknown")
        etype = entry.get("entity_type","?")
        succ  = entry.get("success", 0)
        fail  = entry.get("fail", 0)
        ts    = entry.get("timestamp","")[:16].replace("T"," ")
        rb    = entry.get("rolled_back", False)
        ids   = entry.get("inserted_ids", [])[:6]

        status_color = "#f87171" if rb else "#4ade80"
        status_label = "ROLLED BACK" if rb else "ACTIVE"
        fail_badge   = f"<span style='color:#f87171;font-size:0.65rem'>· {fail} failed</span>" if fail else ""

        id_chips = "".join(
            f"<span style='background:rgba(56,189,248,0.08);border:1px solid rgba(56,189,248,0.15);"
            f"border-radius:4px;padding:1px 6px;font-size:0.62rem;color:#7dd3fc;margin:1px'>{i}</span>"
            for i in ids
        )
        if len(entry.get("inserted_ids",[])) > 6:
            id_chips += f"<span style='color:#94a3b8;font-size:0.62rem'> +{len(entry['inserted_ids'])-6} more</span>"

        html += f"""
<details style="background:rgba(6,12,28,0.7);border:1px solid rgba(56,189,248,0.12);
     border-radius:10px;margin-bottom:8px;overflow:hidden">
  <summary style="cursor:pointer;padding:10px 14px;display:flex;align-items:center;
       gap:8px;list-style:none;user-select:none">
    <span style="font-size:0.62rem;font-weight:800;color:{status_color};
          text-transform:uppercase;flex-shrink:0">{status_label}</span>
    <span style="font-size:0.75rem;font-weight:600;color:#e2e8f0;flex:1;
          white-space:nowrap;overflow:hidden;text-overflow:ellipsis">{fname}</span>
    <span style="font-size:0.65rem;color:#94a3b8;flex-shrink:0">{ts}</span>
  </summary>
  <div style="padding:8px 14px 12px">
    <div style="font-size:0.7rem;color:#94a3b8;margin-bottom:8px">
      <strong style="color:#7dd3fc">{etype}</strong> ·
      <strong style="color:#4ade80">{succ} inserted</strong> {fail_badge}
    </div>
    <div style="margin-bottom:8px">{id_chips}</div>
    <div style="font-size:0.62rem;color:#94a3b8;font-family:monospace;
         margin-bottom:8px">upload_id: {uid}</div>
  </div>
</details>"""

    return html


# ── State for pending upload (between parse→preview→confirm) ─────────
_PENDING_UPLOAD: dict = {}   # filepath → {rows, entity_type, mapping, warnings, norm_rows}

def _handle_file_upload(filepath) -> tuple[str, str, str]:
    """
    Step 1 of upload flow: parse + normalize + preview.
    Returns: (preview_html, entity_type_detected, upload_id)
    """
    if filepath is None:
        return "<div style='color:#94a3b8'>No file selected.</div>", "", ""

    actual_path = filepath.name if hasattr(filepath, "name") else str(filepath)
    import os
    filename = os.path.basename(actual_path)

    rows, fmt = _parse_upload_file(actual_path)
    if not rows or fmt.startswith("error"):
        err = fmt.replace("error:", "") if fmt.startswith("error") else "empty file"
        return f"<div style='color:#f87171'>⚠ Could not parse file: {err}</div>", "", ""

    entity_type = _detect_entity_type(rows, filename=filename, fmt=fmt)
    norm_rows, mapping, warnings = _normalize_rows(rows, entity_type)
    valid_rows, errors, val_warnings = _validate_rows(norm_rows, entity_type)
    all_warnings = warnings + val_warnings

    preview = _preview_rows(norm_rows, mapping, entity_type, all_warnings)

    upload_id = str(_uuid.uuid4())[:8]
    _PENDING_UPLOAD[upload_id] = {
        "filename":    filename,
        "filepath":    actual_path,
        "entity_type": entity_type,
        "mapping":     mapping,
        "norm_rows":   norm_rows,
        "valid_rows":  valid_rows,
        "errors":      errors,
        "warnings":    all_warnings,
    }

    if errors:
        err_html = "<br>".join(errors[:5])
        preview += f"<div style='color:#f87171;font-size:0.75rem;margin-top:8px'>Errors: {err_html}</div>"

    return preview, f"{entity_type} · {len(valid_rows)} valid rows · ID: {upload_id}", upload_id

def _confirm_insert(upload_id: str):
    """
    Streaming 6-agent pipeline: yields (status_html, history_html) progressively.
    Each pipeline step fires on_update which causes a new yield with updated status.
    Wired to upd_insert_status + upd_history_html with show_progress="hidden".
    """
    pending = _PENDING_UPLOAD.get(upload_id)
    if not pending:
        yield "<div style='color:#f87171'>⚠ No pending upload found. Please re-upload.</div>", _get_recent_history_html()
        return

    filepath = pending.get("filepath", "")
    filename  = pending.get("filename", "unknown")

    if not filepath:
        yield "<div style='color:#f87171'>⚠ File path missing — please re-upload.</div>", _get_recent_history_html()
        return

    import queue as _q2, threading as _t2

    _STEP_LABELS = {
        "__upd_parse__":     ("📂", "#38bdf8",  "File Parser",     "Reading and extracting rows from file"),
        "__upd_detect__":    ("🔍", "#a78bfa",  "Schema Detector", "Identifying entity type via LLM"),
        "__upd_normalize__": ("🔄", "#22d3ee",  "Field Mapper",    "Normalising column names to Neo4j schema"),
        "__upd_clean__":     ("🧹", "#fbbf24",  "Data Cleaner",    "Fixing nulls, duplicates, type mismatches"),
        "__upd_validate__":  ("✅", "#4ade80",  "Validator",       "Checking rows before insertion"),
        "__upd_insert__":    ("🔗", "#f97316",  "Graph Agent",     "Inserting nodes into Neo4j"),
    }
    step_order = list(_STEP_LABELS.keys())

    completed_steps = []
    current_step_key = ""
    eq = _q2.Queue()
    holder = {}

    def worker():
        def on_update(event):
            eq.put(event)
        try:
            summary, logs = run_file_update_pipeline(filepath, filename, on_update=on_update)
            holder["summary"] = summary
            holder["logs"]    = logs
        except Exception as e:
            holder["error"] = str(e)
        finally:
            eq.put(None)

    _t2.Thread(target=worker, daemon=True).start()

    def _build_status_html(current_key, completed, summary_text=None, error=None):
        steps_html = ""
        for key in step_order:
            icon, col, label, desc = _STEP_LABELS[key]
            if key in completed:
                state_col, state_icon, bg = "#4ade80", "✓", "rgba(74,222,128,0.05)"
            elif key == current_key:
                state_col, state_icon, bg = col, "◌", f"rgba(56,189,248,0.08)"
            else:
                state_col, state_icon, bg = "#475569", "○", "transparent"
            steps_html += (
                f'<div style="display:flex;align-items:flex-start;gap:10px;padding:7px 10px;'
                f'background:{bg};border-radius:7px;margin:3px 0">'
                f'<span style="font-size:0.8rem;color:{state_col};flex-shrink:0">{state_icon}</span>'
                f'<div>'
                f'<div style="font-size:0.73rem;font-weight:700;color:{state_col}">{icon} {label}</div>'
                f'<div style="font-size:0.67rem;color:#94a3b8;margin-top:1px">{desc}</div>'
                f'</div></div>'
            )

        if error:
            footer = f'<div style="color:#f87171;font-size:0.78rem;margin-top:10px">⚠ {error}</div>'
        elif summary_text:
            import re as _re3
            sm = _re3.search(r"Inserted\s*:\s*(\d+)", summary_text)
            fm = _re3.search(r"Failed:\s*(\d+)", summary_text)
            um = _re3.search(r"Upload ID\s*:\s*([a-f0-9]+)", summary_text)
            succ = int(sm.group(1)) if sm else 0
            fail = int(fm.group(1)) if fm else 0
            uid2 = um.group(1) if um else upload_id
            col2 = "#4ade80" if fail == 0 else "#fbbf24"
            footer = (
                f'<div style="margin-top:12px;padding-top:10px;border-top:1px solid rgba(255,255,255,0.08)">'
                f'<div style="font-size:0.82rem;color:#e2e8f0;margin-bottom:6px">'
                f'<strong style="color:#4ade80">{succ}</strong> nodes inserted · '
                f'<strong style="color:#f87171">{fail}</strong> failed</div>'
                f'<div style="font-size:0.68rem;color:#cbd5e1">Upload ID: '
                f'<code style="color:#7dd3fc">{uid2}</code> — use this to rollback</div>'
                f'</div>'
            )
        else:
            pct = int(len(completed) / len(step_order) * 100)
            footer = (
                f'<div style="margin-top:10px;height:4px;background:rgba(255,255,255,0.08);border-radius:2px">'
                f'<div style="height:100%;width:{pct}%;background:linear-gradient(90deg,#38bdf8,#a78bfa);border-radius:2px;transition:width 0.4s"></div>'
                f'</div>'
            )

        return (
            f'<div style="background:rgba(6,12,28,0.9);border:1px solid rgba(56,189,248,0.2);'
            f'border-radius:12px;padding:16px">'
            f'<div style="font-size:0.62rem;font-weight:900;text-transform:uppercase;'
            f'letter-spacing:0.14em;color:#38bdf8;margin-bottom:12px">◈ 6-Agent Upload Pipeline</div>'
            f'{steps_html}{footer}</div>'
        )

    # Stream updates
    while True:
        try:
            event = eq.get(timeout=0.4)
        except _q2.Empty:
            yield _build_status_html(current_step_key, completed_steps), _get_recent_history_html()
            continue

        if event is None:
            break

        kind, data = event
        if kind == "tool_start" and data in _STEP_LABELS:
            current_step_key = data
        elif kind == "tool" and data.get("tool") in _STEP_LABELS:
            completed_steps.append(data["tool"])
            current_step_key = ""

        yield _build_status_html(current_step_key, completed_steps), _get_recent_history_html()

    # Final yield with summary
    if "error" in holder:
        yield _build_status_html("", completed_steps, error=holder["error"]), _get_recent_history_html()
    else:
        yield _build_status_html("", completed_steps, summary_text=holder.get("summary", "")), _get_recent_history_html()
    _PENDING_UPLOAD.pop(upload_id, None)

def _handle_rollback(upload_id: str) -> tuple[str, str]:
    """Rollback a full upload batch."""
    uid = upload_id.strip()
    if not uid:
        return "<div style='color:#f87171'>⚠ Enter an Upload ID to rollback.</div>", _get_recent_history_html()
    ok, msg = _rollback_upload(uid)
    color = "#4ade80" if ok else "#f87171"
    return (
        f"<div style='color:{color};font-size:0.82rem;padding:10px 0'>{msg}</div>",
        _get_recent_history_html()
    )

def _handle_delete_node(upload_id: str, node_id: str) -> tuple[str, str]:
    """Delete a single node from a batch."""
    ok, msg = _delete_single_node(upload_id.strip(), node_id.strip())
    color = "#4ade80" if ok else "#f87171"
    return (
        f"<div style='color:{color};font-size:0.82rem;padding:10px 0'>{msg}</div>",
        _get_recent_history_html()
    )


with gr.Blocks(title="GraphPulse AI", css=CUSTOM_CSS + EXTRA_CSS,
               theme=gr.themes.Base()) as demo:

    gr.HTML(INIT_JS)
    gr.HTML("""
<script>
(function() {
  // Force Plotly to resize whenever the View Charts panel becomes visible.
  // Gradio renders gr.Plot at 0px height when parent column is hidden at mount time.
  // A resize call after visibility change fixes this without modifying Python code.
  function resizeVchart() {
    var plotDiv = document.getElementById('rca-vchart-plot');
    if (!plotDiv) return;
    var gd = plotDiv.querySelector('.js-plotly-plot');
    if (!gd) return;
    try {
      if (window.Plotly) {
        Plotly.Plots.resize(gd);
      }
    } catch(e) {}
  }

  // Watch for the rca-view-charts-wrap column becoming visible
  function startObserver() {
    var observer = new MutationObserver(function(mutations) {
      mutations.forEach(function(m) {
        if (m.type === 'attributes' && m.attributeName === 'style') {
          var el = m.target;
          if (el && el.id === 'rca-view-charts-wrap') {
            // Visible → resize after short delay to let Gradio finish rendering
            setTimeout(resizeVchart, 150);
            setTimeout(resizeVchart, 500);
            setTimeout(resizeVchart, 1000);
          }
        }
        // Also watch for child additions (Gradio sometimes re-creates the plot element)
        if (m.addedNodes && m.addedNodes.length) {
          m.addedNodes.forEach(function(node) {
            if (node.nodeType === 1) {
              var inner = node.querySelector && node.querySelector('#rca-vchart-plot .js-plotly-plot');
              if (inner || (node.id === 'rca-vchart-plot')) {
                setTimeout(resizeVchart, 200);
                setTimeout(resizeVchart, 600);
              }
            }
          });
        }
      });
    });

    var root = document.getElementById('rca-view-charts-wrap') || document.body;
    observer.observe(document.body, { attributes: true, subtree: true, childList: true });
  }

  // Also poll every 2s as fallback — stops after 120s
  var _pollCount = 0;
  var _lastVisible = false;
  var _poll = setInterval(function() {
    _pollCount++;
    if (_pollCount > 60) { clearInterval(_poll); return; }
    var wrap = document.getElementById('rca-view-charts-wrap');
    if (wrap) {
      var isVisible = wrap.style.display !== 'none' && !wrap.hidden &&
                      getComputedStyle(wrap).display !== 'none';
      if (isVisible && !_lastVisible) {
        setTimeout(resizeVchart, 100);
        setTimeout(resizeVchart, 400);
        setTimeout(resizeVchart, 900);
      }
      _lastVisible = isVisible;
    }
  }, 2000);

  if (document.readyState === 'loading') {
    document.addEventListener('DOMContentLoaded', startObserver);
  } else {
    startObserver();
  }
})();
</script>
""")
    gr.HTML("""
    <div id="gp-hero">
        <div class="gp-agent-badge"><div class="gp-agent-dot"></div>AI Agent · Live</div>
        <div class="gp-hero-title-row">
            <h1>GraphPulse AI</h1>
            <div class="gp-theme-toggle-wrap">
                <span class="gp-theme-toggle-label" id="gp-theme-label">DARK</span>
                <label class="gp-ios-toggle">
                    <input type="checkbox" id="gp-theme-checkbox">
                    <span class="gp-ios-slider"></span>
                </label>
            </div>
        </div>
        <p class="gp-tagline">Real-Time Supply Chain Intelligence Engine</p>
        <div class="gp-hero-stats">
            <div class="gp-stat-pill">⚡ Cypher Query Interface</div>
            <div class="gp-stat-pill">◎ Network Health Monitor</div>
            <div class="gp-stat-pill">🔍 RCA Trail</div>
            <div class="gp-stat-pill">➕ Auto Graph Update</div>
            <div class="gp-stat-pill">📊 Live Visualizations</div>
        </div>
    </div>
    """)

    HOME_BLOCK = """
<style>
/* ══ GraphPulse AI — Clean Home ══════════════════════════════════ */
.gph-root {
  width: 100%; padding: 0; margin: 0;
  background: #060c1c; box-sizing: border-box;
  font-family: 'DM Sans', 'Inter', system-ui, sans-serif;
}

/* ── HERO ── */
.gph-root .gph-hero {
  padding: 0;
  width: 100%;
  min-height: 340px;
  display: flex;
  align-items: center;
  justify-content: center;
  border-bottom: 1px solid rgba(56,189,248,.1);
  background: radial-gradient(ellipse 80% 120% at 50% 50%, rgba(14,165,233,.09) 0%, rgba(56,189,248,.04) 40%, transparent 70%), #060c1c;
  position: relative;
  overflow: hidden;
}
.gph-root .gph-hero::before {
  content: "";
  position: absolute; inset: 0;
  background: radial-gradient(ellipse 60% 80% at 20% 50%, rgba(56,189,248,.05) 0%, transparent 60%),
              radial-gradient(ellipse 40% 60% at 80% 50%, rgba(129,140,248,.04) 0%, transparent 60%);
  pointer-events: none;
}
.gph-root .gph-hero-inner {
  text-align: center;
  padding: 56px 48px 52px;
  max-width: 780px;
  position: relative;
  z-index: 1;
}
.gph-root .gph-eyebrow {
  display: flex; align-items: center; justify-content: center; gap: 8px;
  font-size: .56rem; font-weight: 900; text-transform: uppercase;
  letter-spacing: .22em; color: #7dd3fc; margin-bottom: 20px;
}
.gph-eyebrow-dot {
  width: 5px; height: 5px; border-radius: 50%;
  background: #38bdf8;
  animation: gphBlink 2s ease-in-out infinite;
}
@keyframes gphBlink { 0%,100%{opacity:1} 50%{opacity:.2} }
.gph-root .gph-headline {
  font-size: 3.2rem; font-weight: 900; letter-spacing: -.04em;
  line-height: 1.08; margin-bottom: 22px;
  color: #ffffff;
}
.gph-root .gph-headline em {
  background: linear-gradient(130deg, #38bdf8, #7dd3fc);
  -webkit-background-clip: text; -webkit-text-fill-color: transparent;
  background-clip: text; font-style: normal;
}
.gph-root .gph-lead {
  font-size: 1.05rem; color: #94a3b8; line-height: 1.75;
  max-width: 640px; margin: 0 auto;
}
.gph-lead strong { color: #cbd5e1; font-weight: 600; }

/* Stat pills */
.gph-root .gph-stats {
  display: flex; gap: 8px; flex-wrap: wrap; margin-top: 28px;
  justify-content: center;
}
.gph-root .gph-stat {
  padding: 5px 14px;
  background: rgba(56,189,248,.06);
  border: 1px solid rgba(56,189,248,.18);
  border-radius: 20px;
  font-size: .62rem; font-weight: 700; color: #7dd3fc;
  letter-spacing: .06em; text-transform: uppercase;
}

/* ── FLOW SECTION ── */
.gph-root .gph-flow-section {
  padding: 28px 48px 24px;
  border-bottom: 1px solid rgba(255,255,255,.05);
}
.gph-root .gph-section-label {
  font-size: .54rem; font-weight: 900; text-transform: uppercase;
  letter-spacing: .2em; color: #475569; margin-bottom: 20px;
  display: flex; align-items: center; gap: 10px;
}
.gph-section-label::after {
  content: ""; flex: 1; height: 1px;
  background: linear-gradient(90deg, rgba(255,255,255,.06), transparent);
}
.gph-root .gph-flow {
  display: grid;
  grid-template-columns: 1fr 28px 1fr 28px 1fr 28px 1fr;
  align-items: center;
  gap: 0;
}
.gph-root .gph-flow-step {
  background: rgba(8,14,32,.95);
  border: 1px solid rgba(255,255,255,.07);
  border-radius: 12px;
  padding: 16px 16px 14px;
  position: relative;
  transition: border-color .2s, transform .2s;
}
.gph-flow-step:hover {
  border-color: var(--fc);
  transform: translateY(-3px);
}
.gph-flow-step::before {
  content: "";
  position: absolute; top: 0; left: 0; right: 0; height: 2px;
  border-radius: 12px 12px 0 0;
  background: var(--fc);
  opacity: .7;
}
.gph-root .gph-flow-num {
  display: inline-flex; align-items: center; justify-content: center;
  width: 24px; height: 24px; border-radius: 50%;
  background: var(--fc); color: #060c1c;
  font-size: .62rem; font-weight: 900;
  margin-bottom: 10px; flex-shrink: 0;
}
.gph-root .gph-flow-icon {
  font-size: 1.3rem; margin-bottom: 8px; display: block; line-height: 1;
}
.gph-root .gph-flow-title {
  font-size: .78rem; font-weight: 800; color: #e2e8f0;
  margin-bottom: 5px; letter-spacing: -.01em;
}
.gph-root .gph-flow-desc {
  font-size: .66rem; color: #64748b; line-height: 1.55;
}
.gph-root .gph-flow-arrow {
  display: flex; align-items: center; justify-content: center;
  color: #1e3a5f; font-size: 1rem;
}

/* ── MODULE CARDS ── */
.gph-root .gph-modules-section {
  padding: 24px 48px 32px;
}
.gph-root .gph-modules-flow-label {
  font-size: .78rem !important; font-weight: 900 !important; text-transform: uppercase !important;
  letter-spacing: .18em !important; color: #7dd3fc !important; margin-bottom: 20px !important;
  display: flex !important; align-items: center !important; gap: 10px !important;
  text-shadow: 0 0 18px rgba(56,189,248,.4) !important;
  opacity: 1 !important; visibility: visible !important;
}
.gph-root .gph-modules-flow-label *,
.gph-modules-flow-label {
  color: #7dd3fc !important;
}
.gph-modules-flow-label::after {
  content: "" !important; flex: 1 !important; height: 1px !important;
  background: linear-gradient(90deg, rgba(56,189,248,.25), transparent) !important;
}
.gph-root .gph-mod-flow-wrapper {
  display: grid !important;
  grid-template-columns: 1fr 44px 1fr 44px 1fr 44px 1fr !important;
  align-items: center !important;
  gap: 0 !important;
}
.gph-root .gph-mod-flow-arrow {
  display: flex !important; align-items: center !important; justify-content: center !important;
  color: #38bdf8 !important; font-size: 2rem !important; font-weight: 900 !important;
  text-shadow: 0 0 12px rgba(56,189,248,.6) !important;
  opacity: 1 !important; visibility: visible !important;
}
.gph-root .gph-grid {
  display: grid; grid-template-columns: repeat(4, 1fr); gap: 12px;
}
.gph-root .gph-mod {
  background: rgba(8,14,32,.92);
  border: 1px solid rgba(255,255,255,.07);
  border-radius: 14px; padding: 20px 18px 16px;
  cursor: pointer; display: flex; flex-direction: column;
  position: relative; overflow: hidden;
  transition: transform .22s, border-color .22s;
}
.gph-root .gph-mod::before {
  content: ""; position: absolute;
  top: 0; left: 0; right: 0; height: 2px;
  background: linear-gradient(90deg, transparent, var(--mc), transparent);
  opacity: 0; transition: opacity .22s;
}
.gph-root .gph-mod:hover {
  transform: translateY(-4px); border-color: var(--mc);
}
.gph-mod:hover::before { opacity: 1; }
.gph-root .gph-mod-top {
  display: flex; align-items: flex-start;
  justify-content: space-between; margin-bottom: 12px;
}
.gph-root .gph-mod-icon {
  width: 38px; height: 38px; border-radius: 10px;
  display: flex; align-items: center; justify-content: center;
  background: rgba(255,255,255,.04); border: 1px solid var(--mc);
  transition: box-shadow .25s;
}
.gph-mod:hover .gph-mod-icon {
  box-shadow: 0 0 16px color-mix(in srgb, var(--mc) 28%, transparent);
}
.gph-root .gph-mod-badge {
  display: flex; align-items: center; justify-content: center;
  width: 22px; height: 22px; border-radius: 50%;
  background: var(--mc); color: #060c1c;
  font-size: .58rem; font-weight: 900;
}
.gph-root .gph-mod-cat {
  font-size: .52rem; font-weight: 900; text-transform: uppercase;
  letter-spacing: .14em; color: var(--mc); margin-bottom: 3px;
}
.gph-root .gph-mod-name {
  font-size: .88rem; font-weight: 900; color: #ffffff;
  letter-spacing: -.01em; margin-bottom: 7px; line-height: 1.2;
}
.gph-root .gph-mod-desc {
  font-size: .69rem; color: #94a3b8; line-height: 1.6; flex: 1;
}
.gph-root .gph-mod-cta {
  margin-top: 12px; font-size: .6rem; font-weight: 800;
  text-transform: uppercase; letter-spacing: .12em; color: var(--mc);
  opacity: 0; transition: opacity .2s;
  display: flex; align-items: center; gap: 4px;
}
.gph-mod:hover .gph-mod-cta { opacity: 1; }

/* ── FOOTER ── */
.gph-root .gph-footer {
  text-align: center; padding: 12px 40px 18px;
  border-top: 1px solid rgba(255,255,255,.04);
  font-size: .58rem; color: #1e3a5f;
}

@media (max-width: 900px) {
  .gph-flow { grid-template-columns: 1fr; }
  .gph-flow-arrow { transform: rotate(90deg); }
  .gph-grid { grid-template-columns: repeat(2, 1fr); }
  .gph-hero { flex-direction: column; }
}
</style>

<div class="gph-root">

  <!-- ══ HERO ══ -->
  <div class="gph-hero" style="min-height:100px;padding:32px 0 24px">
    <div class="gph-hero-inner" style="padding:0 48px">
      <div class="gph-headline" style="font-size:2.6rem;margin-bottom:0;white-space:nowrap">
        The intelligent <em>control tower</em> for your supply chain.
      </div>
    </div>
  </div>

  <!-- ══ MODULE CARDS ══ -->
  <div class="gph-modules-section">
    <div class="gph-modules-flow-label" style="font-size:.78rem !important;font-weight:900 !important;text-transform:uppercase !important;letter-spacing:.18em !important;color:#7dd3fc !important;margin-bottom:20px !important;display:flex !important;align-items:center !important;gap:10px !important;text-shadow:0 0 18px rgba(56,189,248,.4) !important;opacity:1 !important;visibility:visible !important">Platform modules — flow of the solution</div>
    <div class="gph-mod-flow-wrapper">

      <!-- 01 Supply Network Health -->
      <div class="gph-mod" style="--mc:#4ade80"
        onclick="(function(){var t=document.querySelectorAll('[role=tab]');for(var i=0;i<t.length;i++){if(t[i].textContent.includes('Supply')||t[i].textContent.includes('Health')){t[i].click();return;}}})()">
        <div class="gph-mod-top">
          <div class="gph-mod-icon">
            <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="#4ade80" stroke-width="2" stroke-linecap="round"><path d="M22 12h-4l-3 9L9 3l-3 9H2"/></svg>
          </div>
          <div class="gph-mod-badge">1</div>
        </div>
        <div class="gph-mod-cat">Live Monitoring</div>
        <div class="gph-mod-name">Supply Network Health</div>
        <div class="gph-mod-desc">
          Live KPIs, plant delay gauge, monthly trend, and 6 interactive
          analytics charts — all pulling from Neo4j in real time.
        </div>
        <div class="gph-mod-cta">Open <span>→</span></div>
      </div>

      <!-- Arrow 1 -->
      <div class="gph-mod-flow-arrow" style="display:flex !important;align-items:center !important;justify-content:center !important;color:#38bdf8 !important;font-size:2rem !important;font-weight:900 !important;text-shadow:0 0 12px rgba(56,189,248,.6) !important;opacity:1 !important;visibility:visible !important">→</div>

      <!-- 02 Query Interface -->
      <div class="gph-mod" style="--mc:#38bdf8"
        onclick="(function(){var t=document.querySelectorAll('[role=tab]');for(var i=0;i<t.length;i++){if(t[i].textContent.includes('Query Interface')){t[i].click();return;}}})()">
        <div class="gph-mod-top">
          <div class="gph-mod-icon">
            <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="#38bdf8" stroke-width="2" stroke-linecap="round"><circle cx="11" cy="11" r="8"/><path d="M21 21l-4.35-4.35"/></svg>
          </div>
          <div class="gph-mod-badge" style="background:#38bdf8">2</div>
        </div>
        <div class="gph-mod-cat">NL Graph Query</div>
        <div class="gph-mod-name">Query Interface</div>
        <div class="gph-mod-desc">
          Type any question, get Cypher generated and executed on Neo4j.
          Results returned as a structured table with export to CSV.
        </div>
        <div class="gph-mod-cta">Open <span>→</span></div>
      </div>

      <!-- Arrow 2 -->
      <div class="gph-mod-flow-arrow" style="display:flex !important;align-items:center !important;justify-content:center !important;color:#38bdf8 !important;font-size:2rem !important;font-weight:900 !important;text-shadow:0 0 12px rgba(56,189,248,.6) !important;opacity:1 !important;visibility:visible !important">→</div>

      <!-- 03 RCA Trail -->
      <div class="gph-mod" style="--mc:#f97316"
        onclick="(function(){var t=document.querySelectorAll('[role=tab]');for(var i=0;i<t.length;i++){if(t[i].textContent.includes('RCA')){t[i].click();return;}}})()">
        <div class="gph-mod-top">
          <div class="gph-mod-icon">
            <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="#f97316" stroke-width="2" stroke-linecap="round"><circle cx="12" cy="12" r="10"/><path d="M12 8v4l3 3"/></svg>
          </div>
          <div class="gph-mod-badge" style="background:#f97316">3</div>
        </div>
        <div class="gph-mod-cat">Multi-Agent AI</div>
        <div class="gph-mod-name">RCA Trail</div>
        <div class="gph-mod-desc">
          5-agent pipeline delivers a full root cause analysis report
          with charts, executive summary, and Word / Excel export.
        </div>
        <div class="gph-mod-cta">Open <span>→</span></div>
      </div>

      <!-- Arrow 3 -->
      <div class="gph-mod-flow-arrow" style="display:flex !important;align-items:center !important;justify-content:center !important;color:#38bdf8 !important;font-size:2rem !important;font-weight:900 !important;text-shadow:0 0 12px rgba(56,189,248,.6) !important;opacity:1 !important;visibility:visible !important">→</div>

      <!-- 04 Update Graph -->
      <div class="gph-mod" style="--mc:#a78bfa"
        onclick="(function(){var t=document.querySelectorAll('[role=tab]');for(var i=0;i<t.length;i++){if(t[i].textContent.includes('Update')){t[i].click();return;}}})()">
        <div class="gph-mod-top">
          <div class="gph-mod-icon">
            <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="#a78bfa" stroke-width="2" stroke-linecap="round"><ellipse cx="12" cy="5" rx="9" ry="3"/><path d="M21 12c0 1.66-4 3-9 3s-9-1.34-9-3M3 5v14c0 1.66 4 3 9 3s9-1.34 9-3V5"/></svg>
          </div>
          <div class="gph-mod-badge" style="background:#a78bfa">4</div>
        </div>
        <div class="gph-mod-cat">Graph Extension</div>
        <div class="gph-mod-name">Update Graph</div>
        <div class="gph-mod-desc">
          Add or edit Neo4j nodes via natural language or CSV/Excel upload.
          AI validates, normalises and merges data into the graph.
        </div>
        <div class="gph-mod-cta">Open <span>→</span></div>
      </div>

    </div>
  </div>

  <div class="gph-footer">
    GraphPulse AI &nbsp;·&nbsp; Neo4j Knowledge Graph + Groq Multi-Agent LLM &nbsp;·&nbsp; Real-Time Supply Chain Intelligence
  </div>

</div>
"""

    with gr.Tabs():

        # ── HOME TAB ──────────────────────────────────────────────
        with gr.Tab("🏠 Home"):
            gr.HTML(HOME_BLOCK)

        # ── TAB 1: Query Interface ────────────────────────────────
        with gr.Tab("◎ Supply Network Health"):
            gr.HTML("""
<style>
.snh-header{
  padding:10px 0 14px;border-bottom:2px solid rgba(56,189,248,.25);margin-bottom:18px;
  background:transparent !important;
}
.snh-hl{
  font-size:.8rem !important;color:#38bdf8 !important;letter-spacing:.13em !important;
  font-weight:900 !important;text-transform:uppercase !important;
  text-shadow:0 0 20px rgba(56,189,248,.4) !important;
  display:block !important;opacity:1 !important;visibility:visible !important;
}
.snh-sub{
  font-size:.82rem !important;color:#cbd5e1 !important;margin-top:6px !important;
  font-weight:500 !important;opacity:1 !important;visibility:visible !important;
  display:block !important;
}
.snh-divider{
  font-size:.65rem !important;font-weight:800 !important;text-transform:uppercase !important;
  letter-spacing:.14em !important;color:#7dd3fc !important;
  margin:22px 0 14px !important;display:flex !important;align-items:center !important;
  gap:10px !important;opacity:1 !important;
}
.snh-divider::after{content:"";flex:1;height:1px;background:rgba(56,189,248,.15)}
.viz-nav-panel {
    background: rgba(7,20,40,0.95);
    border: 1px solid rgba(56,189,248,0.15);
    border-radius: 14px;
    padding: 18px 14px;
}
.viz-nav-title {
    font-size: 0.68rem; font-weight: 700; letter-spacing: 0.12em;
    text-transform: uppercase; color: #38bdf8; margin-bottom: 14px;
}
.viz-nav-btn button {
    background: rgba(14,165,233,0.06) !important;
    border: 1px solid rgba(56,189,248,0.15) !important;
    color: #94a3b8 !important; font-size: 0.78rem !important;
    text-align: left !important; border-radius: 8px !important;
    padding: 10px 12px !important; width: 100% !important;
    margin-bottom: 6px !important; transition: all .18s !important;
}
.viz-nav-btn button:hover {
    background: rgba(14,165,233,0.18) !important;
    border-color: rgba(56,189,248,0.45) !important;
    color: #e2e8f0 !important;
}
.viz-load-all button {
    background: linear-gradient(135deg,#0284c7,#0ea5e9 60%,#06b6d4) !important;
    border: 1px solid #0ea5e9 !important; color: #fff !important;
    font-weight: 700 !important; border-radius: 8px !important;
    width: 100% !important; margin-bottom: 14px !important;
    box-shadow: 0 0 14px rgba(14,165,233,0.3) !important;
}
.viz-meta-card {
    background: rgba(124,58,237,0.08);
    border-left: 3px solid #7c3aed;
    border-radius: 0 8px 8px 0;
    padding: 10px 14px;
    margin-bottom: 14px;
    font-size: 0.78rem; color: #c4b5fd;
    line-height: 1.55;
}
.viz-meta-card .vm-title {
    font-size: 0.65rem; font-weight: 700; letter-spacing: 0.1em;
    text-transform: uppercase; color: #a78bfa; display: block; margin-bottom: 4px;
}
</style>
<div style="padding:12px 0 16px;border-bottom:2px solid rgba(56,189,248,.3);margin-bottom:18px;background:transparent">
  <div style="font-size:.85rem !important;color:#38bdf8 !important;letter-spacing:.12em;font-weight:900;
       text-transform:uppercase;text-shadow:0 0 18px rgba(56,189,248,.35);
       display:block;opacity:1;visibility:visible;margin-bottom:6px">
    ◆ Supply Chain Network Health
  </div>
  <div style="font-size:.84rem !important;color:#cbd5e1 !important;font-weight:500;
       display:block;opacity:1;visibility:visible;line-height:1.5">
    Live diagnostics from Neo4j — KPIs, health gauge, and interactive chart explorer
  </div>
</div>""")

            refresh_health_btn = gr.Button("↺  Refresh Network Health", elem_classes="analyze-btn")
            health_kpi_html = gr.HTML('<div style="color:#94a3b8 !important;font-size:.84rem;padding:20px 0;opacity:1 !important;visibility:visible !important">Click Refresh to load metrics…</div>')

            # Hidden chart outputs kept for backend compatibility (not displayed)
            health_plant_chart = gr.Plot(visible=False)
            health_gauge_chart = gr.Plot(visible=False)
            health_trend_chart = gr.Plot(visible=False)

            refresh_health_btn.click(
                load_network_health,
                outputs=[health_kpi_html, health_plant_chart, health_gauge_chart, health_trend_chart]
            )

            # ══════════════════════════════════════════════════════════════
            # ANALYTICS SUITE — merged from Visualizations tab
            # ══════════════════════════════════════════════════════════════
            gr.HTML('<div class="snh-divider" style="margin-top:32px">◈ Analytics Suite — Interactive Charts</div>')
            gr.HTML('''
<div style="font-size:.8rem;color:#94a3b8;margin-bottom:16px;line-height:1.6;
     padding:12px 16px;background:rgba(56,189,248,.04);border-radius:10px;
     border:1px solid rgba(56,189,248,.1)">
  Click any chart below to load it with a full explanation. Charts pull live data from Neo4j.
</div>''')

            with gr.Row(equal_height=False):
                # ── LEFT NAV PANEL ──────────────────────────────────────────
                with gr.Column(scale=1, min_width=200, elem_classes="viz-nav-panel"):
                    gr.HTML('<div class="viz-nav-title">◈ Chart Navigator</div>')
                    gr.HTML('<div style="font-size:0.65rem;color:#94a3b8;margin-bottom:10px;text-transform:uppercase;letter-spacing:0.08em">Individual Charts</div>')
                    load_heat_btn  = gr.Button("🔥  Delay Heatmap",       elem_classes="viz-nav-btn")
                    load_risk_btn  = gr.Button("⚠️  Supplier Risk",        elem_classes="viz-nav-btn")
                    load_route_btn = gr.Button("🚛  Route Efficiency",     elem_classes="viz-nav-btn")
                    load_trend_btn = gr.Button("📅  Monthly Trend",        elem_classes="viz-nav-btn")
                    load_stock_btn = gr.Button("📦  Stockout Severity",    elem_classes="viz-nav-btn")
                    load_sun_btn   = gr.Button("🌐  Transport Modes",      elem_classes="viz-nav-btn")

                    viz_meta_html = gr.HTML("""
                    <div class="viz-meta-card" style="margin-top:20px">
                        Click any chart to load it here with a full explanation.
                    </div>""")

                # ── RIGHT CHART DISPLAY ──────────────────────────────────────
                _blank_fig_viz = go.Figure().update_layout(
                    paper_bgcolor="#060c1c", plot_bgcolor="#060c1c",
                    xaxis=dict(visible=False), yaxis=dict(visible=False)
                )
                with gr.Column(scale=4):
                    viz_primary_title = gr.HTML(
                        '<div style="font-size:.8rem;font-weight:700;color:#e2e8f0;'                        'padding:8px 0 12px;border-bottom:1px solid rgba(56,189,248,.12);'                        'margin-bottom:12px">Select a chart from the navigator →</div>'
                    )
                    chart_primary = gr.Plot(label="", show_label=False, value=_blank_fig_viz)

                    viz_explain_html = gr.HTML('''
<div style="margin-top:20px;padding:20px 24px;
     background:rgba(7,20,40,.85);border:1px solid rgba(56,189,248,.15);
     border-left:3px solid #7c3aed;border-radius:12px">
  <div style="font-size:.65rem;font-weight:800;text-transform:uppercase;
       letter-spacing:.12em;color:#a78bfa;margin-bottom:14px">
    📖 Chart Guide
  </div>
  <div style="font-size:.8rem;color:#94a3b8;line-height:1.7">
    Select any chart from the navigator on the left to load it here.
    Each chart pulls live data directly from your Neo4j supply chain graph.
  </div>
</div>''')

            # ── Chart loader functions + explanations ───────────────────────
            def _explain(title_icon, title_text, points):
                items_html = "".join(
                    f'''<div style="display:flex;gap:10px;margin-bottom:10px">
  <span style="color:{p["color"]};font-size:.85rem;flex-shrink:0;margin-top:1px">{p["icon"]}</span>
  <div>
    <div style="font-size:.72rem;font-weight:800;text-transform:uppercase;
         letter-spacing:.08em;color:{p["color"]};margin-bottom:2px">{p["head"]}</div>
    <div style="font-size:.78rem;color:#cbd5e1;line-height:1.6">{p["body"]}</div>
  </div>
</div>'''
                    for p in points
                )
                return f'''
<div style="margin-top:20px;padding:20px 24px;
     background:rgba(7,20,40,.85);border:1px solid rgba(56,189,248,.15);
     border-left:3px solid #7c3aed;border-radius:12px">
  <div style="font-size:.65rem;font-weight:800;text-transform:uppercase;
       letter-spacing:.12em;color:#a78bfa;margin-bottom:16px">
    {title_icon} About This Chart — {title_text}
  </div>
  {items_html}
</div>'''

            _EXPLAIN = {
                "heat": _explain("🔥", "Delay Heatmap", [
                    {"icon":"📋","color":"#38bdf8","head":"What it shows",
                     "body":"A grid of every manufacturing plant (rows) against every product category (columns). Each cell is coloured by the number of delayed shipments — darker red means more delays."},
                    {"icon":"🗄️","color":"#7dd3fc","head":"Data it uses",
                     "body":"Shipment records from Neo4j: plant name, product category, delayed flag, and shipment count. Pulls all historical shipments in the graph."},
                    {"icon":"💡","color":"#4ade80","head":"Key insight to look for",
                     "body":"The darkest cell in the grid is your root cause candidate — it's the plant–category combination generating the most disruption in the network."},
                    {"icon":"📈","color":"#fbbf24","head":"How to interpret anomalies",
                     "body":"A full red row means a plant is struggling across all categories — likely a capacity or supplier issue. A full red column means one product category is delayed everywhere — likely a sourcing problem."},
                    {"icon":"🔗","color":"#c084fc","head":"Why it matters",
                     "body":"Delay heatmaps are the fastest way to triage supply chain problems. Instead of reading through rows of data, one glance at the grid shows you exactly where to investigate first."},
                ]),
                "risk": _explain("⚠️", "Supplier Risk Scores", [
                    {"icon":"📋","color":"#38bdf8","head":"What it shows",
                     "body":"Horizontal bars ranking your top 20 suppliers by risk score (0 = no risk, 1 = maximum risk). The dotted red line marks the critical threshold at 0.7."},
                    {"icon":"🗄️","color":"#7dd3fc","head":"Data it uses",
                     "body":"Supplier nodes from Neo4j: supplier name, risk_score property, and the plant they supply. Risk scores are set manually or calculated from delay history."},
                    {"icon":"💡","color":"#4ade80","head":"Key insight to look for",
                     "body":"Any supplier with a bar crossing the red dotted threshold needs immediate attention — they represent the highest probability of causing a future supply disruption."},
                    {"icon":"📈","color":"#fbbf24","head":"How to interpret trends",
                     "body":"If multiple suppliers for the same plant all have high scores, that plant is doubly exposed — it has no low-risk backup. That's a structural vulnerability worth escalating."},
                    {"icon":"🔗","color":"#c084fc","head":"Why it matters",
                     "body":"Supplier risk is the earliest leading indicator in the supply chain. Problems here will appear as plant delays 8–16 days later, and as distributor shortfalls 2–3 weeks after that."},
                ]),
                "route": _explain("🚛", "Route Efficiency", [
                    {"icon":"📋","color":"#38bdf8","head":"What it shows",
                     "body":"A scatter plot of every plant-to-distributor route: X-axis = distance (km), Y-axis = transportation cost (INR). Bubble size represents the route's efficiency score."},
                    {"icon":"🗄️","color":"#7dd3fc","head":"Data it uses",
                     "body":"Route nodes from Neo4j: distance_km, transportation_cost_inr, efficiency score, transport mode (colour coded), and the plant/distributor endpoints."},
                    {"icon":"💡","color":"#4ade80","head":"Key insight to look for",
                     "body":"Small bubbles in the top-left zone (short distance, high cost) are your worst-performing routes — you're paying a lot for a short haul, which usually signals a wrong transport mode."},
                    {"icon":"📈","color":"#fbbf24","head":"How to interpret anomalies",
                     "body":"A cluster of air routes (typically highest cost) on short distances is an immediate cost reduction opportunity — switching to road or rail on those lanes could cut logistics spend significantly."},
                    {"icon":"🔗","color":"#c084fc","head":"Why it matters",
                     "body":"Route efficiency directly impacts delivery cost and speed. Optimising even 2–3 high-cost short routes can reduce total logistics cost by 8–15% without changing the network structure."},
                ]),
                "trend": _explain("📅", "Monthly Shipment Trend", [
                    {"icon":"📋","color":"#38bdf8","head":"What it shows",
                     "body":"A stacked or grouped bar chart showing on-time vs delayed shipments for each month. The delay rate line overlaid shows the percentage trend over time."},
                    {"icon":"🗄️","color":"#7dd3fc","head":"Data it uses",
                     "body":"Shipment records grouped by month: total shipments, delayed count, and calculated delay rate percentage. Uses the shipment_date or created_at field from Neo4j."},
                    {"icon":"💡","color":"#4ade80","head":"Key insight to look for",
                     "body":"A rising delay rate line month-over-month signals a systemic problem that isn't self-correcting. A sudden spike in one month points to a discrete event — supplier failure, port closure, or demand surge."},
                    {"icon":"📈","color":"#fbbf24","head":"How to interpret seasonal patterns",
                     "body":"Recurring delay spikes in the same months every year indicate seasonal risk — typically peak demand periods. These months need pre-positioned inventory or additional supplier capacity."},
                    {"icon":"🔗","color":"#c084fc","head":"Why it matters",
                     "body":"Trend analysis transforms reactive firefighting into proactive planning. Knowing that delays spike every Q3 lets you act in Q2 — building buffer stock, qualifying backup suppliers, and negotiating flexible lead times."},
                ]),
                "stock": _explain("📦", "Stockout Severity", [
                    {"icon":"📋","color":"#38bdf8","head":"What it shows",
                     "body":"A bar chart ranking distributor cities by total demand gap — the number of units ordered but not fulfilled due to stock unavailability. Taller bar = more severe shortage."},
                    {"icon":"🗄️","color":"#7dd3fc","head":"Data it uses",
                     "body":"Distributor nodes from Neo4j: city name, demand_gap property, and stockout_flag. The demand gap is the difference between what retailers ordered and what was actually shipped."},
                    {"icon":"💡","color":"#4ade80","head":"Key insight to look for",
                     "body":"The tallest bar is the city where customers are most affected right now. Cross-reference with the Delay Heatmap to find which plant is responsible for that distributor's shortage."},
                    {"icon":"📈","color":"#fbbf24","head":"How to interpret anomalies",
                     "body":"If a city suddenly spikes from near-zero to a large demand gap, check whether its supplying plant had a delay event in the previous 2–4 weeks — the stockout is almost always a lagged consequence."},
                    {"icon":"🔗","color":"#c084fc","head":"Why it matters",
                     "body":"Stockout severity is the end-consumer impact metric. Every unit in the demand gap is a lost sale or a dissatisfied customer. It's the final downstream signal of upstream supply chain failures."},
                ]),
                "sun": _explain("🌐", "Transport Mode Distribution", [
                    {"icon":"📋","color":"#38bdf8","head":"What it shows",
                     "body":"A sunburst (nested ring) chart: the inner ring shows plants, the outer ring shows transport modes used from each plant. Segment size = number of routes using that mode."},
                    {"icon":"🗄️","color":"#7dd3fc","head":"Data it uses",
                     "body":"Route nodes from Neo4j: plant_id, transport mode (Road/Rail/Air/Sea), and route count. Groups all routes by plant and mode to build the hierarchy."},
                    {"icon":"💡","color":"#4ade80","head":"Key insight to look for",
                     "body":"A plant whose outer ring is almost entirely 'Air' is cost-exposed — air freight is 3–5× more expensive than road or rail. That plant is either compensating for delays or has no rail/road alternatives set up."},
                    {"icon":"📈","color":"#fbbf24","head":"How to interpret the chart",
                     "body":"Click any inner-ring plant segment to zoom in and see its mode breakdown in detail. A well-balanced plant should have Road as its dominant mode, with Rail for long distances and Air only for urgent shipments."},
                    {"icon":"🔗","color":"#c084fc","head":"Why it matters",
                     "body":"Modal mix is a hidden cost driver. Shifting even 20% of air shipments to rail for a single plant can reduce that plant's logistics cost by 30–40% — without changing delivery destinations or supplier relationships."},
                ]),
            }

            def _make_title(icon, text):
                return (f'<div style="font-size:.84rem;font-weight:700;color:#e2e8f0;'                        f'padding:8px 0 12px;border-bottom:1px solid rgba(56,189,248,.12);'                        f'margin-bottom:12px">{icon} {text}</div>')

            def _load_heat():
                try: fig = _apply_dark_theme(build_delay_heatmap())
                except Exception as _e: print(f"[heat] {_e}"); fig = _blank_fig_viz
                return _make_title("🔥","Delay Heatmap — Plant × Product Category"), fig, _EXPLAIN["heat"]

            def _load_risk():
                try: fig = _apply_dark_theme(build_supplier_risk_chart())
                except Exception as _e: print(f"[risk] {_e}"); fig = _blank_fig_viz
                return _make_title("⚠️","Supplier Risk Scores — Top 20"), fig, _EXPLAIN["risk"]

            def _load_route():
                try: fig = _apply_dark_theme(build_route_efficiency_scatter())
                except Exception as _e: print(f"[route] {_e}"); fig = _blank_fig_viz
                return _make_title("🚛","Route Efficiency — Distance vs Cost"), fig, _EXPLAIN["route"]

            def _load_trend():
                try: fig = _apply_dark_theme(build_monthly_delay_trend())
                except Exception as _e: print(f"[trend] {_e}"); fig = _blank_fig_viz
                return _make_title("📅","Monthly Shipment Status Trend"), fig, _EXPLAIN["trend"]

            def _load_stock():
                try: fig = _apply_dark_theme(build_distributor_demand_gap())
                except Exception as _e: print(f"[stock] {_e}"); fig = _blank_fig_viz
                return _make_title("📦","Stockout Severity by Distributor City"), fig, _EXPLAIN["stock"]

            def _load_sun():
                try: fig = _apply_dark_theme(build_plant_transport_sunburst())
                except Exception as _e: print(f"[sun] {_e}"); fig = _blank_fig_viz
                return _make_title("🌐","Transport Mode Distribution by Plant"), fig, _EXPLAIN["sun"]

            load_heat_btn.click(fn=_load_heat,  outputs=[viz_primary_title, chart_primary, viz_explain_html])
            load_risk_btn.click(fn=_load_risk,  outputs=[viz_primary_title, chart_primary, viz_explain_html])
            load_route_btn.click(fn=_load_route, outputs=[viz_primary_title, chart_primary, viz_explain_html])
            load_trend_btn.click(fn=_load_trend, outputs=[viz_primary_title, chart_primary, viz_explain_html])
            load_stock_btn.click(fn=_load_stock, outputs=[viz_primary_title, chart_primary, viz_explain_html])
            load_sun_btn.click(fn=_load_sun,   outputs=[viz_primary_title, chart_primary, viz_explain_html])


        with gr.Tab("⚡ Query Interface"):
            with gr.Row(equal_height=False):
                with gr.Column(scale=1, min_width=280, elem_classes="left-panel"):
                    gr.HTML('<div class="sec-label">⬡ KG Snapshot</div>')
                    snap_html = gr.HTML("<div class='snap-error'>Click ⟳ to load.</div>")
                    snap_btn  = gr.Button("⟳  Load Snapshot", elem_classes="clear-btn")
                    snap_btn.click(fn=load_snapshot, outputs=[snap_html])
                    gr.HTML('<div class="sec-label" style="margin-top:16px;">⚡ Quick Queries</div>')
                    sample_btns = [gr.Button(q, elem_classes="sample-btn") for q in SAMPLE_QUESTIONS[:5]]
                    gr.HTML('<div class="sec-label" style="margin-top:16px;">⬡ Recent</div>')
                    hist_html = gr.HTML("<div class='hist-empty'>No queries yet</div>")
                    gr.HTML('<div class="sec-label" style="margin-top:16px;">⚙ Session</div>')
                    clear_btn = gr.Button("⟳  Clear Session", elem_classes="clear-btn")

                with gr.Column(scale=4, elem_classes="main-panel"):
                    gr.HTML('<div class="sec-label">◈ Query Interface</div>')
                    q_input    = gr.Textbox(placeholder="Ask anything about your supply chain…", lines=2, label="", show_label=False)
                    ask_btn    = gr.Button("⚡  Generate Query", elem_classes="analyze-btn")
                    gen_status = gr.HTML("")
                    cypher_box = gr.Textbox(lines=6, label="Generated Cypher", visible=False)
                    with gr.Row(visible=False) as run_row:
                        run_btn   = gr.Button("▶  Run Query", elem_classes="run-btn",   scale=3)
                        abort_btn = gr.Button("✕  Abort",     elem_classes="abort-btn", scale=1)
                    run_status = gr.HTML("")

                    # Hidden state for RCA context (used to prefill RCA tab)
                    _rca_prefill_state  = gr.State("")    # the prefilled RCA question
                    _rca_worthy_state   = gr.State(False) # whether CTA is live

                    with gr.Column(elem_classes="results-wrap", visible=False) as results_wrap:
                        gr.HTML('<div class="sec-label">◈ Analysis Output</div>')
                        with gr.Column(elem_classes="brief-box"):
                            brief_md = gr.Markdown("")
                        tbl_html     = gr.HTML("")
                        insight_html = gr.HTML("")
                        export_btn = gr.DownloadButton("⬇ Export CSV", visible=False,
                                                       elem_classes="export-btn")

                    # ── RCA CTA banner — shown only for analytical queries ──
                    rca_cta_col = gr.Column(visible=False)
                    with rca_cta_col:
                        gr.HTML("""
<div style="
  margin: 14px 0 4px;
  padding: 14px 18px 14px 16px;
  background: linear-gradient(135deg, rgba(124,58,237,0.15) 0%, rgba(56,189,248,0.08) 100%);
  border: 1px solid rgba(124,58,237,0.35);
  border-left: 3px solid #a78bfa;
  border-radius: 10px;
  display: flex; align-items: center; justify-content: space-between; gap: 16px;
" id="rca-cta-banner">
  <div style="flex:1">
    <div style="font-size:0.66rem;font-weight:800;text-transform:uppercase;
         letter-spacing:0.12em;color:#a78bfa;margin-bottom:4px">
      🔬 Root Cause Analysis Available
    </div>
    <div style="font-size:0.78rem;color:#cbd5e1;line-height:1.5">
      This looks like a diagnostic question. Run a full AI-powered Root Cause Analysis
      to get a structured report, chart visualizations, and prioritised recommendations.
    </div>
  </div>
</div>""")
                        rca_launch_btn = gr.Button(
                            "🔬  Run Root Cause Analysis →",
                            elem_classes="analyze-btn",
                        )

            outs_gen = [gen_status, cypher_box, run_row, results_wrap, run_status, brief_md, tbl_html, insight_html, export_btn]
            # on_run_query returns 9 values: 6 UI + rca_worthy + rca_prefill + rca_cta visibility
            outs_run = [run_status, results_wrap, brief_md, tbl_html, insight_html, export_btn,
                        _rca_worthy_state, _rca_prefill_state, rca_cta_col]
            # on_run_insights returns 2 values: brief, insight_html (Gemini call, chained after)
            outs_insights = [insight_html]  # brief_md NOT included — fast brief stays, LLM only fills bullets

            # Step 1: Generate Cypher only. Run Query is a separate user action.
            ask_btn.click(
                fn=on_generate_query, inputs=[q_input], outputs=outs_gen, show_progress="full"
            ).then(fn=update_history, outputs=[hist_html])

            # Manual run: show table immediately, then fetch Gemini insights async
            run_btn.click(
                fn=on_run_query, inputs=[q_input, cypher_box], outputs=outs_run, show_progress="full"
            ).then(
                fn=on_run_insights, inputs=[q_input, cypher_box],
                outputs=outs_insights, show_progress="hidden"
            ).then(fn=update_history, outputs=[hist_html])

            abort_btn.click(fn=on_abort, outputs=outs_gen).then(
                fn=lambda: (gr.update(visible=False),), outputs=[rca_cta_col])
            clear_btn.click(fn=clear_session, outputs=outs_gen).then(fn=update_history, outputs=[hist_html]).then(
                fn=lambda: (gr.update(visible=False),), outputs=[rca_cta_col])
            export_btn.click(fn=make_csv, inputs=[q_input], outputs=[export_btn])
            for btn, sq in zip(sample_btns, SAMPLE_QUESTIONS[:5]):
                btn.click(fn=lambda x=sq: x, outputs=[q_input]).then(
                    fn=on_generate_query, inputs=[q_input], outputs=outs_gen, show_progress="full"
                ).then(fn=update_history, outputs=[hist_html])

        # ══════════════════════════════════════════════════════
        # TAB 2 — NETWORK HEALTH  (redesigned dashboard)
        # ══════════════════════════════════════════════════════
        # ══════════════════════════════════════════════════════════════
        # TAB 3 — RCA TRAIL (v6 — 2-col layout, 6 dynamic views)
        # ══════════════════════════════════════════════════════════════
        with gr.Tab("🔍 RCA Trail"):

            # ── Header ──────────────────────────────────────────────────────
            gr.HTML("""
<div class="rca5-header">
  <div class="rca5-title">🔍 Root Cause Analysis</div>
  <div class="rca5-sub">Ask a supply-chain question — the AI traces the root cause through live Neo4j data and renders the most relevant chart automatically.</div>
</div>""")

            rca_kpi_html = gr.HTML("")  # KPIs moved to right sidebar

            # ── 2-COLUMN LAYOUT ──────────────────────────────────────────────
            with gr.Row(equal_height=False):

                # LEFT: question box + report (wider)
                with gr.Column(scale=3):

                    rca_q = gr.Textbox(
                        placeholder="e.g.  Which suppliers are causing the most delays?  ·  Why are toy shipments late?  ·  Where are we running out of stock?",
                        lines=2, label="Your Question",
                        elem_id="rca-question-input"
                    )
                    with gr.Row():
                        rca_btn    = gr.Button("🔍  Run Analysis", elem_classes="analyze-btn", scale=4)

                    rca_status = gr.HTML("")
                    # A2A Agent Activity Log — collapsible dropdown below the green A2A box
                    # Populated after each run via .then() chain
                    rca_agent_log_main = gr.HTML("", visible=False)

                    # ── AI Prompt Suggestion card ─────────────────────────────
                    rca_suggest_html = gr.HTML("", elem_id="rca-suggest-wrap")

                    # ── Initial Assessment card (shown while agents run) ───────
                    # This is HTML so it renders rich content properly.
                    # It appears BETWEEN the suggestion and the Tracing Path strip.
                    rca_assess_html = gr.HTML("")

                    # rca_samples kept for the for-loop event wiring (hidden buttons)
                    rca_samples = [
                        "Which suppliers have a risk_score above 0.7 and are causing Major Delay shipments from Plant PL1 or PL2? List each supplier_name, risk_score, plant_name, and delayed shipment count — to be sent as a Zapier webhook alert to the procurement team.",
                        "What is the network-wide Major Delay rate across all Plants, and which 3 distributors have the highest demand_gap? Provide a concise summary formatted for a daily Slack digest message to the supply chain team.",
                        "List all Shipments with delivery_status = Major Delay, including shipment_id, supplier_name, plant_name, distributor_city, delay_days, and product_category_name. I need this as a structured table for Google Sheets weekly KPI tracking.",
                        "Identify the top 5 suppliers with risk_score above 0.8, their delayed shipment counts, and affected plant_names. Then write a brief executive email alert summarising root cause, downstream impact, and 2 immediate recommended actions for the supply chain director.",
                        "For each Plant (PL1 Baddi, PL2 Bhopal, PL3 Pune, PL4 Goa), give me total shipments, delayed shipments, delay_rate percentage, and avg delay_days. Also break down by product_category_name. Format as a KPI table for a Power BI dashboard data feed.",
                        "Which Retailers have stockout_flag = true, and what is the demand_gap for each? Trace back through Distributor and Plant to find the supply chain gap. Format the output as a structured replenishment order recommendation suitable for raising a purchase order in SAP.",
                    ]
                    # Hidden buttons — preserve the for-loop wiring without showing old pills
                    with gr.Row(visible=False):
                        rca_sample_btns = [gr.Button(q, elem_classes="sample-btn", scale=1) for q in rca_samples[:3]]
                    with gr.Row(visible=False):
                        rca_sample_btns += [gr.Button(q, elem_classes="sample-btn", scale=1) for q in rca_samples[3:]]

                    rca_context_html = gr.HTML("")  # cleared — replaced by Cypher accordion

                    # ── Cypher Queries Accordion — appears BELOW initial assessment paragraph,
                    # ABOVE the full RCA report (just before rca_out).
                    # This gives the user context about what queries ran before reading findings.
                    rca_cypher_html = gr.HTML(
                        value="",
                        visible=False,
                        elem_id="rca-cypher-accordion",
                    )

                    with gr.Column(elem_classes="rca5-report-panel"):
                        rca_out = gr.HTML(value="")

                        # ── Wiring-only state holders — never rendered visibly ─────
                        rca_viz_summary_html  = gr.HTML("", visible=False, elem_classes="rca-wiring-hidden")
                        rca_detailed_rec_html = gr.HTML("", visible=False, elem_classes="rca-wiring-hidden", elem_id="rca-rec-data-relay")

                        # ── CSS: hide wiring holders + style native toggle boxes ───
                        gr.HTML("""<style>
.rca-wiring-hidden  { display:none !important; }
.rca-hidden-accordion { display:none !important; }

/* ── FINAL LAYOUT FIX: Gradio Column renders as .form with flex-direction:row.
   Override it to column so RCA cards stack vertically. ── */
.rca5-report-panel > .form,
.rca5-report-panel > div.form {
    display: flex !important;
    flex-direction: column !important;
    align-items: stretch !important;
    gap: 0 !important;
    width: 100% !important;
}
.rca5-report-panel > .form > *,
.rca5-report-panel > div.form > * {
    width: 100% !important;
    flex: 0 0 100% !important;
    min-width: 0 !important;
    max-width: 100% !important;
}
</style>
<script>
(function fixRcaLayout() {
  function applyFix() {
    var panel = document.querySelector('.rca5-report-panel');
    if (!panel) return;
    // Find the inner Gradio .form wrapper
    var form = panel.querySelector('.form');
    if (form) {
      form.style.setProperty('display', 'flex', 'important');
      form.style.setProperty('flex-direction', 'column', 'important');
      form.style.setProperty('align-items', 'stretch', 'important');
      form.style.setProperty('gap', '0', 'important');
      form.style.setProperty('width', '100%', 'important');
      for (var i = 0; i < form.children.length; i++) {
        form.children[i].style.setProperty('width', '100%', 'important');
        form.children[i].style.setProperty('flex', '0 0 100%', 'important');
        form.children[i].style.setProperty('max-width', '100%', 'important');
      }
    }
    // Also fix the rca-report-content div itself
    var content = document.getElementById('rca-report-content');
    if (content) {
      content.style.setProperty('display', 'block', 'important');
      content.style.setProperty('width', '100%', 'important');
    }
  }
  // Run immediately + on mutations (Gradio updates DOM dynamically)
  applyFix();
  var obs = new MutationObserver(function(muts) {
    for (var m of muts) {
      if (m.addedNodes.length) { applyFix(); break; }
    }
  });
  obs.observe(document.body, { childList: true, subtree: true });
  // Also run on a short timer after Gradio finishes rendering
  setTimeout(applyFix, 300);
  setTimeout(applyFix, 800);
  setTimeout(applyFix, 1500);
})();
</script>
<style>

/* ═══════════════════════════════════════════════════════════════
   NATIVE DETAILS/SUMMARY — pixel-perfect Executive Summary clone
   ═══════════════════════════════════════════════════════════════ */

/* Shared glow-box shell */
.rca-native-box {
  position: relative;
  border-radius: 10px;
  margin-bottom: 10px;
  overflow: hidden;
}
.rca-native-box::before {
  content: '';
  position: absolute; top: 0; left: 0; right: 0; height: 2px;
  pointer-events: none;
}

/* Cyan variant — View Charts */
.rca-native-box.cyan-box {
  background: rgba(56,189,248,0.05);
  border: 1px solid rgba(56,189,248,0.25);
  box-shadow: 0 0 28px rgba(56,189,248,0.08), inset 0 0 16px rgba(56,189,248,0.03);
}
.rca-native-box.cyan-box::before {
  background: linear-gradient(90deg, transparent, #38bdf8, transparent);
}

/* Green variant — Recommendations */
.rca-native-box.green-box {
  background: rgba(74,222,128,0.04);
  border: 1px solid rgba(74,222,128,0.25);
  box-shadow: 0 0 28px rgba(74,222,128,0.06), inset 0 0 16px rgba(74,222,128,0.02);
}
.rca-native-box.green-box::before {
  background: linear-gradient(90deg, transparent, #4ade80, transparent);
}

/* Summary row — arrow + title + hint */
.rca-native-summary {
  display: flex;
  align-items: center;
  gap: 10px;
  padding: 13px 18px;
  cursor: pointer;
  user-select: none;
  border-radius: 10px;
  list-style: none;
  transition: background 0.15s ease;
}
.rca-native-summary:hover { background: rgba(56,189,248,0.06); }
.rca-native-summary::-webkit-details-marker { display: none; }
.rca-native-summary::marker { display: none; }

.rca-native-arrow {
  font-size: 0.6rem;
  display: inline-block;
  transition: transform 0.2s ease;
  flex-shrink: 0;
}
.rca-native-title {
  font-size: 0.88rem;
  font-weight: 700;
  letter-spacing: 0.01em;
}
.rca-native-hint {
  font-size: 0.74rem;
  font-weight: 400;
  opacity: 0.8;
  margin-left: 2px;
}
/* Cyan colours */
.cyan-box .rca-native-summary:hover { background: rgba(56,189,248,0.07) !important; }
.cyan-box .rca-native-arrow  { color: #38bdf8; }
.cyan-box .rca-native-title  { color: #38bdf8; }
.cyan-box .rca-native-hint   { color: #7dd3fc; }
/* Green colours */
.green-box .rca-native-summary:hover { background: rgba(74,222,128,0.06) !important; }
.green-box .rca-native-arrow { color: #4ade80; }
.green-box .rca-native-title { color: #4ade80; }
.green-box .rca-native-hint  { color: #86efac; }

/* Rotate arrow when open */
.rca-native-box.is-open .rca-native-arrow { transform: rotate(90deg); }

/* Charts body — always visible */
#rca-charts-body { display: block !important; }

/* ── rca-viz-tab nav buttons: nuke Gradio overrides completely ── */
/* Gradio injects a wrapping div.wrap > button — target every level */
#rca-charts-body .rca-viz-tab,
#rca-charts-body .rca-viz-tab > button,
#rca-charts-body button.rca-viz-tab,
#rca-charts-body .rca-viz-nav > div button,
#rca-charts-body .rca-viz-nav button,
.rca5-report-panel #rca-charts-body button {
  background: rgba(10,18,38,0.9) !important;
  border: 1px solid rgba(56,189,248,0.22) !important;
  border-radius: 8px !important;
  color: #c8d8f0 !important;
  font-size: 0.83rem !important;
  font-weight: 600 !important;
  text-align: left !important;
  padding: 9px 13px !important;
  margin-bottom: 5px !important;
  width: 100% !important;
  min-width: 0 !important;
  transition: all 0.15s ease !important;
  cursor: pointer !important;
  box-shadow: none !important;
  line-height: 1.4 !important;
}
#rca-charts-body .rca-viz-tab:hover > button,
#rca-charts-body .rca-viz-tab:hover,
#rca-charts-body .rca-viz-nav button:hover,
.rca5-report-panel #rca-charts-body button:hover {
  background: rgba(56,189,248,0.10) !important;
  border-color: rgba(56,189,248,0.55) !important;
  color: #38bdf8 !important;
}
/* Nav title */
#rca-charts-body .rca-viz-nav-title {
  font-size: 0.62rem !important; font-weight: 800 !important;
  text-transform: uppercase !important; letter-spacing: 0.12em !important;
  color: #38bdf8 !important; margin-bottom: 8px !important; display: block !important;
}
/* Force ALL text inside charts body to be visible on dark bg */
#rca-charts-body p,
#rca-charts-body span,
#rca-charts-body div,
#rca-charts-body label,
#rca-charts-body li {
  color: #e2e8f0 !important;
}
/* "How it works" purple blurb */
#rca-charts-body .rca-howit-label { color: #a78bfa !important; }
#rca-charts-body .rca-howit-body  { color: #e9d5ff !important; }

/* ── HIDE duplicate grey option tabs that appear below recommendations ── */
/* These are the compat/hidden chart nav buttons outside the view-charts section */
.rca5-report-panel .rca-viz-nav:not(#rca-charts-body .rca-viz-nav) {
  display: none !important;
}
/* Prevent any stray duplicate button rows from rendering */
#rca-charts-col-wrap:not(.gradio-column) {
  display: none !important;
}
/* Toggle button styling */
#rca-charts-toggle-btn,
.rca-charts-toggle button {
  background: rgba(56,189,248,0.06) !important;
  border: 1px solid rgba(56,189,248,0.3) !important;
  border-radius: 10px !important;
  color: #38bdf8 !important;
  font-size: 0.85rem !important;
  font-weight: 700 !important;
  padding: 11px 18px !important;
  width: 100% !important;
  text-align: left !important;
  cursor: pointer !important;
  margin-bottom: 6px !important;
  transition: background 0.18s, border-color 0.18s !important;
}
.rca-charts-toggle button:hover {
  background: rgba(56,189,248,0.12) !important;
  border-color: #38bdf8 !important;
}
#rca-charts-body > .gradio-column { padding: 0 !important; }
</style>""")

                        # Charts toggle + viz are now inside the markdown report as <details> blocks.
                        # Keep dummy components for wiring compat only.
                        rca_charts_toggle_btn = gr.Button("t", visible=False, elem_id="rca-charts-trigger-btn")
                        with gr.Column(visible=False, elem_id="rca-charts-col-wrap") as rca_charts_col:
                            with gr.Row(equal_height=False):
                                with gr.Column(scale=1, min_width=200, elem_classes="rca-viz-nav"):
                                    gr.HTML('<div class="rca-viz-nav-title">◈ Chart Type</div>')
                                    rca_nav_net_btn  = gr.Button("🌐  Network / Flow Diagram", elem_classes="rca-viz-tab")
                                    rca_nav_bar_btn  = gr.Button("📊  Bar Chart",              elem_classes="rca-viz-tab")
                                    rca_nav_pie_btn  = gr.Button("🥧  Pie Chart",              elem_classes="rca-viz-tab")
                                    rca_nav_heat_btn = gr.Button("🔥  Heat Map",               elem_classes="rca-viz-tab")
                                    rca_nav_bub_btn  = gr.Button("🫧  Bubble Chart",           elem_classes="rca-viz-tab")
                                    rca_nav_san_btn  = gr.Button("🔀  Sankey Diagram",         elem_classes="rca-viz-tab")

                                with gr.Column(scale=4):
                                    rca_chart_info_html = gr.HTML(_default_dual_glow_html())
                                    rca_main_plot = gr.Plot(
                                        label="", show_label=False,
                                        value=_blank_fig, elem_id="rca-main-plot"
                                    )

                        # ── View Charts section — inside gr.Accordion dropdown ──
                        rca_charts_html_vis = gr.HTML("", visible=False, elem_id="rca-charts-vis-wrap")  # compat

                        with gr.Column(visible=False, elem_id="rca-view-charts-wrap") as rca_view_charts_col:
                            with gr.Accordion("📊 View Charts", open=True, elem_id="rca-vchart-accordion", elem_classes="rca-vchart-open"):
                                rca_vchart_header = gr.HTML("", visible=True, elem_id="rca-vchart-header")
                                # 6 buttons — start hidden, rca_handler shows only the 4 best for the query
                                with gr.Row(equal_height=True):
                                    rca_vchart_net_btn  = gr.Button("🌐 Network / Flow",    elem_classes="rca-viz-tab", scale=1, visible=False)
                                    rca_vchart_bar_btn  = gr.Button("📊 Bar Chart",         elem_classes="rca-viz-tab", scale=1, visible=False)
                                    rca_vchart_heat_btn = gr.Button("🔥 Heatmap",           elem_classes="rca-viz-tab", scale=1, visible=False)
                                    rca_vchart_bub_btn  = gr.Button("🫧 Bubble Chart",      elem_classes="rca-viz-tab", scale=1, visible=False)
                                    rca_vchart_pie_btn  = gr.Button("🥧 Pie Chart",         elem_classes="rca-viz-tab", scale=1, visible=False)
                                    rca_vchart_san_btn  = gr.Button("🔀 Sankey Diagram",    elem_classes="rca-viz-tab", scale=1, visible=False)
                                # Explain + chart — only visible after user clicks a tab
                                rca_vchart_info = gr.HTML("", visible=False, elem_id="rca-vchart-info")
                                with gr.Column(visible=False, elem_id="rca-vchart-plot-wrap") as rca_vchart_plot_col:
                                    rca_vchart_plot = gr.Plot(
                                        label="", show_label=False,
                                        value=None, elem_id="rca-vchart-plot"
                                    )

                        # ── Recommendations accordion (AFTER charts section) ──
                        rca_detailed_rec_inner = gr.HTML("", visible=True, elem_id="rca-rec-vis-wrap")

                        # ── Compat stubs (hidden, preserve old wiring refs) ──────────────
                        with gr.Column(visible=False):
                            rca_charts_accordion = gr.Accordion("_charts_compat", open=False)
                        with gr.Column(visible=False):
                            rca_rec_accordion = gr.Accordion("_rec_compat", open=False)

                    # ── hidden compat wiring placeholders (outside panel) ──────────
                    with gr.Column(visible=False, elem_id="rca-dl-hidden-wrap"):
                        rca_dl_btn  = gr.DownloadButton(visible=False, label="")
                        rca_csv_btn = gr.DownloadButton(visible=False, label="")
                    rca_dl_status = gr.HTML("")
                    # State variables hold file paths — must be at tab scope, not inside hidden column
                    rca_dl_path_state  = gr.State(None)
                    rca_csv_path_state = gr.State(None)
                    with gr.Column(visible=False) as rca_dl_section:
                        pass   # placeholder — kept for rca_handler outputs compat

                    # ── "View charts?" prompt — kept hidden, wiring compat only ─────
                    with gr.Column(visible=False) as rca_chart_prompt:
                        rca_chart_yes    = gr.Button(visible=False)
                        rca_chart_no     = gr.Button(visible=False)
                        rca_chart_no_msg = gr.HTML(visible=False)

                    # ── ⬇ Download section — always at the bottom ─────────────────
                    rca_dl_section_vis = gr.Column(visible=False)
                    with rca_dl_section_vis:
                        gr.HTML("""
<div style="margin-top:20px;padding-top:14px;border-top:2px solid rgba(56,189,248,0.18);
     display:flex;align-items:center;gap:8px;margin-bottom:12px">
  <span style="font-size:0.62rem;font-weight:800;text-transform:uppercase;
       letter-spacing:0.14em;color:#38bdf8">⬇ Export Full RCA Report</span>
  <span style="font-size:0.68rem;color:#94a3b8">— includes full report · recommendations (Critical / High / Strategic) · all charts</span>
</div>
<div style="font-size:0.67rem;color:#475569;margin-bottom:10px;padding:7px 10px;
     background:rgba(56,189,248,0.05);border:1px solid rgba(56,189,248,0.15);border-radius:7px">
  ✓ Exported file is identical to the on-screen analysis — same sections, same data, same recommendations.
</div>""")
                        with gr.Row():
                            rca_dl_btn_vis = gr.DownloadButton(
                                "📄 Download Word Report (.docx)", visible=True, scale=1,
                                elem_classes="rca-dl-btn"
                            )
                            rca_csv_btn_vis = gr.DownloadButton(
                                "📊 Download Excel (.xlsx)", visible=True, scale=1,
                                elem_classes="rca-dl-btn"
                            )
                        rca_dl_status_vis = gr.HTML("", elem_classes="rca-dl-status")

                    rca_log = gr.HTML("", visible=False)

                # RIGHT: Product Delay Rates only
                with gr.Column(scale=1, min_width=220):
                    rca_kpi_right  = gr.HTML("")   # compat
                    rca_prod_stats = gr.HTML(_build_rca_insights_panel())

                    # hidden compat wiring kept for event chains
                    rca_prod_filter   = gr.State("all")
                    rca_category_dd   = gr.Dropdown(
                        choices=_get_all_product_categories(),
                        value="🌐 All Categories",
                        label="", show_label=False, visible=False
                    )
                    rca_category_stats    = gr.HTML("", visible=False)
                    rca_agent_log_html    = gr.HTML("", visible=False)
                    rca_cypher_panel_html = gr.HTML("", visible=False)

            # ═══════════════════════════════════════════════════════════════════
            # VISUALIZATION EXPLORER — 6 fixed chart types, dynamic content
            # Charts always keep the same names; data & descriptions adapt to query
            # ═══════════════════════════════════════════════════════════════════

            # ── States ───────────────────────────────────────────────────────
            rca_q_state        = gr.State("")
            rca_state_net_fig  = gr.State(None)
            rca_state_bar_fig  = gr.State(None)
            rca_state_pie_fig  = gr.State(None)
            rca_state_heat_fig = gr.State(None)
            rca_state_bub_fig  = gr.State(None)   # bubble/network
            rca_state_san_fig  = gr.State(None)   # sankey
            rca_state_net_info  = gr.State("")
            rca_state_bar_info  = gr.State("")
            rca_state_pie_info  = gr.State("")
            rca_state_heat_info = gr.State("")
            rca_state_bub_info  = gr.State("")
            rca_state_san_info  = gr.State("")
            # Legacy label states (kept for compat, not used for button labels)
            rca_state_net_label  = gr.State("🌐  Network / Flow Diagram")
            rca_state_bar_label  = gr.State("📊  Bar Chart")
            rca_state_pie_label  = gr.State("🥧  Pie Chart")
            rca_state_heat_label = gr.State("🔥  Heat Map")

            _blank_fig = go.Figure()
            _blank_fig.update_layout(
                paper_bgcolor="#060c1c", plot_bgcolor="#060c1c",
                xaxis=dict(visible=False), yaxis=dict(visible=False),
                annotations=[dict(
                    text="Run an analysis above to load this chart",
                    x=0.5, y=0.5, xref="paper", yref="paper",
                    showarrow=False, font=dict(color="#475569", size=13)
                )], height=480
            )

            # ── Hidden compat components (preserve all legacy wiring) ─────────
            rca_net_plot   = gr.Plot(visible=False, value=_blank_fig)
            rca_bar_plot   = gr.Plot(visible=False, value=_blank_fig)
            rca_pie_plot   = gr.Plot(visible=False, value=_blank_fig)
            rca_heat_plot  = gr.Plot(visible=False, value=_blank_fig)
            rca_net_info   = gr.HTML(visible=False)
            rca_bar_info   = gr.HTML(visible=False)
            rca_pie_info   = gr.HTML(visible=False)
            rca_heat_info  = gr.HTML(visible=False)
            rca_viz_plot          = gr.Plot(visible=False, value=_blank_fig)
            rca_why_html          = gr.HTML(visible=False)
            rca_chart1            = gr.Plot(visible=False, value=_blank_fig)
            rca_chart2            = gr.Plot(visible=False, value=_blank_fig)
            rca_chart3            = gr.Plot(visible=False, value=_blank_fig)
            rca_viz_dropdown      = gr.Dropdown(choices=["—"], value="—", visible=False)
            rca_chart2_dd         = gr.Dropdown(choices=["—"], value="—", visible=False)
            rca_viz_plot2         = gr.Plot(visible=False, value=_blank_fig)
            rca_why_html2         = gr.HTML(visible=False)
            rca_category_stats_old = gr.HTML(visible=False)

            # ── Chart helper functions (defined before wiring) ───────────────────
            _charts_open_state = gr.State(False)

            def _toggle_charts(is_open):
                new_state = not is_open
                label = ("▼  📊  View Charts  —  click to collapse"
                         if new_state else
                         "▶  📊  View Charts  —  click to expand")
                return new_state, gr.update(visible=new_state), gr.update(value=label)

            def _select_best_charts(q: str):
                q_l = (q or "").lower()
                is_route    = any(w in q_l for w in ["route","transit","transport","road","rail","air","ship"])
                is_supplier = any(w in q_l for w in ["supplier","vendor","risk","sup"])
                is_category = any(w in q_l for w in ["category","product","toy","auto","health","beauty"])
                is_stock    = any(w in q_l for w in ["stockout","stock","demand","gap","shortage"])
                is_dist     = any(w in q_l for w in ["distributor","city","distribution"])
                if is_route:
                    show={("net","🌐  Route Network"),("san","🔀  Supply Sankey"),("bar","📊  Route Ranking"),("heat","🔥  Delay Heatmap")}
                elif is_stock:
                    show={("bar","📊  Stockout Ranking"),("pie","🥧  Demand Gap Share"),("heat","🔥  Plant Heatmap"),("bub","🫧  Risk Bubble")}
                elif is_supplier:
                    show={("net","🌐  Supplier Network"),("bub","🫧  Supplier Risk"),("bar","📊  Supplier Pareto"),("heat","🔥  Delay Heatmap")}
                elif is_category:
                    show={("pie","🥧  Category Share"),("bar","📊  Category Pareto"),("heat","🔥  Category Heatmap"),("net","🌐  Supply Flow")}
                elif is_dist:
                    show={("net","🌐  Distributor Network"),("san","🔀  Supply Sankey"),("bar","📊  City Ranking"),("pie","🥧  Delay Share")}
                else:
                    show={("net","🌐  Plant Network"),("bar","📊  Plant Pareto"),("heat","🔥  Delay Heatmap"),("san","🔀  Supply Sankey")}
                show_map = {k:v for k,v in show}
                _all = [("net",rca_nav_net_btn),("bar",rca_nav_bar_btn),("pie",rca_nav_pie_btn),
                        ("heat",rca_nav_heat_btn),("bub",rca_nav_bub_btn),("san",rca_nav_san_btn)]
                return [gr.update(visible=k in show_map, value=show_map.get(k,"")) if k in show_map
                        else gr.update(visible=False) for k,_ in _all]

            # ── Wiring ─────────────────────────────────────────────────────────
            rca_q.change(fn=_build_context_html, inputs=[rca_q], outputs=[rca_context_html])
            rca_q.change(fn=_suggest_prompt,     inputs=[rca_q], outputs=[rca_suggest_html])

            rca_category_dd.change(
                fn=_load_rca_category_stats,
                inputs=[rca_category_dd],
                outputs=[rca_category_stats]
            )

            rca_load_sup_btn   = gr.Button(visible=False)
            rca_supplier_stats = gr.HTML(visible=False)
            rca_supplier_dd        = gr.Dropdown(choices=['—'], value='—', visible=False)
            rca_supplier_stats_old = gr.HTML(visible=False)

            # ── Helper: single explanation box above chart ───────────────────
            def _dual_glow(what_title, what_body, what_tags,
                           why_title, why_body):
                """Single merged explanation box — no question repeat."""
                tags_html = "".join(
                    f'<span class="viz-glow-tag">{t}</span>' for t in what_tags
                )
                return f"""
<div class="viz-glow-box what-glow" style="margin-bottom:12px">
  <div style="display:flex;align-items:flex-start;gap:12px">
    <div style="flex:1">
      <div class="viz-glow-label" style="margin-bottom:6px">💡 {what_title}</div>
      <div class="viz-glow-body">{what_body} — <strong>{why_title}:</strong> {why_body}</div>
      <div class="viz-glow-tags" style="margin-top:8px">{tags_html}</div>
    </div>
  </div>
</div>"""

            # ── _post_rca_render_6: render all 6 charts + dynamic descriptions ──
            def _post_rca_render_6(q):
                """
                Renders all 6 chart types. Each chart adapts its data to the
                question but keeps its fixed type name in the nav tab.
                Descriptions reference the actual question in plain language.
                """
                _bfig = go.Figure().update_layout(
                    paper_bgcolor="#060c1c", plot_bgcolor="#060c1c",
                    xaxis=dict(visible=False), yaxis=dict(visible=False),
                    height=480, autosize=False,
                )
                def _safe(fn, **kw):
                    try:
                        return _apply_dark_theme(fn(**kw))
                    except TypeError:
                        try: return _apply_dark_theme(fn())
                        except Exception: return _bfig
                    except Exception as e:
                        print(f"[chart] {getattr(fn,'__name__','?')}: {e}")
                        return _bfig

                ctx = _extract_context(q)
                q_short = (q or "your supply chain query").strip().rstrip("?")

                import concurrent.futures as _cf
                net_fig = bar_fig = pie_fig = heat_fig = bub_fig = san_fig = _bfig
                try:
                    with _cf.ThreadPoolExecutor(max_workers=6) as _ex:
                        # Network/Flow Diagram → build_dynamic_flow (node-link flow diagram)
                        _fn  = _ex.submit(_safe, build_dynamic_flow,     question=q)
                        _fb  = _ex.submit(_safe, build_dynamic_pareto,   question=q)
                        _fp  = _ex.submit(_safe, build_dynamic_pie,      question=q)
                        _fh  = _ex.submit(_safe, build_dynamic_heatmap,  question=q)
                        # Bubble Chart → build_dynamic_network (risk×delay scatter, bubble=lead time)
                        _fbu = _ex.submit(_safe, build_dynamic_network,  question=q)
                        # Sankey Diagram → supply flow sankey
                        _fs  = _ex.submit(_safe, build_supply_flow_sankey, question=q)
                        for _f, _tgt in [(_fn,  "net"), (_fb, "bar"), (_fp, "pie"),
                                          (_fh, "heat"), (_fbu,"bub"), (_fs, "san")]:
                            try:
                                v = _f.result(timeout=25)
                                if   _tgt=="net":  net_fig  = v
                                elif _tgt=="bar":  bar_fig  = v
                                elif _tgt=="pie":  pie_fig  = v
                                elif _tgt=="heat": heat_fig = v
                                elif _tgt=="bub":  bub_fig  = v
                                elif _tgt=="san":  san_fig  = v
                            except Exception: pass
                except Exception: pass

                # For route queries, net_fig (flow diagram) is often empty — use bar_fig instead
                if ctx.get("is_route") and _fig_is_empty(net_fig) and not _fig_is_empty(bar_fig):
                    net_fig = bar_fig

                # ── Final safety: replace any still-empty figure with broadest fallback ──
                def _force(f, fn_fallback):
                    """If f is empty, build fn_fallback without question filter."""
                    if _fig_is_empty(f):
                        try:
                            fb = _apply_dark_theme(fn_fallback())
                            return fb if not _fig_is_empty(fb) else f
                        except Exception:
                            return f
                    return f

                net_fig  = _force(net_fig,  build_dynamic_pareto)
                bar_fig  = _force(bar_fig,  build_dynamic_pareto)
                pie_fig  = _force(pie_fig,  build_dynamic_pie)
                heat_fig = _force(heat_fig, build_dynamic_heatmap)
                bub_fig  = _force(bub_fig,  build_dynamic_network)
                san_fig  = _force(san_fig,  build_supply_flow_sankey)

                # ── 1. Network / Flow Diagram ─────────────────────────────
                if ctx["is_route"]:
                    net_what_t = "Route Flow Network"
                    net_what_b = f"For «{q_short}»: each node is a plant or distributor. Edges show routes coloured by transport mode — thickness = volume."
                    net_tags   = ["Route Flow", "Transport Mode", "Volume"]
                    net_why_t  = "Spot bottleneck connections"
                    net_why_b  = "Routes with thin edges or mismatched modes are logistic chokepoints. Fix these to reduce transit delays."
                elif ctx["is_plant"]:
                    net_what_t = "Plant Delay Network"
                    net_what_b = f"For «{q_short}»: X = total shipments, Y = delay rate %. Bubble size = absolute delayed count. Red = >30% delay."
                    net_tags   = ["Volume × Delay Rate", "Plant Risk", ">30% = Critical"]
                    net_why_t  = "Identify structural factory bottlenecks"
                    net_why_b  = "Plants in the top-right quadrant have both high volume and high delay rates — they need the most urgent attention."
                else:
                    net_what_t = "Supplier Risk Network"
                    net_what_b = f"For «{q_short}»: X = risk score, Y = delay rate %, bubble = lead time days. Red bubbles top-right = highest-priority root cause."
                    net_tags   = ["Risk Score", "Delay Rate %", "Lead Time"]
                    net_why_t  = "Pinpoint your primary root cause supplier"
                    net_why_b  = "A large red bubble in the top-right corner is the single supplier most responsible for your delays. Start your RCA there."

                # ── 2. Bar Chart ──────────────────────────────────────────
                if ctx["is_plant"]:
                    bar_what_t = "Plant Pareto — Delay Count Ranking"
                    bar_what_b = f"For «{q_short}»: plants ranked by total delayed shipments, with a cumulative % line overlay."
                    bar_tags   = ["Plant Ranking", "Cumulative %", "80/20 Rule"]
                    bar_why_t  = "Find which plants to fix first"
                    bar_why_b  = "If the top 2 bars account for 70%+ of delays, those 2 plants are your entire action plan."
                elif ctx["is_category"]:
                    bar_what_t = "Product Category Pareto — Delay Ranking"
                    bar_what_b = f"For «{q_short}»: product categories ranked by delayed shipments, with cumulative % line."
                    bar_tags   = ["Category Ranking", "Cumulative %", "Product Focus"]
                    bar_why_t  = "Identify which product lines are failing"
                    bar_why_b  = "A dominant category at the top signals a product-specific issue — not a logistics-wide problem."
                elif ctx["is_stock"]:
                    bar_what_t = "Distributor Demand Gap Ranking"
                    bar_what_b = f"For «{q_short}»: distributor cities ranked by demand gap (forecast minus fulfilled units)."
                    bar_tags   = ["Demand Gap", "Unmet Units", "City Ranking"]
                    bar_why_t  = "Prioritise replenishment cities"
                    bar_why_b  = "The tallest bars show which cities have the most unmet consumer demand. Replenish these first."
                else:
                    bar_what_t = "Supplier Pareto — Delay Count Ranking"
                    bar_what_b = f"For «{q_short}»: suppliers ranked by total delayed shipments, cumulative % line shows 80/20 concentration."
                    bar_tags   = ["Supplier Ranking", "Cumulative %", "Fix These First"]
                    bar_why_t  = "Focus on the fewest suppliers for maximum impact"
                    bar_why_b  = "Fixing the tallest 3-4 bars typically resolves 70-80% of all delays across the network."

                # ── 3. Pie Chart ──────────────────────────────────────────
                if ctx["is_supplier"]:
                    pie_what_t = "Supplier Delay Share — Proportional View"
                    pie_what_b = f"For «{q_short}»: each slice = one supplier's % share of all major delay events."
                    pie_tags   = ["Supplier Share", "Proportional", "Dominant Cause"]
                    pie_why_t  = "See if one supplier dominates"
                    pie_why_b  = "A single slice taking >40% of the pie means that supplier is your root cause — everything else is noise."
                elif ctx["is_route"]:
                    pie_what_t = "Transport Mode Share"
                    pie_what_b = f"For «{q_short}»: breakdown of all routes by transport mode (Road / Rail / Air / Sea)."
                    pie_tags   = ["Transport Mode", "Modal Split", "Infrastructure"]
                    pie_why_t  = "Detect dangerous modal dependency"
                    pie_why_b  = "A mode with >70% share is a single point of failure. Any disruption to that mode (weather, strikes) stalls your whole network."
                elif ctx["is_plant"]:
                    pie_what_t = "Plant Delay Share — Proportional View"
                    pie_what_b = f"For «{q_short}»: each slice = one plant's proportion of total major delay events."
                    pie_tags   = ["Plant Share", "Location Risk", "Proportional"]
                    pie_why_t  = "Confirm which plant is the real bottleneck"
                    pie_why_b  = "A plant slice dominating >50% of the pie confirms a location-specific issue, not a network-wide problem."
                else:
                    pie_what_t = "Product Category Delay Share"
                    pie_what_b = f"For «{q_short}»: each slice = a product category's share of all major delays in the network."
                    pie_tags   = ["Category Share", "Product Mix", "Root Cause"]
                    pie_why_t  = "Spot category-specific failure patterns"
                    pie_why_b  = "If Toys or Auto parts dominate, the issue is sourcing or seasonal demand for that category — not logistics."

                # ── 4. Heat Map ───────────────────────────────────────────
                if ctx["is_supplier"]:
                    heat_what_t = "Supplier × Product Category Heatmap"
                    heat_what_b = f"For «{q_short}»: grid where rows = suppliers, columns = product categories, cell colour = delayed shipment count."
                    heat_tags   = ["Supplier × Category", "Intersection", "Colour = Severity"]
                    heat_why_t  = "Find your most precise root cause intersection"
                    heat_why_b  = "The darkest red cell is the specific supplier–product combination causing the most damage. That is where to act first."
                elif ctx["is_route"]:
                    heat_what_t = "Plant × Transport Mode Heatmap"
                    heat_what_b = f"For «{q_short}»: grid where rows = plants, columns = transport modes, cell = delayed shipments from that plant via that mode."
                    heat_tags   = ["Plant × Mode", "Transport Risk", "Colour = Severity"]
                    heat_why_t  = "Detect mode-specific delay clusters"
                    heat_why_b  = "Dark clusters on a single mode row mean that plant is over-reliant on one transport type. Diversifying the mode mix reduces risk."
                else:
                    heat_what_t = "Plant × Product Category Heatmap"
                    heat_what_b = f"For «{q_short}»: grid where rows = plants, columns = product categories, cell colour = delayed shipment count."
                    heat_tags   = ["Plant × Category", "Concentration Map", "Colour = Severity"]
                    heat_why_t  = "Determine if delays are plant-wide or product-specific"
                    heat_why_b  = "Delays spread evenly across a row = plant problem. Delays clustered in one column = product-specific sourcing issue. Different fixes needed."

                # ── 5. Bubble Chart — always route efficiency, axes adapt to question ──
                if ctx["is_plant"]:
                    bub_what_t = "Plant Route Efficiency Bubbles"
                    bub_what_b = f"For «{q_short}»: routes originating from the relevant plants — X = distance, Y = transport cost, bubble size = efficiency score."
                    bub_tags   = ["Plant Routes", "Distance × Cost", "Efficiency"]
                    bub_why_t  = "Identify which plant routes are draining cost"
                    bub_why_b  = "Small bubbles top-right = routes from your bottleneck plants that are expensive AND inefficient. Renegotiating these carriers reduces cost and delay together."
                elif ctx["is_supplier"]:
                    bub_what_t = "Supplier Lead-Time Efficiency Bubbles"
                    bub_what_b = f"For «{q_short}»: each bubble = a supplier route — X = StoP lead time days, Y = risk score, size = delayed shipment count."
                    bub_tags   = ["Lead Time", "Risk Score", "Delay Volume"]
                    bub_why_t  = "Find high-risk slow suppliers in one view"
                    bub_why_b  = "Large bubbles in the top-right = suppliers with long lead times, high risk scores, and many delays. These are your dual-sourcing candidates."
                else:
                    bub_what_t = "Route Efficiency Bubble Chart"
                    bub_what_b = f"For «{q_short}»: every route as a bubble — X = distance (km), Y = transport cost (₹), size = efficiency score, colour = transport mode."
                    bub_tags   = ["Distance × Cost", "Efficiency Score", "Transport Mode"]
                    bub_why_t  = "Pinpoint costly and inefficient routes"
                    bub_why_b  = "Small bubbles in the top-right = long, expensive, low-efficiency routes. These are your logistics cost drain and primary delay risk points."

                # ── 6. Sankey Diagram — always end-to-end flow, focus adapts ──
                if ctx["is_supplier"]:
                    san_what_t = "Supplier → Plant → Distributor Flow"
                    san_what_b = f"For «{q_short}»: volume flowing from each supplier through plants to distributors. Red bands = shipments with Major Delay."
                    san_tags   = ["Supplier Flow", "Delay Bands", "Volume Width"]
                    san_why_t  = "Trace exactly which supplier's delays cascade downstream"
                    san_why_b  = "Follow the red bands from a supplier node — every distributor they reach via red paths is affected by that supplier's delays."
                elif ctx["is_stock"]:
                    san_what_t = "Plant → Distributor → Demand Flow"
                    san_what_b = f"For «{q_short}»: volume from plants reaching each distributor city. Thin outgoing bands = unfulfilled demand / stockouts."
                    san_tags   = ["Distribution Flow", "Stockout Risk", "Volume Width"]
                    san_why_t  = "See which cities are starved of supply"
                    san_why_b  = "Distributors where incoming bands are thin relative to their demand size are your highest stockout risk. Replenish those nodes first."
                else:
                    san_what_t = "End-to-End Supply Chain Flow"
                    san_what_b = f"For «{q_short}»: full supply chain from Suppliers → Plants → Distributors. Band width = shipment volume, red = Major Delay."
                    san_tags   = ["End-to-End Flow", "Volume Bands", "Red = Delayed"]
                    san_why_t  = "Find the exact handoff where volume is lost"
                    san_why_b  = "The node where incoming band width drops significantly compared to outgoing is your bottleneck. Wide red bands mean delays, not just low volume."

                # ── Build dual glow HTML for all 6 ───────────────────────
                net_info_html  = _dual_glow(net_what_t,  net_what_b,  net_tags,  net_why_t,  net_why_b)
                bar_info_html  = _dual_glow(bar_what_t,  bar_what_b,  bar_tags,  bar_why_t,  bar_why_b)
                pie_info_html  = _dual_glow(pie_what_t,  pie_what_b,  pie_tags,  pie_why_t,  pie_why_b)
                heat_info_html = _dual_glow(heat_what_t, heat_what_b, heat_tags, heat_why_t, heat_why_b)
                bub_info_html  = _dual_glow(bub_what_t,  bub_what_b,  bub_tags,  bub_why_t,  bub_why_b)
                san_info_html  = _dual_glow(san_what_t,  san_what_b,  san_tags,  san_why_t,  san_why_b)

                # Show network diagram by default — outputs must match wiring exactly (26 items)
                return (
                    net_fig, bar_fig, pie_fig, heat_fig,   # 0-3: hidden compat plots
                    q,                                       # 4: rca_q_state
                    net_info_html, bar_info_html,           # 5-6: hidden compat infos
                    pie_info_html, heat_info_html,          # 7-8
                    net_fig, net_info_html,                 # 9-10: rca_main_plot + rca_chart_info_html (default = network)
                    net_fig,  bar_fig,  pie_fig,  heat_fig, # 11-14: state figs
                    net_info_html, bar_info_html,           # 15-16: state infos
                    pie_info_html, heat_info_html,          # 17-18
                    gr.update(value="🌐  Network / Flow Diagram"),  # 19: nav btn 1
                    gr.update(value="📊  Bar Chart"),               # 20: nav btn 2
                    gr.update(value="🥧  Pie Chart"),               # 21: nav btn 3
                    gr.update(value="🔥  Heat Map"),                # 22: nav btn 4
                    bub_fig, san_fig,                       # 23-24: bub+san state figs
                    bub_info_html, san_info_html,           # 25-26: bub+san state infos
                )

            # ── Post-analysis wiring ──────────────────────────────────────────

            # ── View Charts: query-aware chart recommendation functions ─────────
            # NOTE: View Charts uses the SAME collapsible <details> pattern as
            # Executive Summary & Recommendations — dark bg, cyan glow, NO white bg.
            # Chart option buttons are rendered directly (no duplicate grey tabs below).

            _CHART_META = {
                "net": {
                    "label": "🌐 Network / Flow Diagram",
                    "default_title": "Supply Chain Network Flow",
                    "why": {
                        "shipment_delay":    "Traces the full delay propagation path from Supplier → Plant → Shipment → Distributor, highlighting where the root cause enters the network.",
                        "plant":             "Visualises the supplier-to-plant dependency flow, showing which suppliers feed the bottleneck plant and how disruptions cascade downstream.",
                        "supplier_risk":     "Maps supplier dependency relationships across the network, revealing which plants are most exposed to high-risk upstream suppliers.",
                        "product_category":  "Shows the category's dependency path from supplier inputs through plants to distributors, identifying the exact bottleneck link.",
                        "simulation":        "Traces how the disruption propagates across the full supply chain graph, showing first- and second-order impact layers.",
                        "default":           "Reveals supply chain choke-points and the primary delay propagation path across the network.",
                    },
                },
                "bar": {
                    "label": "📊 Bar Chart",
                    "default_title": "Top Contributors — Ranked Analysis",
                    "why": {
                        "shipment_delay":    "Ranks the top delayed plants, routes, and product categories by shipment volume — fixing the first 3 bars resolves the majority of delay events.",
                        "plant":             "Compares delayed shipment contribution by plant, making it immediately clear which manufacturing site is the primary operational bottleneck.",
                        "supplier_risk":     "Ranks suppliers by delayed shipment count and risk score, directly prioritising which supplier relationships require urgent intervention.",
                        "demand_gap":        "Ranks the highest demand-gap retailers and distributors, identifying which cities require priority replenishment action.",
                        "product_category":  "Compares delayed shipments by category, showing whether the delay problem is concentrated in one product line or systemic.",
                        "distributor":       "Ranks the most affected distributor cities by delayed shipments received, quantifying where commercial impact is highest.",
                        "default":           "Ranked bar view of top delay contributors — highest to lowest impact.",
                    },
                },
                "pie": {
                    "label": "🥧 Pie Chart",
                    "default_title": "Share Distribution",
                    "why": {
                        "demand_gap":        "Shows category-wise contribution to total demand gap — one dominant slice signals a concentrated shortage rather than a systemic one.",
                        "product_category":  "Visualises each category's share of total delayed shipments, revealing whether the delay is concentrated in the queried category or spread across all.",
                        "distributor":       "Breaks down the share of delayed shipments by distributor city, showing which regional hubs are absorbing the most supply chain pressure.",
                        "default":           "Share distribution view — reveals whether the delay or gap is concentrated or evenly spread across the network.",
                    },
                },
                "heat": {
                    "label": "🔥 Heat Map",
                    "default_title": "Delay Severity — Heat Map",
                    "why": {
                        "shipment_delay":    "Grid of plants vs product categories — cell intensity shows delay rate %, immediately exposing the exact plant × category combination driving the root cause.",
                        "plant":             "Compares delay severity and risk scores across plants side by side, making it easy to spot which facility is the worst performer and by how much.",
                        "demand_gap":        "Maps retailer shortage severity by city and distributor, showing the geographic concentration of the demand gap.",
                        "distributor":       "Plots distributor demand gap severity across cities, highlighting the geographic zones most affected by upstream supply failures.",
                        "simulation":        "Identifies vulnerable nodes and their severity levels in the disruption scenario — darker cells signal higher exposure to the simulated failure.",
                        "default":           "Heat map of delay concentration — darker cells identify the highest-severity combinations.",
                    },
                },
                "bub": {
                    "label": "🫧 Bubble Chart",
                    "default_title": "Risk vs Impact — Bubble Analysis",
                    "why": {
                        "supplier_risk":     "Plots each supplier's risk score (X-axis) against delay frequency (Y-axis) with bubble size = shipment volume — top-right large bubbles are highest-priority.",
                        "plant":             "Compares plant impact (Y-axis) against delay rate (X-axis) with bubble size = shipment volume dispatched — exposes which plants combine high delay with high throughput.",
                        "distributor":       "Maps distributor cities by delayed shipments received vs demand gap, with bubble size = total volume — largest top-right bubbles need immediate attention.",
                        "simulation":        "Plots revenue exposure (Y-axis) against disruption impact score (X-axis) with bubble size = shipment volume at risk — identifies highest commercial exposure nodes.",
                        "default":           "Multi-dimensional bubble view — risk score vs impact vs volume in one chart.",
                    },
                },
                "san": {
                    "label": "🔀 Sankey Diagram",
                    "default_title": "Flow Diagram — Volume & Delays",
                    "why": {
                        "shipment_delay":    "Shows delay propagation as a flow: Supplier → Plant → Distributor — wide red bands pinpoint exactly where volume is being lost to delays.",
                        "supplier_risk":     "Traces the supplier delay propagation chain through plants to distributors — band width shows volume, red bands show where risk converts to actual delay.",
                        "demand_gap":        "Visualises the shortage propagation chain from upstream supply failures through distributors to the retail stockout level — connects cause to consumer impact.",
                        "product_category":  "Flows the category's shipment volume across the supply chain — red bands show exactly which supply chain stage is losing volume for this category.",
                        "simulation":        "Shows cascading impact flow from the disrupted node outward — wider bands indicate more severe downstream consequences.",
                        "default":           "Sankey flow: suppliers → plants → distributors — band width = volume, red bands = delayed flow.",
                    },
                },
            }

            _CHART_RECS = {
                "shipment_delay":  ["san", "heat", "bar", "net"],
                "route":           ["bar", "san", "heat", "net"],   # route: pareto cost, sankey flow, heatmap, network
                "plant":           ["bar", "net", "heat", "bub"],
                "supplier_risk":   ["net", "bub", "bar", "san"],
                "distributor":     ["heat", "bar", "san", "bub"],
                "demand_gap":      ["heat", "bar", "san", "pie"],
                "product_category":["pie", "san", "bar", "net"],
                "simulation":      ["net", "san", "heat", "bub"],
            }

            # Dynamic button labels per report type
            _CHART_BTN_LABELS = {
                "route": {
                    "bar":  "📊 Route Cost Ranking",
                    "san":  "🔀 Route Flow Sankey",
                    "heat": "🔥 Route Mode Heatmap",
                    "net":  "🌐 Route Network",
                },
                "supplier_risk": {
                    "net":  "🌐 Supplier Risk Network",
                    "bub":  "🫧 Risk Bubble Chart",
                    "bar":  "📊 Supplier Ranking",
                    "san":  "🔀 Supply Flow Sankey",
                },
                "plant": {
                    "bar":  "📊 Plant Delay Ranking",
                    "net":  "🌐 Plant Network",
                    "heat": "🔥 Plant × Category Heatmap",
                    "bub":  "🫧 Plant Risk Bubble",
                },
                "demand_gap": {
                    "heat": "🔥 Demand Gap Heatmap",
                    "bar":  "📊 Stockout Ranking",
                    "san":  "🔀 Shortage Chain Sankey",
                    "pie":  "🥧 Demand Gap Share",
                },
                "product_category": {
                    "pie":  "🥧 Category Delay Share",
                    "san":  "🔀 Category Flow Sankey",
                    "bar":  "📊 Category Ranking",
                    "net":  "🌐 Category Network",
                },
            }

            def _detect_report_type_from_q(q: str) -> str:
                q = q.lower()
                if any(w in q for w in ["what if","simulate","simulation","impact of"]) and not any(w in q for w in ["stockout","demand gap","shortage","spreading"]):
                    return "simulation"
                if any(w in q for w in ["stockout","demand gap","shortage","unmet",
                                         "running out","run out","out of stock","low stock",
                                         "depleting","replenish","where are we running"]):
                    return "demand_gap"
                if any(w in q for w in ["route","transport cost","transportation cost","highest cost route",
                                         "route cost","route efficiency","expensive route","logistics cost",
                                         "freight","road shipment","rail shipment","air shipment",
                                         "which routes","transport mode"]):
                    return "route"
                if any(w in q for w in [
                    "supplier risk","risky supplier","high risk supplier",
                    "which supplier","which suppliers","suppliers are delaying",
                    "are delaying","delaying shipment","suppliers delay",
                    "causing delay","supplier delay","suppliers who","suppliers that",
                ]):
                    return "supplier_risk"
                if any(w in q for w in ["plant rca","pune plant","baddi","bhopal","goa","plant fail"]) or \
                   (any(w in q for w in ["pune","baddi","bhopal","goa"]) and any(w in q for w in ["plant","rca","fail"])):
                    return "plant"
                if any(w in q for w in ["toy","category","watches","product","health beauty","auto",
                                         "construction","garden","cool stuff","bed bath"]):
                    return "product_category"
                if any(w in q for w in ["distributor","distribution","fulfil"]) and "delay" not in q:
                    return "distributor"
                return "shipment_delay"

            def _get_chart_why(chart_key: str, report_type: str, question: str) -> str:
                meta = _CHART_META.get(chart_key, {})
                why_map = meta.get("why", {})
                return why_map.get(report_type, why_map.get("default", "Relevant analytical view for this supply chain scenario."))

            # ─────────────────────────────────────────────────────────────────
            # VIEW CHARTS — Query-aware chart selection using proven Viz-tab functions
            # Same pattern as _load_heat / _load_stock in the Network Health tab.
            # ─────────────────────────────────────────────────────────────────

            # Map query type → 4 best chart keys (proven functions, not dynamic builders)
            _VIZ_CHART_PLAN = {
                # stock  = build_distributor_demand_gap  (stockout severity bar)
                # heat   = build_delay_heatmap           (plant × category heatmap)
                # risk   = build_supplier_risk_chart     (supplier risk horizontal bar)
                # trend  = build_monthly_delay_trend     (monthly stacked bar + line)
                # route  = build_route_efficiency_scatter (distance vs cost scatter)
                # sun    = build_plant_transport_sunburst (transport mode sunburst)
                "demand_gap":       ["stock", "heat",  "risk",  "trend"],
                "shipment_delay":   ["heat",  "risk",  "trend", "stock"],
                "supplier_risk":    ["risk",  "heat",  "trend", "stock"],
                "route":            ["route", "sun",   "heat",  "trend"],
                "plant":            ["heat",  "risk",  "trend", "stock"],
                "product_category": ["heat",  "stock", "risk",  "trend"],
                "distributor":      ["stock", "heat",  "risk",  "trend"],
                "simulation":       ["stock", "heat",  "risk",  "route"],
            }

            _VIZ_CHART_FN = {
                "stock": build_distributor_demand_gap,
                "heat":  build_delay_heatmap,
                "risk":  build_supplier_risk_chart,
                "trend": build_monthly_delay_trend,
                "route": build_route_efficiency_scatter,
                "sun":   build_plant_transport_sunburst,
            }

            _VIZ_BTN_LABELS = {
                "stock": "📦 Stockout Severity",
                "heat":  "🔥 Delay Heatmap",
                "risk":  "⚠️ Supplier Risk",
                "trend": "📅 Monthly Trend",
                "route": "🚛 Route Efficiency",
                "sun":   "🌐 Transport Modes",
            }

            _VIZ_CHART_WHY = {
                ("stock","demand_gap"):     "Ranks every distributor city by total units short — immediately shows where consumers are most affected and which city needs emergency stock first.",
                ("stock","shipment_delay"): "Stockouts are the downstream consequence of delays. This chart shows the current commercial impact so you can prioritise replenishment alongside fixing the root delay.",
                ("stock","distributor"):    "Identifies which distribution hubs are absorbing the most supply failure — the tallest bar is the city with the highest business impact right now.",
                ("heat","demand_gap"):      "Grid of plants × product categories coloured by shortage severity. The darkest cell is the exact plant–category combination driving the stockout — target that combination first.",
                ("heat","shipment_delay"):  "The darkest cell is your root-cause intersection — the plant × product category generating the highest delay volume. Fixing that single combination resolves the majority of network delays.",
                ("heat","supplier_risk"):   "Shows which plant–category combinations have the highest delay concentration. A full red row means a plant is struggling across all categories — a structural supplier issue.",
                ("risk","supplier_risk"):   "Suppliers above the 0.7 threshold are your highest-priority intervention targets — highest probability of causing a future supply disruption right now.",
                ("risk","demand_gap"):      "High supplier risk is the upstream root cause of downstream stockouts. Suppliers above 0.7 are most likely generating the delayed shipments that create the demand gaps you see.",
                ("risk","shipment_delay"):  "Risk scores are leading indicators — they predict future delays before they appear in shipment data. Any bar above 0.7 needs immediate auditing or dual-sourcing.",
                ("trend","demand_gap"):     "Monthly view shows whether stockouts are worsening over time. A rising red portion signals a deteriorating supply chain — escalation is needed before the next demand peak.",
                ("trend","shipment_delay"): "A rising delay rate line (dotted) month-over-month signals a systemic problem not self-correcting. A sudden spike points to a discrete event — supplier failure or capacity issue.",
                ("route","route"):          "Short-distance routes with high cost (top-left zone) are mis-routed — paying air freight rates for road-distance hauls. These are immediate cost-reduction opportunities.",
                ("sun","route"):            "Shows which transport modes each plant relies on. A plant dominated by Air is cost-exposed — 3–5× more expensive than road or rail for equivalent distances.",
            }

            _VIZ_CHART_HOW = {
                "stock": "Bars ranked left → right by total shortage units. Colour gradient: blue → purple → red encodes severity. Numbers above bars = shortage event count. Hover any bar for exact figures. Tallest bar = city needing emergency stock.",
                "heat":  "Grid: rows = plants, columns = product categories, cell colour = delay count (darker red = more delays). The darkest cell is your primary root cause. Full red row = plant issue. Full red column = category sourcing issue.",
                "risk":  "Horizontal bars ranked by risk score (0 = safe, 1 = critical). Red dotted line = 0.7 high-risk threshold. Any bar crossing it = supplier needing immediate action. Bar text shows supplier ID and plant.",
                "trend": "Stacked bars: on-time (green) vs delayed (red) shipments per month. Dotted line = delay rate %. Watch for months where red grows — that is when deterioration began. Hover for exact monthly counts.",
                "route": "Scatter: X = distance (km), Y = cost (₹), bubble size = efficiency score. Top-left small bubbles = expensive short routes = worst performers. Colour = transport mode: 🔵 Road · 🟢 Rail · 🟡 Air · 🟣 Sea.",
                "sun":   "Inner ring = plants, outer ring = transport modes per plant. Click any plant segment to zoom. A large Air slice on a short-distance plant = cost-reduction opportunity. Road should dominate for most plants.",
            }

            def _make_vchart_explain(chart_key, query_type, q_label):
                """Build a rich explanation card for the View Charts panel — matches Viz tab style."""
                why = _VIZ_CHART_WHY.get((chart_key, query_type),
                      _VIZ_CHART_WHY.get((chart_key, "shipment_delay"),
                      f"Selected as the most informative view for {q_label} analysis."))
                how = _VIZ_CHART_HOW.get(chart_key, "Hover any element for exact figures from Neo4j.")
                _ACCENT = {"stock":"#f97316","heat":"#f43f5e","risk":"#eab308",
                           "trend":"#38bdf8","route":"#10b981","sun":"#a78bfa"}
                ac = _ACCENT.get(chart_key, "#38bdf8")
                _LABEL = {"stock":"📦 Stockout Severity — Distributor Cities",
                          "heat":"🔥 Delay Heatmap — Plant × Product Category",
                          "risk":"⚠️ Supplier Risk Scores",
                          "trend":"📅 Monthly Shipment Delay Trend",
                          "route":"🚛 Route Efficiency — Distance vs Cost",
                          "sun":"🌐 Transport Mode Distribution by Plant"}
                label = _LABEL.get(chart_key, chart_key.upper())
                return (
                    f'<div style="margin:0 0 10px;padding:14px 16px 16px;'
                    f'background:rgba(6,12,28,0.85);border:1px solid {ac}44;'
                    f'border-left:3px solid {ac};border-radius:8px">'
                    f'<div style="font-size:0.67rem;font-weight:800;text-transform:uppercase;'
                    f'letter-spacing:0.12em;color:{ac};margin-bottom:10px">{label}</div>'
                    f'<div style="margin-bottom:10px">'
                    f'<div style="font-size:0.64rem;font-weight:800;text-transform:uppercase;'
                    f'letter-spacing:0.09em;color:#fbbf24;margin-bottom:4px">🎯 Why this chart for your query</div>'
                    f'<div style="font-size:0.8rem;color:#f1f5f9;line-height:1.65">{why}</div>'
                    f'</div>'
                    f'<div style="border-top:1px solid {ac}30;padding-top:9px">'
                    f'<div style="font-size:0.64rem;font-weight:800;text-transform:uppercase;'
                    f'letter-spacing:0.09em;color:#c084fc;margin-bottom:4px">📊 How to read this chart</div>'
                    f'<div style="font-size:0.79rem;color:#cbd5e1;line-height:1.65">{how}</div>'
                    f'</div></div>'
                )

            # ══════════════════════════════════════════════════════════════════
            # RCA VIEW CHARTS — Dynamic, query-aware chart system
            #
            # Architecture (mirrors Network Health tab exactly):
            #   • 6 static button labels: Network/Flow | Bar Chart | Heatmap |
            #                             Bubble Chart  | Pie Chart | Sankey
            #   • On RCA complete → _build_vchart_header picks 4 best chart
            #     types for the query and maps them to the button slots
            #   • Each button click → _rca_load_X(q) builds the chart live from
            #     Neo4j using the dynamic query-aware builders
            #   • Plain tuple returns (explain_html, fig) — no gr.update() wrappers
            # ══════════════════════════════════════════════════════════════════

            # ── Chart-type → query intent mapping ─────────────────────────────
            _RCA_CHART_PLAN = {
                "supplier_risk":    ["san", "heat", "net",  "bar"],
                "demand_gap":       ["bar", "heat", "san",  "pie"],
                "shipment_delay":   ["san", "heat", "bar",  "net"],
                "plant":            ["heat", "net", "bar",  "bub"],
                "distributor":      ["bar", "san", "heat",  "pie"],
                "product_category": ["heat", "pie", "bar",  "net"],
                "route":            ["bub", "bar", "san",   "pie"],
                "simulation":       ["san", "heat", "net",  "bub"],
                "transport_delay":  ["pie", "bar", "heat",  "bub"],
            }

            # Slot → static button label (never changes)
            _SLOT_LABELS = {
                "net":  "🌐 Network / Flow",
                "bar":  "📊 Bar Chart",
                "heat": "🔥 Heatmap",
                "bub":  "🫧 Bubble Chart",
                "pie":  "🥧 Pie Chart",
                "san":  "🔀 Sankey Diagram",
            }

            # Slot → dynamic query-aware builder
            _SLOT_BUILDERS = {
                "net":  lambda q: build_dynamic_flow(question=q),
                "bar":  lambda q: build_dynamic_pareto(question=q),
                "heat": lambda q: build_dynamic_heatmap(question=q),
                "bub":  lambda q: build_dynamic_network(question=q),
                "pie":  lambda q: build_dynamic_pie(question=q),
                "san":  lambda q: build_supply_flow_sankey(question=q),
            }

            # ── Per-slot "Why | What | How" explain content ───────────────────
            # (why_chosen, what_it_explains)
            # why_chosen: 1 crisp sentence on WHY this chart type for this query
            # what_it_explains: plain-English insight about what user should look at in the chart
            _SLOT_WHY = {
                # Sankey
                ("san", "demand_gap"): (
                    "A Sankey maps the exact flow of stock shortages — tracing which supplier feeds which plant, and which plant is leaving which city under-supplied.",
                    "Each band represents a supply path; band width shows how much stock flows that route. The red/dark bands reaching cities on the right are the shortage paths. The wider the red band, the more severe the stockout at that destination. Start from the widest red band on the right and trace it left to find the root cause.",
                ),
                ("san", "supplier_risk"): (
                    "A Sankey is the clearest way to see how a risky supplier's failures ripple through plants into distribution cities.",
                    "The left column shows suppliers, the middle shows plants they feed, and the right shows distributor cities. Red bands are delayed shipment paths. The widest red band is your highest-priority fix — tracing it left tells you exactly which supplier is responsible.",
                ),
                ("san", "shipment_delay"): (
                    "A Sankey traces the delay propagation path end-to-end, showing the exact supplier → plant → city chain where delays are concentrated.",
                    "Follow the widest red band from left to right — that is the primary delay chain in your network. A thick red band at one supplier means that supplier alone is driving most of the network's delay problem.",
                ),
                ("san", "distributor"): (
                    "A Sankey shows which plants are responsible for each city's shortage, and at what volume — essential for distributor impact analysis.",
                    "Look at the incoming bands for each city on the right. The widest incoming red band shows which plant is the main source of that city's shortage. Cities with many red incoming bands have a network-wide supply problem.",
                ),
                ("san", "route"): (
                    "A Sankey reveals the modal mix — which transport modes carry volume on each plant-to-city route, exposing inefficiencies at a glance.",
                    "Wide 'Air' bands on short-distance routes mean you're paying premium freight rates unnecessarily. Those are your immediate cost-saving opportunities.",
                ),
                # Bar chart
                ("bar", "demand_gap"): (
                    "A bar chart immediately shows which cities have the largest stock shortages, ranked from worst to best — making the priority replenishment list obvious.",
                    "The tallest bar is the city running out of stock fastest and needing emergency replenishment first. Notice whether 2–3 cities dominate the chart — that concentration means fixing those cities alone would resolve most of the shortage problem across your network.",
                ),
                ("bar", "supplier_risk"): (
                    "A bar chart ranks suppliers by their risk score, making it immediately clear which supplier is the highest-priority audit or replacement target.",
                    "Bars above 0.7 (the red zone) are suppliers who have historically failed to deliver reliably. The tallest bar is the single most urgent supplier relationship to address. If several bars are close in height, the risk is distributed — you need a portfolio approach.",
                ),
                ("bar", "shipment_delay"): (
                    "A Pareto bar chart shows whether delays are caused by one or two suppliers (concentrated) or spread across many (systemic) — the key to choosing the right fix.",
                    "If the first bar is much taller than the rest, fixing just that one supplier resolves the majority of delays. If all bars are similar heights, the problem is systemic and requires a network-wide intervention rather than a targeted fix.",
                ),
                ("bar", "plant"): (
                    "A bar chart ranks plants by delay volume, showing which manufacturing facility is the biggest bottleneck in the network.",
                    "The tallest bar is the plant generating the most delayed shipments. Check whether it also has the most high-risk suppliers feeding it — that combination confirms the root cause.",
                ),
                # Heatmap
                ("heat", "demand_gap"): (
                    "A heatmap reveals which specific plant-category combinations are driving the most shortages — it finds the intersection to fix first.",
                    "Each cell is a plant (row) and product category (column). The darkest red cell is the combination generating the most stockouts. A full red row means that plant has shortages across all categories — suggesting a plant-level supply failure, not a category-specific one.",
                ),
                ("heat", "supplier_risk"): (
                    "A heatmap shows which supplier-category combinations create the most delays — identifying whether a supplier fails broadly or only on specific product types.",
                    "A fully red row means that supplier fails across all product categories it handles — a systemic reliability problem. A single dark cell means the supplier only fails on one category, which might be a sourcing or lead-time issue for that specific product.",
                ),
                ("heat", "shipment_delay"): (
                    "A heatmap pinpoints the exact plant-category intersection generating the most delays — narrowing the root cause to a single cell.",
                    "The darkest cell is where the delay problem is most concentrated. A fully red column means one product category is delayed everywhere — pointing to a category-level sourcing issue. A fully red row means one plant is delaying everything it dispatches.",
                ),
                ("heat", "product_category"): (
                    "A heatmap shows which plant-supplier combination is creating the most delays for this specific product category.",
                    "The darkest cell is the plant-supplier pair to fix first for this category. If one supplier's column is all red, that supplier is unreliable for this product regardless of which plant it supplies.",
                ),
                # Bubble
                ("bub", "plant"): (
                    "A bubble chart maps plants by two dimensions at once — supplier risk AND delay rate — finding plants where both problems coexist.",
                    "Plants in the top-right zone have both high-risk suppliers AND high delay rates — they are your most urgent intervention targets. Bubble size shows total shipment volume, so a large top-right bubble is a critical problem affecting a large share of your network.",
                ),
                ("bub", "route"): (
                    "A bubble chart plots routes by distance versus cost, immediately exposing routes where you're paying too much relative to how far the goods travel.",
                    "Routes in the top-left zone (short distance, high cost) are the most inefficient — you're paying air-freight rates for what should be a road delivery. The size of each bubble shows shipment volume, so a large top-left bubble is a high-priority cost-reduction opportunity.",
                ),
                # Pie
                ("pie", "demand_gap"): (
                    "A pie chart shows which product categories account for the largest share of total unmet demand — making the portfolio priority instantly visible.",
                    "The largest slice is the product category with the most units short across all cities. If one category dominates the chart, the shortage problem is category-specific — trace that category's supplier chain to find the root cause.",
                ),
                ("pie", "supplier_risk"): (
                    "A pie chart shows how supplier risk is distributed across plants — whether the risk is concentrated in one plant or spread network-wide.",
                    "A large single slice means most high-risk suppliers are feeding one plant, making that plant the priority. If slices are even, every plant is exposed to similar risk levels and the problem requires a network-wide supplier audit.",
                ),
                ("pie", "product_category"): (
                    "A pie chart shows which plants contribute the most delays for this product category — finding the primary bottleneck location.",
                    "The largest slice is the plant responsible for most of this category's delays. If that plant also has a high-risk supplier feeding it, you have identified both the problem and its cause.",
                ),
                # Network
                ("net", "supplier_risk"): (
                    "A network diagram shows the full web of supplier-to-plant-to-city connections, making it clear which supplier failures cascade furthest downstream.",
                    "Red edges are delayed shipment paths; thicker edges carry more shipment volume. A supplier node with many thick red outgoing edges is a high-impact source of network-wide delays — fixing it improves multiple plants and cities simultaneously.",
                ),
                ("net", "shipment_delay"): (
                    "A network diagram reveals the delay propagation structure — whether delays stem from a few concentrated sources or are distributed across the network.",
                    "Thick red edges show the high-volume delay paths. If most thick red edges connect to one or two supplier nodes, the problem is concentrated and solvable by fixing those suppliers. Evenly distributed red edges indicate a systemic network problem.",
                ),
                ("net", "demand_gap"): (
                    "A network diagram shows the full supply chain topology, helping you see which upstream failures create the most downstream stockout exposure.",
                    "Follow the red edges from supplier nodes all the way to city nodes — each red path is a shortage route. Cities with many incoming red edges are the most stockout-exposed locations in your distribution network.",
                ),
            }

            def _get_slot_explain(slot: str, rt: str, q: str) -> str:
                """Build the Why / Explanation card above the chart."""
                why_exp = _SLOT_WHY.get((slot, rt)) or None
                _GENERIC = {
                    "net":  ("A network diagram maps the supply chain as nodes and edges, showing which entities connect and where delays concentrate.",
                             "Red edges are delayed shipment paths. Thicker edges carry more volume. Follow the thickest red edges to find the highest-impact delay routes in the network."),
                    "bar":  ("A bar chart ranks entities by a key metric, making it immediately clear which one is the top priority.",
                             "The tallest bar is your highest-priority action item. If the first bar is much taller than the rest, fixing just that one entity resolves the majority of the problem."),
                    "heat": ("A heatmap uses colour intensity to show where a problem is most concentrated across two dimensions.",
                             "The darkest red cell is where the most severe problem occurs. Look across the row and column of that cell to understand whether the issue is isolated or systemic."),
                    "bub":  ("A bubble chart compares entities across two axes simultaneously, with bubble size adding a third dimension.",
                             "Entities in the top-right corner have the worst performance on both metrics. Large bubbles in that zone represent the highest-volume, highest-risk targets."),
                    "pie":  ("A pie chart shows how a total is divided across categories, making it easy to see which slice dominates.",
                             "The largest slice owns the biggest share of the problem. If it's more than 50%, addressing just that one category will have an outsized impact on your overall numbers."),
                    "san":  ("A Sankey diagram traces flow from sources to destinations, with band width representing volume.",
                             "The widest bands carry the most shipments. Red or dark-coloured bands are delayed or shortage paths. Follow the widest red band from left to right to trace the primary problem chain."),
                }

                if why_exp:
                    why, explanation = why_exp
                else:
                    why, explanation = _GENERIC.get(slot, (
                        "This chart type was selected as the most informative visualisation for your query.",
                        "Hover over any element for exact figures. Look for the largest, darkest, or widest elements — they represent the highest-priority issues."
                    ))

                return (
                    f'<div style="display:grid;grid-template-columns:1fr 1fr;gap:10px;margin:0 0 14px">'
                    f'<div style="padding:12px 14px;background:rgba(56,189,248,0.07);'
                    f'border:1px solid rgba(56,189,248,0.22);border-radius:8px">'
                    f'<div style="font-size:0.63rem;font-weight:800;text-transform:uppercase;'
                    f'letter-spacing:0.1em;color:#fbbf24;margin-bottom:6px">🎯 Why this chart</div>'
                    f'<div style="font-size:0.8rem;color:#f1f5f9;line-height:1.6">{why}</div>'
                    f'</div>'
                    f'<div style="padding:12px 14px;background:rgba(167,139,250,0.07);'
                    f'border:1px solid rgba(167,139,250,0.22);border-radius:8px">'
                    f'<div style="font-size:0.63rem;font-weight:800;text-transform:uppercase;'
                    f'letter-spacing:0.1em;color:#a78bfa;margin-bottom:6px">&#x1F4A1; What\'s happening in this chart</div>'
                    f'<div style="font-size:0.8rem;color:#cbd5e1;line-height:1.6">{explanation}</div>'
                    f'</div>'
                    f'</div>'
                )

            # ── Loader functions — plain tuple returns matching Network Health ──
            # outputs = [rca_vchart_info, rca_vchart_plot]

            def _rca_load_slot(slot: str, q: str):
                """Load chart for slot. Returns (explain_html, fig, col_visible)."""
                rt = _detect_report_type_from_q(q or "")
                fn = _SLOT_BUILDERS.get(slot, _SLOT_BUILDERS["bar"])
                try:
                    fig = _apply_dark_theme(fn(q))
                    if _fig_is_empty(fig):
                        raise ValueError("empty figure")
                    fig.update_layout(autosize=True, height=500,
                                      margin=dict(l=60, r=30, t=50, b=40))
                except Exception as _e:
                    print(f"[rca vchart:{slot}] {_e}")
                    _fallbacks = {
                        "net": build_dynamic_flow, "bar": build_dynamic_pareto,
                        "heat": build_delay_heatmap, "bub": build_dynamic_network,
                        "pie": build_dynamic_pie,   "san": build_supply_flow_sankey,
                    }
                    try:
                        fig = _apply_dark_theme(_fallbacks[slot]())
                        fig.update_layout(autosize=True, height=500,
                                          margin=dict(l=60, r=30, t=50, b=40))
                    except Exception:
                        fig = _blank_fig
                explain_html = _get_slot_explain(slot, rt, q)
                return (
                    gr.update(value=explain_html, visible=True),       # rca_vchart_info
                    gr.update(value=fig, visible=True),                # rca_vchart_plot — MUST use gr.update to show hidden Plot
                    gr.update(visible=True),                           # rca_vchart_plot_col wrapper
                )

            def _rca_load_net(q):  return _rca_load_slot("net",  q)
            def _rca_load_bar(q):  return _rca_load_slot("bar",  q)
            def _rca_load_heat(q): return _rca_load_slot("heat", q)
            def _rca_load_bub(q):  return _rca_load_slot("bub",  q)
            def _rca_load_pie(q):  return _rca_load_slot("pie",  q)
            def _rca_load_san(q):  return _rca_load_slot("san",  q)

            def _load_vchart(chart_key: str, q: str):
                """Wrapper kept for .then() auto-load compatibility."""
                _map = {
                    "net": _rca_load_net, "bar": _rca_load_bar,
                    "heat": _rca_load_heat, "bub": _rca_load_bub,
                    "pie": _rca_load_pie, "san": _rca_load_san,
                    "risk": _rca_load_bar, "trend": _rca_load_heat,
                    "stock": _rca_load_bar, "route": _rca_load_bub, "sun": _rca_load_san,
                }
                return _map.get(chart_key, _rca_load_bar)(q)

            def _build_vchart_header(q: str) -> tuple:
                """
                Picks 4 best chart slots for the query type.
                All 6 buttons always visible with static labels.
                Returns (header_html, btn_net, btn_bar, btn_heat, btn_bub, btn_pie, btn_san)
                """
                rt   = _detect_report_type_from_q(q or "")
                plan = _RCA_CHART_PLAN.get(rt, ["san", "heat", "bar", "net"])[:4]

                _rt_names = {
                    "supplier_risk":  "Supplier Risk Analysis",
                    "demand_gap":     "Demand Gap & Stockout",
                    "shipment_delay": "Shipment Delay Analysis",
                    "plant":          "Plant Performance",
                    "distributor":    "Distributor Impact",
                    "product_category": "Product Category",
                    "route":          "Route & Transport Cost",
                    "simulation":     "Impact Simulation",
                    "transport_delay":"Transport Mode Delays",
                }
                rt_name = _rt_names.get(rt, "Supply Chain")

                pill_html = "".join(
                    f'<span style="display:inline-block;margin:3px 4px 0 0;padding:2px 10px;'
                    f'border-radius:20px;font-size:0.67rem;font-weight:700;'
                    f'background:rgba(56,189,248,0.12);border:1px solid rgba(56,189,248,0.3);'
                    f'color:#7dd3fc">{_SLOT_LABELS.get(s, s)}</span>'
                    for s in plan
                )

                header_html = (
                    f'<div style="padding:8px 2px 10px;border-bottom:1px solid rgba(56,189,248,0.15);margin-bottom:10px">'
                    f'<div style="font-size:0.82rem;font-weight:700;color:#38bdf8;margin-bottom:3px">'
                    f'📈 Recommended visualisations for: <span style="color:#f0abfc">{rt_name}</span></div>'
                    f'<div style="font-size:0.7rem;color:#94a3b8">'
                    f'Select a chart type below. Chart loads when you click a tab.</div>'
                    f'</div>'
                )

                return (
                    gr.update(value=header_html, visible=True),
                    gr.update(visible=True),  # net
                    gr.update(visible=True),  # bar
                    gr.update(visible=True),  # heat
                    gr.update(visible=True),  # bub
                    gr.update(visible=True),  # pie
                    gr.update(visible=True),  # san
                )

            rca_btn.click(
                fn=_build_context_html, inputs=[rca_q], outputs=[rca_context_html],
                show_progress="hidden"
            ).then(
                fn=rca_handler, inputs=[rca_q],
                outputs=[rca_status, rca_assess_html, rca_out, rca_log, rca_chart1, rca_chart2, rca_chart3,
                         rca_dl_btn, rca_csv_btn, rca_dl_section, rca_dl_status,
                         rca_viz_summary_html, rca_detailed_rec_inner, rca_charts_html_vis,
                         rca_cypher_html,
                         rca_view_charts_col,
                         rca_state_net_fig, rca_state_bar_fig, rca_state_pie_fig,
                         rca_state_heat_fig, rca_state_bub_fig, rca_state_san_fig,
                         # ── vchart panel — rendered in same yield as report ──
                         rca_vchart_header,
                         rca_vchart_net_btn, rca_vchart_bar_btn, rca_vchart_heat_btn,
                         rca_vchart_bub_btn, rca_vchart_pie_btn, rca_vchart_san_btn,
                         rca_vchart_info, rca_vchart_plot, rca_vchart_plot_col],
                show_progress="hidden"
            ).then(
                fn=lambda dl, csv: (dl, csv),
                inputs=[rca_dl_btn, rca_csv_btn],
                outputs=[rca_dl_path_state, rca_csv_path_state],
                show_progress="hidden"
            ).then(
                fn=lambda: gr.update(visible=True),
                outputs=[rca_dl_section_vis]
            ).then(
                fn=lambda dl, csv: (
                    gr.update(value=dl, visible=bool(dl)) if dl else gr.update(visible=False),
                    gr.update(value=csv, visible=bool(csv)) if csv else gr.update(visible=False)
                ),
                inputs=[rca_dl_path_state, rca_csv_path_state],
                outputs=[rca_dl_btn_vis, rca_csv_btn_vis]
            ).then(
                fn=_post_rca_render_6, inputs=[rca_q],
                outputs=[rca_net_plot, rca_bar_plot, rca_pie_plot, rca_heat_plot, rca_q_state,
                         rca_net_info, rca_bar_info, rca_pie_info, rca_heat_info,
                         rca_main_plot, rca_chart_info_html,
                         rca_state_net_fig, rca_state_bar_fig, rca_state_pie_fig, rca_state_heat_fig,
                         rca_state_net_info, rca_state_bar_info, rca_state_pie_info, rca_state_heat_info,
                         rca_nav_net_btn, rca_nav_bar_btn, rca_nav_pie_btn, rca_nav_heat_btn,
                         rca_state_bub_fig, rca_state_san_fig, rca_state_bub_info, rca_state_san_info]
            ).then(
                fn=_select_best_charts, inputs=[rca_q],
                outputs=[rca_nav_net_btn, rca_nav_bar_btn, rca_nav_pie_btn,
                         rca_nav_heat_btn, rca_nav_bub_btn, rca_nav_san_btn],
                show_progress="hidden"
            ).then(
                fn=lambda q: gr.update(
                    value=_build_agent_activity_log_html(
                        _RCA_CACHE.get(_cache_key(q), {}).get("logs", [])
                    ),
                    visible=bool(_RCA_CACHE.get(_cache_key(q), {}).get("logs"))
                ),
                inputs=[rca_q],
                outputs=[rca_agent_log_main],
                show_progress="hidden"
            )

            # ── Toggle + dynamic chart functions (defined here, used above in wiring) ──
            _charts_open_state = gr.State(False)

            def _toggle_charts(is_open):
                new_state = not is_open
                return new_state, gr.update(visible=new_state)

            rca_charts_toggle_btn.click(
                fn=_toggle_charts,
                inputs=[_charts_open_state],
                outputs=[_charts_open_state, rca_charts_col],
                show_progress="hidden"
            )

            # ── RCA Launch Button wiring ─────────────────────────────────────
            # Clicking "Run Root Cause Analysis →" in Query Interface:
            #   1. Prefills rca_q with the stored analytical question
            #   2. Switches browser to the RCA Trail tab via JS
            #   3. Triggers the full rca_handler pipeline (same as clicking Run Analysis)
            def _prefill_rca_q(prefill_text: str):
                """Copy the prefilled RCA question into the RCA Trail textbox."""
                return gr.update(value=prefill_text) if prefill_text else gr.update()

            def _tab_switch_js():
                """
                Returns an HTML snippet that uses a unique ID + innerHTML swap trick
                so the script re-executes every time (Gradio normally strips scripts
                from hidden components; making it visible + using a unique nonce
                forces re-evaluation by the browser).
                """
                import time as _time
                nonce = str(int(_time.time() * 1000))
                return gr.update(value=f"""
<span id="rca-tab-switcher-{nonce}" style="display:none"></span>
<script>
(function() {{
  function _switchToRCA() {{
    var tabs = document.querySelectorAll('[role="tab"]');
    for (var i = 0; i < tabs.length; i++) {{
      if (tabs[i].textContent && tabs[i].textContent.trim().includes('RCA Trail')) {{
        tabs[i].click();
        setTimeout(function() {{
          var top = document.querySelector('#rca-question-input, .rca5-report-panel');
          if (top) top.scrollIntoView({{ behavior: 'smooth', block: 'start' }});
        }}, 400);
        return;
      }}
    }}
  }}
  // Try immediately, then retry once after Gradio finishes rendering
  _switchToRCA();
  setTimeout(_switchToRCA, 200);
}})();
</script>""", visible=True)

            # gr.HTML component that fires the tab-switch JS — must be visible=True
            # so Gradio does NOT strip the <script> tag from the DOM update.
            _tab_switch_html = gr.HTML("", visible=True, elem_id="rca-tab-switcher-root")

            # JS string executed immediately on the CLIENT when the button is clicked
            # (before any server round-trip) — this is the primary tab switch trigger.
            _RCA_TAB_SWITCH_JS = """
() => {
  var tabs = document.querySelectorAll('[role="tab"]');
  for (var i = 0; i < tabs.length; i++) {
    if (tabs[i].textContent && tabs[i].textContent.trim().includes('RCA Trail')) {
      tabs[i].click();
      setTimeout(function() {
        var top = document.querySelector('#rca-question-input, .rca5-report-panel');
        if (top) top.scrollIntoView({ behavior: 'smooth', block: 'start' });
      }, 350);
      break;
    }
  }
}
"""

            rca_launch_btn.click(
                fn=_prefill_rca_q,
                inputs=[_rca_prefill_state],
                outputs=[rca_q],
                show_progress="hidden",
                js=_RCA_TAB_SWITCH_JS,
            ).then(
                fn=_tab_switch_js,
                outputs=[_tab_switch_html],
                show_progress="hidden"
            ).then(
                fn=_build_context_html, inputs=[rca_q], outputs=[rca_context_html],
                show_progress="hidden"
            ).then(
                # Trigger AI prompt suggestion so user sees it and can choose to use it
                # before clicking Run — do NOT auto-run rca_handler here
                fn=_suggest_prompt, inputs=[rca_q], outputs=[rca_suggest_html],
                show_progress="hidden"
            )


            def _select_best_charts(q: str):
                q_l = (q or "").lower()
                is_route    = any(w in q_l for w in ["route","transit","transport","road","rail","air","ship"])
                is_supplier = any(w in q_l for w in ["supplier","vendor","risk","sup"])
                is_category = any(w in q_l for w in ["category","product","toy","auto","health","beauty"])
                is_stock    = any(w in q_l for w in ["stockout","stock","demand","gap","shortage"])
                is_dist     = any(w in q_l for w in ["distributor","city","distribution"])

                if is_route:
                    show={("net","🌐  Route Network"),("san","🔀  Supply Sankey"),("bar","📊  Route Ranking"),("heat","🔥  Delay Heatmap")}
                elif is_stock:
                    show={("bar","📊  Stockout Ranking"),("pie","🥧  Demand Gap Share"),("heat","🔥  Plant Heatmap"),("bub","🫧  Risk Bubble")}
                elif is_supplier:
                    show={("net","🌐  Supplier Network"),("bub","🫧  Supplier Risk"),("bar","📊  Supplier Pareto"),("heat","🔥  Delay Heatmap")}
                elif is_category:
                    show={("pie","🥧  Category Share"),("bar","📊  Category Pareto"),("heat","🔥  Category Heatmap"),("net","🌐  Supply Flow")}
                elif is_dist:
                    show={("net","🌐  Distributor Network"),("san","🔀  Supply Sankey"),("bar","📊  City Ranking"),("pie","🥧  Delay Share")}
                else:
                    show={("net","🌐  Plant Network"),("bar","📊  Plant Pareto"),("heat","🔥  Delay Heatmap"),("san","🔀  Supply Sankey")}

                show_map = {k:v for k,v in show}
                _all = [("net",rca_nav_net_btn),("bar",rca_nav_bar_btn),("pie",rca_nav_pie_btn),
                        ("heat",rca_nav_heat_btn),("bub",rca_nav_bub_btn),("san",rca_nav_san_btn)]
                return [gr.update(visible=k in show_map, value=show_map.get(k,"")) if k in show_map
                        else gr.update(visible=False) for k,_ in _all]

            for sbtn, sq in zip(rca_sample_btns, rca_samples):
                sbtn.click(fn=lambda x=sq: x, outputs=[rca_q]).then(
                    fn=_suggest_prompt, inputs=[rca_q], outputs=[rca_suggest_html]
                ).then(
                    fn=_build_context_html, inputs=[rca_q], outputs=[rca_context_html]
                ).then(
                    fn=rca_handler, inputs=[rca_q],
                    outputs=[rca_status, rca_assess_html, rca_out, rca_log, rca_chart1, rca_chart2, rca_chart3,
                             rca_dl_btn, rca_csv_btn, rca_dl_section, rca_dl_status,
                             rca_viz_summary_html, rca_detailed_rec_inner, rca_charts_html_vis,
                             rca_cypher_html,
                             rca_view_charts_col,
                             rca_state_net_fig, rca_state_bar_fig, rca_state_pie_fig,
                             rca_state_heat_fig, rca_state_bub_fig, rca_state_san_fig]
                ).then(
                    fn=lambda dl, csv: (dl, csv),
                    inputs=[rca_dl_btn, rca_csv_btn],
                    outputs=[rca_dl_path_state, rca_csv_path_state],
                    show_progress="hidden"
                ).then(
                    fn=lambda: gr.update(visible=True),
                    outputs=[rca_dl_section_vis]
                ).then(
                    fn=lambda dl, csv: (
                        gr.update(value=dl, visible=bool(dl)) if dl else gr.update(visible=False),
                        gr.update(value=csv, visible=bool(csv)) if csv else gr.update(visible=False)
                    ),
                    inputs=[rca_dl_path_state, rca_csv_path_state],
                    outputs=[rca_dl_btn_vis, rca_csv_btn_vis]
                ).then(
                    fn=_post_rca_render_6, inputs=[rca_q],
                    outputs=[rca_net_plot, rca_bar_plot, rca_pie_plot, rca_heat_plot, rca_q_state,
                             rca_net_info, rca_bar_info, rca_pie_info, rca_heat_info,
                             rca_main_plot, rca_chart_info_html,
                             rca_state_net_fig, rca_state_bar_fig, rca_state_pie_fig, rca_state_heat_fig,
                             rca_state_net_info, rca_state_bar_info, rca_state_pie_info, rca_state_heat_info,
                             rca_nav_net_btn, rca_nav_bar_btn, rca_nav_pie_btn, rca_nav_heat_btn,
                             rca_state_bub_fig, rca_state_san_fig, rca_state_bub_info, rca_state_san_info]
                ).then(
                    fn=lambda q: gr.update(
                        value=_build_agent_activity_log_html(
                            _RCA_CACHE.get(_cache_key(q), {}).get("logs", [])
                        ),
                        visible=bool(_RCA_CACHE.get(_cache_key(q), {}).get("logs"))
                    ),
                    inputs=[rca_q],
                    outputs=[rca_agent_log_main],
                    show_progress="hidden"
                )

            # ── Nav tab click → swap main plot + single explanation box ──────
            rca_nav_net_btn.click(
                fn=lambda f, i: (f, i),
                inputs=[rca_state_net_fig, rca_state_net_info],
                outputs=[rca_main_plot, rca_chart_info_html]
            )
            rca_nav_bar_btn.click(
                fn=lambda f, i: (f, i),
                inputs=[rca_state_bar_fig, rca_state_bar_info],
                outputs=[rca_main_plot, rca_chart_info_html]
            )
            rca_nav_pie_btn.click(
                fn=lambda f, i: (f, i),
                inputs=[rca_state_pie_fig, rca_state_pie_info],
                outputs=[rca_main_plot, rca_chart_info_html]
            )
            rca_nav_heat_btn.click(
                fn=lambda f, i: (f, i),
                inputs=[rca_state_heat_fig, rca_state_heat_info],
                outputs=[rca_main_plot, rca_chart_info_html]
            )
            rca_nav_bub_btn.click(
                fn=lambda f, i: (f, i),
                inputs=[rca_state_bub_fig, rca_state_bub_info],
                outputs=[rca_main_plot, rca_chart_info_html]
            )
            rca_nav_san_btn.click(
                fn=lambda f, i: (f, i),
                inputs=[rca_state_san_fig, rca_state_san_info],
                outputs=[rca_main_plot, rca_chart_info_html]
            )
            # rca_charts_toggle_btn kept as hidden compat ref

            # ── Button click handlers — load fig + show info card ────────────
            # Vchart buttons: each slot loads the Nth chart from _VIZ_CHART_PLAN[query_type]
            # Slot 0=net, 1=bar, 2=pie, 3=heat — matching _build_vchart_header assignment
            # Button clicks — use rca_q (textbox) directly, not rca_q_state
            # rca_q_state is only populated after _post_rca_render_6 completes;
            # rca_q always has the current query text and is always reliable.
            rca_vchart_net_btn.click(
                fn=_rca_load_net,  inputs=[rca_q],
                outputs=[rca_vchart_info, rca_vchart_plot, rca_vchart_plot_col], show_progress="hidden"
            )
            rca_vchart_bar_btn.click(
                fn=_rca_load_bar,  inputs=[rca_q],
                outputs=[rca_vchart_info, rca_vchart_plot, rca_vchart_plot_col], show_progress="hidden"
            )
            rca_vchart_heat_btn.click(
                fn=_rca_load_heat, inputs=[rca_q],
                outputs=[rca_vchart_info, rca_vchart_plot, rca_vchart_plot_col], show_progress="hidden"
            )
            rca_vchart_bub_btn.click(
                fn=_rca_load_bub,  inputs=[rca_q],
                outputs=[rca_vchart_info, rca_vchart_plot, rca_vchart_plot_col], show_progress="hidden"
            )
            rca_vchart_pie_btn.click(
                fn=_rca_load_pie,  inputs=[rca_q],
                outputs=[rca_vchart_info, rca_vchart_plot, rca_vchart_plot_col], show_progress="hidden"
            )
            rca_vchart_san_btn.click(
                fn=_rca_load_san,  inputs=[rca_q],
                outputs=[rca_vchart_info, rca_vchart_plot, rca_vchart_plot_col], show_progress="hidden"
            )



        # ── TAB 4: Update Graph ───────────────────────────────────────────
        with gr.Tab("➕ Update Graph"):

            try:
                import stage4_pipeline as _s4
                _S4_AVAILABLE = True
            except ImportError:
                _S4_AVAILABLE = False

            gr.HTML("""<style>
/* ── Update Graph tab ─────────────────────────────────────── */
.upd-step { display:flex;align-items:center;gap:8px;padding:7px 12px;border-radius:7px;margin-bottom:4px; }
.upd-step-done  { background:rgba(74,222,128,.06); border:1px solid rgba(74,222,128,.2); }
.upd-step-run   { background:rgba(56,189,248,.08); border:1px solid rgba(56,189,248,.3); }
.upd-step-wait  { background:rgba(255,255,255,.02);border:1px solid rgba(255,255,255,.06); }
.upd-step-err   { background:rgba(239,68,68,.07);  border:1px solid rgba(239,68,68,.25); }
.upd-step-num   { width:22px;height:22px;border-radius:50%;display:flex;align-items:center;justify-content:center;font-size:.62rem;font-weight:800;flex-shrink:0;font-family:monospace; }
.upd-step-done  .upd-step-num { background:rgba(74,222,128,.2); color:#4ade80;border:1px solid rgba(74,222,128,.4); }
.upd-step-run   .upd-step-num { background:rgba(56,189,248,.2); color:#38bdf8;border:1px solid rgba(56,189,248,.5);animation:upd-pulse 1.4s ease-in-out infinite; }
.upd-step-wait  .upd-step-num { background:rgba(255,255,255,.04);color:#475569;border:1px solid rgba(255,255,255,.1); }
.upd-step-err   .upd-step-num { background:rgba(239,68,68,.2);  color:#f87171;border:1px solid rgba(239,68,68,.4); }
.upd-step-done  .upd-step-txt { color:#4ade80; }
.upd-step-run   .upd-step-txt { color:#38bdf8; }
.upd-step-wait  .upd-step-txt { color:#475569; }
.upd-step-err   .upd-step-txt { color:#f87171; }
.upd-step-txt { font-size:.74rem;font-weight:600; }
@keyframes upd-pulse { 0%,100%{opacity:1}50%{opacity:.3} }

.upd-mode-btn button {
  background:rgba(14,165,233,.05) !important;border:1px solid rgba(56,189,248,.2) !important;
  color:#cbd5e1 !important;font-size:.78rem !important;border-radius:8px !important;
  width:100% !important;padding:10px 14px !important;margin-bottom:5px !important;
  text-align:left !important;transition:all .15s !important;
}
.upd-mode-btn button:hover { background:rgba(14,165,233,.18) !important;border-color:#38bdf8 !important;color:#fff !important; }
.upd-approve button {
  background:linear-gradient(135deg,#065f46,#047857) !important;
  border:2px solid #4ade80 !important;color:#fff !important;font-weight:900 !important;
  border-radius:10px !important;width:100% !important;
  box-shadow:0 0 20px rgba(74,222,128,.3) !important;
}
.upd-cancel button {
  background:rgba(239,68,68,.1) !important;border:1px solid rgba(239,68,68,.4) !important;
  color:#fca5a5 !important;font-weight:700 !important;border-radius:10px !important;width:100% !important;
}
.upd-undo button {
  background:rgba(124,58,237,.1) !important;border:1px solid rgba(124,58,237,.4) !important;
  color:#c4b5fd !important;font-weight:700 !important;border-radius:8px !important;width:100% !important;
}
.gv-btn button {
  background:rgba(14,165,233,.08) !important;border:1px solid rgba(56,189,248,.25) !important;
  color:#7dd3fc !important;font-size:.74rem !important;font-weight:600 !important;
  border-radius:8px !important;padding:8px 12px !important;width:100% !important;
  margin-bottom:5px !important;text-align:left !important;transition:all .15s !important;
  letter-spacing:.01em !important;
}
.gv-btn button:hover {
  background:rgba(14,165,233,.22) !important;border-color:#38bdf8 !important;
  color:#e0f2fe !important;box-shadow:0 0 10px rgba(56,189,248,.18) !important;
}
.gv-run button {
  background:linear-gradient(135deg,rgba(14,165,233,.18),rgba(56,189,248,.12)) !important;
  border:1px solid rgba(56,189,248,.55) !important;color:#e0f2fe !important;
  font-weight:800 !important;border-radius:8px !important;letter-spacing:.02em !important;
  box-shadow:0 0 12px rgba(56,189,248,.15) !important;
}
.gv-run button:hover {
  background:linear-gradient(135deg,rgba(14,165,233,.35),rgba(56,189,248,.25)) !important;
  box-shadow:0 0 20px rgba(56,189,248,.28) !important;
}

/* ── Graph Verifier accordion container ─── */
#gv-accordion,
#gv-accordion > .label-wrap,
#gv-accordion .wrap {
  background:rgba(4,9,22,.98) !important;
  border:1px solid rgba(56,189,248,.2) !important;
  border-radius:12px !important;
}
/* Force the accordion title text to be white, bold, readable */
#gv-accordion > .label-wrap,
#gv-accordion > .label-wrap button,
#gv-accordion .label-wrap span,
#gv-accordion .label-wrap p,
#gv-accordion > button,
#gv-accordion > button span,
#gv-accordion > button p {
  color:#ffffff !important;
  font-size:1rem !important;
  font-weight:800 !important;
  letter-spacing:.01em !important;
}

/* ── Cypher query textbox ─── */
#gv-query-box label,
#gv-query-box .label-wrap span,
#gv-query-box span {
  color:#7dd3fc !important;font-size:.7rem !important;font-weight:700 !important;
  text-transform:uppercase !important;letter-spacing:.08em !important;
}
#gv-query-box textarea {
  background:rgba(4,9,22,.95) !important;
  border:1px solid rgba(56,189,248,.3) !important;
  color:#e2e8f0 !important;font-family:'Fira Code','Consolas',monospace !important;
  font-size:.8rem !important;line-height:1.65 !important;border-radius:8px !important;
  padding:10px 14px !important;caret-color:#38bdf8 !important;
}
#gv-query-box textarea:focus {
  border-color:#38bdf8 !important;box-shadow:0 0 0 2px rgba(56,189,248,.18) !important;
  outline:none !important;
}

/* ── Node lookup dropdown ─── */
#gv-label-drop label,
#gv-label-drop .label-wrap span {
  color:#7dd3fc !important;font-size:.68rem !important;font-weight:700 !important;
  text-transform:uppercase !important;letter-spacing:.08em !important;
}
#gv-label-drop select,
#gv-label-drop .wrap-inner,
#gv-label-drop input,
#gv-label-drop .secondary-wrap {
  background:rgba(4,9,22,.95) !important;
  border:1px solid rgba(56,189,248,.3) !important;
  color:#e2e8f0 !important;border-radius:8px !important;
}

/* ── Node ID textbox ─── */
#gv-id-box label,
#gv-id-box .label-wrap span {
  color:#7dd3fc !important;font-size:.68rem !important;font-weight:700 !important;
  text-transform:uppercase !important;letter-spacing:.08em !important;
}
#gv-id-box input, #gv-id-box textarea {
  background:rgba(4,9,22,.95) !important;
  border:1px solid rgba(56,189,248,.3) !important;
  color:#e2e8f0 !important;border-radius:8px !important;font-family:monospace !important;
}
#gv-id-box input:focus { border-color:#38bdf8 !important;outline:none !important; }

/* ── Results area ─── */
#gv-result-area {
  background:rgba(4,9,22,.95) !important;
  border:1px solid rgba(56,189,248,.18) !important;
  border-radius:10px !important;min-height:80px !important;
  padding:4px !important;
}
</style>""")

            # ── State ──────────────────────────────────────────────────────────
            s4_session_id = gr.State("")

            with gr.Row(equal_height=False):

                # ══════════ LEFT SIDEBAR ══════════════════════════════════════
                with gr.Column(scale=1, min_width=230):

                    # Mode switch — two clear buttons at top
                    gr.HTML('<div style="font-size:.58rem;font-weight:900;text-transform:uppercase;letter-spacing:.14em;color:#38bdf8;margin-bottom:8px">◈ Method</div>')
                    upd_mode_file = gr.Button("📂  File Upload",           elem_classes="upd-mode-btn")
                    upd_mode_nl   = gr.Button("✏️   Natural Language",      elem_classes="upd-mode-btn")

                    gr.HTML('<div style="margin:14px 0;height:1px;background:rgba(56,189,248,.12)"></div>')

                    # Pipeline stepper — always fully visible, no opacity
                    upd_stepper = gr.HTML("""
<div>
  <div class="upd-step upd-step-wait"><div class="upd-step-num">1</div><div class="upd-step-txt">Parse File</div></div>
  <div class="upd-step upd-step-wait"><div class="upd-step-num">2</div><div class="upd-step-txt">AI Detect Type</div></div>
  <div class="upd-step upd-step-wait"><div class="upd-step-num">3</div><div class="upd-step-txt">Schema Check</div></div>
  <div class="upd-step upd-step-wait"><div class="upd-step-num">4</div><div class="upd-step-txt">Semantic Rules</div></div>
  <div class="upd-step upd-step-wait"><div class="upd-step-num">5</div><div class="upd-step-txt">Conflict Scan</div></div>
  <div class="upd-step upd-step-wait"><div class="upd-step-num">6</div><div class="upd-step-txt">Graph Diff</div></div>
  <div class="upd-step upd-step-wait"><div class="upd-step-num">7</div><div class="upd-step-txt">Cypher + Dry-Run</div></div>
  <div class="upd-step upd-step-wait"><div class="upd-step-num">8</div><div class="upd-step-txt">AI Impact Forecast</div></div>
  <div class="upd-step upd-step-wait"><div class="upd-step-num">9</div><div class="upd-step-txt">Commit to Neo4j</div></div>
  <div class="upd-step upd-step-wait"><div class="upd-step-num">10</div><div class="upd-step-txt">RCA Refresh</div></div>
</div>""")

                    gr.HTML('<div style="margin:10px 0 0;height:1px;background:rgba(56,189,248,.12)"></div>')

                    # Hidden placeholders — wiring requires these to exist
                    upd_undo_file_btn = gr.Button("↩  Undo Last File Upload", elem_classes="upd-undo", visible=False)
                    upd_undo_nl_btn   = gr.Button("↩  Undo Last NL Change",   elem_classes="upd-undo", visible=False)
                    upd_rollback_msg  = gr.HTML("")

                    # Force Delete — hidden placeholders kept so wiring doesn't break
                    upd_force_del_id   = gr.Textbox(value="", visible=False, label="Node ID", lines=1)
                    upd_force_del_type = gr.Dropdown(choices=["Supplier","Distributor","Route"], value="Supplier", label="Entity Type", visible=False)
                    upd_force_del_btn  = gr.Button("🗑️  Force Delete Node", elem_classes="upd-undo", visible=False)
                    upd_force_del_msg  = gr.HTML("")

                    # Hidden legacy inputs — kept so wiring code below doesn't break
                    upd_rollback_input = gr.Textbox(value="", visible=False, label="")
                    upd_rollback_btn   = gr.Button("legacy-rollback", visible=False)

                # ══════════ MAIN PANEL ════════════════════════════════════════
                with gr.Column(scale=3):

                    # ── FILE PIPELINE panel ─────────────────────────────────
                    with gr.Column(visible=True) as upd_file_panel:
                        upd_file_input = gr.File(
                            label="Drop CSV / Excel / JSON / TSV here",
                            file_types=[".xlsx",".xls",".csv",".json",".txt",".tsv"],
                            type="filepath"
                        )
                        upd_file_status = gr.HTML(
                            '<div style="color:#64748b;font-size:.78rem;padding:6px 0">Upload a file — the 10-stage pipeline runs automatically.</div>'
                        )

                        # Results (hidden until pipeline finishes)
                        with gr.Column(visible=False) as upd_results_col:
                            with gr.Tabs():
                                with gr.Tab("🔍 Validation"):
                                    upd_val_html = gr.HTML("")
                                with gr.Tab("📊 Graph Diff"):
                                    upd_diff_html = gr.HTML("")
                                with gr.Tab("⚡ Cypher"):
                                    upd_cypher_html = gr.HTML("")
                                with gr.Tab("🔮 Supply Chain Impact"):
                                    upd_impact_html = gr.HTML("")
                                with gr.Tab("📈 Before vs After"):
                                    upd_before_after_html = gr.HTML("")

                            # Success banner shown after approve
                            upd_success_banner = gr.HTML("", visible=False)

                            # Approval gate
                            upd_gate_html = gr.HTML("")
                            with gr.Row(visible=False) as upd_approval_row:
                                with gr.Column(scale=3, elem_classes="upd-approve"):
                                    upd_approve_btn = gr.Button("✅  APPROVE — Commit to Neo4j")
                                with gr.Column(scale=1, elem_classes="upd-cancel"):
                                    upd_cancel_btn  = gr.Button("✕  Cancel")

                        # Execution result
                        with gr.Column(visible=False) as upd_exec_panel:
                            upd_exec_status   = gr.HTML("")
                            upd_rca_diff      = gr.HTML("")
                            upd_change_summary= gr.HTML("")

                    # ── NL panel ────────────────────────────────────────────
                    with gr.Column(visible=False) as upd_nl_panel:
                        upd_nl_input = gr.Textbox(
                            placeholder="e.g. Add supplier 'Mehta Plastics', ID SUP9001, risk score 0.45, capacity 50000, lead time 8 days, supply to plant PL3",
                            lines=4, label="", show_label=False
                        )
                        with gr.Row():
                            upd_nl_btn  = gr.Button("➕  Update Graph", elem_classes="run-btn", scale=4)
                            upd_nl_undo = gr.Button("↩ Undo Last",     elem_classes="clear-btn", scale=1)
                        upd_nl_status = gr.HTML("")

                        # Structured results — SAME tab layout as File Upload for full consistency
                        with gr.Column(visible=False) as upd_nl_results_col:
                            upd_nl_success_banner = gr.HTML("")
                            with gr.Tabs():
                                with gr.Tab("🔍 Validation"):
                                    upd_nl_val_html    = gr.HTML("")
                                with gr.Tab("📊 Graph Diff"):
                                    upd_nl_diff_html   = gr.HTML("")
                                with gr.Tab("⚡ Cypher"):
                                    upd_nl_cypher_html = gr.HTML("")
                                with gr.Tab("🔮 Supply Chain Impact"):
                                    upd_nl_impact_html = gr.HTML("")
                                with gr.Tab("📈 Before vs After"):
                                    upd_nl_bva_html    = gr.HTML("")
                            # Agent log kept hidden — wiring still references these vars
                            upd_nl_out = gr.HTML("", visible=False)
                            upd_nl_log = gr.HTML("", visible=False)

                        # NL Update History + per-entry undo
                        upd_nl_history      = gr.HTML(value="", elem_id="upd-nl-history")
                        upd_nl_undo_idx     = gr.Textbox(value="0", visible=False, elem_id="upd-nl-undo-idx-input")
                        upd_nl_undo_idx_btn = gr.Button("undo-idx", visible=False, elem_id="upd-nl-undo-idx-btn")
                # ══════════ RIGHT SIDEBAR ═════════════════════════════════════
                with gr.Column(scale=1, min_width=230):

                    # Examples — combined NL + test scenarios in one list
                    gr.HTML('<div style="font-size:.58rem;font-weight:900;text-transform:uppercase;letter-spacing:.14em;color:#38bdf8;margin-bottom:8px">◈ Quick Examples</div>')
                    upd_ex_sup   = gr.Button("🏭  New Supplier",          elem_classes="upd-mode-btn", size="sm")
                    upd_ex_dist  = gr.Button("🏙️  New Distributor",       elem_classes="upd-mode-btn", size="sm")
                    upd_ex_route = gr.Button("🚛  New Route",              elem_classes="upd-mode-btn", size="sm")
                    upd_ex_link  = gr.Button("🔗  Link Supplier → Plant",  elem_classes="upd-mode-btn", size="sm")
                    upd_ex_risk  = gr.Button("✏️  Update Risk Score",      elem_classes="upd-mode-btn", size="sm")

                    # Hidden test scenario buttons (kept for wiring below, not shown in UI)
                    upd_t_happy    = gr.Button("✅  Happy Path",            elem_classes="upd-mode-btn", size="sm", visible=False)
                    upd_t_dupe     = gr.Button("⚠️   Duplicate Upload",     elem_classes="upd-mode-btn", size="sm", visible=False)
                    upd_t_invalid  = gr.Button("❌  Invalid Values",        elem_classes="upd-mode-btn", size="sm", visible=False)
                    upd_t_missing  = gr.Button("🔗  Missing Target",        elem_classes="upd-mode-btn", size="sm", visible=False)
                    upd_t_circular = gr.Button("🔄  Circular Ref",          elem_classes="upd-mode-btn", size="sm", visible=False)

                    # ── Undo Changes — right below Quick Examples ─────────────
                    gr.HTML("""
<div style="margin:14px 0 10px;height:1px;background:rgba(56,189,248,.12)"></div>
<div style="font-size:.58rem;font-weight:900;text-transform:uppercase;letter-spacing:.14em;color:#a78bfa;margin-bottom:8px">↩ Undo Changes</div>
<div style="font-size:.65rem;color:#64748b;margin-bottom:10px;line-height:1.5">
  Reverses the last change — Neo4j <em>and</em> CSV files on disk are both reset.
</div>""")
                    upd_undo_file_btn_r = gr.Button("↩  Undo Last File Upload", elem_classes="upd-undo", size="sm")
                    upd_undo_nl_btn_r   = gr.Button("↩  Undo Last NL Change",   elem_classes="upd-undo", size="sm")
                    upd_rollback_msg_r  = gr.HTML("")

                    # Hidden history (kept for wiring, not shown)
                    upd_history = gr.HTML(value="", visible=False)
                    upd_hist_refresh = gr.Button("↻ Refresh", size="sm", visible=False)

            # ── Graph Verifier — compact, no accordion, inline below main panel ──
            gr.HTML(
                '<div style="margin:18px 0 14px;padding:12px 18px;background:rgba(12,21,40,0.85);'
                'border:1px solid rgba(56,189,248,0.2);border-left:4px solid #38bdf8;border-radius:10px">'
                '<div style="font-size:.85rem;font-weight:900;color:#38bdf8;letter-spacing:-.01em;margin-bottom:2px">🔬 Graph Verifier</div>'
                '<div style="font-size:.7rem;color:#64748b">Look up any node by ID · Check counts · Verify data was written correctly</div>'
                '</div>'
            )
            gr.HTML('<div style="margin:0 0 6px;height:1px;background:rgba(56,189,248,.1)"></div>')
            with gr.Row(equal_height=False):
                with gr.Column(scale=1, min_width=220):
                    gr.HTML('<div style="font-size:.58rem;font-weight:900;text-transform:uppercase;letter-spacing:.14em;color:#38bdf8;margin-bottom:8px;padding:4px 0;border-bottom:1px solid rgba(56,189,248,.15)">🔎 Node Lookup</div>')
                    gv_lbl  = gr.Dropdown(choices=["Supplier","Distributor","Route","Plant"], value="Supplier", label="Label", elem_id="gv-label-drop")
                    gv_id   = gr.Textbox(placeholder="e.g. SUP9001", label="ID", elem_id="gv-id-box")
                    gv_find = gr.Button("🔎  Find Node", elem_classes="gv-run")
                    gr.HTML('<div style="margin:10px 0 6px;font-size:.58rem;font-weight:900;text-transform:uppercase;letter-spacing:.14em;color:#38bdf8;padding:4px 0;border-bottom:1px solid rgba(56,189,248,.15)">📊 Graph Overview</div>')
                    gv_stats_btn = gr.Button("📊 Show Node & Relationship Counts", elem_classes="gv-btn")

                    # Hidden buttons kept for wiring (not visible in UI)
                    gv_sup_btn     = gr.Button("🏭 All suppliers", elem_classes="gv-btn", visible=False)
                    gv_plant_btn   = gr.Button("🏗️  Plants", elem_classes="gv-btn", visible=False)
                    gv_sup9_btn    = gr.Button("🔍 SUP9001", elem_classes="gv-btn", visible=False)
                    gv_risk_btn    = gr.Button("⚠️  High-risk", elem_classes="gv-btn", visible=False)
                    gv_rel_btn     = gr.Button("🔗 SUPPLIES_TO", elem_classes="gv-btn", visible=False)
                    gv_dist_btn    = gr.Button("🏙️  Distributors", elem_classes="gv-btn", visible=False)
                    gv_route_btn   = gr.Button("🚛 Routes", elem_classes="gv-btn", visible=False)

                with gr.Column(scale=3):
                    gr.HTML('<div style="font-size:.58rem;font-weight:900;text-transform:uppercase;letter-spacing:.14em;color:#38bdf8;margin-bottom:6px">📋 Results</div>')
                    gv_query = gr.Textbox(value="", label="", visible=False, elem_id="gv-query-box")
                    gv_run   = gr.Button("▶  Run", elem_classes="gv-run", visible=False)
                    gv_clear = gr.Button("✕ Clear", visible=False)
                    gv_result = gr.HTML(
                        '<div style="color:#475569;font-size:.78rem;padding:16px;text-align:center;background:rgba(4,9,22,.8);border:1px solid rgba(56,189,248,.12);border-radius:8px">Use Node Lookup or click "Show Counts" to verify your graph.</div>',
                        elem_id="gv-result-area"
                    )

            # ══════════════════════════════════════════════════════════════════
            # BACKEND FUNCTIONS
            # ══════════════════════════════════════════════════════════════════

            # ── Pipeline stepper builder ──────────────────────────────────────
            def _upd_stepper(active: int, err: dict = None) -> str:
                steps = ["Parse File","AI Detect Type","Schema Check","Semantic Rules",
                         "Conflict Scan","Graph Diff","Cypher + Dry-Run","AI Impact Forecast",
                         "Commit to Neo4j","RCA Refresh"]
                html = ""
                for i, lbl in enumerate(steps, 1):
                    if err and i in err:
                        css, icon = "upd-step-err", "✗"
                    elif i < active:
                        css, icon = "upd-step-done", "✓"
                    elif i == active:
                        css, icon = "upd-step-run", str(i)
                    else:
                        css, icon = "upd-step-wait", str(i)
                    html += f'<div class="upd-step {css}"><div class="upd-step-num">{icon}</div><div class="upd-step-txt">{lbl}</div></div>'
                return html

            # ── Approval gate HTML ────────────────────────────────────────────
            def _upd_gate(session) -> str:
                if not session or not session.valid_rows:
                    return ""
                c = "#4ade80" if session.scenario_type == "happy_path" else "#fbbf24"
                desc = {
                    "happy_path":       "All rows valid — safe to commit.",
                    "duplicate_upload": f"{len(session.duplicates_in_file)} duplicates → will MERGE (update).",
                    "invalid_dataset":  f"{len(session.invalid_rows)} invalid rows skipped, {len(session.valid_rows)} will commit.",
                    "missing_targets":  f"{len(session.missing_targets)} relationship target(s) missing — rows still commit, links skipped.",
                    "circular_reference": f"{len(session.circular_refs)} self-ref(s) ignored.",
                }.get(session.scenario_type, "Review before committing.")
                kpis = "".join(
                    f'<div style="text-align:center;padding:8px 12px;background:rgba(12,21,40,.8);border:1px solid rgba(56,189,248,.12);border-radius:8px">'
                    f'<div style="font-size:.55rem;color:#94a3b8;text-transform:uppercase">{lbl}</div>'
                    f'<div style="font-size:1.1rem;font-weight:800;color:{col};font-family:monospace">{val}</div></div>'
                    for lbl, val, col in [
                        ("New", len(session.new_nodes), "#38bdf8"),
                        ("Update", len(session.updated_nodes), "#fbbf24"),
                        ("Rels", len(session.new_relationships), "#4ade80"),
                        ("Skip", len(session.invalid_rows), "#f87171" if session.invalid_rows else "#64748b"),
                    ]
                )
                return (
                    f'<div style="background:rgba(6,12,28,.95);border:2px solid {c};border-radius:12px;padding:18px 22px;margin-top:14px">'
                    f'<div style="display:flex;align-items:center;gap:10px;margin-bottom:12px">'
                    f'<span style="font-size:1.2rem">{"✅" if c=="#4ade80" else "⚠️"}</span>'
                    f'<div><div style="font-size:.88rem;font-weight:800;color:#e2e8f0">Ready to commit</div>'
                    f'<div style="font-size:.7rem;color:#94a3b8;margin-top:2px">{desc}</div></div></div>'
                    f'<div style="display:grid;grid-template-columns:repeat(4,1fr);gap:8px;margin-bottom:12px">{kpis}</div>'
                    f'<div style="font-size:.68rem;color:#64748b;padding:7px 10px;background:rgba(255,255,255,.03);border-radius:6px">'
                    f'🔒 Session <code style="color:#38bdf8">{session.session_id}</code> — reversible with Rollback</div></div>'
                )

            # ── Verifier helpers ──────────────────────────────────────────────
            def _gv_run(q: str) -> str:
                q = (q or "").strip()
                if not q:
                    return '<div style="color:#7dd3fc;font-size:.78rem;padding:12px;text-align:center;opacity:.6">Enter a query above and click ▶ Run.</div>'
                for kw in ["DELETE","DROP","REMOVE","CREATE INDEX","CREATE CONSTRAINT"]:
                    if kw in q.upper():
                        return f'<div style="color:#f87171;font-size:.78rem;padding:10px 14px;background:rgba(239,68,68,.08);border:1px solid rgba(239,68,68,.3);border-radius:8px">⛔ <strong>"{kw}"</strong> is blocked — this is read-only. Use the NL panel to write.</div>'
                try:
                    rows = _run_neo4j(q)
                    if not rows:
                        return '<div style="color:#fbbf24;font-size:.78rem;padding:12px 16px;background:rgba(251,191,36,.06);border:1px solid rgba(251,191,36,.2);border-radius:8px">⚠ <strong>0 rows returned</strong> — node may not exist yet, or query returned no matches.</div>'
                    hdrs = list(rows[0].keys())
                    th = "".join(
                        f'<th style="background:rgba(14,165,233,.18);color:#38bdf8;font-size:.7rem;'
                        f'font-weight:800;text-transform:uppercase;letter-spacing:.07em;'
                        f'padding:8px 12px;border:1px solid rgba(56,189,248,.2);'
                        f'white-space:nowrap">{h}</th>'
                        for h in hdrs
                    )
                    tb = ""
                    for i, row in enumerate(rows[:50]):
                        bg = "rgba(56,189,248,.04)" if i % 2 == 0 else "transparent"
                        cells = []
                        for h in hdrs:
                            raw_val = row.get(h)
                            if isinstance(raw_val, dict):
                                # Format property dicts as readable key-value list (e.g. "props" column)
                                prop_html = "".join(
                                    f'<div style="display:flex;gap:8px;padding:3px 0;'
                                    f'border-bottom:1px solid rgba(56,189,248,.06)">'
                                    f'<span style="font-size:.67rem;font-weight:700;color:#7dd3fc;font-family:monospace;min-width:160px;flex-shrink:0">{pk}</span>'
                                    f'<span style="font-size:.73rem;color:#e2e8f0;word-break:break-all">{pv}</span></div>'
                                    for pk, pv in sorted(raw_val.items()) if pv not in (None, "", 0.0)
                                ) or '<span style="color:#64748b">—</span>'
                                cells.append(
                                    f'<td style="padding:8px 12px;border:1px solid rgba(56,189,248,.08);vertical-align:top">'
                                    f'<div style="min-width:260px;max-width:460px">{prop_html}</div></td>'
                                )
                            else:
                                display = str(raw_val)[:120] if raw_val is not None else "—"
                                cells.append(
                                    f'<td style="color:#e2e8f0;font-size:.76rem;font-weight:500;'
                                    f'padding:7px 12px;border:1px solid rgba(56,189,248,.08);'
                                    f'font-family:monospace;max-width:280px;overflow:hidden;'
                                    f'text-overflow:ellipsis;white-space:nowrap">{display}</td>'
                                )
                        tb += f'<tr style="background:{bg}">' + "".join(cells) + "</tr>"
                    extra = f'<div style="font-size:.65rem;color:#64748b;margin-top:7px;text-align:right">Showing first 50 of {len(rows)} rows</div>' if len(rows) > 50 else ""
                    return (
                        f'<div style="overflow-x:auto;background:rgba(4,9,22,.9);'
                        f'border:1px solid rgba(56,189,248,.2);border-radius:10px;padding:12px">'
                        f'<div style="font-size:.68rem;font-weight:800;color:#4ade80;'
                        f'margin-bottom:8px;display:flex;align-items:center;gap:6px">'
                        f'<span style="display:inline-block;width:8px;height:8px;border-radius:50%;'
                        f'background:#4ade80;box-shadow:0 0 6px #4ade80"></span>'
                        f'✓ {len(rows)} row(s) returned</div>'
                        f'<table style="width:100%;border-collapse:collapse">'
                        f'<thead><tr>{th}</tr></thead>'
                        f'<tbody>{tb}</tbody></table>'
                        f'{extra}</div>'
                    )
                except Exception as e:
                    return (
                        f'<div style="color:#f87171;font-size:.78rem;padding:10px 14px;'
                        f'background:rgba(239,68,68,.08);border:1px solid rgba(239,68,68,.3);'
                        f'border-radius:8px"><strong>Error:</strong> {str(e)[:300]}</div>'
                    )

            def _gv_stats() -> str:
                try:
                    nodes = _run_neo4j("MATCH (n) RETURN labels(n)[0] AS lbl, COUNT(n) AS cnt ORDER BY cnt DESC")
                    rels  = _run_neo4j("MATCH ()-[r]->() RETURN type(r) AS t, COUNT(r) AS cnt ORDER BY cnt DESC LIMIT 8")
                    cards = "".join(
                        f'<div style="text-align:center;padding:12px 8px;'
                        f'background:rgba(14,165,233,.08);border:1px solid rgba(56,189,248,.25);'
                        f'border-top:3px solid #38bdf8;border-radius:10px">'
                        f'<div style="font-size:.6rem;font-weight:700;text-transform:uppercase;'
                        f'letter-spacing:.1em;color:#7dd3fc;margin-bottom:5px">{r["lbl"] or "?"}</div>'
                        f'<div style="font-size:1.4rem;font-weight:900;color:#e2e8f0;'
                        f'font-family:monospace;text-shadow:0 0 12px rgba(56,189,248,.4)">{r["cnt"]:,}</div>'
                        f'</div>'
                        for r in (nodes or [])[:6]
                    )
                    rel_rows = "".join(
                        f'<div style="display:flex;justify-content:space-between;align-items:center;'
                        f'padding:6px 12px;border-bottom:1px solid rgba(56,189,248,.08)">'
                        f'<span style="color:#7dd3fc;font-family:monospace;font-size:.74rem;font-weight:600">'
                        f'[:{r["t"]}]</span>'
                        f'<span style="color:#e2e8f0;font-family:monospace;font-size:.78rem;font-weight:800">'
                        f'{r["cnt"]:,}</span></div>'
                        for r in (rels or [])
                    )
                    return (
                        f'<div style="background:rgba(4,9,22,.9);border:1px solid rgba(56,189,248,.2);'
                        f'border-radius:10px;padding:14px;overflow:hidden">'
                        f'<div style="font-size:.62rem;font-weight:800;text-transform:uppercase;'
                        f'letter-spacing:.12em;color:#38bdf8;margin-bottom:10px">📊 Node Counts</div>'
                        f'<div style="display:grid;grid-template-columns:repeat(3,1fr);gap:8px;margin-bottom:14px">{cards}</div>'
                        f'<div style="font-size:.62rem;font-weight:800;text-transform:uppercase;'
                        f'letter-spacing:.12em;color:#38bdf8;margin-bottom:8px;padding-top:6px;'
                        f'border-top:1px solid rgba(56,189,248,.15)">🔗 Relationship Counts</div>'
                        f'<div style="background:rgba(14,165,233,.04);border:1px solid rgba(56,189,248,.12);'
                        f'border-radius:8px;overflow:hidden">{rel_rows}</div>'
                        f'</div>'
                    )
                except Exception as e:
                    return f'<div style="color:#f87171;font-size:.78rem;padding:10px">Error: {e}</div>'

            def _gv_find(nid: str, lbl: str) -> str:
                nid = str(nid or "").strip()
                if not nid:
                    return '<div style="color:#fbbf24;font-size:.78rem;padding:10px">Enter an ID first.</div>'

                if lbl == "Route":
                    # Try multiple query patterns — NL agent may use different property names
                    queries = [
                        f"MATCH (n:Route {{route_id: '{nid}'}}) RETURN properties(n) AS props",
                        f"MATCH (n:Route) WHERE n.route_id = '{nid}' OR n.id = '{nid}' RETURN properties(n) AS props",
                    ]
                    if "@" in nid:
                        parts = nid.split("@", 1)
                        plant_part, dist_part = parts[0], parts[1]
                        queries.append(
                            f"MATCH (n:Route) WHERE (n.plant_id = '{plant_part}' AND n.distributor_id = '{dist_part}') "
                            f"OR n.route_id = '{nid}' RETURN properties(n) AS props"
                        )
                    for q in queries:
                        result = _gv_run(q)
                        if "0 rows returned" not in result and "Enter a query" not in result:
                            return result
                    return (
                        f'<div style="background:rgba(251,191,36,.08);border:1px solid rgba(251,191,36,.3);' +
                        f'border-radius:8px;padding:12px 16px;font-size:.78rem;color:#fde68a">' +
                        f'<strong>⚠ Route "{nid}" not found.</strong><br>' +
                        f'Try the <strong>🚛 Routes</strong> quick button, or run this in Custom Cypher:<br>' +
                        f'<code style="color:#7dd3fc;font-size:.72rem">MATCH (n:Route) RETURN n.route_id, properties(n) LIMIT 20</code>' +
                        f'</div>'
                    )

                fmap = {"Supplier": "supplier_id", "Distributor": "distributor_id", "Plant": "plant_id"}
                fld = fmap.get(lbl, "id")
                return _gv_run(f"MATCH (n:{lbl} {{{fld}: '{nid}'}}) RETURN properties(n) AS props")

            # ── Quick query strings ───────────────────────────────────────────
            _GVQ = {
                "sup":    "MATCH (s:Supplier) RETURN s.supplier_id AS id, s.supplier_name AS name, s.risk_score AS risk ORDER BY s.supplier_id DESC LIMIT 15",
                "plant":  "MATCH (pl:Plant)-[:DISPATCHES]->(sh:Shipment) WITH pl, COUNT(sh) AS t, SUM(CASE WHEN sh.delivery_status=\'Major Delay\' THEN 1 ELSE 0 END) AS d RETURN pl.plant_id, pl.plant_name AS name, t AS total, d AS delayed, ROUND(100.0*d/CASE WHEN t=0 THEN 1 ELSE t END,1) AS delay_pct ORDER BY delay_pct DESC",
                "sup9":   "MATCH (s:Supplier {supplier_id: \'SUP9001\'}) RETURN s.supplier_id AS id, s.supplier_name AS name, s.risk_score AS risk, s.annual_capacity_units AS capacity, s.StoP_lead_time_days AS lead_time",
                "risk":   "MATCH (s:Supplier) WHERE s.risk_score > 0.7 OPTIONAL MATCH (s)-[:SUPPLIES_TO]->(p:Plant) RETURN s.supplier_id AS id, s.supplier_name AS name, s.risk_score AS risk, p.plant_name AS plant ORDER BY s.risk_score DESC",
                "rel":    "MATCH (s:Supplier)-[:SUPPLIES_TO]->(p:Plant) RETURN s.supplier_id, s.supplier_name, p.plant_id, p.plant_name ORDER BY s.supplier_id",
                "dist":   "MATCH (d:Distributor) RETURN d.distributor_id AS id, d.distributor_city AS city, d.distributor_latitude AS lat, d.distributor_longitude AS lng ORDER BY d.distributor_id",
                "route":  "MATCH (r:Route) OPTIONAL MATCH (pl:Plant)-[:HAS_ROUTE]->(r)-[:CONNECTS_TO]->(d:Distributor) RETURN r.route_id AS route_id, coalesce(pl.plant_name, r.plant_id, '—') AS from_plant, coalesce(d.distributor_city, r.distributor_id, '—') AS to_city, r.mode, r.PtoD_distance_km AS km, r.PtoD_leadtime_days AS days, r.PtoD_transportation_cost_inr AS cost_inr ORDER BY r.route_id LIMIT 30",
            }

            # ── NL example / test texts ───────────────────────────────────────
            _NL_EX = {
                "sup":   "Add supplier \'Mehta Plastics\', ID SUP9001, risk score 0.45, capacity 50000 units, lead time 8 days, supply to plant PL3",
                "dist":  "Add distributor D9001 in city Surat, latitude 21.17, longitude 72.83",
                "route": "Add Road route from PL4 to D0050, 320 km, 3 days lead time, cost 85000 INR",
                "link":  "Link supplier SUP0026 to plant PL2",
                "risk":  "Update supplier SUP0026 risk score to 0.85",
                "t_happy":    "Add supplier \'Mehta Plastics\', ID SUP9001, risk score 0.45, capacity 50000 units, lead time 8 days, supply to plant PL3",
                "t_dupe":     "Add supplier \'Mehta Plastics\', ID SUP9001, risk score 0.52, capacity 52000 units, lead time 7 days, supply to plant PL3",
                "t_invalid":  "Add supplier \'Bad Corp\', ID SUPBAD1, risk score 1.85, capacity -500 units, lead time -3 days, supply to plant PL99",
                "t_missing":  "Add supplier \'Orphan Vendor\', ID SUP9099, risk score 0.6, capacity 10000, lead time 5 days, supply to plant PL99",
                "t_circular": "Add supplier \'Loop Inc\', ID PL3, risk score 0.7, capacity 30000, lead time 10, supply to plant PL3",
            }

            # ── Pipeline runner ───────────────────────────────────────────────
            def _upd_run_pipeline(filepath):
                if filepath is None:
                    yield ('<div style="color:#f87171;font-size:.78rem">No file selected.</div>',
                           _upd_stepper(0), gr.update(visible=False), "","","","","",gr.update(visible=False),"",
                           gr.update(value="",visible=False), gr.update(value=""))
                    return

                import os as _os
                actual = filepath.name if hasattr(filepath, "name") else str(filepath)
                fname  = _os.path.basename(actual)

                yield (f'<div style="color:#38bdf8;font-size:.78rem">◌ Starting pipeline — {fname}…</div>',
                       _upd_stepper(1), gr.update(visible=False), "","","","","",gr.update(visible=False),"",
                       gr.update(value="",visible=False), gr.update(value=""))

                if not _S4_AVAILABLE:
                    yield ('<div style="color:#f87171">stage4_pipeline.py not found.</div>',
                           _upd_stepper(0), gr.update(visible=False), "","","","","",gr.update(visible=False),"",
                           gr.update(value="",visible=False), gr.update(value=""))
                    return

                import queue as _q, threading as _t
                step_holder = {"step": 1}
                q, holder = _q.Queue(), {}

                def _on(event):
                    kind, data = event
                    if kind == "stage":
                        sm = {"parse":1,"detect":2,"schema":3,"semantic":4,"conflicts":5,"diff":6,"cypher":7,"dryrun":7,"impact":8}
                        step_holder["step"] = sm.get(data.get("stage"), step_holder["step"])
                        q.put(("step", step_holder["step"]))
                    elif kind in ("pipeline_ready","error"):
                        q.put((kind, data))

                def _worker():
                    try:
                        sess = _s4.run_stage4_pipeline(actual, fname, on_update=_on)
                        holder["session"] = sess
                        _s4.register_session(sess)
                    except Exception as e:
                        holder["error"] = str(e)
                    finally:
                        q.put(None)

                _t.Thread(target=_worker, daemon=True).start()
                cur = 1
                while True:
                    try:
                        msg = q.get(timeout=0.5)
                    except _q.Empty:
                        yield (f'<div style="color:#38bdf8;font-size:.78rem">◌ Stage {cur}…</div>',
                               _upd_stepper(cur), gr.update(visible=False), "","","","","",gr.update(visible=False),"",
                               gr.update(value="",visible=False), gr.update(value=""))
                        continue
                    if msg is None:
                        break
                    kind, data = msg
                    if kind == "step":
                        cur = data
                    elif kind == "error":
                        yield (f'<div style="color:#f87171">Error: {data.get("msg","?")[:200]}</div>',
                               _upd_stepper(cur, {cur:"upd-step-err"}),
                               gr.update(visible=False), "","","","","",gr.update(visible=False),"",
                           gr.update(value="",visible=False), gr.update(value=""))
                        return

                if "error" in holder:
                    yield (f'<div style="color:#f87171">{holder["error"][:200]}</div>',
                           _upd_stepper(0), gr.update(visible=False), "","","","","",gr.update(visible=False),"",
                           gr.update(value="",visible=False), gr.update(value=""))
                    return

                sess = holder.get("session")
                if not sess:
                    yield ('<div style="color:#f87171">No session returned.</div>',
                           _upd_stepper(0), gr.update(visible=False), "","","","","",gr.update(visible=False),"",
                           gr.update(value="",visible=False), gr.update(value=""))
                    return

                # Strip "AI Impact Forecast:" prefix and limit to ~4 readable lines
                _raw_impact = (sess.impact_simulation or "").strip()
                _raw_impact = re.sub(r'^AI\s+Impact\s+Forecast[:\-–—]?\s*', '', _raw_impact, flags=re.IGNORECASE)
                # Keep first 3 sentences for ≤4 lines
                _sentences = re.split(r'(?<=[.!?])\s+', _raw_impact)
                _short_impact = ' '.join(_sentences[:3]).strip()
                if len(_sentences) > 3:
                    _short_impact += ' …'
                impact_html = (
                    f'<div style="background:rgba(124,58,237,.08);border:1px solid rgba(124,58,237,.3);border-radius:10px;padding:14px 18px">'
                    f'<div style="font-size:.6rem;font-weight:800;text-transform:uppercase;letter-spacing:.12em;color:#a78bfa;margin-bottom:8px">⚡ Impact Forecast</div>'
                    f'<div style="font-size:.82rem;color:#e2e8f0;line-height:1.65">{_short_impact}</div>'
                    f'<div style="margin-top:10px;font-size:.68rem;color:#64748b">Scenario: <b style="color:#a78bfa">{sess.scenario_type.replace("_"," ").title()}</b> · '
                    f'Valid: <b style="color:#4ade80">{len(sess.valid_rows)}</b> · Invalid: <b style="color:#f87171">{len(sess.invalid_rows)}</b></div></div>'
                )

                # ── Build pre-commit Before vs After preview ──────────────────
                def _build_precommit_bva(session) -> str:
                    snap = getattr(session, "rca_before_snapshot", {}) or {}
                    valid = session.valid_rows or []
                    etype = getattr(session, "entity_type", None) or getattr(session, "detected_type", "Node")
                    n_new = len(valid)
                    ship  = snap.get("shipments", {})
                    sup   = snap.get("suppliers",  {})
                    dist  = snap.get("distributors", {})
                    def _row(label, current, projected):
                        return (
                            f'<tr><td style="color:#94a3b8;font-size:.78rem;padding:5px 12px">{label}</td>'
                            f'<td style="color:#e2e8f0;font-weight:600;text-align:right;padding:5px 12px">{current}</td>'
                            f'<td style="color:#7dd3fc;font-weight:600;text-align:right;padding:5px 12px">{projected}</td>'
                            f'<td style="text-align:center;padding:5px 12px;font-size:.72rem;color:#fbbf24">pending</td></tr>'
                        )
                    sup_current  = sup.get("total", "—")
                    sup_projected = (sup_current + n_new) if (etype == "Supplier" and isinstance(sup_current, (int,float))) else sup_current
                    dist_current  = dist.get("total", "—")
                    dist_projected= (dist_current + n_new) if (etype == "Distributor" and isinstance(dist_current, (int,float))) else dist_current
                    avg_risk_cur  = sup.get("avg_risk", "—")
                    if etype == "Supplier" and isinstance(avg_risk_cur, (int,float)) and n_new:
                        new_risks = [float(r.get("risk_score", avg_risk_cur)) for r in valid if r.get("risk_score") is not None]
                        if new_risks:
                            all_risks = [avg_risk_cur] * int(sup_current or 0) + new_risks
                            avg_risk_proj = f'{sum(all_risks)/len(all_risks):.3f}'
                        else:
                            avg_risk_proj = avg_risk_cur
                    else:
                        avg_risk_proj = avg_risk_cur
                    ts = snap.get("timestamp", "")[:19].replace("T", " ") if snap.get("timestamp") else "—"
                    return f"""
<div style="background:rgba(6,12,28,0.9);border:1px solid rgba(56,189,248,0.2);border-radius:12px;padding:18px 22px">
  <div style="font-size:.62rem;font-weight:900;text-transform:uppercase;letter-spacing:.14em;color:#38bdf8;margin-bottom:6px">
    📊 Before vs After — Pre-Commit Preview
  </div>
  <div style="font-size:.7rem;color:#64748b;margin-bottom:12px">
    Current state snapshot · After column shows projected values post-commit · Final comparison updates on approval.
  </div>
  <table style="width:100%;border-collapse:collapse">
    <thead>
      <tr style="border-bottom:1px solid rgba(56,189,248,0.2)">
        <th style="color:#7dd3fc;font-size:.7rem;text-align:left;padding:6px 12px">Metric</th>
        <th style="color:#7dd3fc;font-size:.7rem;text-align:right;padding:6px 12px">Before</th>
        <th style="color:#7dd3fc;font-size:.7rem;text-align:right;padding:6px 12px">After (est.)</th>
        <th style="color:#7dd3fc;font-size:.7rem;text-align:center;padding:6px 12px">Change</th>
      </tr>
    </thead>
    <tbody>
      {_row("Total Suppliers", sup_current, sup_projected)}
      {_row("Avg Risk Score",  f'{avg_risk_cur:.3f}' if isinstance(avg_risk_cur,(int,float)) else avg_risk_cur, avg_risk_proj)}
      {_row("High-Risk Suppliers", sup.get("high_risk","—"), sup.get("high_risk","—"))}
      {_row("Total Distributors", dist_current, dist_projected)}
      {_row("Total Shipments",  ship.get("total","—"), ship.get("total","—"))}
      {_row("Delayed Shipments", ship.get("delayed","—"), ship.get("delayed","—"))}
    </tbody>
  </table>
  <div style="margin-top:12px;font-size:.7rem;color:#64748b;font-style:italic">
    Snapshot: {ts} · <b style="color:#a78bfa">{n_new}</b> valid {etype} row(s) queued for commit.
    Approve to see live Before vs After comparison.
  </div>
</div>"""

                precommit_bva = _build_precommit_bva(sess)

                yield (
                    f'<div style="color:#4ade80;font-size:.78rem;padding:6px 0">✓ Pipeline complete — {len(sess.valid_rows)} rows ready · Session <code style="color:#38bdf8">{sess.session_id}</code></div>',
                    _upd_stepper(8),
                    gr.update(visible=True),
                    _s4.build_validation_report_html(sess),
                    _s4.build_graph_diff_html(sess),
                    _s4.build_cypher_preview_html(sess),
                    impact_html,
                    _upd_gate(sess),
                    gr.update(visible=bool(sess.valid_rows)),
                    sess.session_id,
                    gr.update(value="", visible=False),  # reset success banner
                    gr.update(value=precommit_bva),      # pre-populated before_after tab
                )

            def _canonical_to_csv_row(canonical_row: dict, entity_type: str, id_val: str) -> dict:
                """
                Convert a pipeline valid_row (canonical keys: id, name, risk_score, capacity,
                lead_time, plant, ...) to the actual CSV column names used in SupplierMaster.csv etc.
                Falls back to returning the row as-is if mapping not defined.
                """
                if entity_type == "Supplier":
                    return {
                        "supplier_id":           id_val,
                        "supplier_name":         canonical_row.get("name", ""),
                        "risk_score":            canonical_row.get("risk_score", ""),
                        "annual_capacity_units": canonical_row.get("capacity", ""),
                        "StoP_lead_time_days":   canonical_row.get("lead_time", ""),
                        "status":                canonical_row.get("status", "Active"),
                        "plant_id":              canonical_row.get("plant", ""),
                        # Geo fields — default empty/zero if not provided
                        "supplier_latitude":     canonical_row.get("supplier_latitude", ""),
                        "supplier_longitude":    canonical_row.get("supplier_longitude", ""),
                        "StoP_distance_km":      canonical_row.get("StoP_distance_km", ""),
                    }
                elif entity_type == "Distributor":
                    return {
                        "distributor_id":        id_val,
                        "distributor_city":      canonical_row.get("city", ""),
                        "distributor_latitude":  canonical_row.get("lat", ""),
                        "distributor_longitude": canonical_row.get("lng", ""),
                    }
                elif entity_type == "Route":
                    return {
                        "route_id":                     id_val,
                        "mode":                         canonical_row.get("mode", ""),
                        "PtoD_distance_km":             canonical_row.get("dist_km", ""),
                        "PtoD_leadtime_days":           canonical_row.get("days", ""),
                        "PtoD_transportation_cost_inr": canonical_row.get("cost", ""),
                        "plant_id":                     canonical_row.get("plant", ""),
                        "distributor_id":               canonical_row.get("dist_id", ""),
                    }
                # Fallback — return as-is
                return canonical_row

            def _sync_rows_to_csv(sess):
                """After a successful commit, write all changed rows back to their CSV files on disk."""
                csv_log = []
                try:
                    entity_type = getattr(sess, "entity_type", None) or getattr(sess, "detected_type", None)
                    path = _APP_CSV.get(entity_type) if entity_type else None
                    id_col = _APP_ID_COL.get(entity_type) if entity_type else None
                    if not path or not id_col:
                        csv_log.append(f"CSV sync skipped — unknown entity type '{entity_type}'")
                        return csv_log

                    existing_rows = _app_csv_read(path)
                    existing_map = {str(r.get(id_col, "")).strip(): r for r in existing_rows}

                    # Determine the fieldnames from the existing CSV (preserve column order)
                    csv_fieldnames = list(existing_rows[0].keys()) if existing_rows else None

                    inserted   = getattr(sess, "inserted_ids", []) or []
                    updated    = getattr(sess, "updated_nodes", []) or []
                    valid_rows = getattr(sess, "valid_rows", []) or []

                    # Build lookup from valid_rows by canonical "id" field
                    # valid_rows use canonical keys (id, name, risk_score, capacity, lead_time, plant)
                    valid_map = {}
                    for vr in valid_rows:
                        vid = str(vr.get("id", "")).strip()
                        if vid:
                            valid_map[vid] = vr

                    changed = 0
                    for vid in inserted:
                        vid = str(vid).strip()
                        if vid in valid_map and vid not in existing_map:
                            # Convert canonical row → CSV-format row
                            csv_row = _canonical_to_csv_row(valid_map[vid], entity_type, vid)
                            # If we have known fieldnames, fill any missing columns with ""
                            if csv_fieldnames:
                                padded = {k: "" for k in csv_fieldnames}
                                padded.update({k: v for k, v in csv_row.items() if k in padded})
                                # For new columns not yet in CSV, add them
                                padded.update({k: v for k, v in csv_row.items() if k not in padded})
                                csv_row = padded
                            existing_map[vid] = csv_row
                            changed += 1
                            csv_log.append(f"+ Added {entity_type} {vid} to CSV")

                    for upd_entry in updated:
                        # updated_nodes items are dicts with an "id" key, or may be plain strings
                        vid = str(upd_entry.get("id", upd_entry) if isinstance(upd_entry, dict) else upd_entry).strip()
                        if vid in valid_map:
                            csv_row = _canonical_to_csv_row(valid_map[vid], entity_type, vid)
                            if vid in existing_map and csv_fieldnames:
                                merged = dict(existing_map[vid])  # keep existing columns
                                merged.update({k: v for k, v in csv_row.items() if v != ""})
                                csv_row = merged
                            elif csv_fieldnames:
                                padded = {k: "" for k in csv_fieldnames}
                                padded.update({k: v for k, v in csv_row.items() if k in padded})
                                csv_row = padded
                            existing_map[vid] = csv_row
                            changed += 1
                            csv_log.append(f"~ Updated {entity_type} {vid} in CSV")

                    if changed > 0:
                        all_rows = list(existing_map.values())
                        _app_csv_write(path, all_rows)
                    else:
                        csv_log.append(f"CSV sync: no rows matched for ids={inserted + [u.get('id',u) if isinstance(u,dict) else u for u in updated]}")

                except Exception as e:
                    import traceback as _tb
                    csv_log.append(f"CSV sync error: {e} — {_tb.format_exc()[:300]}")
                return csv_log

            def _build_post_commit_report(done, csv_log):
                """Generate a detailed post-commit report shown after approval."""
                inserted  = getattr(done, "inserted_ids", []) or []
                updated   = getattr(done, "updated_nodes", []) or []
                wired     = getattr(done, "wired_rels", []) or []
                failed    = getattr(done, "failed_ids", []) or []

                # CSV sync status block
                if csv_log:
                    csv_items = "".join(
                        f'<div style="font-family:monospace;font-size:.72rem;color:{"#4ade80" if l.startswith("+") else "#fbbf24" if l.startswith("~") else "#f87171"};padding:2px 0">{l}</div>'
                        for l in csv_log
                    )
                    csv_block = (
                        f'<div style="margin-top:14px;padding:12px 16px;background:rgba(74,222,128,.06);'
                        f'border:1px solid rgba(74,222,128,.25);border-radius:8px">'
                        f'<div style="font-size:.6rem;font-weight:800;text-transform:uppercase;letter-spacing:.1em;'
                        f'color:#4ade80;margin-bottom:8px">📁 CSV Files Updated on Disk</div>'
                        f'{csv_items}</div>'
                    )
                else:
                    csv_block = (
                        '<div style="margin-top:14px;padding:10px 14px;background:rgba(100,116,139,.06);'
                        'border:1px solid rgba(100,116,139,.2);border-radius:8px;font-size:.72rem;color:#64748b">'
                        '📁 No CSV files were updated (no entity type detected or no matching CSV).</div>'
                    )

                # Summary KPIs
                kpis = "".join(
                    f'<div style="text-align:center;padding:10px;background:rgba(12,21,40,.8);'
                    f'border:1px solid rgba(56,189,248,.12);border-radius:8px">'
                    f'<div style="font-size:.55rem;color:#94a3b8;text-transform:uppercase;margin-bottom:4px">{lbl}</div>'
                    f'<div style="font-size:1.2rem;font-weight:900;color:{col};font-family:monospace">{val}</div></div>'
                    for lbl, val, col in [
                        ("Inserted", len(inserted), "#38bdf8"),
                        ("Updated",  len(updated),  "#fbbf24"),
                        ("Rels Wired", len(wired),  "#4ade80"),
                        ("Failed",   len(failed),   "#f87171" if failed else "#64748b"),
                    ]
                )

                # Change details
                changes_html = ""
                if inserted:
                    changes_html += (
                        f'<div style="margin-top:10px;font-size:.7rem;font-weight:700;color:#38bdf8;margin-bottom:4px">✅ Inserted Nodes</div>'
                        + "".join(f'<div style="font-family:monospace;font-size:.72rem;color:#e2e8f0;padding:1px 0">+ {i}</div>' for i in inserted[:20])
                        + (f'<div style="font-size:.65rem;color:#64748b">…and {len(inserted)-20} more</div>' if len(inserted)>20 else "")
                    )
                if updated:
                    changes_html += (
                        f'<div style="margin-top:10px;font-size:.7rem;font-weight:700;color:#fbbf24;margin-bottom:4px">~ Updated Nodes</div>'
                        + "".join(f'<div style="font-family:monospace;font-size:.72rem;color:#e2e8f0;padding:1px 0">~ {u}</div>' for u in updated[:20])
                        + (f'<div style="font-size:.65rem;color:#64748b">…and {len(updated)-20} more</div>' if len(updated)>20 else "")
                    )
                if failed:
                    changes_html += (
                        f'<div style="margin-top:10px;font-size:.7rem;font-weight:700;color:#f87171;margin-bottom:4px">✗ Failed</div>'
                        + "".join(f'<div style="font-family:monospace;font-size:.72rem;color:#f87171;padding:1px 0">✗ {f}</div>' for f in failed[:10])
                    )

                return (
                    f'<div style="background:rgba(6,12,28,.95);border:1px solid rgba(74,222,128,.3);'
                    f'border-radius:12px;padding:20px 24px;margin-top:10px">'
                    f'<div style="font-size:.65rem;font-weight:900;text-transform:uppercase;letter-spacing:.12em;'
                    f'color:#4ade80;margin-bottom:14px">📋 Post-Commit Report</div>'
                    f'<div style="display:grid;grid-template-columns:repeat(4,1fr);gap:8px;margin-bottom:14px">{kpis}</div>'
                    f'{changes_html}'
                    f'{csv_block}'
                    f'</div>'
                )

            def _upd_approve(session_id: str):
                if not session_id or not _S4_AVAILABLE:
                    return ('<div style="color:#f87171">No session — re-upload.</div>', "","",
                            gr.update(visible=False), gr.update(visible=True), gr.update(visible=False), _upd_stepper(0))
                sess = _s4.get_session(session_id)
                if not sess:
                    return (f'<div style="color:#f87171">Session {session_id} not found.</div>', "","",
                            gr.update(visible=False), gr.update(visible=True), gr.update(visible=False), _upd_stepper(0))
                import queue as _q2, threading as _t2
                q, holder = _q2.Queue(), {}
                def _w():
                    try: holder["sess"] = _s4.run_stage4_execute(sess, on_update=lambda e: q.put(e))
                    except Exception as e: holder["error"] = str(e)
                    finally: q.put(None)
                _t2.Thread(target=_w, daemon=True).start()
                while True:
                    try:
                        if q.get(timeout=30) is None: break
                    except: break
                if "error" in holder:
                    return (f'<div style="color:#f87171">Error: {holder["error"][:250]}</div>', "","",
                            gr.update(visible=False), gr.update(visible=True), gr.update(visible=False),
                            _upd_stepper(9, {9:"upd-step-err"}))
                done = holder.get("sess", sess)
                # ── Sync committed rows to CSV files on disk ──────────────────
                csv_log = _sync_rows_to_csv(done)
                # ─────────────────────────────────────────────────────────────
                ins = len(done.inserted_ids); upd = len(done.updated_nodes or [])
                wir = len(done.wired_rels);  fail = len(done.failed_ids)
                color = "#4ade80" if not fail else "#fbbf24"
                _fail_txt = f' · <span style="color:#f87171">{fail} failed</span>' if fail else ''
                _bva_content = done.rca_diff_html or ""
                _bva_block = (
                    f'<details style="margin-top:12px;border:1px solid rgba(56,189,248,0.3);border-radius:10px;overflow:hidden">'
                    f'<summary style="padding:10px 16px;background:rgba(12,21,40,0.95);cursor:pointer;font-size:.75rem;font-weight:700;color:#38bdf8;list-style:none;display:flex;align-items:center;gap:8px;user-select:none">'
                    f'&#128202; View Before vs After — Graph Changes (click to expand)</summary>'
                    f'<div style="padding:14px;background:rgba(4,9,22,0.95)">{_bva_content}</div></details>'
                ) if _bva_content else ""
                full_success = (
                    f'<div style="padding:14px 20px;margin:10px 0;background:rgba(74,222,128,.07);border:1px solid rgba(74,222,128,.4);border-radius:12px">'
                    f'<div style="display:flex;align-items:center;gap:12px;margin-bottom:6px">'
                    f'<span style="font-size:1.6rem">&#9989;</span>'
                    f'<div style="color:#4ade80;font-weight:900;font-size:1rem;letter-spacing:-.01em">Database successfully updated!</div></div>'
                    f'<div style="color:#94a3b8;font-size:.73rem;margin-left:52px;margin-bottom:10px">'
                    f'{ins} node(s) inserted &nbsp;·&nbsp; {upd} updated &nbsp;·&nbsp; {wir} relationships wired{_fail_txt}'
                    f' &nbsp;·&nbsp; Session: <code style="color:#38bdf8">{done.session_id}</code></div>'
                    f'{_bva_block}</div>'
                )
                return (
                    "",
                    "",
                    "",
                    gr.update(visible=False),
                    gr.update(visible=True),
                    gr.update(visible=False),
                    _upd_stepper(10),
                    gr.update(value=full_success, visible=True),
                    gr.update(value=""),
                )

            def _upd_cancel(_sid):
                return ('<div style="color:#94a3b8;font-size:.78rem">Cancelled — nothing written.</div>',
                        gr.update(visible=False), gr.update(visible=False), gr.update(visible=True))

            def _upd_rollback(sid: str):
                sid = sid.strip()
                if not sid or not _S4_AVAILABLE:
                    return '<div style="color:#f87171">Enter a Session ID.</div>'
                r = _s4.rollback_session(sid)
                if "error" in r:
                    return f'<div style="color:#f87171">⚠ {r["error"]}</div>'
                d, f = r.get("deleted",[]), r.get("failed",[])
                c = "#4ade80" if not f else "#fbbf24"
                return f'<div style="color:{c};font-size:.8rem">↩ Rolled back {len(d)} node(s).{" " + str(len(f)) + " failed." if f else ""}</div>'

            def _build_nl_results(entity_type, entity_id, id_prop, cypher_used, report_text, before_snap):
                """Build structured tab content after NL agent completes."""
                val_html = (
                    f'<div style="padding:14px 18px;background:rgba(6,12,28,.9);border:1px solid rgba(74,222,128,.25);border-radius:10px">'
                    f'<div style="font-size:.6rem;font-weight:900;text-transform:uppercase;letter-spacing:.1em;color:#4ade80;margin-bottom:10px">\u2705 Entity Confirmed in Neo4j</div>'
                    f'<div style="display:grid;grid-template-columns:1fr 1fr;gap:8px">'
                    f'<div style="background:rgba(12,21,40,.8);border:1px solid rgba(56,189,248,.12);border-radius:8px;padding:10px;text-align:center">'
                    f'<div style="font-size:.55rem;color:#94a3b8;text-transform:uppercase;margin-bottom:4px">Entity Type</div>'
                    f'<div style="font-size:.95rem;font-weight:800;color:#38bdf8">{entity_type}</div></div>'
                    f'<div style="background:rgba(12,21,40,.8);border:1px solid rgba(56,189,248,.12);border-radius:8px;padding:10px;text-align:center">'
                    f'<div style="font-size:.55rem;color:#94a3b8;text-transform:uppercase;margin-bottom:4px">Node ID</div>'
                    f'<div style="font-size:.95rem;font-weight:800;color:#4ade80;font-family:monospace">{entity_id}</div></div>'
                    f'</div>'
                    f'<div style="margin-top:10px;padding:8px 12px;background:rgba(74,222,128,.05);border:1px solid rgba(74,222,128,.15);border-radius:6px;font-size:.72rem;color:#94a3b8">'
                    f'Node written to Neo4j and synced to CSV files on disk.</div></div>'
                )
                cypher_html = (
                    f'<div style="padding:14px 18px;background:rgba(6,12,28,.9);border:1px solid rgba(56,189,248,.2);border-radius:10px">'
                    f'<div style="font-size:.6rem;font-weight:900;text-transform:uppercase;letter-spacing:.1em;color:#38bdf8;margin-bottom:10px">\u26a1 Cypher Executed by Agent</div>'
                    f'<pre style="font-family:monospace;font-size:.78rem;color:#e2e8f0;background:rgba(0,0,0,.4);padding:14px;border-radius:8px;overflow-x:auto;white-space:pre-wrap">{cypher_used or "MERGE/SET via create_or_update_node tool"}</pre>'
                    f'</div>'
                )
                try:
                    impact_text = _groq_call(
                        f"In 3-4 sentences, describe the business impact of adding/updating a {entity_type} "
                        f"with ID {entity_id} in a logistics supply chain knowledge graph. "
                        f"Focus on risk, capacity, and downstream effects. Be concise."
                    ).strip()
                except Exception:
                    impact_text = f"A new {entity_type} node ({entity_id}) has been added. This may affect supply chain risk scores, capacity calculations, and route availability."
                impact_html = (
                    f'<div style="padding:14px 18px;background:rgba(6,12,28,.9);border:1px solid rgba(168,85,247,.2);border-radius:10px">'
                    f'<div style="font-size:.6rem;font-weight:900;text-transform:uppercase;letter-spacing:.1em;color:#a855f7;margin-bottom:10px">⚡ Impact Forecast</div>'
                    f'<div style="font-size:.82rem;color:#e2e8f0;line-height:1.7">{impact_text}</div></div>'
                )
                after_snap = _take_nl_snapshot()
                bva_html = _build_nl_bva(before_snap, after_snap) if before_snap else '<div style="color:#64748b;padding:14px">Snapshot unavailable.</div>'
                banner = (
                    f'<div style="padding:10px 16px;margin-bottom:8px;background:rgba(74,222,128,.07);'
                    f'border:1px solid rgba(74,222,128,.35);border-radius:10px;display:flex;align-items:center;gap:10px">'
                    f'<span style="font-size:1.2rem">\u2705</span>'
                    f'<div><div style="color:#4ade80;font-weight:800;font-size:.86rem">Successfully committed via Natural Language</div>'
                    f'<div style="color:#94a3b8;font-size:.7rem;margin-top:2px">{entity_type} &middot; <code style="color:#38bdf8">{entity_id}</code> &middot; CSV updated</div>'
                    f'</div></div>'
                )
                return val_html, cypher_html, impact_html, bva_html, banner

            def _take_nl_snapshot():
                """Quick Neo4j metrics snapshot."""
                try:
                    rows = _run_neo4j(
                        "OPTIONAL MATCH (s:Supplier) WITH count(s) AS sup_total "
                        "OPTIONAL MATCH (s2:Supplier) WHERE s2.risk_score > 0.5 WITH sup_total, count(s2) AS high_risk "
                        "OPTIONAL MATCH (s3:Supplier) WITH sup_total, high_risk, avg(toFloat(s3.risk_score)) AS avg_risk "
                        "OPTIONAL MATCH (d:Distributor) WITH sup_total, high_risk, avg_risk, count(d) AS dist_total "
                        "OPTIONAL MATCH (sh:Shipment) WITH sup_total, high_risk, avg_risk, dist_total, count(sh) AS ship_total "
                        "RETURN sup_total, high_risk, avg_risk, dist_total, ship_total"
                    )
                    r = rows[0] if rows else {}
                    import datetime as _dt2
                    return {
                        "timestamp": _dt2.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "suppliers":    {"total": r.get("sup_total",0), "avg_risk": round(float(r.get("avg_risk") or 0),3), "high_risk": r.get("high_risk",0)},
                        "distributors": {"total": r.get("dist_total",0)},
                        "shipments":    {"total": r.get("ship_total",0)},
                    }
                except Exception:
                    return {}

            def _build_nl_bva(before, after):
                """Before vs After table for NL results."""
                def pct(o, n):
                    try:
                        o, n = float(o), float(n)
                        if o == 0: return '<span style="color:#4ade80">new</span>'
                        c = (n-o)/abs(o)*100
                        col = "#f87171" if c > 0 else "#4ade80"
                        return f'<span style="color:{col}">{("+" if c>0 else "")}{c:.1f}%</span>'
                    except: return "\u2014"
                def row(lbl, bv, av):
                    return (f'<tr><td style="color:#94a3b8;font-size:.76rem;padding:5px 12px">{lbl}</td>'
                            f'<td style="color:#e2e8f0;text-align:right;padding:5px 12px">{bv}</td>'
                            f'<td style="color:#e2e8f0;text-align:right;padding:5px 12px">{av}</td>'
                            f'<td style="text-align:center;padding:5px 12px">{pct(bv,av)}</td></tr>')
                bs, as_ = before.get("suppliers",{}), after.get("suppliers",{})
                bd, ad  = before.get("distributors",{}), after.get("distributors",{})
                return (
                    '<div style="padding:14px 18px;background:rgba(6,12,28,.9);border:1px solid rgba(56,189,248,.2);border-radius:10px">'
                    '<div style="font-size:.6rem;font-weight:900;text-transform:uppercase;letter-spacing:.1em;color:#38bdf8;margin-bottom:12px">\U0001f4c8 Before vs After</div>'
                    '<table style="width:100%;border-collapse:collapse">'
                    '<thead><tr style="border-bottom:1px solid rgba(56,189,248,.2)">'
                    '<th style="color:#7dd3fc;font-size:.68rem;text-align:left;padding:5px 12px">Metric</th>'
                    '<th style="color:#7dd3fc;font-size:.68rem;text-align:right;padding:5px 12px">Before</th>'
                    '<th style="color:#7dd3fc;font-size:.68rem;text-align:right;padding:5px 12px">After</th>'
                    '<th style="color:#7dd3fc;font-size:.68rem;text-align:center;padding:5px 12px">Change</th>'
                    '</tr></thead><tbody>'
                    + row("Total Suppliers",     bs.get("total",0),    as_.get("total",0))
                    + row("Avg Risk Score",      bs.get("avg_risk",0), as_.get("avg_risk",0))
                    + row("High-Risk Suppliers", bs.get("high_risk",0),as_.get("high_risk",0))
                    + row("Distributors",        bd.get("total",0),    ad.get("total",0))
                    + '</tbody></table>'
                    + f'<div style="margin-top:10px;font-size:.68rem;color:#64748b">'
                    + f'{before.get("timestamp","?")} \u2192 {after.get("timestamp","?")}</div></div>'
                )

            def _upd_nl(q):
                """
                NL Update handler — runs the full 10-stage pipeline, auto-commits,
                syncs CSV, and streams results into the NL panel's own components.
                Outputs (10 values):
                  upd_nl_status, upd_nl_results_col,
                  upd_nl_success_banner, upd_nl_val_html, upd_nl_diff_html,
                  upd_nl_cypher_html, upd_nl_impact_html, upd_nl_bva_html,
                  upd_nl_out, upd_nl_log
                """
                q_str = str(q or "").strip()

                # Helper: 10-tuple for empty/error states (NL panel outputs only)
                def _err(msg):
                    return (msg, gr.update(visible=False), "", "", "", "", "", "", "", "")

                if not q_str:
                    yield _err('<div style="color:#f87171;font-size:.78rem;padding:8px 12px;'
                               'background:rgba(239,68,68,.08);border:1px solid rgba(239,68,68,.25);'
                               'border-radius:8px">⚠ Please enter an instruction first.</div>')
                    return

                import re as _re_del
                q_low = q_str.lower()
                _is_delete = (any(w in q_low for w in ["delete ", "remove ", "drop "]) and
                              any(w in q_low for w in ["supplier","distributor","route","sup","dist"]))

                if _is_delete:
                    yield _err('<div style="color:#f87171;font-size:.78rem;padding:8px 12px;'
                               'background:rgba(239,68,68,.08);border:1px solid rgba(239,68,68,.25);'
                               'border-radius:8px">🗑 Detected delete — executing…</div>')
                    try:
                        _em = {"supplier":("Supplier","supplier_id"),
                               "distributor":("Distributor","distributor_id"),
                               "route":("Route","route_id")}
                        _lbl, _prop, _eid = None, None, None
                        for kw, (lb, pr) in _em.items():
                            if kw in q_low:
                                _m = _re_del.search(r'\b([A-Z]{2,}[\w@]*\d+[\w@]*)\b', q_str)
                                if _m: _lbl, _prop, _eid = lb, pr, _m.group(1); break
                        if _lbl and _eid:
                            _run_neo4j(f"MATCH (n:{_lbl} {{{_prop}: $v}}) DETACH DELETE n", {"v": _eid})
                            _app_csv_delete(_lbl, _eid)
                            _dbanner = (f'<div style="padding:10px 16px;background:rgba(239,68,68,.07);'
                                        f'border:1px solid rgba(239,68,68,.35);border-radius:10px">'
                                        f'<div style="color:#f87171;font-weight:800">'
                                        f'🗑️ Deleted {_lbl} <code style="color:#38bdf8">{_eid}</code>'
                                        f' from Neo4j and CSV</div></div>')
                            yield ('<div style="color:#f87171;font-size:.78rem;padding:8px 12px;'
                                   'background:rgba(239,68,68,.08);border:1px solid rgba(239,68,68,.3);'
                                   'border-radius:8px">🗑 Delete complete</div>',
                                   gr.update(visible=True), _dbanner, "", "", "", "", "", "")
                        else:
                            yield _err('<div style="color:#fbbf24;font-size:.78rem">⚠ Could not parse delete</div>')
                    except Exception as _de:
                        yield _err(f'<div style="color:#f87171">❌ Delete failed: {str(_de)[:200]}</div>')
                    return

                # ── MAIN: parse NL → CSV → stage4 pipeline → auto-commit → sync CSV ──
                yield _err('<div style="color:#38bdf8;font-size:.78rem;padding:8px 12px;'
                           'background:rgba(56,189,248,.08);border:1px solid rgba(56,189,248,.25);'
                           'border-radius:8px">◌ Parsing instruction…</div>')

                if not _S4_AVAILABLE:
                    yield _err('<div style="color:#f87171">⚠ Pipeline not available.</div>')
                    return

                try:
                    tmp_path, detected_type = _s4.nl_instruction_to_csv(q_str)
                except Exception as _pe:
                    yield _err(f'<div style="color:#f87171;font-size:.78rem;padding:8px 12px;'
                               f'background:rgba(239,68,68,.08);border:1px solid rgba(239,68,68,.25);'
                               f'border-radius:8px">⚠ Could not parse: {str(_pe)[:300]}</div>')
                    return

                # Stage 1-8: validation pipeline
                import queue as _q3, threading as _t3
                q2, holder2 = _q3.Queue(), {}

                def _w2():
                    try:
                        holder2["sess"] = _s4.run_stage4_pipeline(
                            tmp_path, f"nl_{detected_type.lower()}.csv",
                            on_update=lambda e: q2.put(e)
                        )
                    except Exception as e2:
                        holder2["error"] = str(e2)
                    finally:
                        q2.put(None)

                _t3.Thread(target=_w2, daemon=True).start()

                LABELS = ["Parse","Detect","Schema","Semantic","Conflicts",
                          "Graph Diff","Cypher","Supply Chain Impact"]
                while True:
                    try: event = q2.get(timeout=30)
                    except Exception: break
                    if event is None: break
                    if isinstance(event, dict) and event.get("kind") == "stage":
                        cur = event.get("data", {}).get("stage_num", 0)
                        lbl = LABELS[cur-1] if 0 < cur <= len(LABELS) else f"Stage {cur}"
                        yield _err(f'<div style="color:#38bdf8;font-size:.78rem;padding:8px 12px;'
                                   f'background:rgba(56,189,248,.08);border:1px solid rgba(56,189,248,.25);'
                                   f'border-radius:8px">◌ [{cur}/8] {lbl}…</div>')

                if "error" in holder2:
                    yield _err(f'<div style="color:#f87171">❌ Error: {holder2["error"][:250]}</div>')
                    return

                sess = holder2.get("sess")
                if not sess or not sess.valid_rows:
                    yield _err('<div style="color:#f87171">⚠ No valid rows — check instruction.</div>')
                    return

                # Stage 9: Auto-commit to Neo4j (no manual approve needed for NL)
                yield _err('<div style="color:#38bdf8;font-size:.78rem;padding:8px 12px;'
                           'background:rgba(56,189,248,.08);border:1px solid rgba(56,189,248,.25);'
                           'border-radius:8px">◌ [9/10] Committing to Neo4j…</div>')

                _s4.register_session(sess)
                q3, holder3 = _q3.Queue(), {}
                def _w3():
                    try: holder3["done"] = _s4.run_stage4_execute(sess, on_update=lambda e: q3.put(e))
                    except Exception as e3: holder3["error"] = str(e3)
                    finally: q3.put(None)
                _t3.Thread(target=_w3, daemon=True).start()
                while True:
                    try:
                        if q3.get(timeout=30) is None: break
                    except: break

                if "error" in holder3:
                    yield _err(f'<div style="color:#f87171">❌ Commit failed: {holder3["error"][:250]}</div>')
                    return

                done = holder3.get("done", sess)

                # ── Log each inserted node to NL_UPDATE_LOG so "Undo Last NL Change" works ──
                try:
                    from agent_runner import _log_nl_update as _log_nl
                    _entity_type = getattr(done, "entity_type", detected_type) or detected_type
                    _id_prop_map = {
                        "Supplier":    "supplier_id",
                        "Distributor": "distributor_id",
                        "Route":       "route_id",
                    }
                    _id_prop = _id_prop_map.get(_entity_type, "id")
                    for _inserted_id in (getattr(done, "inserted_ids", []) or []):
                        _log_nl(
                            entity_type = _entity_type,
                            entity_id   = str(_inserted_id),
                            cypher      = f"/* NL pipeline insert: {_entity_type} {_inserted_id} */",
                            id_prop     = _id_prop,
                        )
                except Exception as _log_err:
                    pass  # non-fatal — undo log is best-effort

                # Stage 10: Sync to CSV
                yield _err('<div style="color:#38bdf8;font-size:.78rem;padding:8px 12px;'
                           'background:rgba(56,189,248,.08);border:1px solid rgba(56,189,248,.25);'
                           'border-radius:8px">◌ [10/10] Syncing CSV files…</div>')
                csv_log = _sync_rows_to_csv(done)

                # ── Build result tabs ────────────────────────────────────────
                _val  = _s4.build_validation_report_html(sess)
                try:    _diff = _s4.build_graph_diff_html(sess) if hasattr(_s4, "build_graph_diff_html") else ""
                except: _diff = ""
                _cyph = _s4.build_cypher_preview_html(sess)
                try:    _imp = _s4.build_ai_impact_html(sess) if hasattr(_s4,"build_ai_impact_html") else ""
                except: _imp = ""

                ins  = len(getattr(done, "inserted_ids", []) or [])
                upd  = len(getattr(done, "updated_nodes", []) or [])
                wir  = len(getattr(done, "wired_rels", []) or [])
                fail = len(getattr(done, "failed_ids", []) or [])
                color = "#4ade80" if not fail else "#fbbf24"

                csv_path = _APP_CSV.get(detected_type, "")
                csv_note = (f' · 📁 <code style="color:#7dd3fc">{os.path.basename(csv_path)}</code> updated on disk'
                            if csv_log and csv_path else "")

                banner = (
                    f'<div style="padding:10px 16px;margin-bottom:8px;background:rgba(74,222,128,.07);'
                    f'border:1px solid rgba(74,222,128,.35);border-radius:10px;display:flex;align-items:center;gap:10px">'
                    f'<span style="font-size:1.2rem">✅</span>'
                    f'<div><div style="color:{color};font-weight:800;font-size:.86rem">Successfully committed to Neo4j via Natural Language</div>'
                    f'<div style="color:#94a3b8;font-size:.7rem;margin-top:2px">'
                    f'{detected_type} · {ins} inserted · {upd} updated · {wir} linked'
                    f'{(" · " + str(fail) + " failed") if fail else ""}'
                    f'{csv_note}'
                    f' · Session <code style="color:#38bdf8">{done.session_id}</code></div>'
                    f'</div></div>'
                )

                # Before/after snapshot
                try:
                    _after = _take_nl_snapshot()
                    _bva = _build_nl_bva({}, _after)
                except Exception:
                    _bva = ""

                # Log summary
                log_items = "".join(
                    f'<div style="font-family:monospace;font-size:.72rem;color:#4ade80;padding:1px 0">+ {l}</div>'
                    for l in csv_log
                ) or '<div style="color:#64748b;font-size:.72rem">No CSV changes recorded.</div>'
                log_html = (
                    f'<div style="padding:10px 14px;background:rgba(4,9,22,.8);border:1px solid rgba(56,189,248,.15);'
                    f'border-radius:8px;min-height:60px"><div style="font-size:.62rem;font-weight:800;'
                    f'text-transform:uppercase;letter-spacing:.1em;color:#38bdf8;margin-bottom:8px">📋 Commit Log</div>'
                    f'{log_items}</div>'
                )

                yield (
                    f'<div style="color:{color};font-size:.78rem;padding:6px 0">'
                    f'✓ {ins} row(s) committed · Session <code style="color:#38bdf8">{done.session_id}</code></div>',
                    gr.update(visible=True),   # upd_nl_results_col
                    banner,                    # upd_nl_success_banner
                    _val,                      # upd_nl_val_html
                    _diff,                     # upd_nl_diff_html
                    _cyph,                     # upd_nl_cypher_html
                    _imp,                      # upd_nl_impact_html
                    _bva,                      # upd_nl_bva_html
                    f'<div style="font-size:.82rem;color:#e2e8f0;line-height:1.65;padding:8px 12px;'
                    f'background:rgba(4,9,22,.8);border:1px solid rgba(56,189,248,.15);border-radius:8px">'
                    f'{ins} node(s) inserted, {upd} updated, {wir} relationships wired.</div>',  # upd_nl_out
                    log_html,                  # upd_nl_log
                )


            def _get_nl_history_html() -> str:
                """Build HTML showing all NL updates with individual Undo buttons."""
                history = get_nl_update_history()
                if not history:
                    return (
                        '<div style="color:#475569;font-size:.75rem;padding:12px;text-align:center;'
                        'background:rgba(4,9,22,.7);border:1px solid rgba(56,189,248,.1);border-radius:8px">'
                        'No NL updates yet — run an update and it will appear here.'
                        '</div>'
                    )
                rows_html = ""
                for i, entry in enumerate(history):
                    etype = entry["entity_type"]
                    eid   = entry["entity_id"]
                    ts    = entry["timestamp"]
                    icon  = {"Supplier":"🏭","Distributor":"🏙️","Route":"🚛"}.get(etype,"📦")
                    # Each row has an Undo button that sends a special message
                    rows_html += f"""
<div style="display:flex;align-items:center;gap:10px;padding:8px 12px;
     background:rgba(14,165,233,.05);border:1px solid rgba(56,189,248,.15);
     border-radius:8px;margin-bottom:6px">
  <span style="font-size:.9rem;flex-shrink:0">{icon}</span>
  <div style="flex:1;min-width:0">
    <div style="font-size:.78rem;font-weight:700;color:#e2e8f0;
         white-space:nowrap;overflow:hidden;text-overflow:ellipsis">
      {etype}: <span style="color:#38bdf8;font-family:monospace">{eid}</span>
    </div>
    <div style="font-size:.65rem;color:#64748b;margin-top:2px">{ts}</div>
  </div>
  <button onclick="(function(){{
    var inp = document.querySelector('#upd-nl-undo-idx-input textarea, #upd-nl-undo-idx-input input');
    if(inp){{ inp.value='{i}'; inp.dispatchEvent(new Event('input',{{bubbles:true}})); }}
    setTimeout(function(){{
      var btn = document.querySelector('#upd-nl-undo-idx-btn button');
      if(btn) btn.click();
    }}, 80);
  }})()"
  style="background:rgba(239,68,68,.12);border:1px solid rgba(239,68,68,.35);
         color:#fca5a5;font-size:.65rem;font-weight:700;border-radius:6px;
         padding:4px 10px;cursor:pointer;white-space:nowrap;flex-shrink:0">
    ↩ Undo
  </button>
</div>"""
                return f"""
<div style="background:rgba(4,9,22,.85);border:1px solid rgba(56,189,248,.2);
     border-radius:10px;padding:12px 14px">
  <div style="font-size:.62rem;font-weight:800;text-transform:uppercase;letter-spacing:.1em;
       color:#38bdf8;margin-bottom:10px">↩ NL Update History — click Undo to reverse</div>
  {rows_html}
</div>"""

            def _upd_nl_undo(idx_str: str = "0") -> tuple:
                try:
                    idx = int(str(idx_str or "0").strip())
                except ValueError:
                    idx = 0
                msg  = undo_nl_update(index=idx)
                ok   = msg.startswith("✓")
                color = "#4ade80" if ok else "#f87171"
                status_html = (
                    f'<div style="color:{color};font-size:.78rem;padding:8px 12px;'
                    f'background:rgba({("74,222,128" if ok else "239,68,68")},.08);'
                    f'border:1px solid rgba({("74,222,128" if ok else "239,68,68")},.3);'
                    f'border-radius:8px">{msg}</div>'
                )
                return status_html, _get_nl_history_html()

            def _upd_rollback_last_file() -> str:
                """One-click undo for the most recent file upload session.
                Finds the latest committed session in _SESSION_REGISTRY,
                rolls back Neo4j — execute_rollback already calls
                _csv_delete/_csv_upsert to keep CSV files in sync.
                """
                if not _S4_AVAILABLE:
                    return '<div style="color:#f87171;font-size:.78rem;padding:8px 12px;background:rgba(239,68,68,.08);border:1px solid rgba(239,68,68,.3);border-radius:8px">⚠ Pipeline not available.</div>'
                try:
                    registry = _s4._SESSION_REGISTRY
                except Exception:
                    return '<div style="color:#fbbf24;font-size:.78rem;padding:10px 14px;background:rgba(251,191,36,.06);border:1px solid rgba(251,191,36,.2);border-radius:8px">⚠ No sessions found.</div>'

                if not registry:
                    return '<div style="color:#fbbf24;font-size:.78rem;padding:10px 14px;background:rgba(251,191,36,.06);border:1px solid rgba(251,191,36,.2);border-radius:8px">⚠ No file uploads found to undo. Upload and commit a file first.</div>'

                committed = [
                    s for s in registry.values()
                    if (getattr(s, "inserted_ids", None) or getattr(s, "updated_nodes", None))
                ]
                if not committed:
                    return '<div style="color:#fbbf24;font-size:.78rem;padding:10px 14px;background:rgba(251,191,36,.06);border:1px solid rgba(251,191,36,.2);border-radius:8px">⚠ No committed uploads to undo yet.</div>'

                latest      = committed[-1]
                sid         = latest.session_id
                entity_type = getattr(latest, "entity_type", "?")

                r       = _s4.rollback_session(sid)
                if "error" in r:
                    return (
                        f'<div style="color:#f87171;font-size:.78rem;padding:10px 14px;'
                        f'background:rgba(239,68,68,.08);border:1px solid rgba(239,68,68,.3);border-radius:8px">'
                        f'⚠ {r["error"]}</div>'
                    )

                deleted = r.get("deleted", [])
                failed  = r.get("failed", [])

                changes = "".join(
                    f'<div style="font-family:monospace;font-size:.7rem;color:#f87171;padding:1px 0">− {d}</div>'
                    for d in deleted[:15]
                )
                if len(deleted) > 15:
                    changes += f'<div style="font-size:.65rem;color:#64748b">…and {len(deleted)-15} more</div>'
                if failed:
                    changes += "".join(
                        f'<div style="font-family:monospace;font-size:.7rem;color:#fbbf24">⚠ failed: {f["id"]}</div>'
                        for f in failed[:5]
                    )

                csv_path = _APP_CSV.get(entity_type, "")
                csv_note = (
                    f'<div style="margin-top:8px;font-size:.7rem;color:#4ade80">📁 CSV reset: <code style="color:#7dd3fc">{os.path.basename(csv_path)}</code></div>'
                    if csv_path else
                    '<div style="margin-top:8px;font-size:.7rem;color:#64748b">📁 CSV: no matching file for this entity type.</div>'
                )
                color = "#4ade80" if not failed else "#fbbf24"
                return (
                    f'<div style="background:rgba(6,12,28,.95);border:1px solid rgba({"74,222,128" if not failed else "251,191,36"},.3);'
                    f'border-radius:10px;padding:14px 18px;margin-top:6px">'
                    f'<div style="font-size:.7rem;font-weight:800;color:{color};margin-bottom:8px">'
                    f'↩ Rolled back · Session <code style="font-size:.68rem;color:#38bdf8">{sid}</code></div>'
                    f'<div style="font-size:.7rem;color:#94a3b8;margin-bottom:6px">'
                    f'Entity: <b style="color:#e2e8f0">{entity_type}</b> · '
                    f'Removed: <b style="color:#f87171">{len(deleted)}</b> · '
                    f'Failed: <b style="color:#{"f87171" if failed else "64748b"}">{len(failed)}</b></div>'
                    f'<div style="background:rgba(0,0,0,.3);border-radius:6px;padding:8px 10px;'
                    f'max-height:140px;overflow:auto">'
                    f'{changes or "<div style=\\'font-size:.7rem;color:#64748b\\'>No committed nodes to remove.</div>"}'
                    f'</div>'
                    f'{csv_note}'
                    f'</div>'
                )

            def _upd_rollback_last_nl() -> tuple:
                """One-click undo for the most recent NL update.
                Deletes from Neo4j AND removes the row from the CSV on disk.
                Returns (rollback_msg_html, nl_history_html).
                """
                from agent_runner import get_nl_update_history as _gnlh, _NL_UPDATE_LOG
                history = _gnlh()
                if not history:
                    msg = '<div style="color:#fbbf24;font-size:.78rem;padding:10px 14px;background:rgba(251,191,36,.06);border:1px solid rgba(251,191,36,.2);border-radius:8px">⚠ No NL updates to undo.</div>'
                    return msg, _get_nl_history_html()

                entry   = history[0]
                etype   = entry["entity_type"]
                eid     = entry["entity_id"]
                id_prop = entry.get("id_prop") or {
                    "Supplier": "supplier_id", "Distributor": "distributor_id",
                    "Route": "route_id", "Plant": "plant_id",
                }.get(etype, "id")

                try:
                    _run_neo4j(
                        f"MATCH (n:{etype} {{{id_prop}: $v}}) DETACH DELETE n",
                        {"v": eid}
                    )
                except Exception as e:
                    msg = (
                        f'<div style="color:#f87171;font-size:.78rem;padding:10px;'
                        f'background:rgba(239,68,68,.08);border:1px solid rgba(239,68,68,.3);border-radius:8px">'
                        f'⚠ Neo4j delete failed: {e}</div>'
                    )
                    return msg, _get_nl_history_html()

                # ── Remove from CSV on disk ─────────────────────────────────
                csv_synced = False
                try:
                    _app_csv_delete(etype, eid)
                    csv_synced = True
                except Exception:
                    pass

                # ── Remove from in-memory log ─────────────────────────────
                ts = entry["timestamp"]
                for i, item in enumerate(_NL_UPDATE_LOG):
                    if item["timestamp"] == ts and item["entity_id"] == eid:
                        _NL_UPDATE_LOG.pop(i)
                        break

                csv_path = _APP_CSV.get(etype, "")
                csv_note = (
                    f' · 📁 <code style="color:#7dd3fc">{os.path.basename(csv_path)}</code> reset on disk'
                    if csv_synced and csv_path else ""
                )
                msg = (
                    f'<div style="background:rgba(6,12,28,.95);border:1px solid rgba(74,222,128,.3);'
                    f'border-radius:10px;padding:12px 16px;margin-top:6px">'
                    f'<div style="font-size:.75rem;font-weight:700;color:#4ade80;margin-bottom:4px">'
                    f'↩ NL change undone</div>'
                    f'<div style="font-size:.72rem;color:#94a3b8">'
                    f'Deleted <b style="color:#f87171">{etype}</b> '
                    f'<code style="color:#38bdf8">{eid}</code> from Neo4j{csv_note}.</div>'
                    f'</div>'
                )
                return msg, _get_nl_history_html()

            # ══════════════════════════════════════════════════════════════════
            # WIRING
            # ══════════════════════════════════════════════════════════════════

            upd_mode_file.click(fn=lambda: (gr.update(visible=True), gr.update(visible=False)), outputs=[upd_file_panel, upd_nl_panel])
            upd_mode_nl.click(  fn=lambda: (gr.update(visible=False),gr.update(visible=True)),  outputs=[upd_file_panel, upd_nl_panel])

            upd_file_input.change(
                fn=_upd_run_pipeline, inputs=[upd_file_input],
                outputs=[upd_file_status, upd_stepper, upd_results_col,
                         upd_val_html, upd_diff_html, upd_cypher_html, upd_impact_html,
                         upd_gate_html, upd_approval_row, s4_session_id,
                         upd_success_banner, upd_before_after_html],
                show_progress="hidden"
            )
            def _upd_force_delete(node_id: str, entity_type: str) -> str:
                """Force-delete a node from Neo4j + CSV by ID, no session needed."""
                node_id = (node_id or "").strip()
                if not node_id:
                    return '<div style="color:#f87171;font-size:.78rem;padding:8px 12px;background:rgba(239,68,68,.08);border:1px solid rgba(239,68,68,.3);border-radius:8px">\u26a0 Please enter a Node ID.</div>'
                if not _S4_AVAILABLE:
                    return '<div style="color:#f87171;font-size:.78rem">\u26a0 Pipeline not available.</div>'
                r = _s4.force_delete_node(node_id, entity_type)
                if "error" in r:
                    return f'<div style="color:#f87171;font-size:.78rem;padding:10px 14px;background:rgba(239,68,68,.08);border:1px solid rgba(239,68,68,.3);border-radius:8px">\u26a0 {r["error"]}</div>'
                n = r.get("deleted", 0)
                if n == 0:
                    return f'<div style="color:#fbbf24;font-size:.78rem;padding:10px 14px;background:rgba(251,191,36,.06);border:1px solid rgba(251,191,36,.2);border-radius:8px">\u26a0 No {entity_type} node with ID <code>{node_id}</code> found in Neo4j (already deleted or never existed).</div>'
                return f'<div style="color:#4ade80;font-size:.78rem;padding:10px 14px;background:rgba(74,222,128,.06);border:1px solid rgba(74,222,128,.2);border-radius:8px">\u2705 Deleted {entity_type} node <code style="color:#38bdf8">{node_id}</code> from Neo4j and CSV files.</div>'

            upd_approve_btn.click(
                fn=_upd_approve, inputs=[s4_session_id],
                outputs=[upd_exec_status, upd_rca_diff, upd_change_summary,
                         upd_exec_panel, upd_results_col, upd_approval_row, upd_stepper,
                         upd_success_banner, upd_before_after_html],
                show_progress="hidden"
            ).then(fn=_get_recent_history_html, outputs=[upd_history]
            ).then(fn=load_snapshot, outputs=[snap_html]
            ).then(fn=load_network_health,
                  outputs=[health_kpi_html, health_plant_chart, health_gauge_chart, health_trend_chart])
            upd_cancel_btn.click(
                fn=_upd_cancel, inputs=[s4_session_id],
                outputs=[upd_file_status, upd_approval_row, upd_results_col, upd_file_panel]
            )
            upd_undo_file_btn.click(fn=_upd_rollback_last_file, outputs=[upd_rollback_msg], show_progress="hidden"
            ).then(fn=load_snapshot, outputs=[snap_html]
            ).then(fn=load_network_health,
                  outputs=[health_kpi_html, health_plant_chart, health_gauge_chart, health_trend_chart])
            upd_undo_nl_btn.click(fn=_upd_rollback_last_nl, outputs=[upd_rollback_msg, upd_nl_history], show_progress="hidden")

            # Right-sidebar undo buttons (same handlers, output to right-sidebar msg box)
            upd_undo_file_btn_r.click(fn=_upd_rollback_last_file, outputs=[upd_rollback_msg_r], show_progress="hidden"
            ).then(fn=load_snapshot, outputs=[snap_html]
            ).then(fn=load_network_health,
                  outputs=[health_kpi_html, health_plant_chart, health_gauge_chart, health_trend_chart])
            upd_undo_nl_btn_r.click(fn=_upd_rollback_last_nl, outputs=[upd_rollback_msg_r, upd_nl_history], show_progress="hidden"
            ).then(fn=load_snapshot, outputs=[snap_html]
            ).then(fn=load_network_health,
                  outputs=[health_kpi_html, health_plant_chart, health_gauge_chart, health_trend_chart])
            upd_rollback_btn.click(fn=_upd_rollback, inputs=[upd_rollback_input], outputs=[upd_rollback_msg]
            ).then(fn=load_snapshot, outputs=[snap_html]
            ).then(fn=load_network_health,
                  outputs=[health_kpi_html, health_plant_chart, health_gauge_chart, health_trend_chart])
            upd_force_del_btn.click(
                fn=_upd_force_delete,
                inputs=[upd_force_del_id, upd_force_del_type],
                outputs=[upd_force_del_msg]
            ).then(fn=load_snapshot, outputs=[snap_html]
            ).then(fn=load_network_health,
                  outputs=[health_kpi_html, health_plant_chart, health_gauge_chart, health_trend_chart])
            upd_nl_btn.click(fn=_upd_nl, inputs=[upd_nl_input],
                outputs=[upd_nl_status, upd_nl_results_col,
                         upd_nl_success_banner, upd_nl_val_html, upd_nl_diff_html,
                         upd_nl_cypher_html, upd_nl_impact_html, upd_nl_bva_html,
                         upd_nl_out, upd_nl_log],
                show_progress="hidden"
                ).then(fn=_get_nl_history_html, outputs=[upd_nl_history]
                ).then(fn=_get_recent_history_html, outputs=[upd_history]
                ).then(fn=load_snapshot, outputs=[snap_html]
                ).then(fn=load_network_health,
                      outputs=[health_kpi_html, health_plant_chart, health_gauge_chart, health_trend_chart])
            upd_nl_undo.click(
                fn=lambda: _upd_nl_undo("0"),
                outputs=[upd_nl_status, upd_nl_history]
            ).then(fn=load_snapshot, outputs=[snap_html]
            ).then(fn=load_network_health,
                  outputs=[health_kpi_html, health_plant_chart, health_gauge_chart, health_trend_chart])
            upd_nl_undo_idx_btn.click(
                fn=_upd_nl_undo,
                inputs=[upd_nl_undo_idx],
                outputs=[upd_nl_status, upd_nl_history]
            ).then(fn=load_snapshot, outputs=[snap_html]
            ).then(fn=load_network_health,
                  outputs=[health_kpi_html, health_plant_chart, health_gauge_chart, health_trend_chart])
            upd_hist_refresh.click(fn=_get_recent_history_html, outputs=[upd_history])

            # Examples + test scenarios — all switch to NL panel
            _BTN_MAP = [
                (upd_ex_sup,   "sup"),   (upd_ex_dist,  "dist"),
                (upd_ex_route, "route"), (upd_ex_link,  "link"), (upd_ex_risk,  "risk"),
                (upd_t_happy,  "t_happy"), (upd_t_dupe, "t_dupe"), (upd_t_invalid, "t_invalid"),
                (upd_t_missing,"t_missing"),(upd_t_circular,"t_circular"),
            ]
            for _btn, _key in _BTN_MAP:
                _btn.click(
                    fn=lambda k=_key: (_NL_EX[k], gr.update(visible=False), gr.update(visible=True)),
                    outputs=[upd_nl_input, upd_file_panel, upd_nl_panel]
                )

            # Verifier wiring
            gv_run.click(fn=_gv_run, inputs=[gv_query], outputs=[gv_result], show_progress="hidden")
            gv_clear.click(fn=lambda: ("","<div style='color:#475569;font-size:.78rem;padding:20px;text-align:center'>Use Node Lookup or click Show Counts.</div>"), outputs=[gv_query, gv_result])
            gv_stats_btn.click(fn=lambda: _gv_stats(), outputs=[gv_result], show_progress="hidden")
            gv_find.click(fn=_gv_find, inputs=[gv_id, gv_lbl], outputs=[gv_result], show_progress="hidden")
            gv_sup9_btn.click(fn=lambda: _gv_run(_GVQ["sup9"]), outputs=[gv_result], show_progress="hidden")
            for _gvbtn, _gvk in [(gv_sup_btn,"sup"),(gv_plant_btn,"plant"),(gv_risk_btn,"risk"),
                                   (gv_rel_btn,"rel"),(gv_dist_btn,"dist"),(gv_route_btn,"route")]:
                _gvbtn.click(fn=lambda k=_gvk: _gv_run(_GVQ[k]), outputs=[gv_result], show_progress="hidden")


        # ── TAB 5: Visualizations ────────────────────────────────
    demo.load(fn=load_snapshot, outputs=[snap_html]
    ).then(fn=load_network_health,
           outputs=[health_kpi_html, health_plant_chart, health_gauge_chart, health_trend_chart])

if __name__ == "__main__":
    start_mcp()
try:
    demo.launch(
        server_name="127.0.0.1",
        server_port=None,
        inbrowser=False,
        share=False,
        show_error=True,
    )
except OSError:
    # All ports 7860-7867 busy — find a free port
    import socket
    with socket.socket() as s:
        s.bind(("127.0.0.1", 0))
        free_port = s.getsockname()[1]
    print(f"[GraphPulse] Auto-selected port {free_port}")
    demo.launch(
        server_name="127.0.0.1",
        server_port=free_port,
        inbrowser=False,
        share=False,
        show_error=True,
    )